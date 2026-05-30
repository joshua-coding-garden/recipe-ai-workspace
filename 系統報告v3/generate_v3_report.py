"""
generate_v3_report.py — RecipeAI AI營養顧問 Tool Calling 技術報告 v3
教學法：抽換概念 → 專有名詞講解 → 深入講解 → 實際例子
"""
import os, pathlib, io
import cairosvg
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUT_DIR = pathlib.Path(__file__).parent
SVG_DIR = OUT_DIR / "svg"
DOCX_PATH = OUT_DIR / "RecipeAI_v3_Tool_Calling_完整技術報告.docx"

doc = Document()

# ── 全域樣式設定 ──
style = doc.styles["Normal"]
style.font.name = "Microsoft JhengHei"
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
style.paragraph_format.line_spacing = 1.35
style.paragraph_format.space_after = Pt(4)

for level in range(1, 5):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Microsoft JhengHei"
    hs.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    hs.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB) if level <= 2 else RGBColor(0x33, 0x33, 0x33)

def h1(text): doc.add_heading(text, level=1)
def h2(text): doc.add_heading(text, level=2)
def h3(text): doc.add_heading(text, level=3)
def h4(text): doc.add_heading(text, level=4)
def p(text): doc.add_paragraph(text)
def b(text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    return para

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    return t

def add_code(text, lang="python"):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    para.paragraph_format.left_indent = Cm(1)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)

PNG_DIR = OUT_DIR / "png"
PNG_DIR.mkdir(exist_ok=True)

def _svg_to_png(svg_filename: str) -> str:
    svg_path = SVG_DIR / svg_filename
    png_filename = svg_filename.replace(".svg", ".png")
    png_path = PNG_DIR / png_filename
    if not svg_path.exists():
        print(f"  ⚠️ SVG 不存在: {svg_path}")
        return ""
    cairosvg.svg2png(url=str(svg_path), write_to=str(png_path), scale=2.0)
    return str(png_path)

def add_svg_ref(filename, caption=""):
    cap_para = doc.add_paragraph()
    run = cap_para.add_run(f"圖：{caption or filename}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_para.paragraph_format.space_after = Pt(2)
    png_path = _svg_to_png(filename)
    if png_path:
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_para.add_run()
        run.add_picture(png_path, width=Inches(6.2))
        img_para.paragraph_format.space_after = Pt(8)
    else:
        fallback = doc.add_paragraph(f"（SVG 檔案缺失：svg/{filename}）")
        fallback.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_pedagogy_box(concept, term, detail, example):
    """教學法四段式：抽換概念→專有名詞→深入→例子"""
    t = doc.add_table(rows=4, cols=2)
    t.style = "Light Grid Accent 1"
    labels = ["🔄 抽換概念", "📖 專有名詞", "🔬 深入講解", "💡 實際例子"]
    contents = [concept, term, detail, example]
    for i in range(4):
        t.rows[i].cells[0].text = labels[i]
        t.rows[i].cells[1].text = contents[i]
        for para in t.rows[i].cells[0].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
        for para in t.rows[i].cells[1].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
    doc.add_paragraph()

# ════════════════════════════════════════════════════════
#  封面
# ════════════════════════════════════════════════════════
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_para.paragraph_format.space_before = Pt(120)
run = title_para.add_run("RecipeAI AI 營養顧問\nTool Calling 完整技術報告")
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("Version 3.0")
r2.font.size = Pt(18)
r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = sub2.add_run("v1 → v3 差異全面解析\n含工具自癒引擎・記憶樹・認知引擎・會話管理")
r3.font.size = Pt(13)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.paragraph_format.space_before = Pt(40)
r4 = date_p.add_run("2026 年 5 月 19 日")
r4.font.size = Pt(14)

doc.add_page_break()

# ════════════════════════════════════════════════════════
#  目錄頁
# ════════════════════════════════════════════════════════
h1("目錄")
toc_items = [
    "第一章　系統架構總覽與 v3 新增元件",
    "第二章　六步驟 LLM Pipeline（v3 增強版）",
    "第三章　SSE 通信協議（v3：15 種事件）",
    "第四章　v3 新增：工具自癒引擎（ReIn）",
    "第五章　v3 新增：記憶樹（TiMem）階層記憶系統",
    "第六章　v3 新增：認知重建引擎（CRE）",
    "第七章　v3 新增：會話管理與 Checkpoint 系統",
    "第八章　13 個工具完整通信鏈與程式碼解析",
    "第九章　工作流程狀態機（v3：含中斷恢復）",
    "第十章　雙記憶系統（v3：12 類結構化存儲）",
    "第十一章　所有 Prompt 完整內容與定位",
    "第十二章　v1 → v3 差異總覽",
    "附錄 A　前端 SSE 客戶端程式碼解析",
    "附錄 B　SVG 圖表索引",
]
for item in toc_items:
    doc.add_paragraph(item, style="List Number")
doc.add_page_break()

# ════════════════════════════════════════════════════════
#  第一章 系統架構
# ════════════════════════════════════════════════════════
h1("第一章　系統架構總覽與 v3 新增元件")

h2("1.1 抽換概念：什麼是「多服務架構」？")
p("想像一間餐廳：有負責接待客人的櫃檯（前端）、負責協調廚房的經理（Control 層）、"
  "負責從倉庫拿食材的倉管（Backend 層）、儲存所有食材的冷藏庫（資料庫），"
  "還有一位專門回答營養問題的顧問（LLM）。每個角色各司其職，透過對講機互相溝通。"
  "這就是我們系統的基本架構——每個「服務」就像餐廳裡的一個角色。")

h2("1.2 專有名詞講解")
add_pedagogy_box(
    "每個角色各有專長，在不同的「房間」（port）工作，透過「對講機」（HTTP）溝通。",
    "微服務架構（Microservice Architecture）：將系統拆成多個獨立運行的服務，每個服務有自己的 port（通訊埠），透過 HTTP API 互相呼叫。Docker 容器就像每個員工的獨立辦公室。",
    "本系統有 7 個服務：4 個跑在 Docker 容器內（PostgreSQL、Backend、Control、Frontend），3 個是獨立進程（Health Vector、GraphRAG API、Gemma LLM）。所有容器使用 --network host 模式，共享主機的網路空間，所以 localhost 就能互通。",
    "使用者在瀏覽器（port 3001）輸入「幫我查鈣含量最高的食物」→ Vite 代理把 /api 請求轉到 Control（port 8000）→ Control 呼叫 Gemma LLM（port 8080）判斷意圖 → Control 呼叫 Backend（port 8001）查資料庫 → 結果一路傳回瀏覽器。"
)

h2("1.3 深入講解：七大服務拓撲")
add_table(
    ["服務", "容器名 / 進程", "Port", "技術棧", "v3 新增角色"],
    [
        ["PostgreSQL", "recipe_ai_db", "5432", "postgres:16-alpine", "新增 memory_tree_nodes, conversation_sessions, conversation_checkpoints, cognitive_facts 表"],
        ["Backend", "recipe-backend", "8001", "FastAPI + asyncpg", "新增 memory_tree, conversation, cognitive 三組 Router"],
        ["Control", "recipe-control", "8000", "FastAPI + httpx + jieba", "新增 recovery.py, cognitive_engine.py, embedding_service.py"],
        ["Frontend", "recipe-frontend", "3001", "React 18 + Vite", "SSE 新增 recovery/session/checkpoint 事件處理"],
        ["Health Vector", "手動進程", "8003", "FastAPI + ChromaDB", "提供嵌入向量 /embed 端點給記憶樹使用"],
        ["GraphRAG API", "手動進程", "8002", "FastAPI + NetworkX", "無變更"],
        ["Gemma LLM", "llama-server", "8080", "Gemma-4-E4B Q8_0", "新增中斷/恢復分類 prompt"],
    ]
)

add_svg_ref("01_system_architecture_v3.svg", "v3 系統架構總覽")

h2("1.4 v3 新增四大子系統")
p("相比 v1，v3 新增了四個關鍵子系統，它們環繞在原有的 6 步驟 Pipeline 周圍，強化了系統的「記憶力」「自我修復能力」「認知能力」和「持久化能力」：")

add_table(
    ["子系統", "代號", "核心檔案", "一句話說明"],
    [
        ["工具自癒引擎", "ReIn", "recovery.py (237行)", "工具執行失敗時自動診斷→制定策略→重試"],
        ["記憶樹", "TiMem", "memory_tree_repository.py (192行)", "5層階層記憶：leaf→session→daily→weekly→profile"],
        ["認知重建引擎", "CRE", "cognitive_engine.py (97行)", "規則優先提取事實三元組（subject-predicate-object）"],
        ["會話管理", "Session/Checkpoint", "conversation_repository.py (152行)", "UUID 會話 + 每輪 Checkpoint 快照"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════
#  第二章 六步驟 LLM Pipeline
# ════════════════════════════════════════════════════════
h1("第二章　六步驟 LLM Pipeline（v3 增強版）")

h2("2.1 抽換概念：AI 顧問的「思考流程」")
p("想像你去看醫生。醫生不會聽你說完「我頭痛」就立刻開藥——他會先確認你的描述（嗯嗯，頭痛是吧），"
  "然後判斷可能的原因（是偏頭痛？還是感冒引起的？），接著決定要做什麼檢查（量血壓、驗血），"
  "然後執行檢查，看完報告後給你分析（你的血壓偏高，可能是...），最後把你的病歷寫進檔案。"
  "我們的 AI 顧問也是這樣「一步一步思考」的，每一步都有明確的任務。")

h2("2.2 專有名詞講解")
add_pedagogy_box(
    "醫生看診的 6 個步驟：確認→判斷→決定檢查→執行檢查→分析報告→寫病歷",
    "LLM Pipeline（大型語言模型處理管線）：將使用者的一句話，拆成 6 個階段依序處理。每個階段呼叫一次 LLM（Gemma-4-E4B），用不同的 Prompt（指令模板）和 Temperature（隨機度）來完成不同任務。",
    "6 個步驟分別是：Echo（即時回應）→ Phase1（意圖判定）→ Phase2（工具指令生成）→ Execute（工具執行）→ Analysis（結果分析）→ DualMemory（記憶更新）。v3 在 Pipeline 前後新增了「長期記憶載入」和「記憶樹/Checkpoint 儲存」。",
    "使用者說「最近很疲勞」→ Echo:「收到，我來查詢疲勞相關的營養素」→ Phase1:判定 intent=search_symptom → Phase2:生成 tool_call={search_symptom, symptom:疲勞} → Execute:呼叫 /health/symptom-search → Analysis:「找到8個相關效果，鐵和B12可能是關鍵」→ DualMemory:更新 narrative+structured_store"
)

h2("2.3 深入講解：每一步的技術細節")

h3("Step 1: Echo — 即時確認（advisor.py L799-818）")
p("目的：讓使用者知道系統收到了訊息，避免等待焦慮。這一步不回答問題，只是「嗯嗯，我聽到了」。")
add_code("""async def _step_echo(message, history, narrative_memory, long_term_context=""):
    # 組裝 memory_block：對話脈絡 + 長期記憶（v3 新增）
    memory_block = f"對話脈絡摘要：{narrative_memory}"
    if long_term_context:
        memory_block += f"\\n\\n長期記憶：\\n{long_term_context}"
    prompt = PROMPT_ECHO.format(message=message, history=..., memory_block=...)
    result = await _call_gemma([{"role":"user","content":prompt}], temperature=0.5)
    return result or "收到，正在處理你的問題..."  """)
p("Temperature=0.5：中等隨機度，讓回應自然但不離題。v3 新增：Echo 現在也能看到長期記憶（記憶樹），"
  "所以如果使用者上週說過「我對花生過敏」，Echo 可能會回「收到，我會注意你的花生過敏喔」。")

h3("Step 2: Phase1 — 意圖判定（advisor.py L833-871）")
p("目的：讀完系統功能說明書（SYSTEM_MANUAL，約 30KB），判斷使用者想做什麼。支援 14 種意圖。")
add_table(
    ["意圖代碼", "說明", "v3 狀態"],
    [
        ["search_symptom", "搜尋症狀相關營養素和食物", "變更：中文直傳"],
        ["search_graphrag", "學術論文知識圖譜問答", "不變"],
        ["nutrient_ranking", "營養素含量排名", "不變"],
        ["get_recipes", "查看已儲存食譜", "不變"],
        ["get_calendar", "行事曆查詢", "不變"],
        ["add_to_calendar", "儲存食譜加入行事曆", "不變"],
        ["delete_calendar", "刪除行事曆紀錄", "v3 新增"],
        ["search_food", "搜尋食物營養成分", "不變"],
        ["browse_food_database", "瀏覽台灣食品資料庫", "不變"],
        ["analyze_recipe", "分析食譜營養", "增強：內建 DRI 缺口"],
        ["browse_literature", "瀏覽學術論文列表", "不變"],
        ["get_profile", "讀取個人資料", "不變"],
        ["update_profile", "更新個人資料", "不變"],
        ["general_chat", "一般對話", "不變"],
    ]
)
p("v3 新增：Phase1 的 prompt 現在額外注入 narrative_memory（對話脈絡）和 long_term_context（記憶樹），"
  "讓意圖判定能參考歷史對話。例如使用者上一輪問了「鈣」，這一輪說「那鐵呢」，"
  "Phase1 能從脈絡推斷出 intent=nutrient_ranking。")

h3("Step 3: Phase2 — 工具指令生成（advisor.py L874-955）")
p("目的：根據意圖和使用者已知資訊，生成具體的工具呼叫指令。")
p("v3 的關鍵設計：「系統優先，LLM 備援」。首先嘗試 _system_generate_toolcalls()（純 Python 確定性邏輯），"
  "如果能直接生成就不呼叫 LLM——省時間又穩定。只有系統無法確定時才呼叫 LLM。")
add_code("""# Phase2 的三層策略（advisor.py L874-955）
async def _step_phase2(...):
    # 層 0: 系統確定性生成（優先）
    sys_calls = _system_generate_toolcalls(phase1_result, profile)
    if sys_calls is not None and sys_calls:
        validated, corrections = _validate_toolcalls(sys_calls, ...)
        if validated:
            return validated, ["system-deterministic"]  # 不需 LLM

    # 層 1: LLM 生成（備援）
    text = await _call_gemma([...], temperature=0.1)
    result = _parse_llm_json(text)
    validated, corrections = _validate_toolcalls(result["tool_calls"], ...)

    # 層 2: LLM 完全失敗時的最終備援
    if not validated:
        sys_calls = _system_generate_toolcalls(phase1_result, profile)
        validated, corrections = _validate_toolcalls(sys_calls, ...)
        return validated, ["LLM-failed, system-fallback"]""")

h3("Step 4: Execute — 工具執行（v3：含自癒引擎）（advisor.py L1062-1086）")
p("v3 最大變更：所有工具執行都包了一層 _execute_tool_with_recovery()。"
  "執行完工具後，自動呼叫 recovery.diagnose() 檢查結果是否正常，不正常就啟動自癒流程。")
add_code("""async def _execute_tool_with_recovery(name, args, auth_token, user_msg, store, _retry=0):
    result = await _execute_tool(name, args, ...)
    if _retry >= _RECOVERY_MAX_RETRIES:  # 最多重試 1 次
        return result, None
    diag = _recovery_diagnose(name, args, result, user_msg)
    if not diag or diag.confidence < 0.5:
        return result, None  # 沒問題或信心不足，不恢復
    plan = _recovery_plan(diag, {"zh_to_field": _NUTRIENT_ZH_TO_FIELD}, synonym_store)
    if not plan:
        return result, None
    recovered = await _recovery_execute(plan, _exec_fn, name, ...)
    recovery_info = {"tool": name, "strategy": plan.strategy, "success": ...}
    return recovered, recovery_info""")

h3("Step 5: Analysis — 結果分析（advisor.py L989-1036）")
p("目的：將工具查詢結果翻譯成人話，告訴使用者有意義的結論。")
p("v3 新增 _trim_tool_results_for_analysis()：症狀搜尋結果可能很大（數百個推薦），"
  "送進 LLM 前先裁剪為摘要（前 10 個推薦、前 5 個語義匹配），節省 token 又不失重點。")
p("特殊處理：如果工具是 search_graphrag，且 answer 不是錯誤，直接原文呈現（不再讓 LLM 改寫），"
  "因為 GraphRAG 的學術分析已經是完整的。")

h3("Step 6: DualMemory — 雙記憶更新（advisor.py L2819-2848）")
p("目的：並行執行兩個 LLM 呼叫：(a) 更新對話脈絡摘要 (b) 提取結構化資料。")
p("v3 新增：Step 6 結束後，還會觸發「記憶樹 leaf 節點儲存」和「Checkpoint 非同步儲存」——"
  "這兩個都是 fire-and-forget（發射後不管），不阻塞回應速度。")

add_svg_ref("02_llm_pipeline_v3.svg", "v3 六步驟 LLM Pipeline")
doc.add_page_break()

# ════════════════════════════════════════════════════════
#  第三章 SSE 通信協議
# ════════════════════════════════════════════════════════
h1("第三章　SSE 通信協議（v3：15 種事件）")

h2("3.1 抽換概念：什麼是「即時推播」？")
p("想像你在看直播——主播說一句話，你立刻在螢幕上看到。你不需要一直按重新整理，"
  "內容自動推送到你的螢幕。這就是 SSE（Server-Sent Events）的原理。"
  "跟你平常在網頁上按「送出」然後等整個結果不同，SSE 讓伺服器可以「一段一段地」把結果推給你。")

h2("3.2 專有名詞講解")
add_pedagogy_box(
    "直播：主播說一句你看一句，不用等整場結束",
    "SSE（Server-Sent Events）：HTTP 長連線技術，伺服器可以持續推送事件到瀏覽器。每個事件有 event type（類型）和 data（資料），用 JSON 格式傳輸。前端用 fetch + ReadableStream 讀取。",
    "與 WebSocket 不同，SSE 是單向的（伺服器→瀏覽器），走標準 HTTP，不需特殊握手。RecipeAI 用 SSE 而非 WebSocket，因為對話模式本身就是「使用者發一次、AI 回一長串」，單向推送完全夠用。",
    "使用者按「送出」→ POST /advisor/chat → 伺服器開始 SSE 串流 → 前端依序收到 phase_status→echo→tool_start→tool_result→tool_done→reply→memory→workflow_state → 串流結束"
)

h2("3.3 深入講解：15 種事件類型")
p("v1 有 9 種事件，v3 擴展到 15 種。新增的 6 種事件支持自癒引擎、會話管理和更細緻的狀態回報。")
add_table(
    ["事件類型", "資料格式", "用途", "v3 狀態"],
    [
        ["phase_status", "{phase, label}", "Pipeline 進度提示", "不變"],
        ["echo", "{content}", "Step 1 即時確認回應", "不變"],
        ["tool_start", "{name, args}", "工具開始執行", "不變"],
        ["tool_result", "{name, data}", "工具執行結果（完整資料）", "不變"],
        ["tool_done", "{name, summary}", "工具執行完成（摘要）", "不變"],
        ["recovery_attempt", "{tool, strategy, explanation}", "自癒引擎嘗試恢復", "v3 新增"],
        ["recovery_result", "{tool, success}", "自癒引擎恢復結果", "v3 新增"],
        ["reply", "{content}", "AI 分析回覆", "不變"],
        ["validation", "{step, corrections}", "驗證修正報告", "不變"],
        ["memory", "{summary}", "記憶摘要（向後相容）", "不變"],
        ["narrative_memory", "{summary}", "對話脈絡摘要", "不變"],
        ["structured_store", "{store}", "結構化存儲更新", "不變"],
        ["workflow_state", "{phase, ...全部狀態欄位}", "工作流程狀態", "不變"],
        ["session_id", "{session_id}", "會話 ID 通知", "v3 新增"],
        ["checkpoint_saved", "{session_id, turn_index}", "Checkpoint 已儲存通知", "v3 新增"],
    ]
)

h2("3.4 實際例子：一次完整的 SSE 串流")
add_code("""# 使用者說「幫我查鈣含量最高的食物」時，前端依序收到：

event: phase_status
data: {"phase": "echo", "label": "理解需求中..."}

event: echo
data: {"content": "收到，我來幫你查詢鈣含量最高的食物。"}

event: phase_status
data: {"phase": "intent", "label": "理解需求中..."}

event: phase_status
data: {"phase": "toolcall", "label": "分析意圖中..."}

event: phase_status
data: {"phase": "execute", "label": "工具調用中..."}

event: tool_start
data: {"name": "get_nutrient_ranking", "args": {"nutrient_field": "calcium_per_100g", "top_n": 10}}

event: tool_result
data: {"name": "get_nutrient_ranking", "data": {"nutrient": "calcium_per_100g", "top_foods": [...]}}

event: tool_done
data: {"name": "get_nutrient_ranking", "summary": "取得前 10 名"}

event: phase_status
data: {"phase": "analysis", "label": "分析結果中..."}

event: reply
data: {"content": "鈣含量最高的食物是小魚乾（2,213mg/100g），其次是..."}

event: workflow_state
data: {"phase": "done", "active_intent": "", ...}

event: session_id
data: {"session_id": "a1b2c3d4-..."}

event: phase_status
data: {"phase": "memory", "label": "更新記憶中..."}

event: narrative_memory
data: {"summary": "使用者查詢鈣含量排名；AI回報前10名食物..."}

event: structured_store
data: {"store": {"ranking_entry": {"nutrient_field": "calcium_per_100g", ...}}}

event: memory
data: {"summary": "..."}

event: checkpoint_saved
data: {"session_id": "a1b2c3d4-...", "turn_index": 1}""")

add_svg_ref("03_sse_protocol_v3.svg", "v3 SSE 通信協議（15 種事件）")
doc.add_page_break()

# ════════════════════════════════════════════════════════
#  第四章 工具自癒引擎 (ReIn)
# ════════════════════════════════════════════════════════
h1("第四章　v3 新增：工具自癒引擎（ReIn）")

h2("4.1 抽換概念：什麼是「自動修錯」？")
p("想像你在超市找「芝士」，但貨架標的是「起司」。你找不到的時候，"
  "聰明的店員會說：「你要找的可能是起司？」然後帶你去正確的貨架。"
  "自癒引擎就是這個「聰明的店員」——當工具查詢失敗時，它會自動嘗試換個說法再找一次。")

h2("4.2 專有名詞講解")
add_pedagogy_box(
    "超市找不到「芝士」→ 店員建議「試試找起司？」→ 找到了",
    "工具自癒引擎（ReIn, Recovery Engine）：一個三步驟的自動修復機制。① diagnose()：診斷失敗原因（空結果？超時？參數不對？後端錯誤？）。② plan_recovery()：制定恢復策略（換同義詞？重試？欄位映射？跳過？）。③ execute_recovery()：執行恢復策略，最多重試 1 次，防止無限遞迴。",
    "recovery.py 是一個 237 行的獨立模組，不依賴 LLM——所有診斷和恢復策略都是確定性的 Python 邏輯。它內建了一個 _COMMON_FOOD_SYNONYMS 辭典（29 組常見食物的中英文別名），以及 _QUANTITY_PREFIXES 和 _MODIFIER_PATTERNS 兩組正規表達式，用來簡化食物名稱。",
    "使用者說「我想查芝士的營養」→ search_food(food_name='芝士') → 回傳 results=[] → diagnose(): empty_result, confidence=0.9 → plan_recovery(): synonym_expand, 替代詞=['起司','乾酪','cheese'] → execute_recovery(): search_food(food_name='起司') → 找到 3 筆結果 → SSE 送出 recovery_attempt + recovery_result 事件"
)

h2("4.3 深入講解：四種診斷類型 × 四種恢復策略")
add_table(
    ["診斷類型", "觸發條件", "對應策略", "策略行為"],
    [
        ["empty_result", "search_food 回傳 0 筆結果", "synonym_expand", "查同義詞辭典+store，逐個嘗試替代詞（最多5個）"],
        ["tool_timeout", "錯誤訊息含 timeout/timed out", "retry_timeout", "直接重試一次"],
        ["param_mismatch", "get_nutrient_ranking 回傳 0 筆 + 無 error", "field_remap", "用 _NUTRIENT_ZH_TO_FIELD 映射表重新對應欄位名"],
        ["backend_error", "HTTP 500 或錯誤訊息", "skip", "跳過此工具（回傳原始結果）"],
    ]
)

h2("4.4 程式碼解析：recovery.py 核心邏輯")
add_code("""# recovery.py — 診斷函式（L71-90）
def diagnose(tool_name, args, result, user_msg=""):
    if "error" in result:
        err_msg = str(result["error"])
        if "timeout" in err_msg.lower():
            return ErrorDiagnosis("tool_timeout", tool_name, args, 0.9)
        if result.get("status_code", 0) >= 500:
            return ErrorDiagnosis("backend_error", tool_name, args, 0.8)
        return ErrorDiagnosis("backend_error", tool_name, args, 0.6)
    if tool_name == "search_food":
        if len(result.get("results", [])) == 0:
            return ErrorDiagnosis("empty_result", tool_name, args, 0.9)
    if tool_name == "get_nutrient_ranking":
        if len(result.get("foods", [])) == 0 and "error" not in result:
            return ErrorDiagnosis("param_mismatch", tool_name, args, 0.7)
    return None  # 正常結果""")

add_code("""# recovery.py — 同義詞擴展函式（L195-218）
def _expand_food_synonyms(food_name, synonym_store=None):
    candidates = []
    # 1. 內建辭典查詢
    if food_name in _COMMON_FOOD_SYNONYMS:
        candidates.extend(_COMMON_FOOD_SYNONYMS[food_name])
    # 2. 動態同義詞 Store 查詢（3,172 canonical / 8,646 synonyms）
    if synonym_store:
        entry = synonym_store.lookup_synonym(food_name)
        if entry:
            for s in entry.get("synonyms", []):
                if s not in candidates and s != food_name:
                    candidates.append(s)
    return unique[:5]  # 最多回傳 5 個替代詞""")

add_svg_ref("04_recovery_system.svg", "v3 工具自癒引擎（ReIn）")
doc.add_page_break()

# ════════════════════════════════════════════════════════
#  第五章 記憶樹 (TiMem)
# ════════════════════════════════════════════════════════
h1("第五章　v3 新增：記憶樹（TiMem）階層記憶系統")

h2("5.1 抽換概念：什麼是「記憶的階層」？")
p("想像你在學校——每天上完課會寫筆記（單堂筆記），學期末會整理成一份期末總結，"
  "然後畢業前會做一個「我這四年學了什麼」的畢業回顧。"
  "記憶也是這樣分層的：最底層是每一輪對話的即時記錄，往上逐層壓縮合併，"
  "最頂層是使用者的長期人物畫像。這就是「記憶樹」——像一棵倒過來的樹，葉子在底部，根在頂部。")

h2("5.2 專有名詞講解")
add_pedagogy_box(
    "學校筆記系統：單堂筆記→期末總結→畢業回顧，越往上越精簡，但保留最重要的資訊",
    "記憶樹（TiMem, Tiered Memory Tree）：5 層階層記憶系統，每層有不同的壽命和精簡度。Level 0（Leaf，單輪摘要）→ Level 1（Session，會話摘要）→ Level 2（Daily，每日）→ Level 3（Weekly，每週）→ Level 4（Profile，長期人物畫像）。使用 pgvector 做語義搜尋，嵌入維度 1024。",
    "技術實現分 3 層：① Backend 層 memory_tree_repository.py 負責 CRUD + 向量搜尋（PostgreSQL + pgvector）。② Control 層 advisor.py 的 _save_memory_leaf() / _consolidate_session_memory() / _get_memory_tree_context() 負責業務邏輯。③ embedding_service.py 呼叫 Health Vector（port 8003）的 /embed 端點取得向量。",
    "使用者第 1 輪說「我25歲男性」→ _save_memory_leaf() 存一個 Level 0 node，content='使用者提供基本資訊：25歲男性'。第 10 輪時 _consolidate_session_memory() 觸發，把 10 個 leaf 合併成 1 個 Level 1 session node。下一次新對話開始時，_get_memory_tree_context(query='疲勞') 會用語義搜尋找到相關的歷史記憶。"
)

h2("5.3 深入講解：記憶樹的三個關鍵操作")

h3("5.3.1 儲存 Leaf 節點（advisor.py L1743-1780）")
add_code("""async def _save_memory_leaf(user_id, session_id, summary, tools_used):
    token_count = len(summary) // 3  # 粗估 token 數
    # 1. 建立 Level 0 節點
    resp = await c.post(f"{backend_url}/memory/nodes", json={
        "user_id": user_id, "session_id": session_id,
        "level": 0, "content": summary,
        "metadata": {"tools": tools_used}, "token_count": token_count,
    })
    node_id = resp.json().get("id")
    # 2. 計算嵌入向量（呼叫 Health Vector port 8003）
    embedding = await _embed_text(summary)  # 1024 維向量
    if embedding:
        await c.post(f"{backend_url}/memory/embedding",
                     json={"node_id": node_id, "embedding": embedding})
    # 3. 檢查是否需要合併（每 10 輪觸發）
    session_leaves = [n for n in leaves if n.get("session_id") == session_id
                      and n.get("parent_id") is None]
    if len(session_leaves) >= 10:
        asyncio.create_task(_consolidate_session_memory(...))""")

h3("5.3.2 會話合併（advisor.py L1783-1815）")
p("當同一個 session 累積了 10 個以上的 leaf 節點，自動觸發合併：")
add_code("""async def _consolidate_session_memory(user_id, session_id, leaves, client):
    combined = "\\n".join(n["content"] for n in leaves)
    prompt = "以下是同一對話中多輪摘要，請合併為一段精簡的會話摘要（3-5句）：\\n\\n" + combined
    session_summary = await _call_gemma([{"role":"user","content":prompt}], temperature=0.2)
    # 建立 Level 1 父節點
    parent_resp = await c.post(f"{backend_url}/memory/nodes", json={
        "user_id": user_id, "session_id": session_id,
        "level": 1, "content": session_summary, ...
    })
    parent_id = parent_resp.json()["id"]
    # 將所有 leaf 指向此父節點
    await c.post(f"{backend_url}/memory/consolidate", json={
        "parent_id": parent_id, "summary": session_summary,
        "child_ids": [n["id"] for n in leaves],
    })""")

h3("5.3.3 語義搜尋取上下文（advisor.py L1818-1849）")
add_code("""async def _get_memory_tree_context(user_id, max_tokens=500, query=""):
    parts = []
    # 1. Token-budget-aware 階層檢索
    resp = await c.get(f"{backend_url}/memory/context",
                       params={"user_id": user_id, "max_tokens": max_tokens})
    ctx = resp.json().get("context", "")
    if ctx: parts.append(ctx)
    # 2. 語義搜尋增強（用使用者當前問題的嵌入向量）
    if query:
        query_emb = await _embed_text(query)
        if query_emb:
            search_resp = await c.post(f"{backend_url}/memory/search", json={
                "user_id": user_id, "query_embedding": query_emb,
                "level_min": 0, "level_max": 3, "limit": 3,
            })
            for r in search_resp.json():
                if r["content"] not in parts:
                    parts.append(f"[相關記憶] {r['content']}")
    return "\\n---\\n".join(parts)""")

h2("5.4 資料庫 Schema")
add_code("""-- migration 009_memory_tree.sql
CREATE TABLE memory_tree_nodes (
    id SERIAL PRIMARY KEY,
    user_id INTEGER NOT NULL REFERENCES users(id),
    session_id UUID REFERENCES conversation_sessions(id),
    level INTEGER NOT NULL CHECK (level BETWEEN 0 AND 4),
    parent_id INTEGER REFERENCES memory_tree_nodes(id),
    content TEXT NOT NULL,
    metadata JSONB DEFAULT '{}',
    token_count INTEGER DEFAULT 0,
    created_at TIMESTAMPTZ DEFAULT NOW(),
    updated_at TIMESTAMPTZ DEFAULT NOW()
);
-- migration 010_memory_vectors.sql
ALTER TABLE memory_tree_nodes ADD COLUMN embedding vector(1024);
CREATE INDEX ON memory_tree_nodes
  USING ivfflat (embedding vector_cosine_ops) WITH (lists = 50);""")

add_svg_ref("05_memory_tree_system.svg", "v3 記憶樹（TiMem）階層記憶系統")
doc.add_page_break()

# ════════════════════════════════════════════════════════
#  第六章 認知重建引擎 (CRE)
# ════════════════════════════════════════════════════════
h1("第六章　v3 新增：認知重建引擎（CRE）")

h2("6.1 抽換概念：什麼是「自動做筆記」？")
p("想像你在跟朋友聊天，聊到一半他說「我25歲，對花生過敏，最近在減脂」。"
  "如果你是一個好聽眾，你的大腦會自動「摘錄」這些事實：年齡=25、過敏=花生、目標=減脂。"
  "下次聊天時你不需要再問一遍。認知引擎做的就是同樣的事——自動從對話中擷取事實。")

h2("6.2 專有名詞講解")
add_pedagogy_box(
    "聊天時自動記住對方說的事實：25歲、花生過敏、減脂",
    "認知重建引擎（CRE, Cognitive Reconstruction Engine）：用「規則優先，LLM 備援」的方式，從對話中提取事實三元組（subject-predicate-object）。例如（使用者, 年齡, 25, 信心度 1.0）。提取後存入 cognitive_facts 資料表。",
    "cognitive_engine.py 只有 97 行，全靠 8 個正規表達式（regex）工作，完全不需要 LLM。8 個 regex 分別匹配：年齡（_AGE_RE）、性別（_GENDER_RE）、過敏（_ALLERGY_RE）、飲食偏好（_DIET_KW）、體重（_WEIGHT_RE）、身高（_HEIGHT_RE）、懷孕/哺乳（_PREGNANT_RE）、疾病（_DISEASE_RE）。",
    "使用者說：「我是25歲男性，170公分65公斤，對花生過敏，正在生酮飲食」→ extract_facts() 回傳 5 個事實：\n(使用者,年齡,25,1.0), (使用者,性別,男,1.0), (使用者,身高,170cm,1.0), (使用者,體重,65kg,1.0), (使用者,過敏,花生,0.9), (使用者,飲食偏好,生酮,0.8)"
)

h2("6.3 深入講解：8 個正規表達式")
add_code("""# cognitive_engine.py — 8 個 regex pattern
_AGE_RE      = re.compile(r"(\\d{1,3})\\s*歲")
_GENDER_RE   = re.compile(r"(男性?|女性?|male|female)", re.IGNORECASE)
_ALLERGY_RE  = re.compile(r"過敏.*?[：:]?\\s*([^，。、\\n]+)")
_DIET_KW     = re.compile(r"(素食|全素|蛋奶素|無麩質|低碳|生酮|地中海飲食|得舒飲食|低鈉|低脂|高蛋白)")
_WEIGHT_RE   = re.compile(r"(\\d{2,3}(?:\\.\\d)?)\\s*(?:公斤|kg)", re.IGNORECASE)
_HEIGHT_RE   = re.compile(r"(\\d{2,3}(?:\\.\\d)?)\\s*(?:公分|cm)", re.IGNORECASE)
_PREGNANT_RE = re.compile(r"(懷孕|孕期|孕婦|哺乳)")
_DISEASE_RE  = re.compile(r"(糖尿病|高血壓|高血脂|痛風|腎臟病|心臟病|貧血|骨質疏鬆)")""")

h2("6.4 context 重建")
add_code("""def reconstruct_context(facts, token_budget=200):
    lines = []
    budget = token_budget
    for f in facts:
        if budget <= 0: break
        line = f"- {f['subject']}的{f['predicate']}：{f['object']}"
        if f.get("confidence", 1.0) < 0.8:
            line += "（待確認）"
        lines.append(line)
        budget -= len(line) // 3
    return "使用者已知資訊：\\n" + "\\n".join(lines)
# 輸出範例：
# 使用者已知資訊：
# - 使用者的年齡：25
# - 使用者的性別：男
# - 使用者的過敏：花生
# - 使用者的飲食偏好：生酮（待確認）""")

add_svg_ref("06_cognitive_engine.svg", "v3 認知重建引擎（CRE）")
doc.add_page_break()

# ════════════════════════════════════════════════════════
#  第七章 會話管理與 Checkpoint
# ════════════════════════════════════════════════════════
h1("第七章　v3 新增：會話管理與 Checkpoint 系統")

h2("7.1 抽換概念：什麼是「存檔和讀檔」？")
p("玩 RPG 遊戲時，你會在重要關卡前「存檔」，萬一掛了可以從存檔點「讀檔」繼續。"
  "會話管理做的就是同樣的事——每輪對話結束後自動存一個「快照」，"
  "包含所有的對話紀錄、記憶、工具結果和狀態。如果瀏覽器重新整理或斷線，"
  "只要帶著 session_id 回來，就能從上次的存檔點繼續。")

h2("7.2 專有名詞講解")
add_pedagogy_box(
    "RPG 存檔：在重要關卡自動存檔，掛了可以讀檔繼續",
    "Session（會話）：用 UUID 識別的一次對話過程。前端首次連線時，後端建立新 session 並回傳 session_id。之後每次請求都帶著這個 ID。Checkpoint（檢查點）：每輪對話結束後的狀態快照，包含 workflow_state、narrative_memory、structured_store、message_history 等。",
    "兩種 Checkpoint 儲存策略：① _save_checkpoint_async()（完整快照）：每輪結束後儲存。② _save_checkpoint_incremental()（增量快照）：工具執行中，每完成一個工具就存一次，記錄 pending_tool_calls 和 completed_results——萬一中途斷線，恢復時只需執行剩餘的工具。",
    "使用者分析食譜（需要 3 步工具呼叫）→ 完成第 1 步 extract → 存增量 checkpoint（pending=[lookup,calculate], completed=[extract]) → 此時瀏覽器斷線 → 重新連線帶 resume_checkpoint=true → 讀取 latest checkpoint → 直接從第 2 步 lookup 繼續 → 不需重跑 extract"
)

h2("7.3 深入講解：會話生命週期")
add_code("""# advisor.py L2695-2706 — 會話建立
session_id = body.session_id
if not session_id:
    resp = await c.post(f"{backend_url}/conversations/sessions",
                        json={"user_id": user_id, "title": body.message[:50]})
    if resp.status_code == 201:
        session_id = resp.json().get("id", "")  # UUID 字串""")

add_code("""# advisor.py L2708-2732 — Checkpoint 恢復
if body.resume_checkpoint and session_id:
    resp = await c.get(f"{backend_url}/conversations/checkpoints/latest",
                       params={"session_id": session_id, "user_id": user_id})
    if resp.status_code == 200:
        cp = resp.json()
        body.history = cp.get("message_history", body.history)
        body.narrative_memory = cp.get("narrative_memory", ...)
        body.structured_store = cp.get("structured_store", ...)
        body.workflow_state = cp.get("workflow_state", ...)
        pending = cp.get("pending_tool_calls") or []
        if pending and cp.get("phase") == "execute":
            # 有未完成的工具！設定 resume 模式
            body.workflow_state = {"phase": "ready"}
            _resume_tool_calls = pending
            _resume_completed = cp.get("completed_results") or []""")

h2("7.4 資料庫 Schema")
add_code("""-- migration 008_conversation_checkpoints.sql
CREATE TABLE conversation_sessions (
    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
    user_id INTEGER NOT NULL REFERENCES users(id),
    title VARCHAR(200),
    is_active BOOLEAN DEFAULT TRUE,
    created_at TIMESTAMPTZ DEFAULT NOW(),
    updated_at TIMESTAMPTZ DEFAULT NOW()
);
CREATE TABLE conversation_checkpoints (
    id SERIAL PRIMARY KEY,
    session_id UUID NOT NULL REFERENCES conversation_sessions(id),
    user_id INTEGER NOT NULL REFERENCES users(id),
    turn_index INTEGER NOT NULL,
    phase VARCHAR(20),
    workflow_state JSONB DEFAULT '{}',
    pending_tool_calls JSONB DEFAULT '[]',
    completed_results JSONB DEFAULT '[]',
    narrative_memory TEXT DEFAULT '',
    structured_store JSONB DEFAULT '{}',
    message_history JSONB DEFAULT '[]',
    created_at TIMESTAMPTZ DEFAULT NOW()
);""")

add_svg_ref("07_session_checkpoint.svg", "v3 會話管理與 Checkpoint 系統")
doc.add_page_break()

# ════════════════════════════════════════════════════════
#  第八章 13 個工具完整通信鏈
# ════════════════════════════════════════════════════════
h1("第八章　13 個工具完整通信鏈與程式碼解析")

h2("8.1 抽換概念：什麼是「工具」？")
p("想像你是一個萬能秘書。老闆說「查一下明天的天氣」，你不是自己變出天氣資料，"
  "而是拿起手機打開天氣 App（工具）去查。老闆說「幫我記到行事曆」，"
  "你打開行事曆 App（另一個工具）去新增。AI 顧問也是這樣——它本身不會查資料庫，"
  "但它會呼叫 13 個「工具」去做各種事情，然後把結果統整回報。")

h2("8.2 工具調度表")
add_table(
    ["#", "工具名", "意圖代碼", "後端路由", "必要參數", "v3 狀態"],
    [
        ["1", "get_nutrient_ranking", "nutrient_ranking", "GET /health/nutrients/{field}/top-foods", "nutrient_field", "不變"],
        ["2", "search_food", "search_food", "GET /food/search?q=", "food_name", "不變"],
        ["3", "browse_taiwan_food", "browse_food_database", "GET /food/taiwan/browse", "（無）", "不變"],
        ["4", "analyze_recipe", "analyze_recipe", "3步：extract→lookup→calculate", "recipe_text", "增強"],
        ["5", "get_literature_papers", "browse_literature", "GET /graphrag-api/papers/enriched", "（無）", "不變"],
        ["6", "search_symptom", "search_symptom", "POST /health/symptom-search", "symptom", "變更"],
        ["7", "search_graphrag", "search_graphrag", "POST graphrag:8002/query-sync", "query", "不變"],
        ["8", "get_saved_recipes", "get_recipes", "GET /recipe/list", "（無）", "不變"],
        ["9", "get_calendar_entries", "get_calendar", "GET /calendar/entries", "start_date, end_date", "不變"],
        ["10", "add_to_calendar", "add_to_calendar", "POST /recipe/save + POST /calendar/entries", "entry_date", "不變"],
        ["11", "delete_calendar_entry", "delete_calendar", "GET→filter→DELETE", "entry_date", "新增"],
        ["12", "get_user_profile", "get_profile", "GET /profile", "（無）", "不變"],
        ["13", "update_user_profile", "update_profile", "PUT /profile", "至少一欄", "不變"],
    ]
)

h2("8.3 v3 變更工具詳解")

h3("8.3.1 search_symptom — 症狀搜尋（v3：中文直傳 + 5 階段管道）")
add_pedagogy_box(
    "以前：使用者說中文→系統翻成英文→去查。現在：使用者說中文→直接查，系統自動處理翻譯和語義搜尋。",
    "v3 的 search_symptom 不再需要前端翻譯，直接把中文症狀詞傳給 POST /health/symptom-search。後端的 symptom_service.py 實現了 5 階段搜尋管道：LLM 解析→雙路搜尋→LLM 篩選→compound→food→去重彙總。",
    "v1：advisor.py 內部查 _SYMPTOM_ZH_TO_EN 辭典翻譯→GET /food/symptom/{英文}→只走 symptom_mapping.json 靜態對照。v3：直接 POST json={keyword: 中文}→symptom_service.py 5 階段→回傳包含 parsed_symptoms（LLM 拆解結果）、vector_matches（語義搜尋匹配，含中文名和相似度%）、effects、recommendations。",
    "使用者：「手肘肌肉酸痛」→ LLM 拆解為 ['肌肉酸痛', '關節疼痛'] → 雙路搜尋：symptom_mapping 找到 3 個 + Health Vector 語義搜尋找到 5 個 → LLM 篩選保留 top 10 → 回傳 vector_matches=[{name:'Muscle pain relief', name_zh:'肌肉疼痛緩解', score:0.87}, ...]"
)

h3("8.3.2 analyze_recipe — 食譜分析（v3：文字前處理 + 內建 DRI 缺口）")
add_pedagogy_box(
    "以前：只算營養素，不告訴你哪些不夠。現在：算完營養素後，自動對照 HPA 第八版 DRI，告訴你哪些夠、哪些不夠。",
    "v3 的 analyze_recipe 增加兩個功能：① 文字前處理（regex 拆分頓號/逗號為換行，去除「早上吃了」等前綴）。② Step 4 DRI 缺口分析：自動讀取使用者 profile（年齡/性別），對照 33 項 DRI 標準，標記 6 種狀態。",
    "6 種 DRI 狀態判定邏輯：ratio = actual/target × 100%。sufficient（≥100%）、slightly_low（80-99%）、deficient（<80%）、excess（超過上限 UL）、near_limit（≥85% UL）、no_data（DB 無此欄位資料）。",
    "使用者貼食譜「雞蛋3顆、牛奶200ml、菠菜100g」→ Step1 前處理拆分 → Step2 extract 得到 3 食材 → Step3 lookup 匹配台灣食品 → Step4 calculate → Step5 DRI 缺口：蛋白質 sufficient(125%)、鈣 slightly_low(85%)、鐵 deficient(45%)、維生素D deficient(30%)"
)

add_code("""# advisor.py L1475-1650 — analyze_recipe（4+1 步驟）
async def _tool_analyze_recipe(args, token, user_msg):
    recipe_text = args.get("recipe_text") or user_msg
    # v3 新增：文字前處理
    recipe_text = re.sub(r"[、，,;；+＋]\\s*", "\\n", recipe_text)  # 頓號→換行
    recipe_text = re.sub(r"(?<=[一-鿿])(?:跟|還有|以及)(?=[一-鿿])", "\\n", recipe_text)
    recipe_text = re.sub(r"(?:早[上餐]|中[午餐]|晚[上餐])(?:吃了?)\\s*", "", recipe_text)

    # Step 1: 食材提取
    resp1 = await c.post(f"{url}/recipe/extract-with-amounts", json={"recipe_text": recipe_text})
    # Step 2: 雙軌匹配（taiwan_foods 優先）
    resp2 = await c.post(f"{url}/recipe/lookup", json={"ingredients": names, "top_n": 20})
    # Step 3: 營養計算
    resp3 = await c.post(f"{url}/recipe/calculate-nutrition", json={"ingredients": calc_items})

    # Step 4 (v3 新增): DRI 缺口分析
    profile = await c.get(f"{url}/profile", headers=_headers(token))
    dri_info = _get_dri_targets(age, sex)
    for key, tgt in dri_info["targets"].items():
        actual = float(total.get(key) or 0)
        ratio_pct = round(actual / tgt * 100, 1) if tgt else 0
        # ... 判定 6 種狀態 ...
    result["dri_analysis"] = {"age_group": ..., "rows": rows, "deficient_count": ...}""")

h3("8.3.3 delete_calendar_entry — 行事曆刪除（v3 新增）")
add_pedagogy_box(
    "以前只能加行事曆，不能刪。現在可以對 AI 說「幫我刪掉昨天的食譜」，它會自動找到並刪除。",
    "delete_calendar_entry 是 v3 完全新增的工具。它用「先查後刪」的策略：先 GET 該日期的所有行事曆條目，然後根據 recipe_name 做模糊匹配（case-insensitive contains），最後逐筆 DELETE。",
    "3 種刪除模式：① entry_id 精確刪（直接 DELETE /calendar/entries/{id}）。② recipe_name 模糊刪（先查全部→name 匹配→只刪匹配的）。③ 全部刪（先查全部→全部 DELETE）。找不到時回傳現有紀錄名稱供使用者確認。",
    "使用者：「幫我刪掉今天的雞胸肉沙拉」→ GET /calendar/entries?date=2026-05-19 → 找到 [{id:42, recipe_name:'雞胸肉沙拉'}, {id:43, recipe_name:'燕麥粥'}] → 匹配'雞胸肉沙拉'→ DELETE /calendar/entries/42 → 回傳 {success:true, deleted_count:1, recipe_name:'雞胸肉沙拉'}"
)

add_svg_ref("09_tool_dispatch_v3.svg", "v3 完整工具調度表（13 個工具）")
add_svg_ref("12_delete_calendar_flow.svg", "v3 行事曆刪除工具流程")
add_svg_ref("13_recipe_analysis_enhanced.svg", "v3 食譜分析增強流程")
add_svg_ref("14_symptom_search_v3.svg", "v3 症狀搜尋流程")
doc.add_page_break()

# ════════════════════════════════════════════════════════
#  第九章 工作流程狀態機
# ════════════════════════════════════════════════════════
h1("第九章　工作流程狀態機（v3：含中斷恢復）")

h2("9.1 抽換概念：什麼是「被打斷後能回來」？")
p("想像你在跟朋友點餐。你說「我要一個漢堡」，店員問「要加蛋嗎？」"
  "這時你突然看到甜點說「那個蛋糕多少錢？」——你被打斷了。"
  "聰明的店員會先回答蛋糕的價格，然後問「那你的漢堡要加蛋嗎？」——他記得你未完成的訂單。"
  "v3 的狀態機就有這個能力：偵測到你被打斷後，記住未完成的任務，等你回來。")

h2("9.2 專有名詞講解")
add_pedagogy_box(
    "點餐被打斷→店員記住未完成訂單→問完蛋糕後提醒「那漢堡要加蛋嗎？」",
    "工作流程狀態機（Workflow State Machine）：用 5 個狀態管理一次工具呼叫的生命週期。idle（空閒）→ collecting（收集參數）→ ready（參數齊全）→ executing（執行中）→ done（完成）。v3 新增 interrupted_state 欄位和 _detect_resume_smart() 函式。",
    "中斷偵測用「LLM + regex 雙軌」：先呼叫 PROMPT_CLASSIFY_REPLY（temp=0.0）讓 Gemma 判斷「這句話是在回答問題，還是在說別的事？」。如果 LLM 失敗，fallback 到 _detect_interruption_regex()。恢復偵測同理：PROMPT_CLASSIFY_RESUME 讓 LLM 判斷「使用者是想回到未完成的任務嗎？」。兩種恢復路徑：顯式（使用者說「回到剛才」）和隱式（使用者直接提供之前缺少的參數）。",
    "使用者：「幫我分析食譜」→ Phase1: analyze_recipe → 缺 recipe_text → 狀態: collecting → AI:「請提供食譜內容」→ 使用者：「等等，先幫我查一下鈣」→ _detect_interruption_smart(): B（不是回答）→ 儲存 interrupted_state={intent:analyze_recipe, missing:[recipe_text]} → 執行 nutrient_ranking → 使用者：「回到剛才」→ _detect_resume_smart(): A（想回去）→ 恢復 collecting → AI:「好的，回到之前的任務。請提供食譜內容。」"
)

h2("9.3 深入講解：狀態轉換圖")
add_table(
    ["當前狀態", "事件/條件", "下一狀態", "動作"],
    [
        ["idle", "Phase1 判定需要參數", "collecting", "記錄 active_intent, 初始化 collected"],
        ["collecting", "使用者提供了所有缺少的參數", "ready", "標記參數齊全"],
        ["collecting", "中斷偵測(_detect_interruption_smart)=True", "idle", "保存 interrupted_state"],
        ["collecting", "turn_count ≥ 3 仍缺參數", "idle", "放棄：「無法取得所需資訊」"],
        ["idle/done", "恢復偵測(_detect_resume_smart)=True", "collecting", "載入 interrupted_state"],
        ["idle/done", "Phase1+Phase2 生成 tool_calls", "ready→executing", "直接跳到執行"],
        ["ready", "開始執行工具", "executing", "逐個呼叫 _execute_tool_with_recovery"],
        ["executing", "所有工具完成", "done", "設定 _tool_results, _tools_used"],
        ["done", "下一輪開始", "idle", "清空狀態"],
    ]
)

h2("9.4 中斷偵測 LLM Prompt")
add_code("""# advisor.py L1993-2001 — PROMPT_CLASSIFY_REPLY
PROMPT_CLASSIFY_REPLY = '''AI 正在執行「{intent_zh}」，等使用者提供「{missing_desc}」。
使用者說：「{message}」
這句話是在回答上述問題嗎？
A = 是（提供了所需的資訊）
B = 否（在說別的事、問新問題、放棄、或換話題）
只輸出 A 或 B：'''

# advisor.py L2003-2011 — PROMPT_CLASSIFY_RESUME
PROMPT_CLASSIFY_RESUME = '''使用者之前有未完成的任務：「{intent_zh}」（還需要：{missing_desc}）。
使用者現在說：「{message}」
這句話是想回到那個未完成的任務嗎？
A = 是
B = 否（在說新的事情）
只輸出 A 或 B：'''""")

add_svg_ref("08_workflow_state_machine_v3.svg", "v3 工作流程狀態機（含中斷恢復）")
doc.add_page_break()

# ════════════════════════════════════════════════════════
#  第十章 雙記憶系統
# ════════════════════════════════════════════════════════
h1("第十章　雙記憶系統（v3：12 類結構化存儲）")

h2("10.1 抽換概念：兩本筆記本")
p("想像你有兩本筆記本。第一本是「日記」（Narrative Memory）——記錄每天發生的事，"
  "寫的是流水帳。第二本是「資料卡」（Structured Store）——像圖書館的索引卡，"
  "每張卡片記一個事實，分門別類整理。日記讓你回憶故事脈絡，資料卡讓你快速查找某個事實。")

h2("10.2 專有名詞講解")
add_pedagogy_box(
    "日記（流水帳，記故事脈絡）+ 資料卡（分類索引，記事實）",
    "雙記憶系統（Dual Memory）：每輪對話結束後，並行執行兩個 LLM 呼叫：① Narrative Summary（temp=0.2）：用 PROMPT_NARRATIVE_SUMMARY 把舊摘要+新對話壓縮成更新的摘要。② Structured Extract（temp=0.1）：用 PROMPT_STRUCTURED_EXTRACT 從對話中提取結構化資料，存入 12 類欄位。",
    "v3 的 structured_store 從 v1 的 ~6 類擴展到 12 類：user_facts、food_context、recipe_store（v3:eager update）、symptom_history（v3:新增 vector_match_count）、nutrient_ranking_history、calendar_context、academic_context、papers_referenced、dri_awareness、action_entry、saved_recipes_summary、recipe_entry。每類都有容量限制（如 symptom_history 保留最近 20 筆）。",
    "使用者查了疲勞的症狀 → Narrative:「使用者查詢疲勞相關營養素；AI 回報鐵和 B12 可能是關鍵...」 → Structured: {symptom_entry:{symptom_zh:'疲勞', parsed_symptoms:['疲勞'], effect_count:8, vector_match_count:5}, action_entry:{intent:'search_symptom', key_result_summary:'找到 8 個效果、5 個語義匹配'}}"
)

h2("10.3 深入講解：_merge_structured_store 合併邏輯")
add_code("""# advisor.py L1898-1984 — 結構化存儲合併
def _merge_structured_store(store, patch):
    now_iso = datetime.now().isoformat()
    if "user_facts" in patch:        # 覆蓋式更新
        uf = store.setdefault("user_facts", {})
        for k, v in patch["user_facts"].items():
            if v is not None: uf[k] = v
    if "recipe_entry" in patch:      # 追加式（帶時間戳）
        entry = patch["recipe_entry"]
        entry.setdefault("timestamp", now_iso)
        store.setdefault("recipe_store", []).append(entry)
    if "symptom_entry" in patch:     # 追加式 + 容量限制 20
        arr = store.setdefault("symptom_history", [])
        arr.append(patch["symptom_entry"])
        store["symptom_history"] = arr[-20:]  # 只保留最近 20 筆
    if "action_entry" in patch:      # 追加式 + 容量限制 30
        arr = store.setdefault("recent_actions", [])
        arr.append(patch["action_entry"])
        store["recent_actions"] = arr[-30:]
    # ... 其餘 8 類同理 ...""")

add_svg_ref("10_dual_memory_v3.svg", "v3 雙記憶系統")
doc.add_page_break()

# ════════════════════════════════════════════════════════
#  第十一章 所有 Prompt 完整內容
# ════════════════════════════════════════════════════════
h1("第十一章　所有 Prompt 完整內容與定位")

h2("11.1 抽換概念：什麼是 Prompt？")
p("Prompt 就像你給新員工的「工作指示」。你不會只說「去做事」，而是會詳細說明："
  "「去倉庫拿 A 類貨物，數量根據訂單表，如果缺貨就回報，不要自己亂補。」"
  "每個 Prompt 都是給 AI 的一份「任務說明書」，寫清楚它該做什麼、不該做什麼、格式要求等。")

h2("11.2 Prompt 總覽")
add_table(
    ["#", "Prompt 名稱", "程式碼位置", "Pipeline 步驟", "Temperature", "用途"],
    [
        ["1", "PROMPT_ECHO", "L45-79", "Step 1", "0.5", "即時回應，不回答問題"],
        ["2", "PROMPT_PHASE1", "L86-166", "Step 2", "0.1", "意圖判定（含系統說明書）"],
        ["3", "_PROMPT_PHASE1_RETRY", "L820-830", "Step 2 重試", "0.1", "Phase1 格式錯誤時的簡化版"],
        ["4", "PROMPT_PHASE2", "L173-231", "Step 3", "0.1", "工具指令生成"],
        ["5", "PROMPT_ANALYSIS", "L238-278", "Step 5", "0.3", "結果分析回覆"],
        ["6", "PROMPT_NARRATIVE_SUMMARY", "L285-307", "Step 6a", "0.2", "對話脈絡摘要"],
        ["7", "PROMPT_STRUCTURED_EXTRACT", "L309-366", "Step 6b", "0.1", "結構化資料提取"],
        ["8", "PROMPT_EXTRACT_PARAMS", "L2265-2273", "參數提取", "0.1", "regex 失敗時的 LLM 備援"],
        ["9", "PROMPT_CLASSIFY_REPLY", "L1993-2001", "中斷偵測", "0.0", "判斷使用者是否在回答問題"],
        ["10", "PROMPT_CLASSIFY_RESUME", "L2003-2011", "恢復偵測", "0.0", "判斷使用者是否想回到未完成任務"],
    ]
)

h2("11.3 各 Prompt 完整內容")

h3("Prompt 1: PROMPT_ECHO（L45-79）")
add_code("""PROMPT_ECHO = \"\"\"請你根據使用者的訊息和對話脈絡，用繁體中文回應。

## 規則
- 簡潔回應
- 如果使用者的訊息需要查詢工具（記錄飲食、查營養、查症狀等）：概括需求，表達「已收到、正在處理」
- 如果是一般對話（打招呼、感謝、閒聊）：直接給出友善回應
- 不要回答營養問題本身，不要提供營養數據或建議
- 語氣親切自然

## 範例
使用者：「營養夠嗎」→ 收到，我先查看你的個人資料，再幫你分析今天的營養攝取是否達標。
使用者：「最近常覺得疲勞」→ 了解，我來查詢與疲勞相關的營養素和推薦食物。
使用者：「你好」→ 你好！我是你的 AI 營養顧問，有什麼飲食或營養問題可以幫你的嗎？

{memory_block}
對話歷史：{history}
使用者訊息：{message}
請直接輸出回應：\"\"\" """)

h3("Prompt 2: PROMPT_PHASE1（L86-166）")
p("最長的 Prompt，包含完整的系統說明書（SYSTEM_MANUAL）。定義 14 種意圖、營養素欄位對照、確認詞處理規則。")
p("v3 新增：注入 {narrative_memory}（對話脈絡）和 {long_term_context}（記憶樹長期記憶）。"
  "新增 delete_calendar intent。symptom 參數改為中文。")

h3("Prompt 5: PROMPT_ANALYSIS（L238-278）")
p("分析 Prompt 的核心規則：「所有數據必須來自工具結果，禁止從訓練資料引用任何營養數字」。"
  "「前端已自動渲染工具結果的表格和卡片，你不需要列出原始數據清單」。"
  "「search_graphrag 的 answer 和 evidence 必須完整呈現，禁止摘要、禁止縮減」。")

h3("Prompt 7: PROMPT_STRUCTURED_EXTRACT（L309-366）")
p("v3 更新：symptom_entry 格式新增 vector_match_count 欄位。明確要求輸出 JSON 物件（禁止 JSON 陣列）。"
  "recipe_entry 的 total_nutrition 明確要求「有多少記多少，寧可多記不可少記」。")

add_svg_ref("11_prompt_positions_v3.svg", "v3 Prompt 定位圖")
doc.add_page_break()

# ════════════════════════════════════════════════════════
#  第十二章 v1→v3 差異總覽
# ════════════════════════════════════════════════════════
h1("第十二章　v1 → v3 差異總覽")

h2("12.1 檔案層級變更")
add_table(
    ["層級", "變更", "說明"],
    [
        ["advisor.py", "1876行 → 2969行 (+1093行)", "核心檔案，新增自癒、記憶樹、認知、session、checkpoint 整合"],
        ["新增 recovery.py", "237 行", "工具自癒引擎（ReIn）"],
        ["新增 cognitive_engine.py", "97 行", "認知重建引擎（CRE）"],
        ["新增 embedding_service.py", "35 行", "嵌入向量代理（呼叫 port 8003）"],
        ["新增 memory_tree_repository.py", "192 行", "記憶樹 CRUD + 向量搜尋"],
        ["新增 conversation_repository.py", "152 行", "會話 + checkpoint 管理"],
        ["新增 cognitive_repository.py", "87 行", "認知事實存儲"],
        ["新增 memory_tree.py (router)", "106 行", "記憶樹 REST API"],
        ["新增 conversation.py (router)", "89 行", "會話管理 REST API"],
        ["新增 cognitive.py (router)", "55 行", "認知引擎 REST API"],
        ["新增 4 份 migration", "107 行", "memory_tree, vector, sessions, cognitive_facts 表"],
        ["更新 advisorApi.js", "+18 行", "新增 4 個 SSE callback (recovery/session/checkpoint)"],
    ]
)

h2("12.2 功能層級變更")
add_table(
    ["功能", "v1", "v3"],
    [
        ["工具數量", "12 個", "13 個（+delete_calendar_entry）"],
        ["SSE 事件數", "9 種", "15 種（+6 種）"],
        ["意圖數量", "13 種", "14 種（+delete_calendar）"],
        ["Prompt 數量", "6 個", "10 個（+RETRY, CLASSIFY_REPLY, CLASSIFY_RESUME, EXTRACT_PARAMS）"],
        ["記憶系統", "Narrative + Structured（前端暫存）", "+ 記憶樹 5 層 + 認知事實 + 語義搜尋"],
        ["工具失敗處理", "回傳 error 訊息", "自動診斷→策略→重試（ReIn 引擎）"],
        ["會話持久化", "無（重整即失）", "UUID session + Checkpoint 快照"],
        ["症狀搜尋", "英文靜態對照", "中文直傳 + 5 階段管道 + 語義搜尋"],
        ["食譜分析", "3 步（extract→lookup→calculate）", "4+1 步（+文字前處理 + DRI 缺口）"],
        ["中斷偵測", "純 regex", "LLM + regex 雙軌"],
        ["恢復偵測", "純 regex（_RETURN_RE）", "LLM + 確定性參數探測"],
        ["工具結果裁剪", "無（完整送入 LLM）", "_trim_tool_results_for_analysis() 節省 token"],
        ["Phase1 上下文", "history + narrative_memory", "+ long_term_context（記憶樹）"],
        ["Phase2 上下文", "structured_store + narrative_memory", "+ long_term_context"],
    ]
)

h2("12.3 架構層級變更")
add_table(
    ["架構面", "v1", "v3"],
    [
        ["Python 服務模組", "1 個（advisor.py）", "4 個（+recovery, cognitive_engine, embedding_service）"],
        ["Backend Repository", "基本 CRUD", "+3 個（memory_tree, conversation, cognitive）"],
        ["Backend Router", "基本路由", "+3 個（memory_tree, conversation, cognitive）"],
        ["資料庫表", "users, recipes, etc.", "+4 表（memory_tree_nodes, conversation_sessions, conversation_checkpoints, cognitive_facts）"],
        ["向量索引", "無", "pgvector ivfflat (1024 維, 50 lists)"],
        ["外部服務依賴", "Gemma LLM", "+Health Vector（port 8003 嵌入向量）"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════
#  附錄 A 前端 SSE
# ════════════════════════════════════════════════════════
h1("附錄 A　前端 SSE 客戶端程式碼解析")

h2("A.1 抽換概念")
p("前端的 SSE 客戶端就像一個「快遞收件員」——它站在門口，每來一個包裹（事件），"
  "就看包裹上的標籤（event type），然後送到正確的部門（callback 函式）。")

h2("A.2 advisorApi.js 完整解析")
add_code("""// advisorApi.js — SSE 串流連線（124 行）
export function streamAdvisorChat(message, history, userProfile, callbacks,
    memory="", workflowState={}, narrativeMemory="", structuredStore={},
    sessionId="", resumeCheckpoint=false) {  // v3 新增最後兩個參數

    const body = JSON.stringify({
        message, history, user_profile: userProfile,
        memory, workflow_state: workflowState,
        narrative_memory: narrativeMemory,
        structured_store: structuredStore,
        session_id: sessionId,              // v3 新增
        resume_checkpoint: resumeCheckpoint, // v3 新增
    });

    fetch("/api/advisor/chat", {
        method: "POST",
        headers: { "Content-Type": "application/json", Authorization: `Bearer ${token}` },
        body, signal: controller.signal,
    }).then(async (response) => {
        const reader = response.body.getReader();
        const decoder = new TextDecoder();
        let buffer = "", eventType = null;

        while (true) {
            const { done, value } = await reader.read();
            if (done) break;
            buffer += decoder.decode(value, { stream: true });
            const lines = buffer.split("\\n");
            buffer = lines.pop() || "";
            for (const line of lines) {
                if (line.startsWith("event: "))
                    eventType = line.slice(7).trim();
                else if (line.startsWith("data: ") && eventType) {
                    const data = JSON.parse(line.slice(6));
                    switch (eventType) {
                        case "thinking": callbacks.onThinking?.(data); break;
                        case "echo": callbacks.onEcho?.(data); break;
                        case "tool_start": callbacks.onToolStart?.(data); break;
                        case "tool_done": callbacks.onToolDone?.(data); break;
                        case "tool_result": callbacks.onToolResult?.(data); break;
                        case "reply": callbacks.onReply?.(data); break;
                        case "memory": callbacks.onMemory?.(data); break;
                        case "workflow_state": callbacks.onWorkflowState?.(data); break;
                        case "phase_status": callbacks.onPhaseStatus?.(data); break;
                        case "narrative_memory": callbacks.onNarrativeMemory?.(data); break;
                        case "structured_store": callbacks.onStructuredStore?.(data); break;
                        // v3 新增 4 個事件：
                        case "session_id": callbacks.onSessionId?.(data); break;
                        case "checkpoint_saved": callbacks.onCheckpointSaved?.(data); break;
                        case "recovery_attempt": callbacks.onRecoveryAttempt?.(data); break;
                        case "recovery_result": callbacks.onRecoveryResult?.(data); break;
                        case "error": callbacks.onError?.(data.message); break;
                    }
                    eventType = null;
                }
            }
        }
        callbacks.onComplete?.();
    });
    return { abort: () => controller.abort() };
}""")

doc.add_page_break()

# ════════════════════════════════════════════════════════
#  附錄 B SVG 圖表索引
# ════════════════════════════════════════════════════════
h1("附錄 B　SVG 圖表索引")
add_table(
    ["#", "檔名", "內容", "章節"],
    [
        ["1", "01_system_architecture_v3.svg", "v3 系統架構總覽", "第一章"],
        ["2", "02_llm_pipeline_v3.svg", "v3 六步驟 LLM Pipeline", "第二章"],
        ["3", "03_sse_protocol_v3.svg", "v3 SSE 通信協議（15 種事件）", "第三章"],
        ["4", "04_recovery_system.svg", "v3 工具自癒引擎（ReIn）", "第四章"],
        ["5", "05_memory_tree_system.svg", "v3 記憶樹（TiMem）", "第五章"],
        ["6", "06_cognitive_engine.svg", "v3 認知重建引擎（CRE）", "第六章"],
        ["7", "07_session_checkpoint.svg", "v3 會話管理與 Checkpoint", "第七章"],
        ["8", "08_workflow_state_machine_v3.svg", "v3 工作流程狀態機", "第九章"],
        ["9", "09_tool_dispatch_v3.svg", "v3 完整工具調度表", "第八章"],
        ["10", "10_dual_memory_v3.svg", "v3 雙記憶系統", "第十章"],
        ["11", "11_prompt_positions_v3.svg", "v3 Prompt 定位圖", "第十一章"],
        ["12", "12_delete_calendar_flow.svg", "v3 行事曆刪除流程", "第八章"],
        ["13", "13_recipe_analysis_enhanced.svg", "v3 食譜分析增強流程", "第八章"],
        ["14", "14_symptom_search_v3.svg", "v3 症狀搜尋流程", "第八章"],
    ]
)

# ══ 儲存 ══
doc.save(str(DOCX_PATH))
print(f"✅ 報告已儲存到: {DOCX_PATH}")
print(f"   檔案大小: {DOCX_PATH.stat().st_size / 1024:.1f} KB")
print(f"   段落數: {len(doc.paragraphs)}")
print(f"   表格數: {len(doc.tables)}")
