#!/usr/bin/env python3
"""
生成 Tool Calling & Agent 機制完整教學 Word 文件
包含 22+ 張 SVG 圖表、逐行程式碼解釋、完整通信鏈
"""
import io, os, re, textwrap
import cairosvg
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

BASE = "/home/aiiauser/JM/0325new"
OUT = os.path.join(BASE, "Tool_Calling_Agent_完整教學.docx")
FONT_ZH = "Noto Sans CJK TC"
FONT_EN = "Arial, Helvetica, sans-serif"
FONT_SVG = f"{FONT_ZH}, {FONT_EN}"

# ═══════════════════════════════════════════════════════════
# Section A: SVG Helpers
# ═══════════════════════════════════════════════════════════

COMMON_DEFS = '''
<marker id="arr-blue" markerWidth="10" markerHeight="7" refX="9" refY="3.5" orient="auto">
  <polygon points="0 0, 10 3.5, 0 7" fill="#3182CE"/></marker>
<marker id="arr-green" markerWidth="10" markerHeight="7" refX="9" refY="3.5" orient="auto">
  <polygon points="0 0, 10 3.5, 0 7" fill="#38A169"/></marker>
<marker id="arr-red" markerWidth="10" markerHeight="7" refX="9" refY="3.5" orient="auto">
  <polygon points="0 0, 10 3.5, 0 7" fill="#E53E3E"/></marker>
<marker id="arr-orange" markerWidth="10" markerHeight="7" refX="9" refY="3.5" orient="auto">
  <polygon points="0 0, 10 3.5, 0 7" fill="#DD6B20"/></marker>
<marker id="arr-purple" markerWidth="10" markerHeight="7" refX="9" refY="3.5" orient="auto">
  <polygon points="0 0, 10 3.5, 0 7" fill="#805AD5"/></marker>
<marker id="arr-teal" markerWidth="10" markerHeight="7" refX="9" refY="3.5" orient="auto">
  <polygon points="0 0, 10 3.5, 0 7" fill="#319795"/></marker>
<marker id="arr-gray" markerWidth="10" markerHeight="7" refX="9" refY="3.5" orient="auto">
  <polygon points="0 0, 10 3.5, 0 7" fill="#718096"/></marker>
<marker id="arrd-blue" markerWidth="7" markerHeight="10" refX="3.5" refY="9" orient="auto">
  <polygon points="0 0, 7 0, 3.5 10" fill="#3182CE"/></marker>
<marker id="arrd-green" markerWidth="7" markerHeight="10" refX="3.5" refY="9" orient="auto">
  <polygon points="0 0, 7 0, 3.5 10" fill="#38A169"/></marker>
<marker id="arrd-red" markerWidth="7" markerHeight="10" refX="3.5" refY="9" orient="auto">
  <polygon points="0 0, 7 0, 3.5 10" fill="#E53E3E"/></marker>
<marker id="arrd-orange" markerWidth="7" markerHeight="10" refX="3.5" refY="9" orient="auto">
  <polygon points="0 0, 7 0, 3.5 10" fill="#DD6B20"/></marker>
<marker id="arrd-purple" markerWidth="7" markerHeight="10" refX="3.5" refY="9" orient="auto">
  <polygon points="0 0, 7 0, 3.5 10" fill="#805AD5"/></marker>
<marker id="arrd-teal" markerWidth="7" markerHeight="10" refX="3.5" refY="9" orient="auto">
  <polygon points="0 0, 7 0, 3.5 10" fill="#319795"/></marker>
<marker id="arrd-gray" markerWidth="7" markerHeight="10" refX="3.5" refY="9" orient="auto">
  <polygon points="0 0, 7 0, 3.5 10" fill="#718096"/></marker>
'''

def svg_wrap(content, w, h):
    return f'''<svg xmlns="http://www.w3.org/2000/svg" width="{w}" height="{h}" viewBox="0 0 {w} {h}">
<defs>{COMMON_DEFS}
<style>
text {{ font-family: {FONT_SVG}; }}
.title {{ font-size: 18px; font-weight: bold; fill: #1a202c; }}
.subtitle {{ font-size: 15px; font-weight: bold; fill: #2d3748; }}
.label {{ font-size: 13px; fill: #2d3748; }}
.label-bold {{ font-size: 13px; fill: #2d3748; font-weight: bold; }}
.label-sm {{ font-size: 11px; fill: #4a5568; }}
.label-xs {{ font-size: 10px; fill: #718096; }}
.code {{ font-family: monospace; font-size: 11px; fill: #2d3748; }}
.white {{ font-size: 13px; fill: #ffffff; font-weight: bold; }}
.white-sm {{ font-size: 11px; fill: #ffffff; }}
.accent {{ font-size: 13px; fill: #e53e3e; font-weight: bold; }}
.good {{ font-size: 13px; fill: #38a169; font-weight: bold; }}
.num {{ font-size: 22px; fill: #ffffff; font-weight: bold; }}
</style>
</defs>
{content}
</svg>'''

def rr(x, y, w, h, fill="#EBF8FF", stroke="#3182CE", rx=8):
    return f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="{rx}" fill="{fill}" stroke="{stroke}" stroke-width="1.5"/>'

def rr_bold(x, y, w, h, fill="#EBF8FF", stroke="#3182CE", rx=8):
    return f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="{rx}" fill="{fill}" stroke="{stroke}" stroke-width="2.5"/>'

def circle(cx, cy, r, fill="#3182CE"):
    return f'<circle cx="{cx}" cy="{cy}" r="{r}" fill="{fill}"/>'

def tc(x, y, txt, cls="label"):
    return f'<text x="{x}" y="{y}" text-anchor="middle" class="{cls}">{txt}</text>'

def tl(x, y, txt, cls="label"):
    return f'<text x="{x}" y="{y}" class="{cls}">{txt}</text>'

def tr(x, y, txt, cls="label"):
    return f'<text x="{x}" y="{y}" text-anchor="end" class="{cls}">{txt}</text>'

def arr_r(x1, y, x2, color="blue", label="", ly=-12):
    parts = [f'<line x1="{x1}" y1="{y}" x2="{x2}" y2="{y}" stroke="{"#3182CE" if color=="blue" else "#38A169" if color=="green" else "#E53E3E" if color=="red" else "#DD6B20" if color=="orange" else "#805AD5" if color=="purple" else "#319795" if color=="teal" else "#718096"}" stroke-width="2" marker-end="url(#arr-{color})"/>']
    if label:
        mx = (x1+x2)/2
        parts.append(tc(mx, y+ly, label, "label-sm"))
    return "\n".join(parts)

def arr_d(x, y1, y2, color="blue", label="", lx=10):
    c = {"blue":"#3182CE","green":"#38A169","red":"#E53E3E","orange":"#DD6B20","purple":"#805AD5","teal":"#319795","gray":"#718096"}.get(color,"#3182CE")
    parts = [f'<line x1="{x}" y1="{y1}" x2="{x}" y2="{y2}" stroke="{c}" stroke-width="2" marker-end="url(#arrd-{color})"/>']
    if label:
        my = (y1+y2)/2
        parts.append(tl(x+lx, my+4, label, "label-sm"))
    return "\n".join(parts)

def arr_l(x1, y, x2, color="green", label="", ly=15):
    c = {"blue":"#3182CE","green":"#38A169","red":"#E53E3E","orange":"#DD6B20","purple":"#805AD5","teal":"#319795","gray":"#718096"}.get(color,"#38A169")
    parts = [f'<line x1="{x1}" y1="{y}" x2="{x2}" y2="{y}" stroke="{c}" stroke-width="2" marker-end="url(#arr-{color})"/>']
    if label:
        mx = (x1+x2)/2
        parts.append(tc(mx, y+ly, label, "label-sm"))
    return "\n".join(parts)

def dash_d(x, y1, y2, color="#718096"):
    return f'<line x1="{x}" y1="{y1}" x2="{x}" y2="{y2}" stroke="{color}" stroke-width="1.5" stroke-dasharray="6,4"/>'

def dash_r(x1, y, x2, color="#718096"):
    return f'<line x1="{x1}" y1="{y}" x2="{x2}" y2="{y}" stroke="{color}" stroke-width="1.5" stroke-dasharray="6,4"/>'

def svg_to_png(svg_str, scale=2):
    return io.BytesIO(cairosvg.svg2png(bytestring=svg_str.encode("utf-8"), scale=scale))


# ═══════════════════════════════════════════════════════════
# Section B: Word Document Helpers
# ═══════════════════════════════════════════════════════════

doc = Document()

section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)

style_normal = doc.styles['Normal']
style_normal.font.name = FONT_ZH
style_normal.font.size = Pt(11)
style_normal.paragraph_format.space_after = Pt(6)
style_normal.paragraph_format.line_spacing = 1.4
style_normal.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_ZH)

for lvl, sz, clr in [('Heading 1', 22, '1A365D'), ('Heading 2', 16, '2B6CB0'), ('Heading 3', 13, '2C7A7B')]:
    s = doc.styles[lvl]
    s.font.size = Pt(sz)
    s.font.color.rgb = RGBColor.from_string(clr)
    s.font.bold = True
    s.font.name = FONT_ZH
    s.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_ZH)
    s.paragraph_format.space_before = Pt(20 if sz > 15 else 12)
    s.paragraph_format.space_after = Pt(8)

code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = 'Consolas'
code_style.font.size = Pt(9)
code_style.paragraph_format.space_before = Pt(2)
code_style.paragraph_format.space_after = Pt(2)
code_style.paragraph_format.left_indent = Cm(0.5)
code_style.paragraph_format.element.append(
    code_style.paragraph_format.element.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'F0F0F0'
    })
)

def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
    shading.append(shd)

def h1(text):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.name = FONT_ZH
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_ZH)

def h2(text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.name = FONT_ZH
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_ZH)

def h3(text):
    h = doc.add_heading(text, level=3)
    for run in h.runs:
        run.font.name = FONT_ZH
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_ZH)

def p(text, bold=False, italic=False, size=11, color=None, align=None, space_after=6):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.name = FONT_ZH
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_ZH)
    if bold: run.bold = True
    if italic: run.italic = True
    if color: run.font.color.rgb = RGBColor(*color)
    para.paragraph_format.space_after = Pt(space_after)
    if align: para.alignment = align
    return para

def pb(text): return p(text, bold=True)
def pi(text): return p(text, italic=True, size=10, color=(0x71, 0x80, 0x96))

def code(text):
    for line in text.strip().split('\n'):
        para = doc.add_paragraph(style='CodeBlock')
        run = para.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

def bullet(text):
    para = doc.add_paragraph(style='List Bullet')
    run = para.runs[0] if para.runs else para.add_run(text)
    if not para.runs:
        pass
    else:
        run.font.name = FONT_ZH
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_ZH)
    para.text = text
    return para

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for par in cell.paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in par.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = FONT_ZH
                run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_ZH)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            cell.text = str(val)
            for par in cell.paragraphs:
                for run in par.runs:
                    run.font.size = Pt(10)
                    run.font.name = FONT_ZH
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_ZH)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

def add_image(svg_str, width=Inches(6.2)):
    png = svg_to_png(svg_str)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(png, width=width)
    doc.add_paragraph()

def tip_box(title, content):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, "EBF8FF")
    pp = cell.paragraphs[0]
    run = pp.add_run(f"💡 {title}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = FONT_ZH
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_ZH)
    p2 = cell.add_paragraph()
    run2 = p2.add_run(content)
    run2.font.size = Pt(10)
    run2.font.name = FONT_ZH
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_ZH)
    doc.add_paragraph()

def warn_box(content):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, "FFF5F5")
    pp = cell.paragraphs[0]
    run = pp.add_run(f"⚠️ {content}")
    run.font.size = Pt(10)
    run.font.name = FONT_ZH
    run.font.color.rgb = RGBColor(0xC5, 0x30, 0x30)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_ZH)
    doc.add_paragraph()

def analogy_box(title, content):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, "F0FFF4")
    pp = cell.paragraphs[0]
    run = pp.add_run(f"🍳 {title}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = FONT_ZH
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_ZH)
    p2 = cell.add_paragraph()
    run2 = p2.add_run(content)
    run2.font.size = Pt(10)
    run2.font.name = FONT_ZH
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_ZH)
    doc.add_paragraph()

def prompt_box(title, content):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, "FAF5FF")
    pp = cell.paragraphs[0]
    run = pp.add_run(f"📋 Prompt: {title}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = FONT_ZH
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_ZH)
    for line in content.strip().split('\n'):
        cp = cell.add_paragraph()
        cr = cp.add_run(line)
        cr.font.size = Pt(9)
        cr.font.name = 'Consolas'
    doc.add_paragraph()

def page_break():
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════
# Section C: SVG Diagrams
# ═══════════════════════════════════════════════════════════

def make_svg_architecture():
    """圖 1: 系統架構全貌"""
    w, h = 920, 560
    content = f'''
    {tc(460, 30, "Recipe-AI 系統架構全貌", "title")}

    <!-- Frontend -->
    {rr_bold(30, 60, 200, 90, "#EBF8FF", "#3182CE")}
    {tc(130, 95, "Frontend", "label-bold")}
    {tc(130, 115, "React + Vite", "label-sm")}
    {tc(130, 132, "Port 3001", "label-xs")}

    <!-- Control -->
    {rr_bold(310, 60, 200, 90, "#F0FFF4", "#38A169")}
    {tc(410, 95, "Control", "label-bold")}
    {tc(410, 115, "FastAPI + Advisor", "label-sm")}
    {tc(410, 132, "Port 8000", "label-xs")}

    <!-- Backend -->
    {rr_bold(590, 60, 200, 90, "#FFFAF0", "#DD6B20")}
    {tc(690, 95, "Backend", "label-bold")}
    {tc(690, 115, "FastAPI + asyncpg", "label-sm")}
    {tc(690, 132, "Port 8001", "label-xs")}

    <!-- Arrows top row -->
    {arr_r(230, 95, 308, "blue", "API + SSE", -14)}
    {arr_r(510, 95, 588, "green", "httpx proxy", -14)}

    <!-- LLM -->
    {rr_bold(310, 200, 200, 80, "#FFF5F5", "#E53E3E")}
    {tc(410, 235, "Gemma LLM", "label-bold")}
    {tc(410, 255, "llama-server :8080", "label-sm")}

    <!-- GraphRAG -->
    {rr_bold(590, 200, 200, 80, "#FAF5FF", "#805AD5")}
    {tc(690, 235, "GraphRAG API", "label-bold")}
    {tc(690, 255, "FastAPI :8002", "label-sm")}

    <!-- Health Vector -->
    {rr_bold(590, 320, 200, 80, "#E6FFFA", "#319795")}
    {tc(690, 355, "Health Vector", "label-bold")}
    {tc(690, 375, "ChromaDB :8003", "label-sm")}

    <!-- PostgreSQL -->
    {rr_bold(30, 200, 200, 80, "#FFF5F5", "#E53E3E", 12)}
    {tc(130, 230, "PostgreSQL", "label-bold")}
    {tc(130, 255, "recipe_ai_db :5432", "label-sm")}

    <!-- Arrows -->
    {arr_d(410, 150, 198, "red", "LLM 呼叫")}
    {arr_d(690, 150, 198, "purple", "學術查詢")}
    {arr_d(690, 280, 318, "teal", "向量搜尋")}
    {arr_d(130, 150, 198, "red", "SQL 查詢")}

    <!-- Analogy Box -->
    {rr(30, 430, 860, 110, "#FFFFF0", "#D69E2E", 10)}
    {tl(50, 460, "🍳 餐廳比喻", "subtitle")}
    {tl(50, 485, "Frontend = 用餐區（客人看到的畫面）    Control = 主廚（負責調度所有工作站）", "label-sm")}
    {tl(50, 505, "Backend = 食材倉庫（存放所有食品資料）  LLM = 主廚的大腦（負責思考和判斷）", "label-sm")}
    {tl(50, 525, "GraphRAG = 圖書館（查學術論文）         Health Vector = 症狀百科（語義搜尋症狀）", "label-sm")}
    '''
    return svg_wrap(content, w, h)


def make_svg_pipeline_overview():
    """圖 2: 六步驟 LLM Pipeline"""
    w, h = 850, 920
    steps = [
        ("1", "Echo", "即時回應", "temp=0.5", "#3182CE", "#EBF8FF", "SSE: echo"),
        ("2", "Phase 1", "意圖判定", "temp=0.1", "#38A169", "#F0FFF4", "SSE: phase_status"),
        ("3", "Phase 2", "工具調用生成", "temp=0.1", "#2C7A7B", "#E6FFFA", "系統優先 → LLM 備援"),
        ("4", "Execute", "工具執行", "—", "#DD6B20", "#FFFAF0", "SSE: tool_start/result/done"),
        ("5", "Analysis", "分析回覆", "temp=0.3", "#805AD5", "#FAF5FF", "SSE: thinking + reply"),
        ("6", "DualMemory", "雙記憶更新", "temp=0.2", "#E53E3E", "#FFF5F5", "背景非同步"),
    ]
    content = tc(425, 30, "六步驟 LLM Pipeline — 使用者每句話的旅程", "title")
    y = 60
    for i, (num, name, desc, temp, stroke, fill, note) in enumerate(steps):
        content += f'''
        {circle(80, y+40, 22, stroke)}
        {tc(80, y+47, num, "num")}
        {rr_bold(120, y+5, 320, 70, fill, stroke)}
        {tl(140, y+35, name, "label-bold")}
        {tl(140, y+55, desc, "label-sm")}
        {rr(470, y+15, 100, 50, "#F7FAFC", "#A0AEC0")}
        {tc(520, y+45, temp, "label-sm")}
        {tl(600, y+45, note, "label-xs")}
        '''
        if i < len(steps) - 1:
            content += arr_d(80, y+62, y+135, "gray")
        y += 130

    content += f'''
    {rr(30, y+20, 790, 70, "#FFFFF0", "#D69E2E", 10)}
    {tl(50, y+50, "🍳 想像你走進一間高級餐廳：你說「我要牛排」→ 服務生說「收到」(Echo) → 判斷你要什麼 (Phase1)", "label-sm")}
    {tl(50, y+70, "   → 派廚師準備 (Phase2) → 實際烹飪 (Execute) → 上菜並說明 (Analysis) → 記在服務記錄上 (Memory)", "label-sm")}
    '''
    return svg_wrap(content, w, h)


def make_svg_phase1_decision():
    """圖 3: Phase1 意圖判定決策樹"""
    w, h = 880, 620
    content = f'''
    {tc(440, 30, "Phase 1: 意圖判定 — 使用者想做什麼？", "title")}

    <!-- Input -->
    {rr_bold(320, 55, 240, 50, "#EBF8FF", "#3182CE")}
    {tc(440, 85, "使用者訊息進入", "label-bold")}

    <!-- Fast Path -->
    {arr_d(440, 105, 140, "blue")}
    {rr(280, 145, 320, 50, "#FFF5F5", "#E53E3E")}
    {tc(440, 175, "Fast Path 正則比對？", "label-bold")}

    {tl(610, 170, "命中 → 跳過 LLM", "good")}
    {arr_r(600, 170, 610, "green")}
    {rr(620, 150, 230, 45, "#F0FFF4", "#38A169")}
    {tl(640, 178, "直接產生工具調用", "label-sm")}

    <!-- RAG -->
    {arr_d(440, 195, 230, "blue", "未命中")}
    {rr(280, 235, 320, 50, "#E6FFFA", "#319795")}
    {tc(440, 265, "Manual RAG 搜尋說明書", "label")}

    <!-- LLM -->
    {arr_d(440, 285, 320, "blue")}
    {rr_bold(280, 325, 320, 60, "#F0FFF4", "#38A169")}
    {tc(440, 350, "LLM 意圖判定", "label-bold")}
    {tc(440, 370, "PROMPT_PHASE1 + RAG 上下文", "label-xs")}

    <!-- Validate -->
    {arr_d(440, 385, 420, "green")}
    {rr(280, 425, 320, 50, "#FFFAF0", "#DD6B20")}
    {tc(440, 455, "驗證 + 別名修正", "label")}

    <!-- Output -->
    {arr_d(440, 475, 510, "orange")}
    {rr_bold(280, 515, 320, 55, "#FAF5FF", "#805AD5")}
    {tc(440, 538, "intents + entities", "label-bold")}
    {tc(440, 555, "輸出 JSON", "label-xs")}

    <!-- Side: 15 intents list -->
    {rr(20, 310, 230, 290, "#F7FAFC", "#A0AEC0", 6)}
    {tl(35, 335, "15 個合法意圖：", "label-bold")}
    {tl(35, 355, "search_symptom    搜尋症狀", "label-xs")}
    {tl(35, 370, "search_graphrag   學術查詢", "label-xs")}
    {tl(35, 385, "nutrient_ranking  營養素排名", "label-xs")}
    {tl(35, 400, "analyze_recipe    分析食譜", "label-xs")}
    {tl(35, 415, "search_food       搜尋食物", "label-xs")}
    {tl(35, 430, "get_calendar      行事曆查詢", "label-xs")}
    {tl(35, 445, "add_to_calendar   加入行事曆", "label-xs")}
    {tl(35, 460, "delete_calendar   刪除行事曆", "label-xs")}
    {tl(35, 475, "delete_recipe     刪除食譜", "label-xs")}
    {tl(35, 490, "get_recipes       查看食譜", "label-xs")}
    {tl(35, 505, "browse_food_db    瀏覽食品庫", "label-xs")}
    {tl(35, 520, "browse_literature 瀏覽論文", "label-xs")}
    {tl(35, 535, "get/update_profile 個人資料", "label-xs")}
    {tl(35, 550, "synthesize_advice 綜合建議", "label-xs")}
    {tl(35, 565, "general_chat      一般聊天", "label-xs")}
    '''
    return svg_wrap(content, w, h)


def make_svg_phase2_cascade():
    """圖 4: Phase2 三層瀑布"""
    w, h = 820, 420
    content = f'''
    {tc(410, 30, "Phase 2: 工具調用生成 — 三層安全網", "title")}

    <!-- Layer 1 -->
    {rr_bold(50, 65, 220, 70, "#F0FFF4", "#38A169")}
    {tc(160, 95, "第一層：系統確定性", "label-bold")}
    {tc(160, 118, "_system_generate_toolcalls()", "code")}

    {arr_r(270, 100, 330, "green", "成功？")}
    {rr(340, 75, 80, 50, "#F0FFF4", "#38A169")}
    {tc(380, 105, "完成", "good")}

    <!-- Layer 2 -->
    {arr_d(160, 135, 175, "gray", "失敗")}
    {rr_bold(50, 180, 220, 70, "#FAF5FF", "#805AD5")}
    {tc(160, 210, "第二層：LLM 備援", "label-bold")}
    {tc(160, 233, "PROMPT_PHASE2", "code")}

    {arr_r(270, 215, 330, "purple", "成功？")}
    {rr(340, 190, 80, 50, "#FAF5FF", "#805AD5")}
    {tc(380, 220, "完成", "good")}

    <!-- Layer 3 -->
    {arr_d(160, 250, 290, "gray", "失敗")}
    {rr_bold(50, 295, 220, 70, "#FFFAF0", "#DD6B20")}
    {tc(160, 325, "第三層：系統備援", "label-bold")}
    {tc(160, 348, "_system_fallback()", "code")}

    <!-- Explanation -->
    {rr(470, 65, 320, 300, "#FFFFF0", "#D69E2E", 10)}
    {tl(490, 95, "🍳 為什麼要三層？", "subtitle")}
    {tl(490, 120, "第一層像「菜單自動點餐機」：", "label-sm")}
    {tl(490, 140, "你說「我要牛排」→ 機器直接下單", "label-xs")}
    {tl(490, 165, "第二層像「服務生幫你判斷」：", "label-sm")}
    {tl(490, 185, "你說「我想吃清淡的」→ LLM 思考", "label-xs")}
    {tl(490, 210, "第三層像「經理出來處理」：", "label-sm")}
    {tl(490, 230, "前兩層都搞不定 → 用最簡單的預設", "label-xs")}
    {tl(490, 265, "好處：", "label-bold")}
    {tl(490, 285, "• 90% 的問題第一層就解決（快！）", "label-xs")}
    {tl(490, 305, "• LLM 掛掉也不怕（有備援）", "label-xs")}
    {tl(490, 325, "• 永遠不會「什麼都做不了」", "label-xs")}
    '''
    return svg_wrap(content, w, h)


def make_svg_dag_execution():
    """圖 5: DAG 並行執行"""
    w, h = 880, 520
    content = f'''
    {tc(440, 30, "工具執行：DAG 並行排程", "title")}

    <!-- User msg -->
    {rr_bold(240, 55, 400, 45, "#EBF8FF", "#3182CE")}
    {tc(440, 83, "使用者：「分析雞胸肉食譜，加到今天行事曆」", "label-sm")}

    <!-- Phase1 output -->
    {arr_d(440, 100, 130, "blue")}
    {rr(300, 135, 280, 40, "#F0FFF4", "#38A169")}
    {tc(440, 160, "intents: [analyze_recipe, add_to_calendar]", "code")}

    <!-- DAG planning -->
    {arr_d(440, 175, 210, "green")}
    {rr(300, 215, 280, 40, "#FFFAF0", "#DD6B20")}
    {tc(440, 240, "_plan_task_dag() 建立相依圖", "label")}

    <!-- Task nodes -->
    {arr_d(350, 255, 295, "orange")}
    {rr_bold(200, 300, 240, 55, "#FFFAF0", "#DD6B20")}
    {tc(320, 325, "analyze_recipe", "label-bold")}
    {tc(320, 342, "（先跑 — 無相依）", "label-xs")}

    {arr_r(440, 327, 530, "orange", "相依：需要食譜結果")}
    {rr_bold(540, 300, 240, 55, "#E6FFFA", "#319795")}
    {tc(660, 325, "add_to_calendar", "label-bold")}
    {tc(660, 342, "（等 recipe 完成後跑）", "label-xs")}

    <!-- Results merge -->
    {arr_d(320, 355, 405, "orange")}
    {arr_d(660, 355, 405, "teal")}
    {rr(300, 410, 280, 40, "#FAF5FF", "#805AD5")}
    {tc(440, 435, "合併結果 → Step 5 Analysis", "label")}

    <!-- Explanation -->
    {rr(30, 470, 820, 40, "#FFFFF0", "#D69E2E", 6)}
    {tc(440, 495, "🍳 先煎牛排（analyze_recipe），牛排好了才能擺盤上桌（add_to_calendar）——有順序的任務排隊，沒相依的同時開工", "label-sm")}
    '''
    return svg_wrap(content, w, h)


def make_svg_dual_memory():
    """圖 6: 雙記憶架構"""
    w, h = 850, 480
    content = f'''
    {tc(425, 30, "Step 6: 雙記憶系統 — 記住對話的兩種方式", "title")}

    <!-- Input -->
    {rr(300, 55, 250, 40, "#EBF8FF", "#3182CE")}
    {tc(425, 80, "本輪對話 + 工具結果", "label")}

    <!-- Fork -->
    {arr_d(350, 95, 140, "blue")}
    {arr_d(500, 95, 140, "teal")}

    <!-- Narrative -->
    {rr_bold(80, 145, 320, 120, "#F0FFF4", "#38A169")}
    {tc(240, 175, "敘事記憶 (Narrative)", "label-bold")}
    {tl(100, 200, "PROMPT_NARRATIVE_SUMMARY", "code")}
    {tl(100, 220, "記「對話流程」：誰問了什麼、", "label-sm")}
    {tl(100, 238, "話題怎麼轉、使用者的態度", "label-sm")}
    {tl(100, 256, "不記具體數字（交給右邊）", "label-xs")}

    <!-- Structured -->
    {rr_bold(450, 145, 370, 120, "#E6FFFA", "#319795")}
    {tc(635, 175, "結構化記憶 (Structured)", "label-bold")}
    {tl(470, 200, "PROMPT_STRUCTURED_EXTRACT", "code")}
    {tl(470, 220, "記「具體數據」：年齡 28 歲、", "label-sm")}
    {tl(470, 238, "鈣攝取 500mg、搜過「疲勞」", "label-sm")}
    {tl(470, 256, "12 類欄位，JSON 格式", "label-xs")}

    <!-- Merge -->
    {arr_d(240, 265, 310, "green")}
    {arr_d(635, 265, 310, "teal")}
    {rr(250, 315, 350, 40, "#FFFAF0", "#DD6B20")}
    {tc(425, 340, "合併到 structured_store + narrative", "label")}

    <!-- Save -->
    {arr_d(425, 355, 390, "orange")}
    {rr(250, 395, 350, 40, "#FAF5FF", "#805AD5")}
    {tc(425, 420, "背景儲存 → Checkpoint + 記憶樹", "label")}

    {rr(80, 445, 690, 30, "#FFFFF0", "#D69E2E", 4)}
    {tc(425, 465, "🍳 敘事記憶 = 服務日誌（今天哪桌點了什麼、客人心情好不好）；結構化記憶 = 庫存清單（剩多少牛排、用了多少鹽）", "label-xs")}
    '''
    return svg_wrap(content, w, h)


def make_svg_tools_overview():
    """圖 7: 15 個工具分類"""
    w, h = 900, 520
    content = f'''
    {tc(450, 30, "15 個工具分類總覽", "title")}

    <!-- Search Tools -->
    {rr_bold(30, 60, 400, 200, "#EBF8FF", "#3182CE")}
    {tc(230, 85, "🔍 搜尋類工具（6 個）", "label-bold")}
    {rr(50, 100, 170, 30, "#DBEAFE", "#3182CE", 4)}
    {tc(135, 120, "search_food", "code")}
    {rr(230, 100, 180, 30, "#DBEAFE", "#3182CE", 4)}
    {tc(320, 120, "browse_taiwan_food", "code")}
    {rr(50, 140, 180, 30, "#DBEAFE", "#3182CE", 4)}
    {tc(140, 160, "get_nutrient_ranking", "code")}
    {rr(240, 140, 170, 30, "#DBEAFE", "#3182CE", 4)}
    {tc(325, 160, "search_symptom", "code")}
    {rr(50, 180, 180, 30, "#DBEAFE", "#3182CE", 4)}
    {tc(140, 200, "search_graphrag", "code")}
    {rr(240, 180, 180, 30, "#DBEAFE", "#3182CE", 4)}
    {tc(330, 200, "get_literature_papers", "code")}

    <!-- Analysis Tools -->
    {rr_bold(470, 60, 400, 100, "#F0FFF4", "#38A169")}
    {tc(670, 85, "📊 分析類工具（2 個）", "label-bold")}
    {rr(490, 105, 170, 30, "#C6F6D5", "#38A169", 4)}
    {tc(575, 125, "analyze_recipe", "code")}
    {rr(680, 105, 170, 30, "#C6F6D5", "#38A169", 4)}
    {tc(765, 125, "synthesize_advice", "code")}

    <!-- CRUD Tools -->
    {rr_bold(470, 190, 400, 170, "#FFFAF0", "#DD6B20")}
    {tc(670, 215, "📝 CRUD 類工具（5 個）", "label-bold")}
    {rr(490, 235, 170, 30, "#FED7AA", "#DD6B20", 4)}
    {tc(575, 255, "get_saved_recipes", "code")}
    {rr(680, 235, 180, 30, "#FED7AA", "#DD6B20", 4)}
    {tc(770, 255, "delete_saved_recipe", "code")}
    {rr(490, 275, 180, 30, "#FED7AA", "#DD6B20", 4)}
    {tc(580, 295, "get_calendar_entries", "code")}
    {rr(680, 275, 170, 30, "#FED7AA", "#DD6B20", 4)}
    {tc(765, 295, "add_to_calendar", "code")}
    {rr(490, 315, 190, 30, "#FED7AA", "#DD6B20", 4)}
    {tc(585, 335, "delete_calendar_entry", "code")}

    <!-- Profile Tools -->
    {rr_bold(30, 290, 400, 80, "#FAF5FF", "#805AD5")}
    {tc(230, 315, "👤 個人資料類工具（2 個）", "label-bold")}
    {rr(50, 335, 170, 30, "#E9D8FD", "#805AD5", 4)}
    {tc(135, 355, "get_user_profile", "code")}
    {rr(240, 335, 180, 30, "#E9D8FD", "#805AD5", 4)}
    {tc(330, 355, "update_user_profile", "code")}

    {rr(30, 400, 840, 105, "#FFFFF0", "#D69E2E", 8)}
    {tl(50, 425, "🍳 餐廳比喻", "subtitle")}
    {tl(50, 448, "搜尋工具 = 查閱食材百科 / 營養手冊 / 學術論文        分析工具 = 請大廚評估這道菜的營養", "label-xs")}
    {tl(50, 468, "CRUD 工具 = 翻食譜本 / 排菜單 / 管理行事曆            個人資料 = 客人的口味偏好卡", "label-xs")}
    {tl(50, 488, "每個工具就像餐廳裡的一個專業工作站——有的負責查資料，有的負責計算，有的負責存檔", "label-xs")}
    '''
    return svg_wrap(content, w, h)


def make_svg_symptom_pipeline():
    """圖 8: 症狀搜尋 5 階段"""
    w, h = 880, 560
    stages = [
        ("1", "LLM 結構化解析", "使用者自由輸入 → Gemma 拆解", "#E53E3E", "#FFF5F5"),
        ("2", "雙路搜尋", "靜態對照表 + 向量語義搜尋", "#3182CE", "#EBF8FF"),
        ("3", "LLM 判斷相關性", "合併候選 → Gemma 篩選 Top 10", "#38A169", "#F0FFF4"),
        ("4", "化合物→食物查詢", "Backend compound→food 管道", "#DD6B20", "#FFFAF0"),
        ("5", "去重彙總", "合併結果 → 回傳前端", "#805AD5", "#FAF5FF"),
    ]
    content = tc(440, 30, "search_symptom 工具：五階段搜尋管道", "title")
    y = 60
    for num, name, desc, stroke, fill in stages:
        content += f'''
        {circle(60, y+30, 20, stroke)}
        {tc(60, y+37, num, "num")}
        {rr_bold(100, y+5, 300, 50, fill, stroke)}
        {tl(115, y+28, name, "label-bold")}
        {tl(115, y+48, desc, "label-xs")}
        '''
        y += 85

    content += f'''
    <!-- Example flow -->
    {rr(480, 60, 370, 440, "#F7FAFC", "#A0AEC0", 8)}
    {tl(500, 85, "範例：使用者說「最近常覺得疲勞」", "label-bold")}
    {tl(500, 110, "① LLM 解析 → [\"疲勞\", \"體力不足\"]", "label-sm")}
    {tl(500, 140, "② 靜態表找到 3 筆 +", "label-sm")}
    {tl(500, 160, "   向量搜尋找到 10 筆語義相似", "label-sm")}
    {tl(500, 190, "③ LLM 從 13 筆挑出最相關 Top 10", "label-sm")}
    {tl(500, 220, "④ 每個 effect → 查化合物 → 查食物", "label-sm")}
    {tl(500, 250, "⑤ 合併去重 → 回傳：", "label-sm")}
    {tl(520, 275, "• 8 個健康效果", "label-xs")}
    {tl(520, 295, "• 15 種相關化合物（如：鐵、B12）", "label-xs")}
    {tl(520, 315, "• 20 種推薦食物（如：牛肝、菠菜）", "label-xs")}

    {rr(500, 340, 330, 60, "#FFFFF0", "#D69E2E", 6)}
    {tl(515, 365, "🍳 就像去醫院：你說「我累」→", "label-sm")}
    {tl(515, 385, "醫生翻書+電腦搜 → 交叉比對 → 開處方", "label-sm")}

    {rr(500, 410, 330, 80, "#FFF5F5", "#E53E3E", 6)}
    {tl(515, 435, "⚠️ 如果 LLM 掛了？", "label-bold")}
    {tl(515, 455, "每一階段都有 fallback：", "label-sm")}
    {tl(515, 475, "LLM 解析失敗 → 直接用原始關鍵字", "label-xs")}
    '''
    return svg_wrap(content, w, h)


def make_svg_recipe_flow():
    """圖 9: 食譜分析兩階段流程"""
    w, h = 860, 480
    content = f'''
    {tc(430, 30, "analyze_recipe：兩階段食譜分析", "title")}

    {rr_bold(280, 55, 300, 40, "#EBF8FF", "#3182CE")}
    {tc(430, 80, "使用者貼食譜文字", "label")}
    {arr_d(430, 95, 125, "blue")}

    <!-- Phase A -->
    {rr_bold(180, 130, 500, 60, "#FFFAF0", "#DD6B20")}
    {tc(430, 155, "階段 A：食材提取 + 使用者確認", "label-bold")}
    {tc(430, 175, "jieba 斷詞 + regex 提取 → confirm_ingredients 卡片", "label-xs")}
    {arr_d(430, 190, 220, "orange")}

    {rr(300, 225, 260, 35, "#FFF5F5", "#E53E3E")}
    {tc(430, 248, "使用者可新增/移除食材", "label-sm")}
    {arr_d(430, 260, 285, "red")}

    {rr(300, 290, 260, 35, "#F0FFF4", "#38A169")}
    {tc(430, 313, "[RECIPE_CONFIRM] confirm", "code")}
    {arr_d(430, 325, 355, "green")}

    <!-- Phase B -->
    {rr_bold(180, 360, 500, 60, "#F0FFF4", "#38A169")}
    {tc(430, 385, "階段 B：營養計算 + DRI 缺口分析", "label-bold")}
    {tc(430, 405, "POST /recipe/calculate-nutrition → DRI gap → 推薦食物", "label-xs")}

    {rr(30, 130, 130, 290, "#FFFFF0", "#D69E2E", 6)}
    {tl(40, 160, "🍳 就像：", "label-bold")}
    {tl(40, 185, "你把菜單", "label-xs")}
    {tl(40, 200, "交給廚師", "label-xs")}
    {tl(40, 225, "廚師先列", "label-xs")}
    {tl(40, 240, "出食材清單", "label-xs")}
    {tl(40, 260, "問你：", "label-xs")}
    {tl(40, 275, "「對嗎？」", "label-xs")}
    {tl(40, 300, "你說「對」", "label-xs")}
    {tl(40, 320, "廚師才開始", "label-xs")}
    {tl(40, 335, "算營養", "label-xs")}
    {tl(40, 360, "最後告訴你", "label-xs")}
    {tl(40, 375, "缺什麼營養", "label-xs")}
    {tl(40, 395, "該補什麼", "label-xs")}
    '''
    return svg_wrap(content, w, h)


def make_svg_delete_confirm():
    """圖 10: 刪除確認流程"""
    w, h = 800, 380
    content = f'''
    {tc(400, 30, "delete_saved_recipe：二步確認機制", "title")}

    {rr_bold(250, 55, 300, 40, "#EBF8FF", "#3182CE")}
    {tc(400, 80, "使用者：「幫我刪除昨天的食譜」", "label-sm")}
    {arr_d(400, 95, 125, "blue")}

    <!-- Step 1 -->
    {rr_bold(200, 130, 400, 55, "#FFF5F5", "#E53E3E")}
    {tc(400, 152, "第一步：預覽（confirmed=false）", "label-bold")}
    {tc(400, 170, "查詢匹配食譜 → 列出清單 → 問使用者確認", "label-xs")}
    {arr_d(400, 185, 220, "red")}

    {rr(250, 225, 300, 35, "#FFFFF0", "#D69E2E")}
    {tc(400, 248, "使用者：「全部刪除」或「只刪第1個」", "label-sm")}
    {arr_d(400, 260, 290, "orange")}

    <!-- Step 2 -->
    {rr_bold(200, 295, 400, 55, "#F0FFF4", "#38A169")}
    {tc(400, 317, "第二步：實際刪除（confirmed=true）", "label-bold")}
    {tc(400, 335, "帶 recipe_ids → DELETE /recipe/id", "label-xs")}
    '''
    return svg_wrap(content, w, h)


def make_svg_prompt_lifecycle():
    """圖 11: Prompt 生命週期"""
    w, h = 900, 560
    prompts = [
        ("PROMPT_ECHO", "L51", "0.5", "Echo", "#3182CE"),
        ("PROMPT_PHASE1", "L92", "0.1", "Phase1", "#38A169"),
        ("PROMPT_PHASE2", "L196", "0.1", "Phase2", "#2C7A7B"),
        ("PROMPT_ANALYSIS", "L262", "0.3", "Analysis", "#805AD5"),
        ("PROMPT_NARRATIVE_SUMMARY", "L317", "0.2", "Memory", "#E53E3E"),
        ("PROMPT_STRUCTURED_EXTRACT", "L341", "0.1", "Memory", "#E53E3E"),
        ("PROMPT_CLASSIFY_REPLY", "L3266", "0.0", "中斷偵測", "#DD6B20"),
        ("PROMPT_CLASSIFY_RESUME", "L3276", "0.0", "恢復偵測", "#DD6B20"),
        ("PROMPT_EXTRACT_PARAMS", "L3586", "0.1", "參數收集", "#DD6B20"),
    ]
    content = tc(450, 30, "9 個 Prompt 範本與其在 Pipeline 中的位置", "title")
    y = 60
    for name, line, temp, step, color in prompts:
        content += f'''
        {rr(30, y, 380, 35, "#F7FAFC", color, 4)}
        {tl(45, y+23, name, "code")}
        {tl(430, y+23, line, "label-xs")}
        {tl(510, y+23, f"temp={temp}", "label-xs")}
        {rr(600, y+3, 120, 30, "#F7FAFC", color, 4)}
        {tc(660, y+23, step, "label-sm")}
        '''
        y += 45

    content += f'''
    {rr(30, y+10, 840, 55, "#FFFFF0", "#D69E2E", 8)}
    {tl(50, y+35, "🍳 Prompt 就像「指令卡」：每張卡告訴 LLM 這一步要做什麼、用什麼格式回答、", "label-sm")}
    {tl(50, y+55, "不能做什麼。溫度（temperature）像「創意旋鈕」：0.1=精準回答，0.5=自然對話", "label-sm")}
    '''
    return svg_wrap(content, w, h)


def make_svg_sse_timeline():
    """圖 12: SSE 事件時間軸"""
    w, h = 900, 800
    events = [
        (0, "phase_status", "{phase:'echo'}", "#718096", "Pipeline 進度指示"),
        (1, "echo", "{content:'收到...'}", "#3182CE", "即時確認回應"),
        (2, "phase_status", "{phase:'intent'}", "#718096", "開始意圖判定"),
        (3, "tool_start", "{name:'search_food'}", "#DD6B20", "工具開始執行"),
        (4, "tool_result", "{name:'search_food',data:{...}}", "#DD6B20", "工具結果（完整資料）"),
        (5, "tool_done", "{name:'search_food'}", "#DD6B20", "工具執行完畢"),
        (6, "phase_status", "{phase:'analysis'}", "#718096", "開始生成分析"),
        (7, "thinking_start", "{step:'analysis'}", "#805AD5", "LLM 開始思考"),
        (8, "thinking_delta", "{content:'根據查詢...'}", "#805AD5", "思考過程串流（多次）"),
        (9, "thinking_end", "{step:'analysis'}", "#805AD5", "思考結束"),
        (10, "reply", "{content:'完整分析回覆'}", "#38A169", "最終回覆"),
        (11, "workflow_state", "{phase:'done',...}", "#718096", "狀態機更新"),
        (12, "narrative_memory", "{summary:'...'}", "#E53E3E", "敘事記憶更新"),
        (13, "structured_store", "{store:{...}}", "#319795", "結構化記憶更新"),
        (14, "checkpoint_saved", "{session_id,turn}", "#718096", "Checkpoint 已儲存"),
    ]
    content = tc(450, 30, "SSE 事件時間軸：一次完整查詢的事件流", "title")
    y = 60
    content += f'<line x1="80" y1="55" x2="80" y2="{55+len(events)*46}" stroke="#CBD5E0" stroke-width="2"/>'
    for i, (_, etype, payload, color, desc) in enumerate(events):
        content += f'''
        {circle(80, y+10, 8, color)}
        {tl(100, y+5, etype, "code")}
        {tl(310, y+5, payload, "label-xs")}
        {tl(640, y+15, desc, "label-sm")}
        '''
        y += 46
    return svg_wrap(content, w, h)


def make_svg_communication_chain():
    """圖 13: 完整通信鏈"""
    w, h = 900, 480
    content = f'''
    {tc(450, 30, "Frontend → Control → Backend 完整通信鏈", "title")}

    <!-- Three columns -->
    {rr_bold(30, 65, 200, 380, "#EBF8FF", "#3182CE")}
    {tc(130, 90, "Frontend :3001", "label-bold")}

    {rr_bold(310, 65, 230, 380, "#F0FFF4", "#38A169")}
    {tc(425, 90, "Control :8000", "label-bold")}

    {rr_bold(620, 65, 250, 380, "#FFFAF0", "#DD6B20")}
    {tc(745, 90, "Backend :8001", "label-bold")}

    <!-- Step 1: User sends message -->
    {tl(50, 130, "1. 使用者打字送出", "label-sm")}
    {arr_r(230, 135, 308, "blue", "POST /advisor/chat (SSE)", -14)}

    <!-- Step 2: Echo -->
    {tl(325, 160, "2. LLM Echo", "label-sm")}
    {arr_l(308, 170, 230, "green", "event: echo", 12)}

    <!-- Step 3: Phase1 -->
    {tl(325, 200, "3. Phase1 意圖判定", "label-sm")}

    <!-- Step 4: Tool call -->
    {tl(325, 230, "4. 執行工具", "label-sm")}
    {arr_r(540, 235, 618, "green", "GET /food/search", -14)}
    {arr_l(618, 255, 540, "orange", "JSON result", 12)}

    <!-- Step 5: Tool result to frontend -->
    {arr_l(308, 275, 230, "green", "event: tool_result", 12)}

    <!-- Step 6: Analysis stream -->
    {tl(325, 300, "5. LLM 分析（串流）", "label-sm")}
    {arr_l(308, 315, 230, "green", "event: thinking_delta ×N", 12)}

    <!-- Step 7: Reply -->
    {arr_l(308, 345, 230, "green", "event: reply", 12)}

    <!-- Step 8: Memory -->
    {tl(325, 370, "6. 背景記憶更新", "label-sm")}
    {arr_r(540, 375, 618, "green", "POST /checkpoint", -14)}
    {arr_l(308, 395, 230, "green", "event: structured_store", 12)}

    <!-- JWT -->
    {rr(50, 415, 820, 25, "#FFFFF0", "#D69E2E", 4)}
    {tc(460, 432, "所有請求都帶 JWT Bearer Token（登入時取得，存在 localStorage，7 天有效）", "label-xs")}
    '''
    return svg_wrap(content, w, h)


def make_svg_six_factions_overview():
    """圖 14: 六個派別總覽"""
    w, h = 920, 580
    factions = [
        ("智慧派", "_detect_interruption_smart", "LLM 三態判斷", "#E53E3E", "#FFF5F5"),
        ("規則派", "_detect_interruption_regex", "正則表達式比對", "#DD6B20", "#FFFAF0"),
        ("態度派", "_classify_collecting_response", "回覆態度分類", "#38A169", "#F0FFF4"),
        ("回歸派", "_detect_resume_smart", "中斷恢復偵測", "#3182CE", "#EBF8FF"),
        ("快捷派", "[RECIPE_CONFIRM]", "食譜確認協議", "#805AD5", "#FAF5FF"),
        ("前線派", "AbortController", "前端串流中斷", "#319795", "#E6FFFA"),
    ]
    content = tc(460, 30, "中斷處理的六個派別", "title")
    content += tc(460, 55, "每個派別負責不同情境下的「打斷」和「恢復」", "label-sm")

    for i, (name, func, desc, stroke, fill) in enumerate(factions):
        col = i % 3
        row = i // 3
        x = 30 + col * 300
        y = 80 + row * 220
        content += f'''
        {rr_bold(x, y, 270, 190, fill, stroke)}
        {circle(x+25, y+25, 18, stroke)}
        {tc(x+25, y+32, str(i+1), "num")}
        {tl(x+55, y+32, name, "subtitle")}
        {tl(x+20, y+55, func, "code")}
        {tl(x+20, y+80, desc, "label-sm")}
        '''
        if i == 0:
            content += f'''
            {tl(x+20, y+105, "三態：answer / interrupt /", "label-xs")}
            {tl(x+20, y+120, "ambiguous（模糊時追問）", "label-xs")}
            {tl(x+20, y+145, "LLM 與 regex 交叉驗證", "label-xs")}
            {tl(x+20, y+165, "分歧 → ambiguous → 問使用者", "label-xs")}
            '''
        elif i == 1:
            content += f'''
            {tl(x+20, y+105, "_QUESTION_RE 問句模式", "label-xs")}
            {tl(x+20, y+120, "_RETURN_RE 回歸模式", "label-xs")}
            {tl(x+20, y+140, "_NEW_INTENT_RE 新意圖模式", "label-xs")}
            {tl(x+20, y+160, "按缺少參數類型做特化判斷", "label-xs")}
            '''
        elif i == 2:
            content += f'''
            {tl(x+20, y+105, "answer → 繼續收集", "label-xs")}
            {tl(x+20, y+120, "soft_rejection → 再給一次機會", "label-xs")}
            {tl(x+20, y+140, "hard_abandon → 立即放棄", "label-xs")}
            {tl(x+20, y+160, "軟拒 2 次 → 自動放棄", "label-xs")}
            '''
        elif i == 3:
            content += f'''
            {tl(x+20, y+105, "隱式：確定性參數提取", "label-xs")}
            {tl(x+20, y+120, "（使用者直接補參數 → 自動恢復）", "label-xs")}
            {tl(x+20, y+140, "顯式：LLM + regex", "label-xs")}
            {tl(x+20, y+160, "（使用者說「回到剛才」）", "label-xs")}
            '''
        elif i == 4:
            content += f'''
            {tl(x+20, y+105, "add / remove / confirm / abort", "label-xs")}
            {tl(x+20, y+120, "前端送特殊前綴訊息", "label-xs")}
            {tl(x+20, y+140, "[RECIPE_CONFIRM]{{json}}", "label-xs")}
            {tl(x+20, y+160, "跳過整個 Pipeline 直接處理", "label-xs")}
            '''
        else:
            content += f'''
            {tl(x+20, y+105, "前端 AbortController", "label-xs")}
            {tl(x+20, y+120, "使用者按「停止」→ 中斷串流", "label-xs")}
            {tl(x+20, y+140, "SSE 連線立即斷開", "label-xs")}
            {tl(x+20, y+160, "abort signal 傳到伺服器", "label-xs")}
            '''
    return svg_wrap(content, w, h)


def make_svg_tristate():
    """圖 15: 三態中斷偵測"""
    w, h = 920, 490
    content = f'''
    {tc(460, 30, "智慧派：三態中斷偵測", "title")}

    {rr_bold(280, 55, 260, 40, "#EBF8FF", "#3182CE")}
    {tc(410, 80, "使用者在收集參數時說了什麼", "label-sm")}
    {arr_d(410, 95, 130, "blue")}

    <!-- Quick check -->
    {rr(280, 135, 260, 40, "#F7FAFC", "#A0AEC0")}
    {tc(410, 160, "是「好/嗯/OK」等確認詞？", "label-sm")}
    {arr_r(540, 155, 620, "green", "是", -10)}
    {rr(625, 140, 100, 35, "#F0FFF4", "#38A169", 4)}
    {tc(675, 162, "→ answer", "good")}

    {arr_d(410, 175, 210, "gray", "不是")}

    <!-- LLM -->
    {rr_bold(240, 215, 340, 50, "#FFF5F5", "#E53E3E")}
    {tc(410, 237, "LLM 判斷（PROMPT_CLASSIFY_REPLY）", "label")}
    {tc(410, 255, "A=回答上述問題 / B=說別的事", "label-xs")}

    {arr_d(410, 265, 300, "red")}

    <!-- Regex -->
    {rr_bold(240, 305, 340, 45, "#FFFAF0", "#DD6B20")}
    {tc(410, 332, "Regex 同步判斷", "label")}

    <!-- Compare -->
    {arr_d(410, 350, 385, "orange")}
    {rr(240, 388, 340, 40, "#FAF5FF", "#805AD5")}
    {tc(410, 413, "LLM 與 Regex 結果比較", "label")}

    <!-- Three outcomes -->
    {arr_r(580, 395, 640, "green")}
    {tl(650, 390, "一致「answer」", "good")}
    {tl(650, 408, "→ 繼續收集參數", "label-xs")}

    {arr_r(580, 408, 640, "red")}
    {tl(650, 425, "一致「interrupt」→ 中斷，開始新任務", "accent")}

    {arr_d(410, 428, 445, "purple")}
    {tc(410, 458, "分歧 → ambiguous → 追問使用者釐清", "label-sm")}
    '''
    return svg_wrap(content, w, h)


def make_svg_collecting_states():
    """圖 16: 收集階段狀態轉換"""
    w, h = 850, 420
    content = f'''
    {tc(425, 30, "態度派：收集階段的三種回應", "title")}

    <!-- answer -->
    {rr_bold(30, 80, 240, 120, "#F0FFF4", "#38A169")}
    {tc(150, 110, "answer（正常回答）", "label-bold")}
    {tl(50, 135, "使用者提供了所需資訊", "label-sm")}
    {tl(50, 155, "例：「我25歲」「疲勞」", "label-xs")}
    {tl(50, 180, "→ 提取參數，繼續流程", "good")}

    <!-- soft_rejection -->
    {rr_bold(310, 80, 240, 120, "#FFFAF0", "#DD6B20")}
    {tc(430, 110, "soft_rejection（軟拒）", "label-bold")}
    {tl(330, 135, "使用者猶豫/不確定", "label-sm")}
    {tl(330, 155, "例：「不知道」「隨便」", "label-xs")}
    {tl(330, 180, "→ 再問一次（最多2次）", "label-sm")}

    <!-- hard_abandon -->
    {rr_bold(590, 80, 240, 120, "#FFF5F5", "#E53E3E")}
    {tc(710, 110, "hard_abandon（放棄）", "label-bold")}
    {tl(610, 135, "使用者明確取消", "label-sm")}
    {tl(610, 155, "例：「算了」「不要了」", "label-xs")}
    {tl(610, 180, "→ 立即取消，回到閒置", "accent")}

    <!-- State flow -->
    {rr(30, 240, 800, 160, "#F7FAFC", "#A0AEC0", 8)}
    {tl(50, 270, "狀態流轉：", "label-bold")}
    {tl(50, 295, "soft_rejection 第 1 次 → 提示範例，再問一次", "label-sm")}
    {tl(50, 318, "soft_rejection 第 2 次 → 自動放棄（「沒關係，想做別的隨時說」）", "label-sm")}
    {tl(50, 345, "整個收集流程最多 3 輪（turn_count >= 3 → 強制放棄）", "label-sm")}
    {tl(50, 375, "🍳 就像服務生問你「要幾分熟？」—— 你說「不知道」他會建議「七分熟不錯喔」；你說「算了不要了」他就收走菜單", "label-xs")}
    '''
    return svg_wrap(content, w, h)


def make_svg_resume_flow():
    """圖 17: 中斷→恢復流程"""
    w, h = 860, 500
    content = f'''
    {tc(430, 30, "回歸派：中斷保存與恢復", "title")}

    <!-- Interrupt save -->
    {rr_bold(30, 65, 370, 130, "#FFF5F5", "#E53E3E")}
    {tc(215, 90, "中斷時：保存狀態", "label-bold")}
    {tl(50, 115, "interrupted_state = {{", "code")}
    {tl(50, 132, "  intent: 'search_symptom',", "code")}
    {tl(50, 149, "  collected: {{...}},", "code")}
    {tl(50, 166, "  missing: ['symptom'],", "code")}
    {tl(50, 183, "}}", "code")}

    {arr_r(400, 130, 460, "gray", "使用者去做別的事...")}

    <!-- Resume -->
    {rr_bold(470, 65, 370, 130, "#F0FFF4", "#38A169")}
    {tc(655, 90, "恢復時：兩條路徑", "label-bold")}
    {tl(490, 115, "隱式：使用者直接補參數", "label-sm")}
    {tl(490, 135, "  例：「疲勞」→ 自動偵測是", "label-xs")}
    {tl(490, 152, "  symptom 參數 → 恢復任務", "label-xs")}
    {tl(490, 175, "顯式：使用者說「回到剛才」", "label-sm")}

    <!-- Hint -->
    {rr(30, 220, 810, 50, "#FFFFF0", "#D69E2E", 6)}
    {tl(50, 250, "💡 中斷後 AI 會提示：「您之前還有未完成的任務：搜尋症狀相關營養素。回覆「回到剛才」即可繼續。」", "label-sm")}

    <!-- Decision flow -->
    {rr_bold(200, 295, 460, 180, "#F7FAFC", "#A0AEC0")}
    {tc(430, 320, "恢復決策流程", "label-bold")}
    {tl(220, 345, "1. 先嘗試確定性參數提取", "label-sm")}
    {tl(240, 365, "→ 成功？隱式恢復（使用者無感）", "label-xs")}
    {tl(220, 390, "2. 再用 LLM (PROMPT_CLASSIFY_RESUME)", "label-sm")}
    {tl(240, 410, "→ A=想恢復，B=做新事情", "label-xs")}
    {tl(220, 435, "3. LLM 失敗 → regex fallback", "label-sm")}
    {tl(240, 455, "→ _RETURN_RE 比對「回到剛才」等", "label-xs")}
    '''
    return svg_wrap(content, w, h)


def make_svg_recipe_confirm():
    """圖 18: 食譜確認快捷派"""
    w, h = 860, 450
    content = f'''
    {tc(430, 30, "快捷派：[RECIPE_CONFIRM] 食譜確認協議", "title")}

    {rr_bold(250, 55, 360, 40, "#EBF8FF", "#3182CE")}
    {tc(430, 80, "前端送出特殊前綴訊息", "label")}
    {arr_d(430, 95, 125, "blue")}

    <!-- Decision -->
    {rr(250, 130, 360, 35, "#F7FAFC", "#A0AEC0")}
    {tc(430, 152, "解析 action 欄位", "label-sm")}

    <!-- Four branches -->
    {rr(30, 200, 180, 75, "#F0FFF4", "#38A169")}
    {tc(120, 225, "add", "label-bold")}
    {tc(120, 250, "新增食材到清單", "label-xs")}
    {tc(120, 265, "更新 pending", "label-xs")}

    {rr(240, 200, 180, 75, "#FFFAF0", "#DD6B20")}
    {tc(330, 225, "remove", "label-bold")}
    {tc(330, 250, "移除指定食材", "label-xs")}
    {tc(330, 265, "更新 pending", "label-xs")}

    {rr(450, 200, 180, 75, "#805AD5", "#805AD5")}
    {tc(540, 225, "confirm", "white")}
    {tc(540, 250, "確認分析", "white-sm")}
    {tc(540, 265, "→ 進入 Phase B", "white-sm")}

    {rr(660, 200, 180, 75, "#FFF5F5", "#E53E3E")}
    {tc(750, 225, "abort", "label-bold")}
    {tc(750, 250, "取消食譜分析", "label-xs")}
    {tc(750, 265, "清除 pending", "label-xs")}

    {rr(30, 310, 810, 120, "#FFFFF0", "#D69E2E", 8)}
    {tl(50, 340, "🍳 餐廳比喻：", "subtitle")}
    {tl(50, 365, "廚師列出食材清單後，你可以：加菜（add）、退菜（remove）、確認開煮（confirm）、不要了（abort）", "label-sm")}
    {tl(50, 390, "格式：message = '[RECIPE_CONFIRM]{{\"action\":\"confirm\",\"ingredients\":[...]}}'", "code")}
    {tl(50, 415, "這條訊息跳過整個 6 步 Pipeline，直接由 _flow_controller 的 Branch R 處理", "label-xs")}
    '''
    return svg_wrap(content, w, h)


def make_svg_compound_dag():
    """圖 19: 複合問題 DAG"""
    w, h = 880, 500
    content = f'''
    {tc(440, 30, "複合問題處理：多意圖 DAG 執行", "title")}

    {rr(200, 55, 480, 35, "#EBF8FF", "#3182CE", 4)}
    {tc(440, 78, "「分析雞胸肉食譜，加到行事曆，告訴我還缺什麼營養」", "label-sm")}
    {arr_d(440, 90, 120, "blue")}

    {rr(260, 125, 360, 35, "#F0FFF4", "#38A169")}
    {tc(440, 148, "Phase1 → 三個意圖：", "label")}
    {arr_d(440, 160, 195, "green")}

    <!-- Level 0 -->
    {rr_bold(280, 200, 200, 45, "#FFFAF0", "#DD6B20")}
    {tc(380, 228, "analyze_recipe", "label-bold")}
    {tl(500, 225, "Level 0（無相依）", "label-xs")}

    <!-- Level 1 -->
    {arr_d(330, 245, 295, "orange")}
    {arr_d(430, 245, 295, "orange")}

    {rr_bold(140, 300, 200, 45, "#E6FFFA", "#319795")}
    {tc(240, 328, "add_to_calendar", "label-bold")}
    {tl(50, 325, "Level 1", "label-xs")}

    {rr_bold(440, 300, 200, 45, "#FAF5FF", "#805AD5")}
    {tc(540, 328, "synthesize_advice", "label-bold")}
    {tl(660, 325, "Level 1", "label-xs")}

    <!-- Merge -->
    {arr_d(240, 345, 390, "teal")}
    {arr_d(540, 345, 390, "purple")}
    {rr(260, 395, 360, 35, "#F0FFF4", "#38A169")}
    {tc(440, 418, "合併所有結果 → Analysis LLM", "label")}

    {rr(50, 450, 780, 40, "#FFFFF0", "#D69E2E", 6)}
    {tc(440, 475, "Level 0 的工具先跑完 → 結果存入 structured_store → Level 1 的工具才能拿到食譜資料開始跑", "label-xs")}
    '''
    return svg_wrap(content, w, h)


def make_svg_recovery_engine():
    """圖 20: 自癒引擎"""
    w, h = 850, 460
    content = f'''
    {tc(425, 30, "自癒引擎（Recovery Engine）", "title")}

    {rr_bold(300, 55, 250, 40, "#FFFAF0", "#DD6B20")}
    {tc(425, 80, "工具執行出問題！", "accent")}
    {arr_d(425, 95, 130, "red")}

    <!-- Diagnose -->
    {rr_bold(230, 135, 390, 55, "#FFF5F5", "#E53E3E")}
    {tc(425, 158, "Step 1: diagnose() 診斷錯誤類型", "label-bold")}
    {tc(425, 178, "empty_result | tool_timeout | param_mismatch | backend_error", "label-xs")}
    {arr_d(425, 190, 225, "red")}

    <!-- Plan -->
    {rr_bold(230, 230, 390, 55, "#FFFAF0", "#DD6B20")}
    {tc(425, 253, "Step 2: plan_recovery() 規劃恢復策略", "label-bold")}
    {tc(425, 273, "synonym_expand | retry_timeout | field_remap | skip", "label-xs")}
    {arr_d(425, 285, 320, "orange")}

    <!-- Execute -->
    {rr_bold(230, 325, 390, 55, "#F0FFF4", "#38A169")}
    {tc(425, 348, "Step 3: execute_recovery() 自動重試", "label-bold")}
    {tc(425, 368, "修改參數後重新呼叫工具（最多重試 1 次）", "label-xs")}

    <!-- Example -->
    {rr(30, 135, 180, 245, "#FFFFF0", "#D69E2E", 6)}
    {tl(45, 165, "範例：", "label-bold")}
    {tl(45, 188, "搜「芝士」→ 0 筆", "label-xs")}
    {tl(45, 208, "診斷：empty_result", "label-xs")}
    {tl(45, 232, "策略：synonym_expand", "label-xs")}
    {tl(45, 252, "嘗試：「起司」→ 3 筆", "label-xs")}
    {tl(45, 275, "成功！回傳結果", "good")}
    {tl(45, 305, "🍳 就像你問", "label-xs")}
    {tl(45, 322, "「有沒有芝士」", "label-xs")}
    {tl(45, 339, "服務生說「你是", "label-xs")}
    {tl(45, 356, "指起司嗎？」", "label-xs")}

    {rr(30, 395, 790, 50, "#F7FAFC", "#A0AEC0", 6)}
    {tl(50, 420, "4 種策略：synonym_expand（換同義詞）、retry_timeout（超時重試）、field_remap（欄位映射修正）、skip（跳過）", "label-xs")}
    '''
    return svg_wrap(content, w, h)


def make_svg_state_machine():
    """圖 21: 工作流程狀態機"""
    w, h = 900, 520
    content = f'''
    {tc(450, 30, "工作流程狀態機", "title")}

    <!-- idle -->
    {rr_bold(50, 100, 150, 60, "#EBF8FF", "#3182CE")}
    {tc(125, 127, "idle", "label-bold")}
    {tc(125, 147, "等待輸入", "label-xs")}

    <!-- collecting -->
    {arr_r(200, 130, 290, "blue", "參數不足")}
    {rr_bold(300, 100, 160, 60, "#FFFAF0", "#DD6B20")}
    {tc(380, 127, "collecting", "label-bold")}
    {tc(380, 147, "收集參數", "label-xs")}
    {tl(310, 175, "最多 3 輪", "label-xs")}

    <!-- ready -->
    {arr_r(460, 130, 540, "orange", "參數齊全")}
    {rr_bold(550, 100, 140, 60, "#F0FFF4", "#38A169")}
    {tc(620, 127, "ready", "label-bold")}
    {tc(620, 147, "準備執行", "label-xs")}

    <!-- executing -->
    {arr_d(620, 160, 225, "green", "開始執行")}
    {rr_bold(550, 230, 140, 60, "#FAF5FF", "#805AD5")}
    {tc(620, 257, "executing", "label-bold")}
    {tc(620, 277, "工具跑中", "label-xs")}

    <!-- done -->
    {arr_d(620, 290, 355, "purple", "全部完成")}
    {rr_bold(550, 360, 140, 60, "#F0FFF4", "#38A169")}
    {tc(620, 387, "done", "label-bold")}
    {tc(620, 407, "分析回覆", "label-xs")}

    <!-- Back to idle -->
    {arr_l(550, 400, 125, "gray")}
    {tc(340, 415, "回到等待", "label-xs")}
    <line x1="125" y1="400" x2="125" y2="160" stroke="#718096" stroke-width="1.5" stroke-dasharray="6,4"/>

    <!-- Interrupted branch -->
    {rr_bold(30, 280, 200, 80, "#FFF5F5", "#E53E3E")}
    {tc(130, 307, "interrupted_state", "label-bold")}
    {tc(130, 327, "保存中斷狀態", "label-xs")}
    {tc(130, 347, "等待恢復指令", "label-xs")}

    <line x1="380" y1="160" x2="380" y2="210" stroke="#E53E3E" stroke-width="1.5" stroke-dasharray="6,4"/>
    <line x1="380" y1="210" x2="230" y2="280" stroke="#E53E3E" stroke-width="1.5" stroke-dasharray="6,4" marker-end="url(#arr-red)"/>
    {tl(250, 230, "使用者中斷", "accent")}

    <line x1="30" y1="340" x2="30" y2="130" stroke="#38A169" stroke-width="1.5" stroke-dasharray="6,4"/>
    <line x1="30" y1="130" x2="48" y2="130" stroke="#38A169" stroke-width="1.5" stroke-dasharray="6,4" marker-end="url(#arr-green)"/>
    {tl(15, 230, "恢復", "good")}

    {rr(50, 440, 800, 55, "#FFFFF0", "#D69E2E", 6)}
    {tl(70, 463, "🍳 idle=等客人 → collecting=問幾分熟 → ready=準備好了 → executing=煎牛排 → done=上菜", "label-sm")}
    {tl(70, 483, "   中途客人改主意 → interrupted_state 記住「本來要七分熟牛排」→ 客人回來時可以繼續", "label-xs")}
    '''
    return svg_wrap(content, w, h)


def make_svg_validation():
    """圖 22: 多層驗證瀑布"""
    w, h = 850, 480
    layers = [
        ("意圖別名修正", "_INTENT_ALIASES", "get_user_profile→get_profile", "#3182CE", "#EBF8FF"),
        ("意圖合法性驗證", "_VALID_INTENTS", "過濾非法意圖名稱", "#38A169", "#F0FFF4"),
        ("工具參數驗證", "_validate_toolcalls()", "類型轉換+去重+必填檢查", "#DD6B20", "#FFFAF0"),
        ("營養素欄位解析", "_NUTRIENT_ZH_TO_FIELD", "「鈣」→ calcium_per_100g", "#805AD5", "#FAF5FF"),
        ("分析品質檢查", "_validate_analysis()", "最低 50 字，不夠重試", "#E53E3E", "#FFF5F5"),
        ("JSON 解析備援", "_parse_llm_json()", "直接→fence→深度配對", "#319795", "#E6FFFA"),
    ]
    content = tc(425, 30, "驗證與自修復管道：六層安全網", "title")
    y = 65
    for name, func, desc, stroke, fill in layers:
        content += f'''
        {rr_bold(100, y, 650, 50, fill, stroke)}
        {tl(120, y+22, name, "label-bold")}
        {tl(300, y+22, func, "code")}
        {tl(520, y+35, desc, "label-sm")}
        '''
        if y < 65 + 5*65:
            content += arr_d(425, y+50, y+63, "gray")
        y += 65
    return svg_wrap(content, w, h)


def make_svg_memory_tree():
    """圖 23: 記憶樹階層"""
    w, h = 850, 480
    content = f'''
    {tc(425, 30, "記憶架構：記憶樹 + 結構化儲存", "title")}

    <!-- Memory Tree -->
    {rr_bold(30, 65, 380, 260, "#F0FFF4", "#38A169")}
    {tc(220, 90, "記憶樹（Memory Tree）", "label-bold")}

    {rr(50, 110, 140, 40, "#C6F6D5", "#38A169", 4)}
    {tc(120, 135, "Leaf 節點", "label-sm")}
    {tl(200, 135, "每輪對話一筆摘要", "label-xs")}

    {arr_d(120, 150, 175, "green")}
    {rr(50, 180, 140, 40, "#C6F6D5", "#38A169", 4)}
    {tc(120, 205, "Session 摘要", "label-sm")}
    {tl(200, 205, "每 10 個 leaf 合併", "label-xs")}

    {arr_d(120, 220, 245, "green")}
    {rr(50, 250, 140, 40, "#C6F6D5", "#38A169", 4)}
    {tc(120, 275, "User 摘要", "label-sm")}
    {tl(200, 275, "跨 session 合併", "label-xs")}

    <!-- Structured Store -->
    {rr_bold(440, 65, 380, 260, "#E6FFFA", "#319795")}
    {tc(630, 90, "結構化儲存（Structured Store）", "label-bold")}

    {tl(460, 115, "user_facts       年齡/性別/活動量", "label-xs")}
    {tl(460, 133, "food_context      搜過的食物", "label-xs")}
    {tl(460, 151, "recipe_entry      分析過的食譜", "label-xs")}
    {tl(460, 169, "symptom_history   查過的症狀", "label-xs")}
    {tl(460, 187, "ranking_history   查過的排名", "label-xs")}
    {tl(460, 205, "calendar_context  行事曆查詢", "label-xs")}
    {tl(460, 223, "academic_context  學術查詢", "label-xs")}
    {tl(460, 241, "recent_actions    最近動作記錄", "label-xs")}
    {tl(460, 259, "dri_awareness     營養缺口意識", "label-xs")}
    {tl(460, 277, "saved_recipes     食譜摘要", "label-xs")}
    {tl(460, 295, "papers_referenced 引用的論文", "label-xs")}
    {tl(460, 313, "action_entry      本輪動作", "label-xs")}

    <!-- Checkpoint -->
    {rr_bold(30, 350, 790, 60, "#FAF5FF", "#805AD5")}
    {tc(425, 375, "Checkpoint 持久化（存到 PostgreSQL）", "label-bold")}
    {tc(425, 395, "session_id + turn_index + workflow_state + narrative + structured_store + tool_results", "label-xs")}

    {rr(30, 425, 790, 40, "#FFFFF0", "#D69E2E", 6)}
    {tc(425, 450, "🍳 記憶樹=服務日誌的年度總結；結構化儲存=客人的口味檔案；Checkpoint=存檔點（斷電也不怕）", "label-xs")}
    '''
    return svg_wrap(content, w, h)


# ═══════════════════════════════════════════════════════════
# Section D: Document Content
# ═══════════════════════════════════════════════════════════

def chapter_0():
    """前言"""
    h1("前言")
    p("這份文件是一份完整的技術教學，帶你了解 Recipe-AI 系統背後的「AI 大腦」是怎麼運作的。")
    p("我們會從最簡單的概念開始，一步步帶你認識：")
    bullet("AI 怎麼「聽懂」你說的話（意圖判定）")
    bullet("AI 怎麼「決定」要做什麼事（工具調用）")
    bullet("AI 怎麼「記住」你之前說過的話（記憶系統）")
    bullet("AI 怎麼處理你「突然改主意」（中斷處理）")
    bullet("整個系統的「溝通方式」（通信協議）")

    p("所有內容都會用日常生活的比喻來解釋，讓沒學過程式設計的人也能理解。專有名詞會在第一次出現時加上解釋。")

    analogy_box("全文核心比喻：高級餐廳",
        "想像 Recipe-AI 是一間高級餐廳。你（使用者）坐在用餐區（前端畫面），告訴服務生你要什麼。"
        "服務生把你的需求轉達給主廚（Control 服務），主廚用他的大腦（LLM 語言模型）判斷該做什麼，"
        "然後指揮不同的工作站（工具）來備料、烹飪。最後，服務生一步步把進度回報給你——"
        "這就是整個系統運作的縮影。")


def chapter_1():
    """系統全貌"""
    h1("第一章：系統全貌")

    p("在深入細節之前，先來看整個系統的「地圖」。Recipe-AI 不是一個單獨的程式，"
      "而是由好幾個獨立的「服務」合作組成的。就像一間餐廳不只有廚房，還有前台、倉庫、圖書館。")

    h2("1.1 系統架構圖")
    add_image(make_svg_architecture())

    h2("1.2 各服務說明")
    add_table(
        ["服務名稱", "角色", "技術", "Port", "餐廳比喻"],
        [
            ["Frontend", "使用者看到的網頁畫面", "React + Vite", "3001", "用餐區"],
            ["Control", "AI 主控台，負責調度一切", "FastAPI + Python", "8000", "主廚"],
            ["Backend", "資料庫操作（食品、食譜、使用者）", "FastAPI + asyncpg", "8001", "食材倉庫"],
            ["Gemma LLM", "語言模型，負責「思考」", "llama-server", "8080", "主廚的大腦"],
            ["GraphRAG API", "學術論文知識圖譜", "FastAPI + ChromaDB", "8002", "圖書館"],
            ["Health Vector", "症狀語義搜尋", "FastAPI + ChromaDB", "8003", "症狀百科"],
            ["PostgreSQL", "關聯式資料庫", "postgres:16", "5432", "檔案室"],
        ]
    )

    tip_box("什麼是「Port」？",
        "Port（埠號）就像公寓的門牌號碼。同一台電腦上可以跑好多個服務，"
        "每個服務佔一個不同的 Port。例如你在瀏覽器輸入 localhost:3001 就是去前端的「門」，"
        "localhost:8000 就是去 Control 的「門」。")

    h2("1.3 資料怎麼流動")
    p("使用者在前端（:3001）打字 → 前端把訊息送到 Control（:8000）→ "
      "Control 叫 LLM（:8080）思考 → Control 呼叫 Backend（:8001）查資料 → "
      "結果一路回傳到前端顯示。")
    p("所有的溝通都經過 Control 這個「主廚」中轉。前端不會直接跟 Backend 對話，"
      "就像客人不會直接衝進倉庫拿食材一樣。")

    tip_box("什麼是 FastAPI？",
        "FastAPI 是一個用 Python 寫的「網路服務框架」。簡單說，它讓你的 Python 程式"
        "能夠接受來自網路的請求（就像接電話），處理後回覆結果。"
        "Control、Backend、GraphRAG 都是用 FastAPI 寫的。")


def chapter_2():
    """六步驟 LLM Pipeline"""
    h1("第二章：六步驟 LLM Pipeline — AI 處理你每句話的流程")

    p("當你在聊天框打一句話（比如「我最近常覺得疲勞」），AI 並不是直接去搜尋答案。"
      "它會經過六個精心設計的步驟，就像高級餐廳處理一張訂單一樣有條不紊。")

    analogy_box("Pipeline 是什麼？",
        "Pipeline（流水線）就是「一連串按順序執行的步驟」。就像工廠的生產線："
        "原料進去 → 第一站加工 → 第二站組裝 → … → 成品出來。"
        "每一站做一件特定的事，做完交給下一站。")

    h2("2.0 全景圖")
    add_image(make_svg_pipeline_overview())

    # ── Step 1: Echo ──
    h2("2.1 Step 1: Echo — 即時回應「收到了！」")

    p("你在餐廳跟服務生說「我要一份牛排」，服務生不會沉默不語地走掉。"
      "他會先說「好的，一份牛排，馬上為您準備！」讓你知道他聽到了。")
    p("Echo 就是這個「收到確認」。AI 收到你的訊息後，會在 0.5 秒內先回一句簡短的確認，"
      "讓你知道它正在處理。但它不會在這個階段回答你的問題——它只是說「我聽到了，正在查」。")

    pb("溫度（Temperature）= 0.5")
    tip_box("什麼是「溫度」？",
        "語言模型在生成文字時，有一個叫「溫度」的參數。想像你在烤箱烤東西：\n"
        "• 溫度低（0.1）= 烤出來的東西很穩定、很規律，每次差不多。用在需要精準回答的場合。\n"
        "• 溫度高（0.5-1.0）= 烤出來的東西比較多變、有創意。用在需要自然對話的場合。\n"
        "Echo 用 0.5 是因為它只是打招呼，稍微有點變化比較自然。")

    pb("完整 Prompt 內容（PROMPT_ECHO，advisor.py 第 51-84 行）：")
    prompt_box("PROMPT_ECHO", """請你根據使用者的訊息和對話脈絡，用繁體中文回應。

## 規則
- 簡潔回應
- 如果使用者的訊息需要查詢工具（記錄飲食、查營養、查症狀等）：概括需求，表達「已收到、正在處理」
- 如果是一般對話（打招呼、感謝、閒聊）：直接給出友善回應
- 不要回答營養問題本身，不要提供營養數據或建議
- 語氣親切自然

## 範例
使用者：「營養夠嗎」→ 收到，我先查看你的個人資料，再幫你分析今天的營養攝取是否達標。
使用者：「為什麼蛋白質對肌肉修復很重要？」→ 好問題，我來搜尋相關的學術文獻幫你找答案。
使用者：「你好」→ 你好！我是你的 AI 營養顧問，有什麼飲食或營養問題可以幫你的嗎？

---
{memory_block}       ← 之前的對話記憶
對話歷史：{history}  ← 最近幾輪對話
使用者訊息：{message} ← 你剛剛打的那句話

請直接輸出回應（不要加引號或其他格式）：""")

    p("這段 Prompt 裡用了 {placeholder}（佔位符），系統會在送出前把它們替換成真正的內容。"
      "例如 {message} 會被替換成你打的那句話，{history} 會被替換成之前的對話記錄。")

    pb("SSE 事件產出：")
    code('event: phase_status\ndata: {"phase": "echo", "label": "理解需求中..."}\n\nevent: echo\ndata: {"content": "了解，我來查詢與疲勞相關的營養素和推薦食物。"}')

    tip_box("什麼是 SSE？",
        "SSE 全名是 Server-Sent Events（伺服器推送事件）。普通的網路請求是「你問一次，我答一次」。\n"
        "但 AI 的回答需要時間，我們不想讓你等到天荒地老才看到結果。\n"
        "SSE 就像服務生一直跑過來跟你報告進度：「食材找到了！」「正在煎了！」「快好了！」「上菜！」\n"
        "它讓伺服器可以「主動」一直送訊息給你的瀏覽器，不用你一直重新問。")

    # ── Step 2: Phase 1 ──
    h2("2.2 Step 2: Phase 1 — 意圖判定「你到底想做什麼？」")

    p("服務生聽到你說的話後，第一件事是判斷你想要什麼。"
      "你說「牛排」很明確，但如果你說「我今天想吃清淡一點」呢？服務生就要動腦判斷了。")
    p("Phase 1 就是讓 AI 判斷你的「意圖」。系統定義了 15 種合法的意圖，"
      "AI 必須從中選擇一個或多個。")

    add_image(make_svg_phase1_decision())

    pb("三條判斷路徑：")
    p("① Fast Path（快速路徑）：用正則表達式（一種文字比對規則）直接判斷。"
      "如果你的訊息很明確，例如「鈣含量最高的食物」，系統不需要問 LLM，直接就知道你要查「營養素排名」。")
    p("② Manual RAG 上下文注入：在送給 LLM 之前，系統會先從「系統說明書」中搜尋跟你問題最相關的章節，"
      "附在 Prompt 裡讓 LLM 參考。這樣 LLM 就不需要靠記憶，而是「翻書」來判斷。")
    p("③ LLM 判定：帶著 RAG 上下文的 Prompt 送給 Gemma LLM，用溫度 0.1（非常精準）判斷意圖。")

    tip_box("什麼是「正則表達式」（Regex）？",
        "正則表達式是一種「文字搜尋的公式」。例如：\n"
        "• 「最高|排名|最多|含量」這個 regex 會比對包含「最高」「排名」「最多」「含量」任一個詞的句子\n"
        "• 如果你說「鈣含量最高的食物」→ 命中！系統知道你要查營養素排名\n"
        "好處是非常快（不需要呼叫 LLM），壞處是只能處理固定模式的句子")

    tip_box("什麼是 RAG？",
        "RAG = Retrieval-Augmented Generation（檢索增強生成）。\n"
        "簡單說就是「開書考試」：AI 在回答之前，先去翻一本書（系統說明書），找到最相關的段落，\n"
        "然後帶著這些「小抄」來回答你的問題。這樣比純靠記憶準確很多。")

    pb("完整 Prompt 內容（PROMPT_PHASE1，advisor.py 第 92-189 行）：")
    p("這個 Prompt 很長，因為它要告訴 LLM 所有 15 種意圖的定義、參數格式、確認詞處理規則等。"
      "以下是關鍵段落：")
    prompt_box("PROMPT_PHASE1（核心段落）", """你是 AI 營養顧問系統的意圖判定器。今天是 {current_date}。
根據使用者訊息、對話脈絡和下方的相關功能說明，判斷需要呼叫哪些工具。

## 與使用者問題最相關的系統功能
{relevant_manual_context}    ← RAG 搜出的系統說明書段落

## 可用工具（intents 欄位只允許下表名稱）
| intent           | 必要參數                    | 說明               |
| search_symptom   | symptom（中文症狀描述）      | 搜尋症狀相關營養素   |
| analyze_recipe   | recipe_text                | 分析食譜營養成分     |
| nutrient_ranking | nutrient_field             | 營養素含量排名       |
| ... (共 15 種)

## 確認詞處理
使用者說「好」「好啊」「可以」「對」「嗯」= 確認上一輪 AI 提議的動作。

## 複合任務
「今天吃了雞胸肉和花椰菜，加到5/20行事曆」→ intents: ["analyze_recipe","add_to_calendar"]

## 輸出格式（嚴格 JSON）
{"intents":["intent_name"], "entities":{...}, "needs_profile":false, "reasoning":"一句話說明"}""")

    # ── Step 3: Phase 2 ──
    h2("2.3 Step 3: Phase 2 — 工具調用生成「派誰去做？」")

    p("Phase 1 告訴我們「你想做什麼」，Phase 2 決定「具體怎麼做」。"
      "就像主廚知道客人要「海鮮拼盤」後，要決定派哪個工作站、用什麼食材、怎麼配量。")
    p("Phase 2 有一個聰明的設計：三層安全網。")

    add_image(make_svg_phase2_cascade())

    pb("第一層：系統確定性生成（_system_generate_toolcalls）")
    p("不需要問 LLM，系統直接根據 Phase 1 的結果和已知資訊生成工具呼叫。"
      "這一層速度極快（幾毫秒）。大約 90% 的情況在這裡就解決了。")

    pb("第二層：LLM 備援（PROMPT_PHASE2）")
    p("當第一層搞不定（例如參數需要從對話中推斷），才呼叫 LLM。溫度 0.1，非常精準。")

    pb("第三層：系統備援（_system_fallback）")
    p("如果 LLM 也失敗了（例如模型掛了），系統會用最簡單的預設值來生成工具呼叫。"
      "這確保「永遠不會什麼都做不了」。")

    # ── Step 4: Execute ──
    h2("2.4 Step 4: Execute — 工具執行「開始烹飪！」")

    p("現在我們知道要用哪些工具、帶什麼參數了。是時候真正「動手」了。")
    p("系統會根據工具之間的相依關係，建立一個「任務有向無環圖」（DAG）。"
      "聽起來很複雜？其實就是：")

    analogy_box("什麼是 DAG？",
        "想像你要做一桌菜。有些菜可以同時做（沙拉和湯可以同時準備），"
        "有些必須按順序（先煎牛排，才能把牛排擺盤）。\n"
        "DAG 就是一張「任務先後順序圖」：沒有相依的任務同時開工，有相依的按順序排隊。")

    add_image(make_svg_dag_execution())

    pb("Recovery 機制：")
    p("如果工具執行失敗了（例如搜尋「芝士」找不到結果），系統不會直接放棄。"
      "它有一個「自癒引擎」會自動診斷問題、嘗試修復。詳見第八章。")

    # ── Step 5: Analysis ──
    h2("2.5 Step 5: Analysis — 分析回覆「為您上菜，並解說」")

    p("工具跑完了，拿到了一堆原始資料。但直接把資料丟給你看太粗暴了——"
      "就像主廚不會把一塊生肉直接端上桌。")
    p("Analysis 步驟讓 LLM 扮演「分析師」：閱讀工具回傳的所有資料，"
      "用你能理解的語言告訴你重點是什麼。")

    pb("溫度 0.3（平衡精準與自然）")

    pb("完整 Prompt 內容（PROMPT_ANALYSIS，advisor.py 第 262-310 行）：")
    prompt_box("PROMPT_ANALYSIS（核心規則）", """請你根據工具查詢結果，用繁體中文提供專業營養分析回覆。

## 核心規則
- 所有數據必須來自下方的「工具執行結果」，禁止從訓練資料引用任何營養數字
- 沒有查到相關數據就明確說「目前沒有查到相關數據」
- 前端已自動渲染工具結果的表格和卡片，你不需要列出原始數據清單
- 你的角色是「分析師」：解讀數據的意義，不是複述數據

### search_graphrag（學術文獻搜尋）— 最高優先規則
你必須將 answer 和 evidence 的完整內容直接呈現，禁止摘要、禁止縮減、禁止改寫。

### 其他搜尋類工具
前端已渲染完整互動式卡片。你的回覆應：
- 一句說找到什麼
- 一句指出最重要的洞見
- 可選一句建議下一步

---
使用者原始問題：{message}
{memory_block}
工具執行結果：{tool_results}
請用繁體中文輸出分析回覆：""")

    p("注意那條規則：「禁止從訓練資料引用任何營養數字」。"
      "這是因為 LLM 的訓練資料可能過時或不準確，所有數字都必須來自剛剛查到的工具結果。"
      "就像主廚不能憑記憶說「這道菜有 500 大卡」，要看實際食材的標籤。")

    # ── Step 6: Dual Memory ──
    h2("2.6 Step 6: Dual Memory — 雙記憶更新「記在服務紀錄上」")

    p("一輪對話結束後，系統會「記住」這次發生了什麼。"
      "但它不是把所有東西都塞進一個大桶——它用兩種不同的方式來記：")

    add_image(make_svg_dual_memory())

    pb("敘事記憶（Narrative Memory）")
    p("記住「對話流程」：誰問了什麼、AI 怎麼回的、話題怎麼轉、使用者的態度。"
      "不記具體數字（那是結構化記憶的工作）。")

    pb("結構化記憶（Structured Store）")
    p("記住「具體數據」：使用者年齡 28 歲、上次查了「疲勞」、鈣攝取 500mg。"
      "用 12 類固定欄位的 JSON 格式存儲。")

    p("這兩種記憶都在「背景」執行——不會讓你等。系統一邊回覆你，一邊在幕後更新記憶。")

    page_break()


def chapter_3():
    """十五個工具全解析"""
    h1("第三章：十五個工具全解析")

    p("上一章講了六步驟 Pipeline 的全貌。這一章我們要把 Step 4「執行」中用到的 15 個工具逐一攤開看。")
    p("每個工具就像餐廳裡的一個專業工作站——有的負責查食材百科，有的負責算營養，有的負責管行事曆。")

    add_image(make_svg_tools_overview())

    h2("3.1 搜尋類工具")

    # search_food
    h3("search_food — 搜尋食物營養成分")
    add_table(
        ["項目", "說明"],
        [
            ["使用者說", "「雞蛋的營養成分是什麼？」"],
            ["意圖名稱", "search_food"],
            ["工具名稱", "search_food"],
            ["必要參數", "food_name: 食物名稱（如「雞蛋」）"],
            ["後端 API", "GET /food/search?q=雞蛋"],
            ["回傳什麼", "台灣食品 + FooDB 雙資料庫的匹配食物清單，含營養成分"],
            ["前端顯示", "食物卡片，可展開查看完整營養資料"],
            ["人工操作等價", "你自己去「食材搜尋」頁面，打字搜尋「雞蛋」"],
        ]
    )

    # search_symptom
    h3("search_symptom — 搜尋症狀相關營養素")
    p("這是最複雜的工具，內部跑了一個五階段管道：")
    add_image(make_svg_symptom_pipeline())
    add_table(
        ["項目", "說明"],
        [
            ["使用者說", "「最近常覺得疲勞」"],
            ["意圖名稱", "search_symptom"],
            ["工具名稱", "search_symptom"],
            ["必要參數", "symptom: 症狀描述（中文，如「疲勞」）"],
            ["後端 API", "POST /health/symptom-search"],
            ["回傳什麼", "健康效果、相關化合物、推薦食物"],
            ["前端顯示", "LLM 解析標籤 + 向量匹配區塊 + 效果卡片"],
            ["人工操作等價", "你自己去「症狀搜尋」頁面，輸入「疲勞」"],
        ]
    )

    # analyze_recipe
    h3("analyze_recipe — 分析食譜營養成分")
    add_image(make_svg_recipe_flow())
    add_table(
        ["項目", "說明"],
        [
            ["使用者說", "「今天吃了雞胸肉 200g 和花椰菜 100g」"],
            ["意圖名稱", "analyze_recipe"],
            ["工具名稱", "analyze_recipe"],
            ["必要參數", "recipe_text: 食譜文字"],
            ["後端 API", "POST /recipe/extract-with-amounts → POST /recipe/calculate-nutrition"],
            ["回傳什麼", "食材清單（Phase A）→ 營養計算 + DRI 缺口（Phase B）"],
            ["前端顯示", "確認食材卡片 → 營養表格 + 缺口分析"],
            ["人工操作等價", "你自己去「整合分析」頁面，貼上食譜文字"],
        ]
    )

    # delete_saved_recipe
    h3("delete_saved_recipe — 刪除已存食譜")
    add_image(make_svg_delete_confirm())
    add_table(
        ["項目", "說明"],
        [
            ["使用者說", "「幫我刪掉昨天的食譜」"],
            ["意圖名稱", "delete_recipe"],
            ["工具名稱", "delete_saved_recipe"],
            ["流程", "先預覽匹配食譜 → 使用者確認 → 實際刪除"],
            ["後端 API", "GET /recipe/list → DELETE /recipe/{id}"],
            ["安全機制", "第一次呼叫 confirmed=false（只列表），確認後 confirmed=true"],
        ]
    )

    h3("其餘工具快速一覽")
    add_table(
        ["工具名稱", "意圖名稱", "後端 API", "說明"],
        [
            ["get_nutrient_ranking", "nutrient_ranking", "GET /health/nutrients/{field}/top-foods", "營養素含量排名"],
            ["browse_taiwan_food", "browse_food_database", "GET /food/taiwan/browse", "瀏覽台灣食品資料庫"],
            ["search_graphrag", "search_graphrag", "POST localhost:8002/query-sync", "學術論文知識圖譜問答"],
            ["get_literature_papers", "browse_literature", "GET localhost:8002/papers/enriched", "瀏覽學術論文"],
            ["synthesize_advice", "synthesize_advice", "（內部計算）", "根據已分析食譜生成營養建議"],
            ["get_saved_recipes", "get_recipes", "GET /recipe/list", "查看已儲存食譜"],
            ["get_calendar_entries", "get_calendar", "GET /calendar/entries", "行事曆查詢 + DRI 缺口"],
            ["add_to_calendar", "add_to_calendar", "POST /recipe/save → POST /calendar/entries", "儲存食譜並加入行事曆"],
            ["delete_calendar_entry", "delete_calendar", "GET/DELETE /calendar/entries/{id}", "刪除行事曆紀錄"],
            ["get_user_profile", "get_profile", "GET /auth/me", "讀取個人資料"],
            ["update_user_profile", "update_profile", "PUT /profile", "更新年齡/性別/身高等"],
        ]
    )
    page_break()


def chapter_4():
    """九個 Prompt 範本"""
    h1("第四章：九個 Prompt 範本")

    p("Prompt（提示詞）就是我們給 LLM 的「指令卡」。每張卡告訴 LLM：")
    bullet("你現在扮演什麼角色")
    bullet("你有哪些工具可以用")
    bullet("你應該用什麼格式回答")
    bullet("哪些事情絕對不能做")

    add_image(make_svg_prompt_lifecycle())

    add_table(
        ["Prompt 名稱", "位置", "步驟", "溫度", "用途"],
        [
            ["PROMPT_ECHO", "L51", "Echo", "0.5", "即時確認「收到了」"],
            ["PROMPT_PHASE1", "L92", "Phase1", "0.1", "判斷使用者的 15 種意圖"],
            ["PROMPT_PHASE2", "L196", "Phase2", "0.1", "生成具體工具呼叫指令"],
            ["PROMPT_ANALYSIS", "L262", "Analysis", "0.3", "根據工具結果生成分析回覆"],
            ["PROMPT_NARRATIVE_SUMMARY", "L317", "Memory", "0.2", "更新對話脈絡摘要"],
            ["PROMPT_STRUCTURED_EXTRACT", "L341", "Memory", "0.1", "提取結構化資料到 JSON"],
            ["PROMPT_CLASSIFY_REPLY", "L3266", "中斷偵測", "0.0", "判斷使用者是回答問題還是換話題"],
            ["PROMPT_CLASSIFY_RESUME", "L3276", "恢復偵測", "0.0", "判斷使用者是否想恢復之前的任務"],
            ["PROMPT_EXTRACT_PARAMS", "L3586", "參數收集", "0.1", "從訊息中提取參數值（LLM 備援）"],
        ]
    )

    p("前四個 Prompt（ECHO、PHASE1、PHASE2、ANALYSIS）在每次正常對話中都會用到。"
      "後面五個是特殊情境才觸發：收集參數時判斷中斷、恢復中斷任務、記憶更新。")
    page_break()


def chapter_5():
    """SSE 通信協議"""
    h1("第五章：SSE 通信協議 — 即時通報進度")

    h2("5.1 為什麼需要 SSE？")
    p("普通的網路請求是「一問一答」：你發一個請求，伺服器算完後回你一個完整的回應。"
      "但 AI 處理一個問題可能要 5-10 秒，如果讓你乾等到最後才看到結果，體驗會很差。")

    analogy_box("SSE vs 普通請求",
        "普通請求：你跟服務生說「我要牛排」→ 服務生消失 10 分鐘 → 突然端出完整的牛排\n"
        "SSE 模式：你跟服務生說「我要牛排」→ 服務生每隔一會兒回來報告：\n"
        "「食材找到了！」→「正在煎了！」→「快好了！」→「上菜！」")

    h2("5.2 SSE 事件時間軸")
    add_image(make_svg_sse_timeline())

    h2("5.3 SSE 格式說明")
    p("SSE 的格式非常簡單，就是一連串的「事件」，每個事件長這樣：")
    code('event: tool_result\ndata: {"name":"search_food","data":{"results":[...]}}\n')
    p("前端的 JavaScript 收到後，根據 event 的類型做不同的事："
      "收到 echo 就顯示確認訊息、收到 tool_result 就渲染工具結果卡片、收到 reply 就顯示最終回覆。")

    h2("5.4 通信鏈全貌")
    add_image(make_svg_communication_chain())

    h2("5.5 完整範例：「雞蛋的營養」")
    p("讓我們追蹤一個簡單查詢的完整事件流：")
    code("""使用者打字：「雞蛋的營養」

→ Frontend 發出 POST /advisor/chat {message:"雞蛋的營養", ...}
← event: phase_status  {"phase":"echo","label":"理解需求中..."}
← event: echo          {"content":"收到，我來查詢雞蛋的營養成分。"}
← event: phase_status  {"phase":"intent","label":"分析意圖中..."}
   (Control 內部：Phase1 → intent=search_food, entities={food_name:"雞蛋"})
   (Control 內部：Phase2 → tool_calls=[{tool:"search_food",args:{food_name:"雞蛋"}}])
← event: tool_start    {"name":"search_food","label":"搜尋食物營養資料"}
   (Control → Backend: GET /food/search?q=雞蛋)
   (Backend 查 taiwan_foods + FooDB → 回傳 5 筆匹配)
← event: tool_result   {"name":"search_food","data":{"results":[...]}}
← event: tool_done     {"name":"search_food"}
← event: phase_status  {"phase":"analysis","label":"分析結果中..."}
← event: thinking_start {"step":"analysis"}
← event: thinking_delta {"content":"根據查詢結果，"}
← event: thinking_delta {"content":"雞蛋含有豐富的蛋白質"}
← event: thinking_delta {"content":"（每100g 約12.5g）..."}
← event: thinking_end   {"step":"analysis"}
← event: reply         {"content":"根據查詢結果，雞蛋含有豐富的蛋白質..."}
← event: workflow_state {"phase":"done",...}
← event: narrative_memory {"summary":"使用者查了雞蛋營養..."}
← event: structured_store {"store":{...}}
← event: checkpoint_saved {"session_id":"abc123","turn_index":1}""")

    page_break()


def chapter_6():
    """中斷處理六個派別"""
    h1("第六章：中斷處理的六個派別")

    p("在對話中，使用者隨時可能「改主意」。AI 需要聰明地判斷：你是在回答我的問題，"
      "還是突然想做別的事？")
    p("Recipe-AI 有六套不同的中斷處理機制，各自負責不同的情境。"
      "我們用「六個派別」來稱呼它們——每個派別有自己的「武器」和「戰場」。")

    add_image(make_svg_six_factions_overview())

    h2("6.1 智慧派 — LLM 三態判斷")
    p("戰場：使用者在「收集參數」階段說了一句話。"
      "AI 正在問你「你有什麼症狀？」，你回了一句話，智慧派要判斷：")
    bullet("answer（回答）：你提供了 AI 問的資訊")
    bullet("interrupt（中斷）：你在說別的事")
    bullet("ambiguous（模糊）：不確定，需要追問釐清")

    add_image(make_svg_tristate())

    p("智慧派的聰明之處：它同時用 LLM 和 Regex 兩種方法判斷，然後比較結果。"
      "如果兩者一致，就採用；如果分歧，就回報「ambiguous」然後問使用者。")

    prompt_box("PROMPT_CLASSIFY_REPLY（L3266）", """AI 正在執行「{intent_zh}」，等使用者提供「{missing_desc}」。

使用者說：「{message}」

這句話是在回答上述問題嗎？
A = 是（提供了所需的資訊）
B = 否（在說別的事、問新問題、放棄、或換話題）

只輸出 A 或 B：""")

    h2("6.2 規則派 — Regex 正則比對")
    p("戰場：跟智慧派一樣，但規則派用的是「固定規則」而非 LLM。"
      "它是智慧派的 fallback（備援）——如果 LLM 掛了，規則派頂上。")
    p("規則派有三把武器（三個正則表達式）：")
    add_table(
        ["武器", "比對什麼", "範例"],
        [
            ["_QUESTION_RE", "問句模式", "「什麼是...？」「有什麼...」"],
            ["_RETURN_RE", "回歸模式", "「回到剛才」「繼續之前的」"],
            ["_NEW_INTENT_RE", "新意圖模式", "「幫我記錄」「幫我查」「幫我刪」"],
        ]
    )
    p("它還會根據目前缺少的參數類型做特化判斷。例如缺少 age 時，"
      "看到數字就認為是在回答年齡（不是中斷）。")

    h2("6.3 態度派 — 收集回覆分類")
    add_image(make_svg_collecting_states())

    p("態度派不判斷「中斷」，而是判斷使用者的「態度」：是認真回答、猶豫不決、還是明確放棄？")

    h2("6.4 回歸派 — 中斷恢復偵測")
    add_image(make_svg_resume_flow())

    p("當使用者中斷後，系統會保存中斷狀態。下次使用者提到相關內容時，"
      "回歸派負責判斷「你是想回到之前的任務嗎？」")

    prompt_box("PROMPT_CLASSIFY_RESUME（L3276）", """使用者之前有未完成的任務：「{intent_zh}」（還需要：{missing_desc}）。

使用者現在說：「{message}」

這句話是想回到那個未完成的任務嗎？
A = 是
B = 否（在說新的事情）

只輸出 A 或 B：""")

    h2("6.5 快捷派 — [RECIPE_CONFIRM] 食譜確認協議")
    add_image(make_svg_recipe_confirm())

    p("快捷派是一種「特殊通道」——前端送出的訊息如果以 [RECIPE_CONFIRM] 開頭，"
      "就跳過整個 6 步 Pipeline，直接由特殊處理器接手。")

    h2("6.6 前線派 — Frontend AbortController")
    p("前線派是唯一在「前端」運作的中斷機制。當使用者按下「停止」按鈕時，"
      "前端會呼叫 AbortController 立即中斷 SSE 連線。")
    code("// 前端 advisorApi.js\nconst controller = new AbortController();\nfetch('/api/advisor/chat', { signal: controller.signal, ... });\n// 使用者按停止 → controller.abort() → SSE 連線中斷")

    page_break()


def chapter_7():
    """複合問題處理"""
    h1("第七章：複合問題處理")

    p("使用者常常在一句話裡提出多個需求。"
      "例如：「分析雞胸肉食譜，加到行事曆，告訴我還缺什麼營養」——這裡有三個意圖。")

    add_image(make_svg_compound_dag())

    h2("7.1 多意圖偵測")
    p("Phase 1 可以回傳多個 intent。系統會分析它們之間的相依關係，"
      "建立一個 DAG 來決定執行順序。")

    h2("7.2 相依規則")
    add_table(
        ["規則", "說明"],
        [
            ["analyze_recipe → add_to_calendar", "必須先分析食譜，才能把食譜加到行事曆"],
            ["analyze_recipe → synthesize_advice", "必須先有食譜資料，才能生成營養建議"],
            ["無相依 → 並行", "search_food + search_symptom 可以同時跑"],
        ]
    )

    h2("7.3 副作用傳播")
    p("當 analyze_recipe 完成後，它的結果會存入 structured_store 的 recipe_store。"
      "下游的 add_to_calendar 和 synthesize_advice 就能從那裡讀取資料。")
    page_break()


def chapter_8():
    """自癒引擎"""
    h1("第八章：自癒引擎（Recovery Engine）")

    p("工具不一定每次都能成功。搜尋可能找不到結果、連線可能逾時、參數可能不對。"
      "自癒引擎就是系統的「急救箱」。")

    add_image(make_svg_recovery_engine())

    h2("8.1 四種錯誤類型")
    add_table(
        ["錯誤類型", "觸發條件", "範例"],
        [
            ["empty_result", "搜尋回傳 0 筆結果", "search_food(「芝士」) → []"],
            ["tool_timeout", "工具執行超時", "search_graphrag 超過 60 秒"],
            ["param_mismatch", "參數不匹配", "nutrient_ranking(「維生素K」) 欄位不存在"],
            ["backend_error", "後端回 500 錯誤", "Backend 服務暫時掛掉"],
        ]
    )

    h2("8.2 四種恢復策略")
    add_table(
        ["策略", "做什麼", "範例"],
        [
            ["synonym_expand", "換同義詞重試", "「芝士」→「起司」→「乾酪」"],
            ["retry_timeout", "重新執行一次", "等一下再試"],
            ["field_remap", "欄位名稱映射", "「維生素K」→ vitamin_k_ug"],
            ["skip", "跳過這個工具", "告訴使用者「這個服務暫時不可用」"],
        ]
    )
    page_break()


def chapter_9():
    """工作流程狀態機"""
    h1("第九章：工作流程狀態機")

    p("狀態機就是一張「現在在哪裡、接下來去哪裡」的地圖。"
      "系統用它來追蹤每次對話的進度。")

    add_image(make_svg_state_machine())

    h2("9.1 WorkflowState 結構")
    code("""class WorkflowState:
    phase: str = "idle"              # 目前階段
    active_intent: str = ""          # 正在處理的意圖
    collected: dict = {}             # 已收集的參數
    missing: list = []               # 還缺的參數
    pending_intents: list = []       # 待處理的其他意圖
    turn_count: int = 0              # 收集輪次（最多 3 輪）
    interrupted_from: str = ""       # 從哪個意圖中斷的
    interrupted_state: dict = {}     # 中斷時保存的狀態""")
    page_break()


def chapter_10():
    """驗證與自修復"""
    h1("第十章：驗證與自修復管道")

    p("LLM 不是完美的——它可能回傳格式錯誤的 JSON、用錯意圖名稱、忘了必要參數。"
      "系統有六層驗證來確保每一步都正確。")

    add_image(make_svg_validation())

    add_table(
        ["驗證層", "做什麼", "修正範例"],
        [
            ["意圖別名修正", "把常見錯名映射到正確名稱", "get_user_profile → get_profile"],
            ["意圖合法性", "過濾不在 15 種之內的意圖", "chat → general_chat"],
            ["工具參數驗證", "型別轉換 + 去重 + 必填檢查", "age=\"28\" → age=28（字串轉數字）"],
            ["營養素欄位", "中文名 → 資料庫欄位名", "「鈣」→ calcium_per_100g"],
            ["分析品質", "最低 50 字，不夠重試", "太短 → 重新呼叫 LLM"],
            ["JSON 解析", "三級 fallback", "直接解析 → 去 fence → 深度配對"],
        ]
    )
    page_break()


def chapter_11():
    """記憶架構"""
    h1("第十一章：記憶架構")

    add_image(make_svg_memory_tree())

    h2("11.1 記憶樹")
    p("每輪對話產生一個 Leaf（葉子）節點。每 10 個 Leaf 合併成一個 Session 摘要。"
      "多個 Session 合併成 User 層級摘要。這樣即使對話很長，記憶也不會無限膨脹。")

    h2("11.2 Checkpoint 持久化")
    p("每輪對話結束後，系統會把當前的完整狀態存到 PostgreSQL 資料庫。"
      "如果突然斷線、瀏覽器關閉，下次回來時可以從上次的存檔點繼續。")
    page_break()


def chapter_12():
    """附錄"""
    h1("第十二章：附錄")

    h2("A. 專有名詞中英對照")
    add_table(
        ["中文", "英文", "簡單解釋"],
        [
            ["意圖", "Intent", "使用者想做什麼事"],
            ["工具", "Tool", "系統用來完成任務的模組"],
            ["溫度", "Temperature", "控制 LLM 回答創意程度的旋鈕"],
            ["管道", "Pipeline", "一連串按順序執行的步驟"],
            ["有向無環圖", "DAG", "任務先後順序圖"],
            ["伺服器推送事件", "SSE", "伺服器主動持續送訊息給瀏覽器"],
            ["提示詞", "Prompt", "給 AI 的指令卡"],
            ["狀態機", "State Machine", "追蹤「現在在哪、下一步去哪」的機制"],
            ["正則表達式", "Regex", "文字搜尋的公式"],
            ["檢索增強生成", "RAG", "AI 先翻書再回答"],
            ["令牌", "Token", "LLM 處理文字的最小單位"],
            ["同義詞擴展", "Synonym Expand", "用其他說法重新搜尋"],
            ["斷詞", "Tokenization/Segmentation", "把一整句話切成一個個詞"],
            ["語義搜尋", "Semantic Search", "用「意思」而非「文字」來搜尋"],
            ["串流", "Streaming", "一邊產生一邊送出，不用等全部完成"],
            ["存檔點", "Checkpoint", "遊戲存檔，下次可以繼續"],
        ]
    )

    h2("B. 所有 SSE 事件類型")
    add_table(
        ["事件名稱", "觸發時機", "資料格式"],
        [
            ["phase_status", "每個步驟開始前", "{phase, label}"],
            ["echo", "Step 1 完成", "{content}"],
            ["thinking_start", "LLM 開始思考", "{step, label}"],
            ["thinking_delta", "LLM 產出一段文字", "{content}"],
            ["thinking_end", "LLM 思考結束", "{step}"],
            ["tool_start", "工具開始執行", "{name, label}"],
            ["tool_result", "工具執行完畢", "{name, data}"],
            ["tool_done", "工具結果已處理", "{name}"],
            ["reply", "最終回覆", "{content}"],
            ["collecting_state", "收集參數階段", "{intent, missing, turn_count}"],
            ["workflow_state", "狀態機更新", "{phase, active_intent, ...}"],
            ["narrative_memory", "敘事記憶更新", "{summary}"],
            ["structured_store", "結構化記憶更新", "{store}"],
            ["session_id", "Session 識別碼", "{session_id}"],
            ["checkpoint_saved", "存檔完成", "{session_id, turn_index}"],
            ["recovery_attempt", "自癒引擎嘗試", "{tool, strategy}"],
            ["recovery_result", "自癒結果", "{tool, success}"],
            ["error", "錯誤通知", "{message}"],
        ]
    )


# ═══════════════════════════════════════════════════════════
# Section E: Assembly
# ═══════════════════════════════════════════════════════════

def build():
    # Cover
    p("Recipe-AI", bold=True, size=28, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    p("Tool Calling & Agent 機制完整教學", bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    p("涵蓋：六步驟 LLM Pipeline ∙ 15 個工具 ∙ 9 個 Prompt ∙ SSE 通信協議 ∙ 六派中斷處理 ∙ 自癒引擎 ∙ 雙記憶系統",
      size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    p("以高中生也能聽懂的方式，深入淺出解析系統的每一個齒輪",
      italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    p(f"生成日期：2026-05-20", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    page_break()

    chapter_0()
    page_break()
    chapter_1()
    page_break()
    chapter_2()
    chapter_3()
    chapter_4()
    chapter_5()
    chapter_6()
    chapter_7()
    chapter_8()
    chapter_9()
    chapter_10()
    chapter_11()
    chapter_12()

    doc.save(OUT)
    print(f"✅ 文件已生成：{OUT}")
    print(f"   大小：{os.path.getsize(OUT) / 1024 / 1024:.1f} MB")


if __name__ == "__main__":
    build()
