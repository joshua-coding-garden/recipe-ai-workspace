#!/usr/bin/env python3
"""
RecipeAI Tool Calling 完整技術報告生成器
生成包含所有通信協議、程式碼解析、Prompt 完整內容的 Word 文件
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

SVG_DIR = os.path.join(os.path.dirname(__file__), "svg")
OUTPUT = os.path.join(os.path.dirname(__file__), "RecipeAI_Tool_Calling_完整技術報告.docx")

doc = Document()

# ═══════════════════════════════════════════════════════════════
# 樣式設定
# ═══════════════════════════════════════════════════════════════

style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft JhengHei'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')

for level in range(1, 5):
    hs = doc.styles[f'Heading {level}']
    hf = hs.font
    hf.name = 'Microsoft JhengHei'
    hf.bold = True
    hf.color.rgb = RGBColor(0x1A, 0x47, 0x8A)
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
    if level == 1:
        hf.size = Pt(22)
    elif level == 2:
        hf.size = Pt(16)
    elif level == 3:
        hf.size = Pt(13)
    else:
        hf.size = Pt(11)

# 程式碼樣式
code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_font = code_style.font
code_font.name = 'Consolas'
code_font.size = Pt(9)
code_font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
code_style.paragraph_format.space_before = Pt(2)
code_style.paragraph_format.space_after = Pt(2)
code_style.paragraph_format.left_indent = Cm(0.5)

def add_code_block(text, label=""):
    """加入程式碼區塊，帶灰色背景"""
    if label:
        p = doc.add_paragraph()
        run = p.add_run(f"📄 {label}")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x4A, 0x90, 0xD9)
    for line in text.strip().split('\n'):
        p = doc.add_paragraph(line, style='CodeBlock')
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
        p.paragraph_format.element.get_or_add_pPr().append(shading)

def add_table(headers, rows, col_widths=None):
    """加入格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表頭
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1A478A" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # 資料列
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F4FA" w:val="clear"/>')
                cell._tc.get_or_add_tcPr().append(shading)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_svg_reference(filename, caption):
    """加入 SVG 圖表參考"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[圖表: {caption}]")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xCC, 0x55, 0x00)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"請參閱附件 SVG 檔案: svg/{filename}")
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run2.italic = True

def add_note(text):
    """加入提示框"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    run = p.add_run(f">>> {text}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x8B, 0x45, 0x13)
    run.italic = True
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFF8E1" w:val="clear"/>')
    p.paragraph_format.element.get_or_add_pPr().append(shading)

# ═══════════════════════════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("RecipeAI 系統")
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Tool Calling 調用機制與通信協議\n完整技術報告")
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x4A, 0x90, 0xD9)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run(f"生成日期：{datetime.date.today().strftime('%Y-%m-%d')}\n涵蓋範圍：整合分析、食譜分析、營養素查詢、症狀搜尋、學術搜尋、AI 顧問等全部功能\n附件：12 張 SVG 通信流程圖")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 目錄提示
# ═══════════════════════════════════════════════════════════════

doc.add_heading('目錄', level=1)
toc_items = [
    "第一章  什麼是 Tool Calling？——從「人工操作」到「AI 自動操作」",
    "第二章  系統架構總覽——五個服務如何分工合作",
    "第三章  通信協議基礎——HTTP、SSE、JSON 是什麼？",
    "第四章  AI 營養顧問的 6 步驟 Pipeline——一句話如何變成答案",
    "第五章  所有 Tool 逐一拆解——11 個工具完整通信鏈",
    "  5.1  get_nutrient_ranking（營養素排名）",
    "  5.2  search_food（食物搜尋）",
    "  5.3  browse_taiwan_food（台灣食品資料庫瀏覽）",
    "  5.4  analyze_recipe（食譜分析——最複雜的工具）",
    "  5.5  search_symptom（症狀搜尋）",
    "  5.6  search_graphrag（學術論文搜尋）",
    "  5.7  get_literature_papers（文獻瀏覽）",
    "  5.8  get_saved_recipes（已存食譜）",
    "  5.9  get_calendar_entries / add_to_calendar（行事曆）",
    "  5.10 get_user_profile / update_user_profile（個人資料）",
    "第六章  整合分析頁——不經 AI 顧問的 7 步驟直接操作流程",
    "第七章  SSE 通信協議完整解析——每一個事件的意義",
    "第八章  LLM 思考邏輯與交叉互動機制",
    "第九章  雙記憶體系統——AI 如何「記住」你說過的話",
    "第十章  所有 Prompt 完整內容與通信鏈對應",
    "第十一章  前端代理路由與 API 調用機制",
    "第十二章  Gemma LLM 整合細節——食材提取與驗證的子 Prompt",
    "附錄 A  所有 API 端點清單",
    "附錄 B  SVG 圖表索引",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if item.startswith("第"):
        for run in p.runs:
            run.bold = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 第一章
# ═══════════════════════════════════════════════════════════════

doc.add_heading('第一章 什麼是 Tool Calling？', level=1)
doc.add_heading('——從「人工操作」到「AI 自動操作」', level=2)

doc.add_heading('1.1 先想像一個場景', level=3)
doc.add_paragraph(
    '假設你今天想知道「哪些食物鈣含量最高？」。'
    '如果沒有 AI，你需要自己做這些事：')
doc.add_paragraph('1. 打開瀏覽器，進入食品營養網站', style='List Number')
doc.add_paragraph('2. 找到「營養素排名」功能', style='List Number')
doc.add_paragraph('3. 從下拉選單選擇「礦物質」分組', style='List Number')
doc.add_paragraph('4. 再選「鈣 (calcium_per_100g)」', style='List Number')
doc.add_paragraph('5. 設定顯示前 10 名', style='List Number')
doc.add_paragraph('6. 點擊搜尋按鈕', style='List Number')
doc.add_paragraph('7. 等待結果表格出現', style='List Number')
doc.add_paragraph('8. 自己閱讀表格、分析哪些食物適合你', style='List Number')

doc.add_paragraph(
    '這就是「人工操作」——每一步都需要你自己點擊、選擇、等待。')

doc.add_paragraph(
    '現在有了 AI 營養顧問，你只需要說一句話：「幫我查鈣含量最高的食物」。'
    'AI 就會自動幫你完成上面所有步驟。但 AI 本身並不「知道」資料庫裡有什麼——'
    '它需要「工具」來幫它查資料。這就是 Tool Calling 的核心概念。')

doc.add_heading('1.2 Tool Calling 的白話解釋', level=3)
doc.add_paragraph(
    'Tool Calling（工具調用）就像是 AI 的「手」。AI 的「腦」（大型語言模型 LLM）'
    '負責理解你的意思、決定要做什麼；但它自己不能直接去資料庫查資料、'
    '不能去計算營養素、不能去搜尋症狀。它需要「呼叫工具」來完成這些動作。')

doc.add_paragraph('打個比方：')
doc.add_paragraph(
    '想像 AI 是一位很聰明的營養師，坐在辦公室裡。它聽得懂你的問題，也知道該怎麼回答。'
    '但是營養資料庫在另一棟大樓（後端伺服器）。所以這位營養師需要打電話（HTTP 請求）'
    '給資料庫管理員（Backend API），說：「幫我查 calcium_per_100g 欄位，排名前 10 的食物。」'
    '等到結果回來後，營養師再用自己的專業知識，幫你解讀這些數據。')

doc.add_paragraph(
    '這整個「營養師決定打哪通電話、說什麼內容、拿到結果後怎麼解讀」的過程，'
    '就是 Tool Calling 機制。')

doc.add_heading('1.3 Tool Calling 的三個核心問題', level=3)
doc.add_paragraph('每次 Tool Calling 都在解決三個問題：')

add_table(
    ['問題', '白話解釋', '系統中對應的步驟'],
    [
        ['要用哪個工具？', 'AI 聽完你的話後，判斷「你想做的事」對應到哪個工具', 'Phase 1：意圖偵測（Intent Detection）'],
        ['工具需要什麼參數？', '每個工具需要不同的「輸入材料」，AI 要準備好', 'Phase 2：工具指令生成（Tool Call Generation）'],
        ['拿到結果後怎麼回覆？', '工具回傳的是冰冷的數據，AI 要翻譯成你聽得懂的話', 'Analysis：綜合分析回覆'],
    ]
)

doc.add_heading('1.4 本系統有哪些工具？', level=3)
doc.add_paragraph(
    '本系統（RecipeAI）共有 11 個工具，分成五大類。'
    '你可以把它們想成 AI 營養師桌上的 11 支不同的電話，每支打給不同的部門：')

add_table(
    ['類別', '工具名稱', '白話功能', '打給誰（後端端點）'],
    [
        ['營養查詢', 'get_nutrient_ranking', '查營養素含量排名', 'GET /health/nutrients/{field}/top-foods'],
        ['營養查詢', 'search_food', '搜尋特定食物的營養成分', 'GET /food/search?q=...'],
        ['營養查詢', 'browse_taiwan_food', '瀏覽台灣食品資料庫', 'GET /food/taiwan/browse'],
        ['食譜分析', 'analyze_recipe', '分析一份食譜的完整營養', '4 步連鎖呼叫（最複雜）'],
        ['食譜分析', 'get_saved_recipes', '查看已儲存的食譜', 'GET /recipe/list'],
        ['健康查詢', 'search_symptom', '查症狀→化合物→推薦食物', 'POST /health/symptom-search'],
        ['學術搜尋', 'search_graphrag', '學術論文知識圖譜問答', 'POST /graphrag-api/query-sync'],
        ['學術搜尋', 'get_literature_papers', '瀏覽論文列表', 'GET /graphrag-api/papers/enriched'],
        ['行事曆', 'get_calendar_entries', '查行事曆', 'GET /calendar/entries'],
        ['行事曆', 'add_to_calendar', '加入行事曆', 'POST /recipe/save + /calendar/entries'],
        ['個人資料', 'get_user_profile', '讀取個人資料', 'GET /profile'],
        ['個人資料', 'update_user_profile', '更新個人資料', 'PUT /profile'],
    ],
    col_widths=[2.5, 4, 4, 6]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 第二章
# ═══════════════════════════════════════════════════════════════

doc.add_heading('第二章 系統架構總覽', level=1)
doc.add_heading('——五個服務如何分工合作', level=2)

add_svg_reference("01_system_architecture.svg", "系統架構總覽圖")

doc.add_heading('2.1 五個服務是什麼？', level=3)
doc.add_paragraph(
    '整個 RecipeAI 系統由五個獨立的「服務」組成。你可以把它們想成一家公司的五個部門，'
    '各自負責不同的工作，透過內部電話（網路請求）互相溝通。')

doc.add_paragraph('讓我們用一家餐廳來比喻：')

add_table(
    ['服務', 'Port', '比喻', '實際職責', '技術棧'],
    [
        ['Frontend\n前端', '3001', '餐廳大廳\n（客人看到的畫面）', '顯示網頁介面\n處理使用者互動\n發送 API 請求', 'React 18 + Vite\n+ nginx'],
        ['Control\n控制層', '8000', '經理\n（協調一切）', '接收前端請求\n呼叫 AI 模型\n轉發給後端\n業務邏輯處理', 'FastAPI + httpx\n+ jieba + argos'],
        ['Backend\n後端', '8001', '廚房\n（處理食材資料）', '直接操作資料庫\n計算營養數據\n管理使用者資料', 'FastAPI + asyncpg\n（無 ORM）'],
        ['PostgreSQL\n資料庫', '5432', '食材倉庫\n（所有資料存這）', '儲存 2213 筆台灣食品\n992 筆 FooDB 食品\n使用者、食譜、日誌', 'PostgreSQL 16'],
        ['Gemma LLM\n語言模型', '8080', '顧問\n（理解語言的大腦）', '理解使用者意圖\n生成自然語言回覆\n判斷工具調用', 'llama.cpp\n+ Gemma-4-31B'],
        ['GraphRAG API\n學術搜尋', '8002', '圖書館員\n（查學術論文）', '向量搜尋論文\n知識圖譜遍歷\n社群報告檢索', 'FastAPI + ChromaDB\n+ NetworkX'],
    ],
    col_widths=[2.5, 1.5, 3, 4, 3]
)

doc.add_heading('2.2 它們之間怎麼溝通？', level=3)
doc.add_paragraph(
    '這些服務之間的溝通有嚴格的規則，就像公司部門之間有固定的溝通管道：')

doc.add_paragraph('1. 前端（瀏覽器）不能直接跟後端說話。所有請求都要經過「控制層」轉發。'
    '這就像客人不能直接衝進廚房，必須透過經理。', style='List Number')
doc.add_paragraph('2. 前端透過「代理」（Proxy）把請求轉到對的地方。'
    '網址以 /api 開頭的 → 轉給 Control（port 8000）；'
    '以 /graphrag-api 開頭的 → 轉給 GraphRAG API（port 8002）。', style='List Number')
doc.add_paragraph('3. Control 和 Backend 之間用 HTTP 通訊。Control 收到請求後，'
    '會用 httpx 庫去呼叫 Backend 的 API。', style='List Number')
doc.add_paragraph('4. Control 和 Gemma LLM 之間也用 HTTP。走的是 OpenAI 相容的 API 格式'
    '（/v1/chat/completions），讓同一套 code 未來可以換不同的語言模型。', style='List Number')

doc.add_heading('2.3 Vite Proxy 代理配置（前端路由的關鍵）', level=3)
doc.add_paragraph(
    '前端的路由轉發設定在 vite.config.js 裡。這是前端開發伺服器的配置檔，'
    '它告訴 Vite：「如果瀏覽器請求的網址符合某個模式，就把它轉發到對應的後端服務。」')

add_code_block('''// vite.config.js — 前端代理配置
export default defineConfig({
  server: {
    host: "10.22.22.187",     // 監聽的 IP 位址
    port: 3001,                // 前端跑在 port 3001
    proxy: {
      "/api": {                // 規則 1：所有 /api 開頭的請求
        target: "http://localhost:8000",  // 轉發到 Control 服務
        changeOrigin: true,
        rewrite: (path) => path.replace(/^\\/api/, ""),
        // 例：/api/advisor/chat → http://localhost:8000/advisor/chat
        // 把 /api 前綴去掉
      },
      "/graphrag-api": {       // 規則 2：所有 /graphrag-api 開頭的
        target: "http://localhost:8002",  // 轉發到 GraphRAG API
        changeOrigin: true,
        rewrite: (path) => path.replace(/^\\/graphrag-api/, ""),
        // 例：/graphrag-api/query → http://localhost:8002/query
      },
    },
  },
});''', "vite.config.js（前端代理配置）")

add_note("這段 code 的意思是：瀏覽器以為所有東西都在同一個網站（localhost:3001），"
    "但實際上 Vite 偷偷把請求轉到不同的後端服務。這叫做「反向代理」（Reverse Proxy）。")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 第三章
# ═══════════════════════════════════════════════════════════════

doc.add_heading('第三章 通信協議基礎', level=1)
doc.add_heading('——HTTP、SSE、JSON 是什麼？', level=2)

doc.add_paragraph(
    '在深入了解 Tool Calling 之前，我們需要先理解三個基礎概念。'
    '這就像學做菜之前，先認識鍋碗瓢盆。')

doc.add_heading('3.1 HTTP 是什麼？（超文本傳輸協議）', level=3)
doc.add_paragraph(
    'HTTP 是瀏覽器和伺服器之間「說話」的規則。就像寫信有固定格式'
    '（收件人、寄件人、正文），HTTP 也有固定的格式。')

doc.add_paragraph('HTTP 最常用的兩種方式：')
doc.add_paragraph(
    'GET — 「拿東西」。就像跟圖書館說「我要借這本書」。'
    '你告訴伺服器你想要什麼資料，它把資料給你。'
    '例如：GET /health/nutrients/calcium_per_100g/top-foods?limit=10'
    '意思是「給我鈣含量最高的前 10 種食物」。', style='List Bullet')
doc.add_paragraph(
    'POST — 「送東西過去」。就像寄包裹，你把資料放在「body」（包裹內容）裡送出去。'
    '例如：POST /recipe/extract-with-amounts，body 裡放食譜文字，'
    '伺服器幫你提取食材。', style='List Bullet')

doc.add_heading('3.2 JSON 是什麼？（JavaScript 物件表示法）', level=3)
doc.add_paragraph(
    'JSON 是一種「資料格式」。它用花括號 {} 和方括號 [] 來組織資料，'
    '讓電腦和人類都能輕鬆閱讀。')

add_code_block('''{
  "food_name": "小魚乾",           ← 食物名稱（文字用雙引號包）
  "calcium_per_100g": 2213,        ← 每 100g 含鈣量（數字不用引號）
  "source": "taiwan_foods",        ← 資料來源
  "nutrients": {                    ← 巢狀物件（物件裡面還有物件）
    "protein_per_100g": 65.0,
    "iron_per_100g": 10.2
  },
  "categories": ["魚類", "乾貨"]   ← 陣列（用方括號，放多個值）
}''', "JSON 範例")

doc.add_heading('3.3 SSE 是什麼？（Server-Sent Events，伺服器推送事件）', level=3)
doc.add_paragraph(
    '一般的 HTTP 就像「你問一句，我答一句」——發一次請求，收一次回應。'
    '但 AI 顧問的回答需要好幾秒才能完成（因為要經過 6 個步驟），'
    '如果等全部做完再一次回覆，使用者會覺得系統卡住了。')

doc.add_paragraph(
    'SSE 解決這個問題。它讓伺服器可以「持續推送」資料給瀏覽器，'
    '就像直播——畫面是一點一點出來的，不是等到整部電影拍完才播放。')

doc.add_paragraph('SSE 的格式長這樣：')

add_code_block('''event: echo
data: {"content": "收到，我來查詢鈣含量最高的食物。"}

event: phase_status
data: {"phase": "intent", "label": "正在分析意圖"}

event: tool_start
data: {"tool": "get_nutrient_ranking", "params": {"nutrient_field": "calcium_per_100g"}}

event: tool_result
data: {"tool": "get_nutrient_ranking", "result": {"foods": [...]}}

event: reply
data: {"content": "查詢結果顯示，鈣含量最高的食物是小魚乾..."}''', "SSE 事件格式範例")

doc.add_paragraph(
    '每個 SSE 事件有兩行：event 行告訴前端「這是什麼類型的事件」，'
    'data 行是實際的資料（JSON 格式）。前端收到不同的 event 類型，'
    '就做不同的事情——echo 事件顯示確認訊息、tool_result 事件渲染工具結果卡片、'
    'reply 事件顯示 AI 的分析回覆。')

add_note("SSE 跟「串流」（Streaming）不一樣。串流是持續的資料流（像水管），"
    "SSE 是一個一個「事件」推送（像快遞一個一個包裹送到）。每個包裹都有標籤（event type），"
    "讓你知道該怎麼處理它。")

doc.add_heading('3.4 JWT 是什麼？（JSON Web Token，認證令牌）', level=3)
doc.add_paragraph(
    'JWT 是系統用來確認「你是誰」的憑證。就像進入大樓需要刷門禁卡一樣，'
    '每次前端發送 API 請求時，都要在 HTTP header 裡附上 JWT：')

add_code_block('''Authorization: Bearer eyJhbGciOiJIUzI1NiIs...
                   ↑ 固定前綴   ↑ 這串很長的就是 JWT（加密的身份證明）''')

doc.add_paragraph(
    'JWT 在登入時發放（POST /auth/login），有效期 7 天。'
    '如果過期了，前端會自動把你踢到登入頁面重新登入。')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 第四章
# ═══════════════════════════════════════════════════════════════

doc.add_heading('第四章 AI 營養顧問的 6 步驟 Pipeline', level=1)
doc.add_heading('——一句話如何變成答案', level=2)

add_svg_reference("02_llm_pipeline.svg", "AI 顧問 6 步驟 LLM Pipeline 流程圖")

doc.add_paragraph(
    '當你在 AI 營養顧問（/advisor 頁面）輸入一句話時，系統不是直接把你的話丟給 AI。'
    '它會經過精心設計的 6 個步驟，就像一條生產線——'
    '每個步驟都有專門的「指令」（Prompt）告訴 AI 該做什麼。')

doc.add_paragraph('讓我們用一個具體的例子來走完整個流程。使用者輸入：')
doc.add_paragraph('"幫我查鈣含量最高的食物"')

doc.add_heading('步驟 1：Echo（即時確認）', level=3)
doc.add_paragraph(
    '目的：讓使用者知道「系統收到了」，避免覺得卡住。'
    '\n就像你去餐廳點餐，服務生會先說「好的，收到您的點餐」，讓你知道他聽到了。')

doc.add_paragraph('這一步做的事：')
doc.add_paragraph('1. 把使用者的訊息 + 對話歷史 + 記憶 組合成一份「指令」', style='List Number')
doc.add_paragraph('2. 發送給 Gemma LLM（溫度 0.5，較有創意）', style='List Number')
doc.add_paragraph('3. LLM 回覆一句簡短的確認', style='List Number')
doc.add_paragraph('4. 透過 SSE 推送 echo 事件給前端', style='List Number')

doc.add_paragraph('對應程式碼：')
add_code_block('''# advisor.py → _step_echo() 函式
# 組裝 Prompt（把使用者的訊息塞進模板）
prompt = PROMPT_ECHO.format(
    memory_block=memory_block,    # 記憶摘要
    history=history_str,          # 對話歷史
    message=message               # 使用者這次說的話
)
# 發送給 Gemma LLM
echo_text = await _call_gemma(
    [{"role": "user", "content": prompt}],
    temperature=0.5    # 0.5 = 中等隨機性，回覆更自然
)
# 透過 SSE 推送給前端
yield _sse_event("echo", {"content": echo_text})''',
"advisor.py L780-795 _step_echo()")

doc.add_paragraph('LLM 可能回覆：「收到，我來查詢鈣含量最高的食物。」')

doc.add_heading('步驟 2：Phase 1 — 意圖偵測（Intent Detection）', level=3)
doc.add_paragraph(
    '目的：理解使用者「想做什麼事」。'
    '\n這是最關鍵的步驟——就像醫生問完症狀後，判斷你需要做哪項檢查。')

doc.add_paragraph(
    '這一步會把一份巨大的「系統功能說明書」（system_manual.md，548 行）'
    '附在 Prompt 裡，讓 LLM 知道系統有哪些功能可以用。'
    '然後要求 LLM 輸出嚴格的 JSON 格式，告訴系統：')

add_table(
    ['JSON 欄位', '意義', '範例值'],
    [
        ['intents', '你想做哪一件事（工具名稱）', '["nutrient_ranking"]'],
        ['entities', '做這件事需要的資訊', '{"nutrient_field": "calcium_per_100g"}'],
        ['needs_profile', '是否需要先讀取個人資料', 'false'],
        ['reasoning', '為什麼這樣判斷（一句話）', '"使用者想查鈣含量排名"'],
    ]
)

doc.add_paragraph('系統定義了 12 種意圖：')
add_table(
    ['意圖名稱', '白話意思', '對應工具'],
    [
        ['search_symptom', '查症狀相關的營養素和食物', 'search_symptom'],
        ['search_graphrag', '學術論文問答', 'search_graphrag'],
        ['nutrient_ranking', '查某個營養素的食物排名', 'get_nutrient_ranking'],
        ['get_recipes', '看已儲存的食譜', 'get_saved_recipes'],
        ['get_calendar', '查行事曆', 'get_calendar_entries'],
        ['add_to_calendar', '加入行事曆', 'add_to_calendar'],
        ['search_food', '搜尋食物', 'search_food'],
        ['browse_food_database', '瀏覽食品資料庫', 'browse_taiwan_food'],
        ['analyze_recipe', '分析食譜營養', 'analyze_recipe'],
        ['browse_literature', '瀏覽論文列表', 'get_literature_papers'],
        ['get_profile', '讀取個人資料', 'get_user_profile'],
        ['update_profile', '更新個人資料', 'update_user_profile'],
        ['general_chat', '一般閒聊', '（不需要工具）'],
    ]
)

add_note("注意：Phase 1 的意圖名稱（如 nutrient_ranking）和 Phase 2 的工具名稱（如 get_nutrient_ranking）不完全一樣！"
    "這是因為意圖是「使用者想做什麼」，工具是「系統要呼叫什麼函式」。"
    "中間有一個映射表做轉換。")

doc.add_heading('步驟 3：Phase 2 — 工具指令生成', level=3)
doc.add_paragraph(
    '目的：把 Phase 1 偵測到的「意圖」轉換成具體的「工具呼叫指令」。'
    '\n就像醫生說「要做血液檢查」，護理師會具體開單：「抽血項目：CBC + 鐵蛋白」。')

doc.add_paragraph(
    '這一步有兩條路徑：')
doc.add_paragraph(
    '路徑 A（優先）：系統自動生成 — _system_generate_toolcalls()。'
    '如果意圖很明確（例如 nutrient_ranking + nutrient_field 已知），'
    '系統直接用程式邏輯生成工具呼叫，不需要再問 LLM。快又穩定。', style='List Bullet')
doc.add_paragraph(
    '路徑 B（備援）：LLM 生成 — 用 PROMPT_PHASE2 問 Gemma。'
    '當情況比較複雜（例如多個意圖、參數不完整）時才走這條路。', style='List Bullet')

doc.add_paragraph('輸出的 JSON 長這樣：')
add_code_block('''{
  "tool_calls": [
    {"tool": "get_nutrient_ranking", "args": {"nutrient_field": "calcium_per_100g", "top_n": 10}}
  ],
  "params_from_store": {},       // 從結構化記憶取得的參數
  "still_missing": [],           // 還缺少的參數（要追問使用者）
  "ask_user": "",                // 追問語句
  "reasoning": "使用者要查鈣排名，參數完整"
}''', "Phase 2 輸出範例")

doc.add_heading('步驟 4：Execute（工具執行）', level=3)
doc.add_paragraph(
    '目的：真正去呼叫後端 API，拿到資料。'
    '\n這一步不需要 LLM——就像快遞員不需要讀懂包裹內容，只需要送到指定地址。')

doc.add_paragraph(
    '系統有一個「分發表」（_TOOL_DISPATCH），把工具名稱對應到實際的函式：')

add_code_block('''# advisor.py L1547-1560
_TOOL_DISPATCH = {
    "get_nutrient_ranking":  _tool_get_nutrient_ranking,   # → 呼叫後端
    "search_food":           _tool_search_food,
    "browse_taiwan_food":    _tool_browse_taiwan_food,
    "analyze_recipe":        _tool_analyze_recipe,         # → 4 步連鎖呼叫！
    "get_literature_papers": _tool_get_literature_papers,
    "search_symptom":        _tool_search_symptom,
    "search_graphrag":       _tool_search_graphrag,        # → 呼叫 GraphRAG API
    "get_saved_recipes":     _tool_get_saved_recipes,
    "get_calendar_entries":  _tool_get_calendar_entries,
    "add_to_calendar":       _tool_add_to_calendar,
    "get_user_profile":      _tool_get_user_profile,
    "update_user_profile":   _tool_update_user_profile,
}

# 執行時的邏輯：
for tc in tool_calls:
    fn = _TOOL_DISPATCH[tc["tool"]]      # 找到對應的函式
    result = await fn(tc["args"], ...)    # 呼叫它
    yield _sse_event("tool_result", {"tool": tc["tool"], "result": result})''',
"advisor.py 工具分發表")

doc.add_paragraph(
    '執行過程中，SSE 會推送三種事件：tool_start（開始執行）→ tool_result（結果到了）→ tool_done（執行完畢）。'
    '前端的 ToolResultCard 元件會根據工具名稱，用不同的樣板來顯示結果。')

doc.add_heading('步驟 5：Analysis（綜合分析回覆）', level=3)
doc.add_paragraph(
    '目的：把工具回傳的冰冷數據，翻譯成使用者聽得懂的話。'
    '\n就像實驗室把血液報告送回來，醫生不會直接把報告遞給你，'
    '而是會解讀：「你的鐵蛋白偏低，建議多吃紅肉和深綠蔬菜。」')

doc.add_paragraph('這一步的核心規則（寫在 PROMPT_ANALYSIS 裡）：')
doc.add_paragraph('所有數據必須來自工具結果，禁止從 LLM 自己的訓練資料編造數字', style='List Bullet')
doc.add_paragraph('前端已經會自動顯示表格和卡片，AI 不需要重複列出原始資料', style='List Bullet')
doc.add_paragraph('AI 的角色是「分析師」——解讀數據的意義，不是複述數據', style='List Bullet')
doc.add_paragraph('GraphRAG 結果特例：直接透傳完整答案，不摘要、不改寫', style='List Bullet')

doc.add_heading('步驟 6：Dual Memory（雙記憶更新）', level=3)
doc.add_paragraph(
    '目的：讓 AI 「記住」這次對話的內容，下次不用重複問。'
    '\n詳見第九章。')

add_svg_reference("09_dual_memory_system.svg", "雙記憶體系統架構圖")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 第五章 — 每個 Tool 的完整通信鏈
# ═══════════════════════════════════════════════════════════════

doc.add_heading('第五章 所有 Tool 逐一拆解', level=1)
doc.add_heading('——11 個工具完整通信鏈', level=2)

add_svg_reference("08_tool_dispatch.svg", "Tool Dispatch 工具分發機制圖")

# --- 5.1 ---
doc.add_heading('5.1 get_nutrient_ranking（營養素排名）', level=2)

add_svg_reference("06_nutrient_ranking_flow.svg", "營養素排名通信鏈圖")

doc.add_heading('如果人在操作（不透過 AI）', level=3)
doc.add_paragraph('使用者會看到的頁面：/nutrient-ranking（營養素排名頁）')
doc.add_paragraph('操作步驟：', style='List Bullet')
doc.add_paragraph('1. 從「分組」下拉選單選一個類別（13 組：基本、糖類、礦物質、維生素 A/D/E/K/B群、SFA/MUFA/PUFA、胺基酸、其他）', style='List Number')
doc.add_paragraph('2. 從「營養素」下拉選單選一個具體營養素（110+ 個選項）', style='List Number')
doc.add_paragraph('3. 設定顯示筆數（預設 10）', style='List Number')
doc.add_paragraph('4. 系統自動發送 API 請求', style='List Number')
doc.add_paragraph('5. 看到排名表格（排名、食物名、含量、來源）', style='List Number')
doc.add_paragraph('6. 點擊任一食物可展開詳細營養資訊（Modal）', style='List Number')

doc.add_paragraph('人操作時的輸入：選單選擇（分組 + 營養素 + 筆數）')
doc.add_paragraph('人操作時的回傳：排名表格 + 可點擊展開的食物卡片')

doc.add_heading('完整通信鏈（程式碼對應）', level=3)

add_code_block('''// ① 前端 — NutrientRankingPage.jsx
// 使用者選擇營養素後觸發
const handleSearch = async () => {
  const data = await getNutrientFoods(selectedField, topN);
  setResults(data);
};

// ② 前端 API 模組 — healthApi.js L33
export const getNutrientFoods = (field, limit = 10) =>
  client.get(`/health/nutrients/${field}/top-foods`, { params: { limit } })
        .then(r => r.data);
// 實際發出: GET /api/health/nutrients/calcium_per_100g/top-foods?limit=10

// ③ Vite 代理 — vite.config.js
// /api/health/nutrients/... → http://localhost:8000/health/nutrients/...

// ④ Control 路由 — control/app/routers/health.py
@router.get("/nutrients/{field}/top-foods")
async def get_top_foods(field: str, limit: int = 10):
    resp = await httpx_client.get(
        f"{BACKEND_URL}/health/nutrients/{field}/top-foods",
        params={"limit": limit}
    )
    return resp.json()
// Control 只是「轉發」，不做額外處理

// ⑤ Backend 路由 — backend/routers/health.py
@router.get("/nutrients/{field}/top-foods")
async def get_top_foods_by_nutrient(field: str, limit: int = 10):
    # 先查白名單（防止 SQL injection）
    field_map = load_json("resources/nutrient_field_map.json")
    if field not in field_map:
        raise HTTPException(400, "Unknown nutrient field")

    config = field_map[field]
    if "taiwan_cols" in config:
        # 合計欄位（如 vitamin_k_ug = k1 + k2mk4 + k2mk7）
        cols = config["taiwan_cols"]
        expr = " + ".join(f"COALESCE({c}, 0)" for c in cols)
        sql = f"SELECT *, ({expr}) AS computed_value FROM taiwan_foods ORDER BY computed_value DESC LIMIT $1"
    else:
        col = config["taiwan_col"]
        sql = f"SELECT * FROM taiwan_foods WHERE {col} IS NOT NULL ORDER BY {col} DESC LIMIT $1"

    rows = await db.fetch(sql, limit)
    return [dict(r) for r in rows]

// ⑥ PostgreSQL 執行 SQL
// SELECT *, calcium_per_100g FROM taiwan_foods
// WHERE calcium_per_100g IS NOT NULL
// ORDER BY calcium_per_100g DESC LIMIT 10

// ⑦ 結果一路回傳：DB → Backend → Control → Frontend
// Frontend 渲染排名表格''',
"營養素排名完整通信鏈")

# --- 5.2 ---
doc.add_heading('5.2 search_food（食物搜尋）', level=2)

doc.add_heading('如果人在操作', level=3)
doc.add_paragraph('頁面：/food-search（食材搜尋頁）')
doc.add_paragraph('1. 在搜尋框輸入食物名稱（如「雞蛋」）', style='List Number')
doc.add_paragraph('2. 系統搜尋兩個資料庫：台灣食品（2213 筆）+ FooDB（992 筆）', style='List Number')
doc.add_paragraph('3. 看到搜尋結果列表（食物名、來源、基本營養）', style='List Number')
doc.add_paragraph('4. 點擊食物→開 Modal 看完整營養素（110+ 欄位）', style='List Number')

doc.add_paragraph('人操作時的輸入：搜尋框文字')
doc.add_paragraph('人操作時的回傳：雙資料庫搜尋結果列表')

add_code_block('''// 通信鏈：
// 前端 foodApi.js: GET /api/food/search?q=雞蛋&lang=auto
// → Vite proxy → Control food.py: GET /food/search
// → Control 判斷語言（中文/英文/自動）
// → Backend food.py: GET /foods/search/zh?terms=雞蛋&limit=20
//   + GET /foods/search/en?terms=egg&limit=20（如果偵測到英文）
// → PostgreSQL: SELECT FROM taiwan_foods WHERE food_name ILIKE '%雞蛋%'
//              + SELECT FROM foods WHERE name ILIKE '%egg%'
// → 合併結果回傳''',
"search_food 通信鏈")

# --- 5.3 ---
doc.add_heading('5.3 browse_taiwan_food（台灣食品資料庫瀏覽）', level=2)

doc.add_heading('如果人在操作', level=3)
doc.add_paragraph('頁面：/food-database（台灣食品資料庫頁）')
doc.add_paragraph('1. 從分類下拉選單選擇（如「乳品類」）', style='List Number')
doc.add_paragraph('2. 可選輸入搜尋關鍵字', style='List Number')
doc.add_paragraph('3. 看到分頁的食品列表', style='List Number')
doc.add_paragraph('4. 點擊食品→TaiwanFoodModal（4 標籤：維生素/礦物質/糖類/脂肪酸）', style='List Number')

add_code_block('''// 通信鏈：
// 前端 taiwanFoodApi.js: GET /api/food/taiwan/browse?category=乳品類&page=1&limit=20
// → Control food.py → Backend food.py: GET /foods/taiwan/browse
// → PostgreSQL: SELECT FROM taiwan_foods WHERE category = '乳品類' LIMIT 20 OFFSET 0
// 分類列表: GET /api/food/taiwan/categories → 回傳所有分類名稱
// 統計: GET /api/food/taiwan/stats → 回傳 {total: 2213, categories: 18}''',
"browse_taiwan_food 通信鏈")

# --- 5.4 ---
doc.add_heading('5.4 analyze_recipe（食譜分析——最複雜的工具）', level=2)

add_svg_reference("04_recipe_analysis_flow.svg", "食譜分析完整通信鏈圖")

doc.add_heading('如果人在操作', level=3)
doc.add_paragraph('頁面：/analysis（整合分析頁，7 步驟流程）或 /recipe-analyzer')
doc.add_paragraph('人需要做的事（很多步驟！）：')
doc.add_paragraph('1. 貼上食譜文字', style='List Number')
doc.add_paragraph('2. 檢查自動提取的食材是否正確（可拖曳修改）', style='List Number')
doc.add_paragraph('3. 設定每種食材的用量（克數）', style='List Number')
doc.add_paragraph('4. 從候選列表中選擇最佳的資料庫匹配', style='List Number')
doc.add_paragraph('5. 確認營養計算結果', style='List Number')
doc.add_paragraph('6. 查看 DRI 缺口分析', style='List Number')
doc.add_paragraph('7. 瀏覽推薦補充食物', style='List Number')

doc.add_paragraph(
    '但如果透過 AI 顧問，使用者只需要說一句話：「幫我分析這個食譜：雞胸肉 200g、花椰菜 150g...」，'
    'AI 就會自動完成上面所有步驟。')

doc.add_heading('AI 顧問的 analyze_recipe 工具做了什麼（4 步連鎖呼叫）', level=3)

add_code_block('''# advisor.py → _tool_analyze_recipe() L1329-1500
# 這是最複雜的工具，內部要連續呼叫 4 個 API

async def _tool_analyze_recipe(args, auth_token, profile_data):
    recipe_text = args["recipe_text"]

    # ── 第 1 步：提取食材 ──
    resp1 = await httpx_client.post(
        f"{BACKEND_URL}/recipe/extract-with-amounts",
        json={"recipe_text": recipe_text},
        headers={"Authorization": f"Bearer {auth_token}"}
    )
    keywords = resp1.json()["keywords"]
    # 這一步內部又會：
    #   a) ai_service.extract_keywords_with_amounts() → 正則解析
    #   b) gemma_extraction_service.refine_extraction() → Gemma LLM 驗證

    # ── 第 2 步：查詢匹配食品 ──
    resp2 = await httpx_client.post(
        f"{BACKEND_URL}/recipe/lookup",
        json={"ingredients": [kw["name"] for kw in keywords]},
        headers={"Authorization": f"Bearer {auth_token}"}
    )
    matches = resp2.json()
    # 這一步內部又會：
    #   a) gemma_verification_service.verify_ingredient_list() → 驗證食材名
    #   b) Backend GET /foods/search/batch → 雙 DB 搜尋
    # AI 自動選第一個匹配（人操作時可以手動選）

    # ── 第 3 步：計算營養 ──
    items = []
    for kw, match in zip(keywords, matches):
        items.append({
            "food_id": match["best_match"]["food_id"],
            "source": match["best_match"]["source"],
            "grams": kw.get("amount_g", 100)
        })
    resp3 = await httpx_client.post(
        f"{BACKEND_URL}/recipe/calculate-nutrition",
        json={"items": items}
    )
    nutrition = resp3.json()  # 110+ 營養素欄位

    # ── 第 4 步：DRI 缺口分析 ──
    if profile_data:
        resp4 = await httpx_client.post(
            f"{BACKEND_URL}/health/dri-gap",
            json={
                "age": profile_data.get("age"),
                "sex": profile_data.get("gender"),
                "daily_intake": nutrition
            }
        )
        dri_gaps = resp4.json()

    return {
        "keywords": keywords,
        "nutrition": nutrition,
        "dri_gaps": dri_gaps,
        "ingredient_count": len(keywords)
    }''',
"analyze_recipe 工具（4 步連鎖呼叫）")

add_note("這個工具最特別的地方：當透過 AI 顧問使用時，食材匹配是自動的（選第一個候選）。"
    "但在整合分析頁（/analysis）手動操作時，使用者可以自己從候選列表中挑選最佳匹配。"
    "兩條路徑共用相同的後端 API。")

# --- 5.5 ---
doc.add_heading('5.5 search_symptom（症狀搜尋）', level=2)

add_svg_reference("05_symptom_search_flow.svg", "症狀搜尋通信鏈圖")

doc.add_heading('如果人在操作', level=3)
doc.add_paragraph('頁面：/symptom-search（症狀搜尋頁）')
doc.add_paragraph('1. 輸入症狀關鍵字（如「疲勞」）', style='List Number')
doc.add_paragraph('2. 系統自動翻譯成英文（fatigue）', style='List Number')
doc.add_paragraph('3. 看到三層結構的結果：', style='List Number')
doc.add_paragraph('   第一層：健康效果（如「缺鐵性貧血」）', style='List Bullet')
doc.add_paragraph('   第二層：展開→化合物（如「鐵」、「維生素 B12」）', style='List Bullet')
doc.add_paragraph('   第三層：展開→含該化合物的食物', style='List Bullet')
doc.add_paragraph('4. 點擊食物名→開 FoodbFoodModal', style='List Number')

add_code_block('''# 通信鏈（AI 顧問路徑）：
# advisor.py → _tool_search_symptom()

async def _tool_search_symptom(args, auth_token, *_):
    symptom = args["symptom"]  # 已是英文（Phase1 翻譯過）

    # POST /health/symptom-search
    resp = await httpx_client.post(
        f"{BACKEND_URL}/health/symptom-search",
        json={"symptom": symptom},
        headers={"Authorization": f"Bearer {auth_token}"}
    )
    return resp.json()
    # 回傳結構：
    # {
    #   "keyword": "fatigue",
    #   "effects": [
    #     {
    #       "effect_name": "Iron deficiency anemia",
    #       "compounds": [
    #         {
    #           "compound_name": "Iron",
    #           "foods": [
    #             {"food_name": "Beef liver", "food_id": 123, "source": "foodb"}
    #           ]
    #         }
    #       ]
    #     }
    #   ]
    # }

# Control 層的處理（health.py）：
# 1. 用 argostranslate 把中文翻英文
# 2. 轉發到 Backend
# Backend 層（health.py）：
# 1. SELECT FROM health_effects WHERE name ILIKE '%fatigue%'
# 2. JOIN effect_compounds → compounds
# 3. JOIN compound_foods → foods
# 4. 組合三層結構回傳''',
"search_symptom 通信鏈")

# --- 5.6 ---
doc.add_heading('5.6 search_graphrag（學術論文搜尋）', level=2)

add_svg_reference("07_graphrag_flow.svg", "GraphRAG 通信鏈圖")

doc.add_heading('如果人在操作', level=3)
doc.add_paragraph('頁面：/graphrag（GraphRAG 知識圖譜問答頁）')
doc.add_paragraph('1. 輸入自然語言問題（如「咖啡因對耐力表現的影響」）', style='List Number')
doc.add_paragraph('2. 系統同時用 4 種方法搜尋：', style='List Number')
doc.add_paragraph('   Method A（Vector RAG）：純向量搜尋，最快', style='List Bullet')
doc.add_paragraph('   Method B（Local Search）：entity → 圖遍歷 → chunk', style='List Bullet')
doc.add_paragraph('   Method C（Global Search）：社群報告 → entity → chunk', style='List Bullet')
doc.add_paragraph('   Method D（Drift Search）：假設答案 → 子問題 → chunk', style='List Bullet')
doc.add_paragraph('3. SSE 串流即時顯示每種方法的結果', style='List Number')
doc.add_paragraph('4. 每個結果包含：Gemma 回答 + 參考文獻 + 搜尋路徑圖 + 論文圖表', style='List Number')

doc.add_heading('AI 顧問路徑的差異', level=3)
doc.add_paragraph(
    'AI 顧問走的是「同步」端點（/query-sync），不是 SSE 串流。'
    '原因：advisor.py 的 LLM pipeline 本身就是 SSE，不能在 SSE 裡套 SSE。'
    '所以 GraphRAG 結果會整包回來後，再由 advisor.py 透過 tool_result 事件推給前端。')

add_code_block('''# advisor.py → _tool_search_graphrag()
async def _tool_search_graphrag(args, *_):
    query = args["query"]  # 英文學術查詢

    # 走同步端點（非 SSE），180 秒超時
    resp = await httpx_client.post(
        "http://localhost:8002/query-sync",
        json={"query": query, "methods": ["B"]},  # 只用 Local Search
        timeout=180.0
    )
    result = resp.json()
    # result 包含 answer（完整學術分析）和 evidence（方法學佐證）

    return result

# 特別規則：PROMPT_ANALYSIS 收到 GraphRAG 結果時，
# 必須「直接透傳」answer 和 evidence 的完整內容，
# 禁止摘要、禁止縮減、禁止改寫。''',
"search_graphrag 通信鏈")

# --- 5.7-5.10 ---
doc.add_heading('5.7 get_literature_papers（文獻瀏覽）', level=2)
doc.add_paragraph('頁面：/literature-review（學術文獻瀏覽頁）')
doc.add_paragraph('人操作：直接看論文列表，點擊可到 PDF 檢視器')
add_code_block('''// 通信鏈：
// GET /graphrag-api/papers/enriched → GraphRAG API (port 8002)
// → 回傳 65 篇論文的元數據：標題、作者、IF/SJR、DOI、摘要''',
"get_literature_papers 通信鏈")

doc.add_heading('5.8 get_saved_recipes（已存食譜）', level=2)
doc.add_paragraph('頁面：/recipes（已儲存食譜頁）')
doc.add_paragraph('人操作：看到食譜列表，可展開查看營養明細，可刪除')
add_code_block('''// 通信鏈：
// GET /api/recipe/list (帶 JWT)
// → Control recipe.py → Backend recipe.py
// → PostgreSQL: SELECT FROM recipes WHERE user_id = $1
// → 回傳食譜列表 [{recipe_name, nutrition, ingredients, created_at}]''',
"get_saved_recipes 通信鏈")

doc.add_heading('5.9 get_calendar_entries / add_to_calendar（行事曆）', level=2)
doc.add_paragraph('頁面：/calendar（每日分析日曆頁）')
doc.add_paragraph('人操作（查詢）：選日期範圍→看到該時段的分析紀錄')
doc.add_paragraph('人操作（新增）：在整合分析完成後點「加入行事曆」')
add_code_block('''// 查詢通信鏈：
// GET /api/calendar/entries?start_date=2026-05-01&end_date=2026-05-31
// → Control → Backend → PostgreSQL: SELECT FROM calendar_entries WHERE date BETWEEN ...

// 新增通信鏈（AI 顧問路徑較複雜）：
// _tool_add_to_calendar():
//   1. POST /recipe/save (先存食譜，拿到 recipe_id)
//   2. POST /calendar/entries (再把 recipe_id 加到指定日期)''',
"行事曆通信鏈")

doc.add_heading('5.10 get_user_profile / update_user_profile（個人資料）', level=2)
doc.add_paragraph('頁面：/profile（個人資料頁）')
doc.add_paragraph('人操作：填寫年齡、性別、身高、體重、活動量、目標→儲存')
add_code_block('''// 讀取通信鏈：
// GET /api/profile (帶 JWT) → Control → Backend → PostgreSQL

// 更新通信鏈：
// PUT /api/profile {age: 25, gender: "male", height_cm: 175, weight_kg: 70, ...}
// → Control profile.py → Backend user.py
// → PostgreSQL: UPDATE user_profiles SET age=$1, gender=$2, ... WHERE user_id=$3''',
"個人資料通信鏈")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 第六章
# ═══════════════════════════════════════════════════════════════

doc.add_heading('第六章 整合分析頁', level=1)
doc.add_heading('——不經 AI 顧問的 7 步驟直接操作流程', level=2)

doc.add_paragraph(
    '整合分析頁（/analysis）是整個系統最核心的功能頁面。'
    '它不透過 AI 顧問，而是讓使用者「一步一步」操作。'
    '每一步都對應不同的 API 呼叫。')

add_table(
    ['步驟', '使用者看到什麼', '使用者做什麼', 'API 呼叫', '回傳什麼'],
    [
        ['Step 1', '空白文字框', '貼上食譜文字', '（尚未呼叫 API）', '（無）'],
        ['Step 2', '自動提取的食材標籤列表\n+ 停用詞拖曳區', '確認食材、可拖曳\n移除/恢復食材', 'POST /recipe/\nextract-with-amounts\n(+Gemma LLM 驗證)', '食材名+用量列表'],
        ['Step 3', '每項食材的用量輸入框', '調整克數', '（前端狀態更新）', '（無）'],
        ['Step 4', '每項食材的候選匹配列表', '選擇最佳匹配', 'POST /recipe/lookup\n(+Gemma 驗證)', '候選食品列表'],
        ['Step 5', '完整營養分析表\n（宏量/微量/胺基酸）', '閱讀結果', 'POST /recipe/\ncalculate-nutrition', '110+ 營養素數值'],
        ['Step 6', 'DRI 缺口表格\n（紅/黃/綠色標）', '閱讀缺口分析', 'POST /health/dri-gap', '33 項營養素缺口'],
        ['Step 7', '推薦補充食物表\n（點擊可展開 Modal）', '瀏覽推薦食物', 'GET /health/nutrients/\n{field}/top-foods\n（每個缺口一次）', '每項缺口的 Top N 食物'],
    ],
    col_widths=[1.5, 4, 3, 4, 3]
)

doc.add_heading('Step 2 的停用詞拖曳機制（雙向 DnD）', level=3)
doc.add_paragraph(
    '這是一個特別的 UX 機制：使用者可以用滑鼠拖曳食材標籤到「停用詞區」，'
    '系統就會記住這個詞不是食材，下次解析時自動過濾掉。')

add_code_block('''// UnifiedAnalysisPage.jsx — 雙向拖曳機制

// === 方向 1：食材 → 停用詞 ===
// 食材標籤設定為可拖曳
<span draggable={true}
  onDragStart={(e) => {
    e.dataTransfer.setData("text/plain", ingredientName);
    // 拖曳時攜帶食材名稱
  }}>
  {ingredientName}
</span>

// 停用詞放置區
<div onDragOver={(e) => e.preventDefault()}
  onDrop={async (e) => {
    const word = e.dataTransfer.getData("text/plain");
    // 呼叫 API 新增停用詞
    await createStopwordsBulk({ words: [word] });
    // POST /api/stopwords/bulk → Control → 寫入 stopwords.json
    removeFromIngredients(word);
  }}>
  拖曳到此新增停用詞
</div>

// === 方向 2：停用詞 → 食材（恢復） ===
<chip draggable={true}
  onDragStart={(e) => {
    e.dataTransfer.setData("application/x-stopword", word);
    // 用不同的 MIME type 標記「這是停用詞方向」
  }}>
  {stopword}
</chip>

// 食材區的接收
<div onDrop={(e) => {
    if (e.dataTransfer.types.includes("application/x-stopword")) {
      const word = e.dataTransfer.getData("application/x-stopword");
      await deleteStopword(word);
      // DELETE /api/stopwords/{word} → 從 stopwords.json 移除
      addBackToIngredients(word);
    }
  }}>''',
"雙向拖曳停用詞機制")

add_note("方向判斷的巧妙設計：食材→停用詞用 'text/plain' MIME type，"
    "停用詞→食材用 'application/x-stopword' MIME type。"
    "同一個 drop zone 靠檢查 e.dataTransfer.types 來區分方向，不會搞混。")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 第七章
# ═══════════════════════════════════════════════════════════════

doc.add_heading('第七章 SSE 通信協議完整解析', level=1)
doc.add_heading('——每一個事件的意義', level=2)

add_svg_reference("03_sse_protocol.svg", "SSE 通信協議序列圖")

doc.add_paragraph(
    '前面提到 SSE 讓伺服器可以「持續推送」事件給瀏覽器。'
    '現在讓我們仔細看每一種事件類型的意義，'
    '以及前端收到後做了什麼。')

add_table(
    ['事件類型', '推送時機', '資料內容', '前端收到後做什麼'],
    [
        ['phase_status', '每個階段開始時', '{"phase":"echo"}', 'ThinkingPanel 更新進度指示器\n（六個圓點依序亮起）'],
        ['echo', 'Step 1 完成時', '{"content":"收到..."}', '顯示 AI 的快速確認訊息'],
        ['tool_start', 'Step 4 工具開始前', '{"tool":"名稱","params":{...}}', 'ToolResultCard 顯示「正在查詢...」loading'],
        ['tool_result', 'Step 4 工具完成後', '{"tool":"名稱","result":{...}}', 'ToolResultCard 渲染工具結果\n（不同工具用不同樣板）'],
        ['tool_done', 'Step 4 所有工具完成', '{}', '前端更新 loading 狀態'],
        ['reply', 'Step 5 分析完成', '{"content":"分析結果..."}', 'MessageBubble 渲染 AI 回覆\n（支援 Markdown 格式）'],
        ['narrative_memory', 'Step 6a 完成', '{"summary":"對話摘要..."}', '更新 localStorage\n"advisor_narrative"'],
        ['structured_store', 'Step 6b 完成', '{"store":{...JSON...}}', '更新 localStorage\n"advisor_structured_store"'],
        ['workflow_state', '流程結束時', '{"phase":"done",...}', '更新工作流狀態機'],
        ['memory', '備用記憶事件', '{"summary":"..."}', '更新 state.memory'],
        ['error', '任何步驟出錯時', '{"message":"錯誤訊息"}', '顯示錯誤提示'],
    ],
    col_widths=[2.5, 3, 4, 5]
)

doc.add_heading('前端如何接收 SSE（advisorApi.js 完整程式碼解析）', level=3)

add_code_block('''// advisorApi.js — 完整的 SSE 客戶端實作

export function streamAdvisorChat(
  message,           // 使用者輸入的文字
  history,           // 對話歷史（陣列）
  userProfile,       // 使用者個人資料
  callbacks,         // 回呼函式物件（每種事件一個）
  memory,            // 舊版記憶（已棄用，保留相容）
  workflowState,     // 工作流狀態機
  narrativeMemory,   // 敘述記憶
  structuredStore    // 結構化存儲
) {
  const token = localStorage.getItem("token");  // 從瀏覽器取 JWT
  const controller = new AbortController();       // 用來中斷連線

  // ① 組裝請求 body
  const body = JSON.stringify({
    message,
    history,
    user_profile: userProfile,
    memory,
    workflow_state: workflowState,
    narrative_memory: narrativeMemory,
    structured_store: structuredStore,
  });

  // ② 發送 POST 請求（不是 EventSource！因為 EventSource 只支援 GET）
  fetch("/api/advisor/chat", {
    method: "POST",
    headers: {
      "Content-Type": "application/json",
      Authorization: `Bearer ${token}`,
    },
    body,
    signal: controller.signal,  // 可以中斷
  })
    .then(async (response) => {
      // ③ 建立 ReadableStream reader
      const reader = response.body.getReader();
      const decoder = new TextDecoder();
      let buffer = "";      // 緩衝區（處理 TCP 分片）
      let eventType = null; // 當前事件類型

      // ④ 持續讀取 SSE 事件
      while (true) {
        const { done, value } = await reader.read();
        if (done) break;  // 伺服器關閉連線

        // ⑤ 解碼二進位資料為文字
        buffer += decoder.decode(value, { stream: true });
        const lines = buffer.split("\\n");
        buffer = lines.pop() || "";  // 最後一行可能不完整

        // ⑥ 逐行解析 SSE 格式
        for (const line of lines) {
          if (line.startsWith("event: ")) {
            // 這行告訴我們「接下來的 data 是什麼類型」
            eventType = line.slice(7).trim();
          } else if (line.startsWith("data: ") && eventType) {
            // 這行是實際資料
            const data = JSON.parse(line.slice(6));

            // ⑦ 根據事件類型呼叫對應的回呼函式
            switch (eventType) {
              case "thinking":    callbacks.onThinking?.(data);    break;
              case "echo":        callbacks.onEcho?.(data);        break;
              case "tool_start":  callbacks.onToolStart?.(data);   break;
              case "tool_result": callbacks.onToolResult?.(data);  break;
              case "reply":       callbacks.onReply?.(data);       break;
              case "memory":      callbacks.onMemory?.(data);      break;
              // ... 其他事件類型
            }
            eventType = null;  // 重置，等下一個 event 行
          }
        }
      }
      callbacks.onComplete?.();  // 全部完成
    });

  // ⑧ 回傳中斷函式（使用者可以隨時取消）
  return { abort: () => controller.abort() };
}''',
"advisorApi.js 完整 SSE 客戶端")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 第八章
# ═══════════════════════════════════════════════════════════════

doc.add_heading('第八章 LLM 思考邏輯與交叉互動機制', level=1)

doc.add_paragraph(
    '本系統使用 Gemma-4-31B-Instruct 作為語言模型（LLM）。'
    '每次使用者發一句話，LLM 會被呼叫 5~6 次，每次的「角色」不同。'
    '這就像同一位營養師，先戴上「接待員」的帽子確認需求，'
    '再戴上「分析師」的帽子解讀結果——每頂帽子就是一份不同的 Prompt。')

doc.add_heading('8.1 LLM 的「溫度」是什麼？', level=3)
doc.add_paragraph(
    '溫度（temperature）控制 LLM 回覆的「隨機性」。'
    '想像你在考試——溫度 0 就是只寫最標準的答案；'
    '溫度 1 就是天馬行空自由發揮。')

add_table(
    ['溫度值', '特性', '用在哪裡', '為什麼'],
    [
        ['0.1', '幾乎確定性\n回覆很穩定', 'Phase 1（意圖偵測）\nPhase 2（工具生成）\n結構化提取', '判斷意圖和生成 JSON\n不能有隨機性\n否則會亂判'],
        ['0.2', '低隨機性', '記憶摘要\n食材提取驗證', '需要一致但允許\n些微變化'],
        ['0.3', '中低隨機性', 'Analysis（分析回覆）', '分析要準確\n但可以有個人風格'],
        ['0.5', '中等隨機性', 'Echo（即時確認）', '確認訊息可以\n多樣化、更自然'],
        ['0.7', '較高隨機性', 'GraphRAG 回答', '學術討論可以\n更全面、更深入'],
    ]
)

doc.add_heading('8.2 多次 LLM 呼叫的交叉互動', level=3)
doc.add_paragraph(
    '每次使用者輸入，LLM 會被呼叫多次，而且後面的呼叫會「用到」前面的結果。'
    '這形成了一條資料流鏈：')

add_code_block('''使用者訊息
    │
    ▼
[Echo] ─── temp=0.5 ──→ 確認文字（推送給前端）
    │
    ▼
[Phase 1] ─── temp=0.1 ──→ {intents, entities}
    │                              │
    │                              ▼
    │                    [Phase 2] ─── temp=0.1 ──→ {tool_calls}
    │                              │                     │
    │                              │                     ▼
    │                              │           [Execute] ──→ tool_results
    │                              │                            │
    │                              │                            ▼
    │                              └──────────→ [Analysis] ─── temp=0.3
    │                                           │    ↑
    │                                           │    └── tool_results
    │                                           │
    │                                           ▼
    │                                      AI 回覆文字
    │                                           │
    ▼                                           ▼
[Narrative Summary] ◄─── 用戶訊息 + AI回覆 + 工具結果 ── temp=0.2
    │
    ▼
[Structured Extract] ◄─── 用戶訊息 + 工具結果 + AI回覆摘要 ── temp=0.1

每一步的輸出都會成為下一步的輸入，形成「思考鏈」。''',
"LLM 呼叫的資料流鏈")

doc.add_heading('8.3 工作流狀態機——處理「參數不足」的情況', level=3)
doc.add_paragraph(
    '有時候使用者說的話不夠明確。例如說「幫我查行事曆」，'
    '但沒有說「哪個日期範圍」。這時系統需要「追問」。')

doc.add_paragraph(
    '系統用一個「狀態機」（State Machine）來管理這個流程。'
    '狀態機就像一個流程圖，記錄現在「走到哪一步了」：')

add_table(
    ['狀態', '意思', '會發生什麼'],
    [
        ['idle', '閒置中，等使用者輸入', '收到訊息後進入 Phase 1'],
        ['collecting', '正在收集缺失的參數', '向使用者追問（最多 3 輪）'],
        ['ready', '參數已備齊', '直接執行工具'],
        ['executing', '工具正在執行中', '等待後端回應'],
        ['done', '本輪完成', '回到 idle 等下一次'],
    ]
)

add_code_block('''# advisor.py — WorkflowState 結構
class WorkflowState:
    phase: str = "idle"              # 目前階段
    active_intent: str = ""          # 當前意圖
    collected: dict = {}             # 已收集的參數
    missing: list = []               # 還缺的參數
    pending_intents: list = []       # 排隊的意圖
    auto_fetch_done: list = []       # 已自動取得的資料
    turn_count: int = 0              # 對話輪次
    interrupted_from: str = ""       # 中斷前的階段

# 狀態轉換範例：
# 使用者：「幫我查行事曆」
# Phase 1 → intents: ["get_calendar"], missing: ["start_date", "end_date"]
# → 狀態：collecting
# AI 追問：「請問你想查哪個日期範圍？」
# 使用者：「這週的」
# Phase 2 → tool_calls: [{tool: "get_calendar_entries", args: {start_date: "2026-05-12", end_date: "2026-05-18"}}]
# → 狀態：executing → done''',
"工作流狀態機")

doc.add_heading('8.4 中斷偵測——使用者突然換話題怎麼辦？', level=3)
doc.add_paragraph(
    '如果 AI 正在追問「你想查哪個日期？」，但使用者突然說「算了，幫我查鈣含量」，'
    '系統不能繼續追問日期。_detect_interruption() 函式會偵測這種情況：')

add_code_block('''# advisor.py → _detect_interruption()
# 邏輯：比較使用者這次的意圖和上次的活躍意圖
# 如果不同 → 中斷舊的，開始新的
# 如果相同 → 繼續收集參數''',
"中斷偵測邏輯")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 第九章
# ═══════════════════════════════════════════════════════════════

doc.add_heading('第九章 雙記憶體系統', level=1)
doc.add_heading('——AI 如何「記住」你說過的話', level=2)

add_svg_reference("09_dual_memory_system.svg", "雙記憶體系統架構圖")

doc.add_paragraph(
    'AI 每次收到你的訊息，處理完後就會「忘記」——因為 LLM 本身沒有記憶功能。'
    '所以系統設計了兩種記憶機制，讓 AI 能「記住」你之前說過的話。')

doc.add_heading('9.1 為什麼需要「兩種」記憶？', level=3)
doc.add_paragraph(
    '想像你在看醫生。醫生會記住兩種東西：')
doc.add_paragraph('1. 對話流程（敘述記憶）：「病人先說頭痛，我建議他做檢查，他同意了」', style='List Number')
doc.add_paragraph('2. 具體資料（結構化存儲）：「病人 25 歲、男性、175cm、70kg、對花生過敏」', style='List Number')

doc.add_paragraph(
    '這兩種記憶用在不同的地方：')
doc.add_paragraph('敘述記憶 → 幫助 Phase 1（意圖偵測）理解對話脈絡', style='List Bullet')
doc.add_paragraph('結構化存儲 → 幫助 Phase 2（工具生成）填入已知參數，避免重複追問', style='List Bullet')

doc.add_heading('9.2 敘述記憶（Narrative Memory）', level=3)
doc.add_paragraph('生成方式：每次對話結束後，用 PROMPT_NARRATIVE_SUMMARY 讓 LLM 摘要')
doc.add_paragraph('存在哪裡：前端 localStorage 的 "advisor_narrative" key')
doc.add_paragraph('格式：用分號分隔的文字摘要')
doc.add_paragraph('範例：')
add_code_block('''"使用者問疲勞原因；AI 查症狀建議補鐵和 B12；使用者接著問哪些食物鐵最多；AI 查營養素排名顯示小魚乾含鐵量最高"''',
"敘述記憶範例")

doc.add_heading('9.3 結構化存儲（Structured Store）', level=3)
doc.add_paragraph('生成方式：每次對話結束後，用 PROMPT_STRUCTURED_EXTRACT 讓 LLM 提取固定欄位')
doc.add_paragraph('存在哪裡：前端 localStorage 的 "advisor_structured_store" key')
doc.add_paragraph('格式：JSON 物件，包含以下類別：')

add_table(
    ['類別', '包含什麼', '用途'],
    [
        ['user_facts', '年齡、性別、身高、體重、活動量、目標、過敏', '下次不用再問「你幾歲？」'],
        ['food_context', '本次搜尋的食物、瀏覽的分類', '追蹤使用者的飲食興趣'],
        ['recipe_entry', '食譜名、食材、110+ 營養素數值', '記住分析過的食譜'],
        ['symptom_entry', '查過的症狀、相關化合物、推薦食物', '避免重複查詢'],
        ['ranking_entry', '查過的營養素排名和 top 3 食物', '快速回顧'],
        ['academic_entry', '學術查詢和關鍵發現', '記住研究結論'],
        ['dri_awareness', '缺乏或過量的營養素', '提供持續建議'],
        ['action_entry', '本輪的動作記錄', '審計追蹤'],
    ]
)

doc.add_heading('9.4 記憶的生命週期', level=3)
add_code_block('''使用者發訊息
  │
  ├─→ [帶著 narrative_memory + structured_store 發送到後端]
  │
  ├─→ [Phase 1 用 narrative_memory 理解上下文]
  ├─→ [Phase 2 用 structured_store 填入已知參數]
  │
  ├─→ [工具執行 + 分析回覆]
  │
  ├─→ [Step 6a: LLM 更新 narrative_memory]
  ├─→ [Step 6b: LLM 更新 structured_store]
  │
  ├─→ [SSE 推送回前端]
  └─→ [前端存入 localStorage]
       │
       └─→ 下次發訊息時，再帶著更新後的記憶一起送出''',
"記憶生命週期")

doc.add_paragraph(
    '對話歷史只保留最近 30 則訊息。但記憶不限制——它會隨著對話越來越豐富，'
    '讓 AI 越來越「了解」你。關閉瀏覽器再開，記憶仍然在（因為存在 localStorage）。')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 第十章 — 所有 Prompt 完整內容
# ═══════════════════════════════════════════════════════════════

doc.add_heading('第十章 所有 Prompt 完整內容', level=1)
doc.add_heading('——與通信鏈對應', level=2)

add_svg_reference("10_prompt_positions.svg", "Prompt 在通信鏈中的位置圖")

doc.add_paragraph(
    'Prompt（指令模板）是系統的「靈魂」——它們告訴 LLM「你現在是誰、要做什麼、'
    '怎麼回答」。以下是系統中所有 Prompt 的完整內容、'
    '在通信鏈中的位置、以及前後文。')

# Prompt 1
doc.add_heading('10.1 PROMPT_ECHO（即時確認指令）', level=2)
doc.add_paragraph('位置：advisor.py L43-77')
doc.add_paragraph('通信鏈位置：使用者訊息進入 → _step_echo() → _call_gemma() → Gemma LLM → SSE echo 事件 → 前端顯示')
doc.add_paragraph('溫度：0.5 | max_tokens：16384')

add_code_block('''請你根據使用者的訊息和對話脈絡，用繁體中文回應。

## 規則
- 簡潔回應
- 如果使用者的訊息需要查詢工具（記錄飲食、查營養、查症狀等）：概括需求，表達「已收到、正在處理」
- 如果是一般對話（打招呼、感謝、閒聊）：直接給出友善回應
- 不要回答營養問題本身，不要提供營養數據或建議
- 語氣親切自然

## 範例
使用者：「營養夠嗎」
→ 收到，我先查看你的個人資料，再幫你分析今天的營養攝取是否達標。

使用者：「最近常覺得疲勞」
→ 了解，我來查詢與疲勞相關的營養素和推薦食物。

使用者：「為什麼蛋白質對肌肉修復很重要？」
→ 好問題，我來搜尋相關的學術文獻幫你找答案。

使用者：「你好」
→ 你好！我是你的 AI 營養顧問，有什麼飲食或營養問題可以幫你的嗎？

使用者：「謝謝」
→ 不客氣！如果之後有其他營養問題，隨時問我。

---

{memory_block}     ← 這裡會插入敘述記憶摘要

對話歷史：
{history}          ← 這裡會插入最近的對話記錄

使用者訊息：{message}   ← 使用者這次說的話

請直接輸出回應（不要加引號或其他格式）：''', "PROMPT_ECHO 完整內容")

# Prompt 2
doc.add_heading('10.2 PROMPT_PHASE1（意圖偵測指令）', level=2)
doc.add_paragraph('位置：advisor.py L84-166')
doc.add_paragraph('通信鏈位置：_step_echo() 完成後 → _step_phase1() → _call_gemma() → Gemma LLM → 解析 JSON')
doc.add_paragraph('溫度：0.1 | max_tokens：16384 | 重試：2 次')
doc.add_paragraph('特別說明：這個 Prompt 最長！它會把 system_manual.md（548 行的完整功能說明書）插入其中。')

add_code_block('''你是 AI 營養顧問系統的意圖判定器。以下是系統的完整功能說明書，請仔細閱讀後，根據使用者訊息、對話脈絡，判斷需要呼叫哪些工具。

[這裡會插入 system_manual.md 的完整 548 行內容]

---

## 可用工具（intents 欄位只允許下表名稱，不可用其他名稱）

| intent | 必要參數 | 說明 |
|---|---|---|
| search_symptom | symptom（英文） | 搜尋症狀相關的營養素和推薦食物 |
| search_graphrag | query（英文學術查詢） | 學術論文知識圖譜問答 |
| nutrient_ranking | nutrient_field | 營養素含量排名 |
| get_recipes | （無） | 查看已儲存的食譜 |
| get_calendar | start_date, end_date | 行事曆查詢 |
| add_to_calendar | entry_date | 儲存食譜到行事曆 |
| search_food | food_name | 搜尋食物營養成分 |
| browse_food_database | （無） | 瀏覽台灣食品資料庫 |
| analyze_recipe | recipe_text | 分析食譜營養 |
| browse_literature | （無） | 瀏覽學術論文 |
| get_profile | （無） | 讀取個人資料 |
| update_profile | 至少一欄 | 更新個人資料 |
| general_chat | — | 一般對話 |

## 症狀翻譯對照
疲勞=fatigue、失眠=insomnia、頭痛=headache、掉髮=hair loss、
便秘=constipation、抽筋=muscle cramp、皮膚差=poor skin health、
貧血=anemia、骨質疏鬆=osteoporosis、免疫力差=weak immunity

## 營養素欄位名對照
蛋白質=protein_per_100g、鈣=calcium_per_100g、鐵=iron_per_100g、
鋅=zinc_per_100g、鎂=magnesium_per_100g、維生素C=vitamin_c_mg

## 確認詞處理
使用者說「好」「可以」「對」= 確認上一輪 AI 提議的動作。
從對話歷史找出上一輪 AI 建議了什麼，轉為對應的 intent。

---

### 對話脈絡摘要
{narrative_memory}     ← 敘述記憶

### 對話歷史
{history}              ← 最近對話

### 使用者最新訊息
{message}

### 輸出格式（嚴格 JSON）
{
  "intents": ["intent_name"],
  "entities": {
    "food_name": "", "recipe_text": "", "symptom_en": "",
    "query_en": "", "nutrient_field": "",
    "profile_fields": {"gender": "", "age": 0, ...},
    "start_date": "", "end_date": ""
  },
  "needs_profile": false,
  "reasoning": "一句話說明"
}''', "PROMPT_PHASE1 完整內容（精簡版，完整含 548 行說明書）")

# Prompt 3
doc.add_heading('10.3 PROMPT_PHASE2（工具指令生成指令）', level=2)
doc.add_paragraph('位置：advisor.py L173-227')
doc.add_paragraph('通信鏈位置：Phase 1 JSON 解析完成 → _step_phase2() → _call_gemma() 或 _system_generate_toolcalls()')
doc.add_paragraph('溫度：0.1 | max_tokens：16384')

add_code_block('''你是工具調用生成模組。根據意圖判定結果和已知的使用者資訊，生成具體的工具呼叫指令。

### 意圖判定結論（來自上一步）
{phase1_json}       ← Phase 1 的輸出 JSON

### 系統已儲存的使用者資訊（結構化資料）
{structured_store}  ← 結構化存儲（之前的對話記住的事實）

### 對話脈絡摘要
{narrative_memory}

### 使用者最新訊息
{message}

### 可用工具與參數
| 工具名 | 必要參數 | 可選參數 | 說明 |
|--------|---------|---------|------|
| search_symptom | symptom(英文) | — | 症狀→化合物→食物 |
| search_graphrag | query(英文) | — | 學術論文搜尋 |
| get_nutrient_ranking | nutrient_field | top_n | 營養素排名 |
| ... （共 12 個工具）

### 輸出格式（嚴格 JSON）
{
  "tool_calls": [{"tool": "名稱", "args": {"參數": "值"}}],
  "params_from_store": {"age": 25},  ← 從記憶中取得的參數
  "still_missing": ["param_name"],   ← 還缺什麼
  "ask_user": "追問語句",            ← 怎麼追問
  "reasoning": "一句話說明"
}

### 規則
- 優先從 structured_store 取值
- still_missing 只列真正無法推斷的參數
- symptom 必須是英文
- 行事曆日期用 YYYY-MM-DD''', "PROMPT_PHASE2 完整內容")

# Prompt 4
doc.add_heading('10.4 PROMPT_ANALYSIS（綜合分析回覆指令）', level=2)
doc.add_paragraph('位置：advisor.py L234-274')
doc.add_paragraph('通信鏈位置：工具執行完成 → _step_analysis() → _call_gemma() → Gemma LLM → SSE reply 事件')
doc.add_paragraph('溫度：0.3 | max_tokens：16384 | 重試：2 次')

add_code_block('''請你根據工具查詢結果，用繁體中文提供專業營養分析回覆。

## 核心規則
- 所有數據必須來自下方的「工具執行結果」，禁止從訓練資料引用任何營養數字
- 沒有查到就明確說「目前沒有查到相關數據」
- 禁止使用「根據一般營養學原理」「通常建議」等包裝語句
- 前端已自動渲染表格和卡片，你不需要列出原始數據清單
- 你的角色是「分析師」：解讀數據的意義，不是複述數據

## 回覆結構

### search_graphrag — 最高優先規則
必須將 answer 和 evidence 完整呈現，禁止摘要、禁止縮減、禁止改寫。

### 其他工具
前端已渲染互動式卡片。你的回覆應：
- 一句說找到什麼
- 一句指出最重要的洞見
- 可選一句建議下一步

---

使用者原始問題：{message}

{memory_block}

工具執行結果：
{tool_results}     ← 所有工具回傳的 JSON 資料

請用繁體中文輸出分析回覆：''', "PROMPT_ANALYSIS 完整內容")

# Prompt 5
doc.add_heading('10.5 PROMPT_NARRATIVE_SUMMARY（敘述記憶指令）', level=2)
doc.add_paragraph('位置：advisor.py L281-303')
doc.add_paragraph('通信鏈位置：分析完成 → _summarize_narrative() → _call_gemma() → SSE narrative_memory 事件')
doc.add_paragraph('溫度：0.2 | max_tokens：16384')

add_code_block('''你是對話摘要器。根據舊摘要和這輪對話，輸出更新後的對話脈絡摘要。

規則：
- 只記對話流程（誰問了什麼、AI 怎麼回的、話題轉折、使用者的情緒和態度）
- 不記具體數據（年齡、營養素數值等已由系統另存）
- 記錄使用者的關注點（例如「使用者特別在意蛋白質攝取」）
- 記錄對話的邏輯鏈（例如「使用者先問了疲勞→AI建議補鐵→使用者接著問含鐵食物」）
- 新資訊覆蓋舊資訊
- 用分號分段

舊摘要：{old_narrative}
本輪使用者說：{user_msg}
本輪 AI 回覆：{ai_reply}
本輪使用的工具：{tools_used}
本輪工具結果摘要：{tool_summary}

更新後的摘要：''', "PROMPT_NARRATIVE_SUMMARY 完整內容")

# Prompt 6
doc.add_heading('10.6 PROMPT_STRUCTURED_EXTRACT（結構化提取指令）', level=2)
doc.add_paragraph('位置：advisor.py L305-362')
doc.add_paragraph('通信鏈位置：與敘述記憶並行 → _extract_structured_data() → _call_gemma() → SSE structured_store 事件')
doc.add_paragraph('溫度：0.1 | max_tokens：16384 | response_format：json_object')

add_code_block('''從這輪對話和工具結果中，提取以下固定欄位的值。
輸出一個 JSON 物件，只包含這一輪有新值的欄位。
⚠️ 必須輸出 JSON 物件（用 { } 包裹），禁止輸出 JSON 陣列。

## 可提取的欄位

### user_facts（使用者基本資訊）
- age, gender, height_cm, weight_kg, activity_level, goal
- health_conditions, dietary_restrictions, allergies

### food_context（食物相關）
- foods_searched, foods_mentioned_today, categories_browsed

### recipe_entry（食譜物件）
total_nutrition 是動態欄位：有多少記多少，寧可多記不可少記。

### symptom_entry / ranking_entry / calendar_update
### academic_entry / papers_referenced / dri_awareness
### action_entry / saved_recipes_summary

本輪使用者說：{user_msg}
本輪工具結果：{tool_results_json}
本輪 AI 回覆摘要：{reply_summary}

JSON 物件：''', "PROMPT_STRUCTURED_EXTRACT 完整內容")

# Prompt 7-8
doc.add_heading('10.7 食材提取與驗證的子 Prompt', level=2)
doc.add_paragraph('這兩個 Prompt 不在主 Pipeline 裡，而是在 analyze_recipe 工具的「子步驟」中。')

doc.add_heading('Extraction Prompt（食材提取驗證）', level=3)
doc.add_paragraph('位置：gemma_extraction_service.py L16-31')
doc.add_paragraph('通信鏈位置：POST /recipe/extract-with-amounts → refine_extraction() → gemma_client.complete_json()')

add_code_block('''You are a food ingredient extraction validator for Traditional Chinese recipes.
You receive a recipe text and a list of ingredients already extracted by a rule-based parser.

Your job:
1. Check if any real food ingredients are MISSING — add them
2. Check if any items are NOT food ingredients (cooking actions, heat levels) — remove them
3. Check for unnecessary suffixes (切絲/切丁/切片) — rename to clean food name
4. Check for DUPLICATES — remove redundant ones

Common false positives: 大火/小火/中火, 備用, 適量

Respond ONLY with JSON:
{"corrections": [
  {"action": "add", "name": "食材名", "amount": "", "reason": ""},
  {"action": "remove", "name": "", "reason": ""},
  {"action": "rename", "original": "", "corrected": "", "reason": ""}
]}''', "Extraction _SYSTEM_PROMPT 完整內容")

doc.add_heading('Verification Prompt（食材驗證）', level=3)
doc.add_paragraph('位置：gemma_verification_service.py L16-28')
doc.add_paragraph('通信鏈位置：POST /recipe/lookup → verify_ingredient_list() → gemma_client.complete_json()')

add_code_block('''You are a food ingredient list validator for a Taiwanese nutrition database.
Review the following ingredient list and:
1. Flag AMBIGUOUS ingredients (e.g. 玉米粉 could be starch or flour)
2. Flag DUPLICATE ingredients (same food, different names)
3. Standardize names to common Taiwanese database names

Do NOT add or remove ingredients — only flag issues.

Respond ONLY with JSON:
{"verified": ["食材1", "食材2"], "warnings": [{"ingredient": "名", "note": "說明"}]}''', "Verification _SYSTEM_PROMPT 完整內容")

# Prompt 9
doc.add_heading('10.8 GraphRAG Prompt（學術搜尋回答）', level=2)
doc.add_paragraph('位置：graph/src/api_server.py L56-80')
doc.add_paragraph('通信鏈位置：POST /graphrag-api/query → _call_gemma_stream() → Gemma LLM')

add_code_block('''# SYSTEM_PROMPT（主回答）
You are a sports nutrition research assistant.
Answer based ONLY on the provided context.

Requirements:
1. Write a comprehensive, well-structured answer.
2. Use markdown headers (##) to organize sections.
3. Cite specific findings: doses, durations, effect sizes, p-values.
4. Compare or synthesize across multiple studies.
5. End with a brief practical takeaway.

Answer in Traditional Chinese (繁體中文).

# EVIDENCE_PROMPT（證據審查，在主回答之後）
Based on the SAME context, provide a detailed evidence review:
- Study design: RCT / meta-analysis / systematic review
- Participants: sample size, population
- Protocol: intervention details, dosages, duration
- Results: measurements, statistical outcomes

Organize by claim. Focus on methodology and data.
Answer in Traditional Chinese (繁體中文).''', "GraphRAG SYSTEM_PROMPT + EVIDENCE_PROMPT 完整內容")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 第十一章
# ═══════════════════════════════════════════════════════════════

doc.add_heading('第十一章 前端代理路由與 API 調用機制', level=1)

add_svg_reference("11_frontend_proxy.svg", "前端代理路由圖")

doc.add_heading('11.1 Axios 客戶端配置', level=3)
doc.add_paragraph(
    '前端所有 API 呼叫（除了 SSE 和 GraphRAG）都透過一個共用的 Axios 客戶端：')

add_code_block('''// client.js — 前端 HTTP 客戶端
import axios from "axios";

const client = axios.create({
  baseURL: "/api",           // 所有請求自動加上 /api 前綴
  timeout: 120000,           // 2 分鐘超時
});

// 請求攔截器：自動加上 JWT
client.interceptors.request.use((config) => {
  const token = localStorage.getItem("token");
  if (token) {
    config.headers.Authorization = `Bearer ${token}`;
  }
  return config;
});

// 回應攔截器：401 自動跳轉登入
client.interceptors.response.use(
  (response) => response,
  (error) => {
    if (error.response?.status === 401) {
      localStorage.removeItem("token");
      localStorage.removeItem("user");
      window.location.href = "/login";
    }
    return Promise.reject(error);
  }
);

export default client;''', "client.js — Axios 客戶端配置")

doc.add_heading('11.2 所有前端 API 模組一覽', level=3)

add_table(
    ['模組檔案', '函式數量', '主要功能'],
    [
        ['advisorApi.js', '1', 'SSE 串流對話（不用 axios）'],
        ['authApi.js', '5', '登入/註冊/取得用戶/讀取Profile/更新Profile'],
        ['recipeApi.js', '10', '食譜提取/查詢/計算/儲存/列表/刪除/更新/Gemma狀態'],
        ['foodApi.js', '1', '食物搜尋'],
        ['taiwanFoodApi.js', '4', '台灣食品瀏覽/分類/統計/詳情'],
        ['foodbFoodApi.js', '1', 'FooDB 食品詳情'],
        ['healthApi.js', '7', '健康效果/化合物/食物/症狀/DRI缺口/營養素排名'],
        ['calendarApi.js', '3', '行事曆查詢/新增/刪除'],
        ['dietApi.js', '3', '飲食日誌查詢/新增/刪除'],
        ['stopwordsApi.js', '5', '停用詞 CRUD + 批次新增'],
        ['synonymDictApi.js', '10', '同義詞 CRUD + 查詢 + 擴展 + Pipeline'],
        ['graphragApi.js', '1', '論文元數據查詢'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 第十二章
# ═══════════════════════════════════════════════════════════════

doc.add_heading('第十二章 Gemma LLM 整合細節', level=1)
doc.add_heading('——食材提取與驗證的子 Prompt', level=2)

doc.add_paragraph(
    '除了主 Pipeline 的 6 個 Prompt 之外，系統還有兩個「子 Prompt」'
    '用於食譜分析過程中的食材提取和驗證。這些子 Prompt 走的是不同的通訊路徑。')

doc.add_heading('12.1 主 Pipeline vs 子 Pipeline 的通訊差異', level=3)

add_table(
    ['項目', '主 Pipeline（advisor.py）', '子 Pipeline（extraction/verification）'],
    [
        ['LLM 客戶端', '_call_gemma() 直接呼叫', 'gemma_client.complete_json() 間接呼叫'],
        ['訊息格式', '單一 {"role":"user"} 訊息', '{"role":"system"} + {"role":"user"} 兩段式'],
        ['JSON 解析', '手動 json.loads()', '3 層 fallback 解析器'],
        ['超時', '120 秒', '提取 10 秒 / 驗證 5 秒'],
        ['失敗處理', '重試 2 次', '直接 fallback 到規則結果'],
        ['溫度', '0.1-0.5', '0.2'],
    ]
)

doc.add_heading('12.2 gemma_client.py — LLM 通訊層', level=3)
doc.add_paragraph(
    '這是系統中「打電話給 Gemma LLM」的共用元件。它的關鍵設計是 3 層 JSON 解析 fallback：')

add_code_block('''# gemma_client.py — JSON 解析 3 層 fallback

async def complete_json(system_prompt, user_content, ...):
    raw_text = await complete(system_prompt, user_content, ...)
    return _extract_json(raw_text)

def _extract_json(text):
    # 第 1 層：直接 parse
    try:
        return json.loads(text)
    except:
        pass

    # 第 2 層：從 Markdown code fence 提取
    # 有時 LLM 會用 ```json ... ``` 包裹
    match = re.search(r"```(?:json)?\\s*(.*?)```", text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(1))
        except:
            pass

    # 第 3 層：找第一組平衡的 {} 或 []
    # 有時 LLM 會在 JSON 前後加多餘的文字
    for start_char, end_char in [("{", "}"), ("[", "]")]:
        start = text.find(start_char)
        if start >= 0:
            depth = 0
            for i in range(start, len(text)):
                if text[i] == start_char: depth += 1
                elif text[i] == end_char: depth -= 1
                if depth == 0:
                    try:
                        return json.loads(text[start:i+1])
                    except:
                        break

    raise GemmaParseError("Cannot parse JSON from LLM output")''',
"gemma_client.py 3 層 JSON 解析 fallback")

add_note("為什麼需要 3 層 fallback？因為 LLM 不是程式——它有時會在 JSON 前面加一句 '好的，以下是結果：'，"
    "或者用 Markdown 的 ``` 包裹。這 3 層 fallback 確保幾乎所有情況都能成功解析。")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 附錄 A
# ═══════════════════════════════════════════════════════════════

doc.add_heading('附錄 A 所有 API 端點清單', level=1)

doc.add_heading('Backend 服務（Port 8001）', level=2)
add_table(
    ['方法', '端點', '用途'],
    [
        ['GET', '/healthz', '健康檢查'],
        ['POST', '/foods/search/batch', '批次搜尋食物（中+英）'],
        ['GET', '/foods/search/zh', '中文食物搜尋'],
        ['GET', '/foods/search/zh/exact', '中文精確搜尋'],
        ['GET', '/foods/search/en', '英文食物搜尋'],
        ['GET', '/foods/search/en/exact', '英文精確搜尋'],
        ['GET', '/foods/names/all', '取得所有食物名稱'],
        ['GET', '/foods/taiwan/browse', '瀏覽台灣食品'],
        ['GET', '/foods/taiwan/categories', '台灣食品分類'],
        ['GET', '/foods/taiwan/stats', '台灣食品統計'],
        ['GET', '/foods/{id}/nutrition', '單筆食物營養'],
        ['GET', '/foods/foodb/{id}', 'FooDB 食品詳情'],
        ['GET', '/foods/{id}', '食物詳情'],
        ['POST', '/foods/aggregate-nutrition', '批次營養計算'],
        ['GET', '/recipes', '列出食譜'],
        ['GET', '/recipes/{id}', '食譜詳情'],
        ['POST', '/recipes', '建立食譜'],
        ['PUT', '/recipes/{id}', '更新食譜'],
        ['DELETE', '/recipes/{id}', '刪除食譜'],
        ['GET', '/health/nutrients/list', '營養素清單'],
        ['GET', '/health/nutrients/search', '搜尋營養素'],
        ['GET', '/health/nutrients/{field}/top-foods', '營養素排名'],
        ['GET', '/health/effects/search', '搜尋健康效果'],
        ['GET', '/health/effects/{id}/compounds', '效果的化合物'],
        ['GET', '/health/compounds/{id}/foods', '化合物的食物'],
        ['GET', '/users/{id}', '取得使用者'],
        ['POST', '/users', '建立使用者'],
        ['GET', '/users/{id}/profile', '取得 Profile'],
        ['PUT', '/users/{id}/profile', '更新 Profile'],
        ['GET', '/diet-logs', '飲食日誌'],
        ['POST', '/diet-logs', '建立日誌'],
        ['GET', '/calendar/entries', '行事曆'],
        ['POST', '/calendar/entries', '建立行事曆'],
        ['DELETE', '/calendar/entries/{id}', '刪除行事曆'],
    ]
)

doc.add_heading('Control 服務（Port 8000）', level=2)
add_table(
    ['方法', '端點', '用途'],
    [
        ['POST', '/auth/login', '登入'],
        ['POST', '/auth/register', '註冊'],
        ['GET', '/auth/me', '取得當前使用者'],
        ['POST', '/advisor/chat', 'AI 顧問 SSE 串流'],
        ['GET', '/recipe/gemma/status', 'Gemma 狀態'],
        ['POST', '/recipe/extract', '食材提取'],
        ['POST', '/recipe/extract-with-amounts', '食材提取（含用量）'],
        ['POST', '/recipe/reverse-search', '反向搜尋（jieba）'],
        ['POST', '/recipe/lookup', '食材查詢匹配'],
        ['POST', '/recipe/calculate-nutrition', '營養計算'],
        ['POST', '/recipe/save', '儲存食譜'],
        ['GET', '/recipe/list', '食譜列表'],
        ['GET', '/food/search', '食物搜尋'],
        ['GET', '/food/taiwan/browse', '台灣食品瀏覽'],
        ['GET', '/food/taiwan/{id}', '台灣食品詳情'],
        ['GET', '/food/foodb/{id}', 'FooDB 詳情'],
        ['POST', '/health/symptom-search', '症狀搜尋'],
        ['POST', '/health/dri-gap', 'DRI 缺口分析'],
        ['GET', '/health/nutrients/{field}/top-foods', '營養素排名'],
        ['GET', '/profile', '讀取 Profile'],
        ['PUT', '/profile', '更新 Profile'],
        ['GET', '/dri/fields', 'DRI 欄位定義'],
        ['GET', '/stopwords', '停用詞列表'],
        ['POST', '/stopwords', '新增停用詞'],
        ['POST', '/stopwords/bulk', '批次新增停用詞'],
        ['DELETE', '/stopwords/{word}', '刪除停用詞'],
        ['GET', '/synonym-dict/entries', '同義詞列表'],
        ['POST', '/synonym-dict/expand', '擴展同義詞'],
        ['GET', '/calendar/entries', '行事曆'],
        ['POST', '/calendar/entries', '建立行事曆'],
    ]
)

doc.add_heading('GraphRAG API 服務（Port 8002）', level=2)
add_table(
    ['方法', '端點', '用途'],
    [
        ['GET', '/health', '健康檢查'],
        ['POST', '/query', '檢索 + LLM SSE 串流'],
        ['POST', '/query-sync', '同步檢索 + LLM'],
        ['POST', '/retrieve-only', '只檢索不呼叫 LLM'],
        ['POST', '/skip', '跳過'],
        ['POST', '/abort', '中止'],
        ['GET', '/papers', '論文元數據'],
        ['GET', '/papers/enriched', '論文豐富元數據'],
        ['GET', '/papers/{id}/pdf', 'PDF 檔案'],
        ['GET', '/papers/{id}/figures/{fig}', '論文圖表'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 附錄 B
# ═══════════════════════════════════════════════════════════════

doc.add_heading('附錄 B SVG 圖表索引', level=1)

add_table(
    ['編號', '檔案名稱', '內容說明', '對應章節'],
    [
        ['01', '01_system_architecture.svg', '系統架構總覽圖', '第二章'],
        ['02', '02_llm_pipeline.svg', 'AI 顧問 6 步驟 LLM Pipeline', '第四章'],
        ['03', '03_sse_protocol.svg', 'SSE 通信協議序列圖', '第七章'],
        ['04', '04_recipe_analysis_flow.svg', '食譜分析完整通信鏈', '第五章 5.4 + 第六章'],
        ['05', '05_symptom_search_flow.svg', '症狀搜尋通信鏈', '第五章 5.5'],
        ['06', '06_nutrient_ranking_flow.svg', '營養素排名通信鏈', '第五章 5.1'],
        ['07', '07_graphrag_flow.svg', 'GraphRAG 通信鏈', '第五章 5.6'],
        ['08', '08_tool_dispatch.svg', 'Tool Dispatch 工具分發機制', '第五章 開頭'],
        ['09', '09_dual_memory_system.svg', '雙記憶體系統架構', '第九章'],
        ['10', '10_prompt_positions.svg', 'Prompt 在通信鏈中的位置', '第十章'],
        ['11', '11_frontend_proxy.svg', '前端代理路由圖', '第十一章'],
    ]
)

doc.add_paragraph()
doc.add_paragraph('— 報告結束 —').alignment = WD_ALIGN_PARAGRAPH.CENTER

# ═══════════════════════════════════════════════════════════════
# 儲存
# ═══════════════════════════════════════════════════════════════

doc.save(OUTPUT)
print(f"報告已生成：{OUTPUT}")
print(f"SVG 圖表目錄：{SVG_DIR}")
