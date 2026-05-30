#!/usr/bin/env python3
"""
Generate SSE Bug Tutorial Word Document
With embedded SVG→PNG diagrams
"""
import os, io, tempfile
import cairosvg
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUT_DIR = "/home/aiiauser/JM/0325new"
FONT_ZH = "Noto Sans CJK TC"
FONT_EN = "Arial, Helvetica, sans-serif"
FONT_SVG = f"{FONT_ZH}, {FONT_EN}"

# ─────────────────────── SVG Helpers ───────────────────────

COMMON_MARKERS = '''
    <marker id="ag" markerWidth="10" markerHeight="7" refX="9" refY="3.5" orient="auto"><polygon points="0 0, 10 3.5, 0 7" fill="#38A169"/></marker>
    <marker id="ab" markerWidth="10" markerHeight="7" refX="9" refY="3.5" orient="auto"><polygon points="0 0, 10 3.5, 0 7" fill="#3182CE"/></marker>
    <marker id="ar" markerWidth="10" markerHeight="7" refX="9" refY="3.5" orient="auto"><polygon points="0 0, 10 3.5, 0 7" fill="#E53E3E"/></marker>
    <marker id="ap1" markerWidth="8" markerHeight="6" refX="7" refY="3" orient="auto"><polygon points="0 0, 8 3, 0 6" fill="#999"/></marker>
    <marker id="ap2" markerWidth="8" markerHeight="6" refX="7" refY="3" orient="auto"><polygon points="0 0, 8 3, 0 6" fill="#38A169"/></marker>
    <marker id="a1" markerWidth="10" markerHeight="7" refX="9" refY="3.5" orient="auto"><polygon points="0 0, 10 3.5, 0 7" fill="#3182CE"/></marker>
    <marker id="a2" markerWidth="10" markerHeight="7" refX="9" refY="3.5" orient="auto"><polygon points="0 0, 10 3.5, 0 7" fill="#38A169"/></marker>
    <marker id="ay" markerWidth="10" markerHeight="7" refX="9" refY="3.5" orient="auto"><polygon points="0 0, 10 3.5, 0 7" fill="#D69E2E"/></marker>
    <marker id="accc" markerWidth="8" markerHeight="6" refX="7" refY="3" orient="auto"><polygon points="0 0, 8 3, 0 6" fill="#ccc"/></marker>
'''

def svg_wrap(content, w, h):
    return f'''<svg xmlns="http://www.w3.org/2000/svg" width="{w}" height="{h}" viewBox="0 0 {w} {h}">
  <defs>
    {COMMON_MARKERS}
    <style>
      text {{ font-family: {FONT_SVG}; }}
      .title {{ font-size: 18px; font-weight: bold; fill: #1a202c; }}
      .label {{ font-size: 14px; fill: #2d3748; }}
      .label-sm {{ font-size: 12px; fill: #4a5568; }}
      .label-xs {{ font-size: 11px; fill: #718096; }}
      .code {{ font-family: monospace, {FONT_ZH}; font-size: 12px; fill: #2d3748; }}
      .accent {{ font-size: 14px; fill: #e53e3e; font-weight: bold; }}
      .good {{ font-size: 14px; fill: #38a169; font-weight: bold; }}
    </style>
  </defs>
  {content}
</svg>'''

def rounded_rect(x, y, w, h, fill="#EBF8FF", stroke="#3182CE", rx=8):
    return f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="{rx}" fill="{fill}" stroke="{stroke}" stroke-width="1.5"/>'

def arrow_right(x1, y1, x2, y2, color="#3182CE", label="", label_y_offset=-10):
    mid_x = (x1 + x2) / 2
    parts = [
        f'<line x1="{x1}" y1="{y1}" x2="{x2}" y2="{y2}" stroke="{color}" stroke-width="2" marker-end="url(#arr-{color.replace("#","")})"/>',
        f'<defs><marker id="arr-{color.replace("#","")}" markerWidth="10" markerHeight="7" refX="9" refY="3.5" orient="auto"><polygon points="0 0, 10 3.5, 0 7" fill="{color}"/></marker></defs>'
    ]
    if label:
        parts.append(f'<text x="{mid_x}" y="{y1 + label_y_offset}" text-anchor="middle" class="label-sm">{label}</text>')
    return "\n".join(parts)

def arrow_down(x1, y1, x2, y2, color="#3182CE"):
    return f'''<defs><marker id="arrd-{color.replace("#","")}" markerWidth="10" markerHeight="7" refX="3.5" refY="5" orient="auto"><polygon points="0 0, 7 0, 3.5 10" fill="{color}"/></marker></defs>
<line x1="{x1}" y1="{y1}" x2="{x2}" y2="{y2}" stroke="{color}" stroke-width="2" marker-end="url(#arrd-{color.replace("#","")})"/>'''

def text_center(x, y, txt, cls="label"):
    return f'<text x="{x}" y="{y}" text-anchor="middle" class="{cls}">{txt}</text>'

def text_left(x, y, txt, cls="label"):
    return f'<text x="{x}" y="{y}" class="{cls}">{txt}</text>'


def svg_to_png(svg_str, scale=2):
    png_bytes = cairosvg.svg2png(bytestring=svg_str.encode("utf-8"), scale=scale)
    return io.BytesIO(png_bytes)

# ─────────────────────── SVG Diagrams ───────────────────────

def make_svg_http_basic():
    """Diagram 1: Basic HTTP - like ordering food"""
    w, h = 700, 320
    content = f'''
    {text_center(350, 30, "網頁通信的基本模式：就像在餐廳點餐", "title")}

    <!-- Browser -->
    {rounded_rect(40, 70, 180, 80, "#EBF8FF", "#3182CE")}
    {text_center(130, 105, "🖥️ 瀏覽器", "label")}
    {text_center(130, 128, "(你，客人)", "label-sm")}

    <!-- Server -->
    {rounded_rect(480, 70, 180, 80, "#F0FFF4", "#38A169")}
    {text_center(570, 105, "🖧 伺服器", "label")}
    {text_center(570, 128, "(廚房)", "label-sm")}

    <!-- Request arrow -->
    <line x1="220" y1="95" x2="478" y2="95" stroke="#3182CE" stroke-width="2.5" marker-end="url(#a1)"/>
    {text_center(350, 83, "① 請求（Request）：我要一份牛排！", "label-sm")}

    <!-- Response arrow -->
    <line x1="480" y1="130" x2="222" y2="130" stroke="#38A169" stroke-width="2.5" marker-end="url(#a2)"/>
    {text_center(350, 155, "② 回應（Response）：這是你的牛排！", "label-sm")}

    <!-- Explanation -->
    {rounded_rect(40, 190, 620, 110, "#FFFAF0", "#DD6B20", 10)}
    {text_left(60, 218, "📌 傳統 HTTP 通信 = 一問一答", "label")}
    {text_left(60, 245, "• 客人（瀏覽器）問一次，廚房（伺服器）答一次", "label-sm")}
    {text_left(60, 268, "• 想知道新消息？只能再問一次（重新發請求）", "label-sm")}
    {text_left(60, 291, "• 就像你每次想知道餐點好了沒，都要再舉手問服務生", "label-sm")}
    '''
    return svg_wrap(content, w, h)


def make_svg_sse_vs_polling():
    """Diagram 2: SSE vs Polling comparison"""
    w, h = 720, 480
    content = f'''
    {text_center(360, 30, "兩種取得最新消息的方式", "title")}

    <!-- Left: Polling -->
    {rounded_rect(20, 55, 320, 35, "#FED7D7", "#E53E3E", 6)}
    {text_center(180, 78, "❌ 傳統方式：一直問（Polling）", "label-sm")}

    {rounded_rect(30, 105, 80, 40, "#EBF8FF", "#3182CE")}
    {text_center(70, 130, "瀏覽器", "label-xs")}
    {rounded_rect(250, 105, 80, 40, "#F0FFF4", "#38A169")}
    {text_center(290, 130, "伺服器", "label-xs")}

    <line x1="110" y1="120" x2="248" y2="120" stroke="#999" stroke-width="1.5" marker-end="url(#ap1)"/>
    {text_center(180, 113, "好了沒？", "label-xs")}

    <line x1="250" y1="135" x2="112" y2="135" stroke="#ccc" stroke-width="1.5" stroke-dasharray="4" marker-end="url(#accc)"/>
    {text_center(180, 155, "還沒", "label-xs")}

    <line x1="110" y1="170" x2="248" y2="170" stroke="#999" stroke-width="1.5" marker-end="url(#ap1)"/>
    {text_center(180, 163, "好了沒？", "label-xs")}

    <line x1="250" y1="185" x2="112" y2="185" stroke="#ccc" stroke-width="1.5" stroke-dasharray="4" marker-end="url(#accc)"/>
    {text_center(180, 205, "還沒", "label-xs")}

    <line x1="110" y1="220" x2="248" y2="220" stroke="#999" stroke-width="1.5" marker-end="url(#ap1)"/>
    {text_center(180, 213, "好了沒？", "label-xs")}

    <line x1="250" y1="235" x2="112" y2="235" stroke="#38A169" stroke-width="1.5" marker-end="url(#ap1)"/>
    {text_center(180, 255, "好了！", "label-xs")}

    {text_left(40, 285, "😫 問了 3 次才拿到答案", "accent")}
    {text_left(40, 305, "浪費網路流量和伺服器資源", "label-xs")}

    <!-- Right: SSE -->
    {rounded_rect(380, 55, 320, 35, "#C6F6D5", "#38A169", 6)}
    {text_center(540, 78, "✅ SSE：伺服器主動推送", "label-sm")}

    {rounded_rect(390, 105, 80, 40, "#EBF8FF", "#3182CE")}
    {text_center(430, 130, "瀏覽器", "label-xs")}
    {rounded_rect(610, 105, 80, 40, "#F0FFF4", "#38A169")}
    {text_center(650, 130, "伺服器", "label-xs")}

    <line x1="470" y1="125" x2="608" y2="125" stroke="#3182CE" stroke-width="1.5" marker-end="url(#ab)"/>
    {text_center(540, 118, "建立連線", "label-xs")}

    <line x1="610" y1="170" x2="472" y2="170" stroke="#38A169" stroke-width="2" marker-end="url(#ap2)"/>
    {text_center(540, 163, "思考中...", "label-xs")}

    <line x1="610" y1="200" x2="472" y2="200" stroke="#38A169" stroke-width="2" marker-end="url(#ap2)"/>
    {text_center(540, 193, "正在查食物資料...", "label-xs")}

    <line x1="610" y1="230" x2="472" y2="230" stroke="#38A169" stroke-width="2" marker-end="url(#ap2)"/>
    {text_center(540, 223, "查到了！表格資料", "label-xs")}

    <line x1="610" y1="260" x2="472" y2="260" stroke="#38A169" stroke-width="2" marker-end="url(#ap2)"/>
    {text_center(540, 253, "最終回覆", "label-xs")}

    {text_left(400, 295, "😊 只問一次，伺服器持續推送", "good")}
    {text_left(400, 315, "即時、省流量、體驗好", "label-xs")}

    <!-- Bottom analogy -->
    {rounded_rect(20, 340, 680, 120, "#FAF5FF", "#805AD5", 10)}
    {text_left(40, 368, "🎯 生活比喻", "label")}
    {text_left(40, 395, "Polling = 每 5 分鐘打電話問外送到了沒（煩！）", "label-sm")}
    {text_left(40, 418, "SSE  = 外送 App 自動推播通知：接單了→出發了→快到了→已送達", "label-sm")}
    {text_left(40, 441, "SSE 的全名是 Server-Sent Events（伺服器傳送事件）", "label-sm")}
    '''
    return svg_wrap(content, w, h)


def make_svg_sse_format():
    """Diagram 3: SSE event format"""
    w, h = 700, 400
    content = f'''
    {text_center(350, 30, "SSE 事件的格式：伺服器寄出的「明信片」", "title")}

    <!-- Single event -->
    {rounded_rect(40, 55, 300, 160, "#F7FAFC", "#A0AEC0", 6)}
    {text_left(55, 80, "一則 SSE 事件", "label")}

    <rect x="55" y="95" width="270" height="100" rx="4" fill="#2D3748"/>
    {text_left(65, 118, 'event: tool_result', "code")}
    <text x="65" y="118" font-family="monospace" font-size="12" fill="#68D391">event: tool_result</text>
    <text x="65" y="142" font-family="monospace" font-size="12" fill="#63B3ED">data: {{"name":"search_food",</text>
    <text x="65" y="162" font-family="monospace" font-size="12" fill="#63B3ED">       "data":{{"results":[...]}}}}</text>
    <text x="65" y="186" font-family="monospace" font-size="12" fill="#A0AEC0">（空行 = 事件結束）</text>

    <!-- Explanation arrows -->
    {rounded_rect(380, 60, 290, 50, "#FEFCBF", "#D69E2E", 6)}
    {text_center(525, 82, "第一行：事件「類型」", "label-sm")}
    {text_center(525, 100, "（這是什麼消息？）", "label-xs")}

    <line x1="380" y1="85" x2="330" y2="110" stroke="#D69E2E" stroke-width="1.5"/>

    {rounded_rect(380, 125, 290, 50, "#BEE3F8", "#3182CE", 6)}
    {text_center(525, 147, "第二行：事件「內容」", "label-sm")}
    {text_center(525, 165, "（JSON 格式的資料）", "label-xs")}

    <line x1="380" y1="150" x2="330" y2="150" stroke="#3182CE" stroke-width="1.5"/>

    <!-- Our system events -->
    {rounded_rect(40, 235, 620, 145, "#EBF8FF", "#3182CE", 10)}
    {text_left(60, 262, "📋 我們系統用到的 9 種事件：", "label")}

    {rounded_rect(60, 275, 110, 28, "#C6F6D5", "#38A169", 4)}
    {text_center(115, 294, "thinking", "label-xs")}

    {rounded_rect(180, 275, 80, 28, "#C6F6D5", "#38A169", 4)}
    {text_center(220, 294, "echo", "label-xs")}

    {rounded_rect(270, 275, 110, 28, "#FEFCBF", "#D69E2E", 4)}
    {text_center(325, 294, "tool_start", "label-xs")}

    {rounded_rect(390, 275, 110, 28, "#FEFCBF", "#D69E2E", 4)}
    {text_center(445, 294, "tool_result", "label-xs")}

    {rounded_rect(510, 275, 110, 28, "#FEFCBF", "#D69E2E", 4)}
    {text_center(565, 294, "tool_done", "label-xs")}

    {rounded_rect(60, 315, 80, 28, "#BEE3F8", "#3182CE", 4)}
    {text_center(100, 334, "reply", "label-xs")}

    {rounded_rect(150, 315, 100, 28, "#E9D8FD", "#805AD5", 4)}
    {text_center(200, 334, "memory", "label-xs")}

    {rounded_rect(260, 315, 140, 28, "#E9D8FD", "#805AD5", 4)}
    {text_center(330, 334, "workflow_state", "label-xs")}

    {rounded_rect(410, 315, 80, 28, "#FED7D7", "#E53E3E", 4)}
    {text_center(450, 334, "error", "label-xs")}

    {text_left(60, 365, "⚡ 關鍵：tool_result 事件帶著表格資料，如果漏掉就看不到表格", "accent")}
    '''
    return svg_wrap(content, w, h)


def make_svg_proxy_chain():
    """Diagram 4: Proxy chain - the relay race"""
    w, h = 720, 440
    content = f'''
    {text_center(360, 30, "訊息的傳遞鏈：一場接力賽跑", "title")}

    <!-- Localhost path (simple) -->
    {rounded_rect(20, 55, 680, 35, "#C6F6D5", "#38A169", 6)}
    {text_center(360, 78, "路徑 A：在自己電腦上測試（localhost）", "label-sm")}

    {rounded_rect(40, 105, 140, 60, "#EBF8FF", "#3182CE")}
    {text_center(110, 132, "🖥️ 瀏覽器", "label-sm")}
    {text_center(110, 150, "(localhost:3001)", "label-xs")}

    <line x1="180" y1="130" x2="308" y2="130" stroke="#38A169" stroke-width="2" marker-end="url(#ag)"/>

    {rounded_rect(310, 105, 140, 60, "#FEFCBF", "#D69E2E")}
    {text_center(380, 132, "📦 前端 Nginx", "label-sm")}
    {text_center(380, 150, "(容器內代理)", "label-xs")}

    <line x1="450" y1="130" x2="538" y2="130" stroke="#38A169" stroke-width="2" marker-end="url(#ag)"/>

    {rounded_rect(540, 105, 140, 60, "#F0FFF4", "#38A169")}
    {text_center(610, 132, "🖧 後端 API", "label-sm")}
    {text_center(610, 150, "(port 8000)", "label-xs")}

    {text_left(295, 180, "只經過 1 個中間人 ✅", "good")}

    <!-- External path (complex) -->
    {rounded_rect(20, 205, 680, 35, "#FED7D7", "#E53E3E", 6)}
    {text_center(360, 228, "路徑 B：從外面連進來（外部 URL）", "label-sm")}

    {rounded_rect(20, 255, 120, 60, "#EBF8FF", "#3182CE")}
    {text_center(80, 282, "🖥️ 瀏覽器", "label-sm")}
    {text_center(80, 300, "(你的手機)", "label-xs")}

    <line x1="140" y1="280" x2="168" y2="280" stroke="#E53E3E" stroke-width="2" marker-end="url(#ar)"/>

    {rounded_rect(170, 255, 120, 60, "#FED7D7", "#E53E3E")}
    {text_center(230, 277, "🌐 外部 Nginx", "label-sm")}
    {text_center(230, 300, "(port 9987)", "label-xs")}

    <line x1="290" y1="280" x2="318" y2="280" stroke="#E53E3E" stroke-width="2" marker-end="url(#ar)"/>

    {rounded_rect(320, 255, 120, 60, "#FED7E2", "#D53F8C")}
    {text_center(380, 277, "🌐 網路傳輸", "label-sm")}
    {text_center(380, 300, "(TCP 封包)", "label-xs")}

    <line x1="440" y1="280" x2="468" y2="280" stroke="#E53E3E" stroke-width="2" marker-end="url(#ar)"/>

    {rounded_rect(470, 255, 110, 60, "#FEFCBF", "#D69E2E")}
    {text_center(525, 277, "📦 前端", "label-sm")}
    {text_center(525, 300, "Nginx", "label-xs")}

    <line x1="580" y1="280" x2="608" y2="280" stroke="#E53E3E" stroke-width="2" marker-end="url(#ar)"/>

    {rounded_rect(610, 255, 90, 60, "#F0FFF4", "#38A169")}
    {text_center(655, 277, "🖧 後端", "label-sm")}
    {text_center(655, 300, "API", "label-xs")}

    {text_left(200, 340, "經過 3 個中間人！每一層都可能出問題 ❌", "accent")}

    <!-- Analogy -->
    {rounded_rect(20, 365, 680, 60, "#FAF5FF", "#805AD5", 10)}
    {text_left(40, 390, "🎯 想像你在傳話遊戲：傳越多人，訊息越容易走樣或漏掉", "label")}
    {text_left(40, 413, "每個「中間人」(代理伺服器) 都可能改變訊息的傳遞方式", "label-sm")}
    '''
    return svg_wrap(content, w, h)


def make_svg_buffering_problem():
    """Diagram 5: Proxy buffering - the secretary problem"""
    w, h = 720, 520
    content = f'''
    {text_center(360, 30, "Bug #1：中間人把消息攢在手裡", "title")}
    {text_center(360, 52, "（Proxy Buffering 代理緩衝問題）", "label-sm")}

    <!-- Bad: Buffered -->
    {rounded_rect(20, 70, 330, 35, "#FED7D7", "#E53E3E", 6)}
    {text_center(185, 93, "❌ 有緩衝：消息被攢在一起", "label-sm")}

    {rounded_rect(30, 120, 80, 40, "#F0FFF4", "#38A169")}
    {text_center(70, 145, "後端", "label-xs")}

    {rounded_rect(160, 120, 80, 40, "#FEFCBF", "#D69E2E")}
    {text_center(200, 145, "Nginx", "label-xs")}

    {rounded_rect(280, 120, 80, 40, "#EBF8FF", "#3182CE")}
    {text_center(320, 145, "瀏覽器", "label-xs")}

    <!-- Events from backend -->
    <line x1="110" y1="125" x2="158" y2="125" stroke="#38A169" stroke-width="1.5" marker-end="url(#ag)"/>
    <text x="134" y="118" font-family="{FONT_SVG}" font-size="10" fill="#38A169">①②③</text>

    <!-- Nginx holds them -->
    <rect x="165" y="165" width="70" height="65" rx="4" fill="#FFF5F5" stroke="#E53E3E" stroke-dasharray="3"/>
    <text x="173" y="183" font-family="{FONT_SVG}" font-size="10" fill="#E53E3E">① thinking</text>
    <text x="173" y="198" font-family="{FONT_SVG}" font-size="10" fill="#E53E3E">② tool_result</text>
    <text x="173" y="213" font-family="{FONT_SVG}" font-size="10" fill="#E53E3E">③ reply</text>
    <text x="175" y="227" font-family="{FONT_SVG}" font-size="10" fill="#E53E3E">等湊滿...</text>

    <!-- All at once -->
    <line x1="240" y1="135" x2="278" y2="135" stroke="#E53E3E" stroke-width="2.5" marker-end="url(#ar)"/>
    <text x="243" y="155" font-family="{FONT_SVG}" font-size="10" fill="#E53E3E">一口氣全送</text>

    {text_left(30, 255, "💥 瀏覽器收到時已經太晚！", "accent")}
    {text_left(30, 275, "程式來不及依序處理每個事件", "label-xs")}

    <!-- Good: Unbuffered -->
    {rounded_rect(380, 70, 320, 35, "#C6F6D5", "#38A169", 6)}
    {text_center(540, 93, "✅ 無緩衝：消息即時轉發", "label-sm")}

    {rounded_rect(390, 120, 80, 40, "#F0FFF4", "#38A169")}
    {text_center(430, 145, "後端", "label-xs")}

    {rounded_rect(520, 120, 80, 40, "#FEFCBF", "#D69E2E")}
    {text_center(560, 145, "Nginx", "label-xs")}

    {rounded_rect(640, 120, 60, 40, "#EBF8FF", "#3182CE")}
    {text_center(670, 145, "瀏覽器", "label-xs")}

    <!-- Event 1 -->
    <line x1="470" y1="125" x2="518" y2="125" stroke="#38A169" stroke-width="1.5" marker-end="url(#ag)"/>
    <line x1="600" y1="125" x2="638" y2="125" stroke="#38A169" stroke-width="1.5" marker-end="url(#ag)"/>
    <text x="530" y="118" font-family="{FONT_SVG}" font-size="10" fill="#38A169">① →</text>

    <!-- Event 2 -->
    <line x1="470" y1="170" x2="518" y2="170" stroke="#38A169" stroke-width="1.5" marker-end="url(#ag)"/>
    <line x1="600" y1="170" x2="638" y2="170" stroke="#38A169" stroke-width="1.5" marker-end="url(#ag)"/>
    <text x="530" y="163" font-family="{FONT_SVG}" font-size="10" fill="#38A169">② →</text>

    <!-- Event 3 -->
    <line x1="470" y1="195" x2="518" y2="195" stroke="#38A169" stroke-width="1.5" marker-end="url(#ag)"/>
    <line x1="600" y1="195" x2="638" y2="195" stroke="#38A169" stroke-width="1.5" marker-end="url(#ag)"/>
    <text x="530" y="188" font-family="{FONT_SVG}" font-size="10" fill="#38A169">③ →</text>

    {text_left(390, 235, "👍 每條消息立刻轉送", "good")}
    {text_left(390, 255, "瀏覽器能即時處理每個事件", "label-xs")}

    <!-- Analogy box -->
    {rounded_rect(20, 300, 680, 100, "#FAF5FF", "#805AD5", 10)}
    {text_left(40, 328, "🎯 比喻：秘書轉交文件的方式", "label")}
    {text_left(40, 355, "❌ 有緩衝 = 秘書收到 10 封信後，一次全部交給你（你可能漏看重要的）", "label-sm")}
    {text_left(40, 378, "✅ 無緩衝 = 秘書每收到 1 封信，立刻交給你（你能即時處理每一封）", "label-sm")}

    <!-- Fix -->
    {rounded_rect(20, 415, 680, 90, "#F0FFF4", "#38A169", 10)}
    {text_left(40, 443, "🔧 修復方法：在 Nginx 設定檔加上以下指令", "label")}
    <rect x="40" y="455" width="640" height="35" rx="4" fill="#2D3748"/>
    <text x="55" y="478" font-family="monospace" font-size="13" fill="#68D391">proxy_buffering off;   proxy_cache off;   chunked_transfer_encoding off;</text>
    '''
    return svg_wrap(content, w, h)


def make_svg_scope_bug():
    """Diagram 6: Variable scope bug - the whiteboard analogy"""
    w, h = 720, 560
    content = f'''
    {text_center(360, 30, "Bug #2：被自動擦掉的白板", "title")}
    {text_center(360, 52, "（Variable Scope 變數作用域問題）", "label-sm")}

    <!-- Analogy section -->
    {rounded_rect(20, 70, 680, 85, "#FAF5FF", "#805AD5", 10)}
    {text_left(40, 95, "🎯 想像你在用白板記錄老師唸的考試題目：", "label")}
    {text_left(40, 118, "老師先說「第一題是選擇題」，你寫在白板上。然後老師唸題目「下列何者為⋯」", "label-sm")}
    {text_left(40, 141, "但如果每次老師繼續唸，白板就自動擦掉！你就不知道這題是選擇題還是申論題了", "label-sm")}

    <!-- Bug illustration -->
    {rounded_rect(20, 170, 330, 35, "#FED7D7", "#E53E3E", 6)}
    {text_center(185, 193, "❌ 錯誤寫法：白板每圈被擦掉", "label-sm")}

    <!-- Loop iteration 1 -->
    {rounded_rect(30, 218, 300, 60, "#FFF5F5", "#FEB2B2")}
    {text_left(40, 238, "第 1 次讀取：收到", "label-xs")}
    <text x="165" y="238" font-family="monospace" font-size="11" fill="#38A169">event: tool_result</text>
    {text_left(40, 258, "白板：eventType =", "label-xs")}
    <text x="175" y="258" font-family="monospace" font-size="11" fill="#3182CE">"tool_result"</text>
    {text_left(225, 268, "✓ 記住了", "label-xs")}

    <!-- Erase -->
    <text x="180" y="295" font-family="{FONT_SVG}" font-size="16" fill="#E53E3E">↓ 進入下一圈，白板被擦掉！</text>

    <!-- Loop iteration 2 -->
    {rounded_rect(30, 305, 300, 65, "#FFF5F5", "#FEB2B2")}
    {text_left(40, 325, "第 2 次讀取：收到", "label-xs")}
    <text x="165" y="325" font-family="monospace" font-size="11" fill="#63B3ED">data: {{"name":...}}</text>
    {text_left(40, 345, "白板：eventType =", "label-xs")}
    <text x="175" y="345" font-family="monospace" font-size="11" fill="#E53E3E">null</text>
    {text_left(215, 345, "← 被重設了！", "label-xs")}
    {text_left(40, 362, "❌ 不知道這是什麼事件 → 直接丟棄", "accent")}

    <!-- Correct version -->
    {rounded_rect(380, 170, 320, 35, "#C6F6D5", "#38A169", 6)}
    {text_center(540, 193, "✅ 正確寫法：白板不會被擦", "label-sm")}

    {rounded_rect(390, 218, 300, 60, "#F0FFF4", "#C6F6D5")}
    {text_left(400, 238, "第 1 次讀取：收到", "label-xs")}
    <text x="525" y="238" font-family="monospace" font-size="11" fill="#38A169">event: tool_result</text>
    {text_left(400, 258, "白板：eventType =", "label-xs")}
    <text x="535" y="258" font-family="monospace" font-size="11" fill="#3182CE">"tool_result"</text>

    <text x="540" y="295" font-family="{FONT_SVG}" font-size="16" fill="#38A169">↓ 進入下一圈，白板保留！</text>

    {rounded_rect(390, 305, 300, 65, "#F0FFF4", "#C6F6D5")}
    {text_left(400, 325, "第 2 次讀取：收到", "label-xs")}
    <text x="525" y="325" font-family="monospace" font-size="11" fill="#63B3ED">data: {{"name":...}}</text>
    {text_left(400, 345, "白板：eventType =", "label-xs")}
    <text x="535" y="345" font-family="monospace" font-size="11" fill="#38A169">"tool_result"</text>
    {text_left(400, 362, "✅ 知道是 tool_result → 正確處理！", "good")}

    <!-- Code comparison -->
    {rounded_rect(20, 390, 330, 155, "#FFF5F5", "#E53E3E", 8)}
    {text_left(35, 413, "❌ 錯誤的程式碼", "accent")}
    <rect x="35" y="420" width="300" height="110" rx="4" fill="#2D3748"/>
    <text x="45" y="440" font-family="monospace" font-size="11" fill="#A0AEC0">while (true) {'{'}</text>
    <text x="45" y="458" font-family="monospace" font-size="11" fill="#FC8181">  let eventType = null; // 每圈重設!</text>
    <text x="45" y="476" font-family="monospace" font-size="11" fill="#A0AEC0">  const {'{'}done, value{'}'} = await reader.read();</text>
    <text x="45" y="494" font-family="monospace" font-size="11" fill="#A0AEC0">  // ... 處理 lines ...</text>
    <text x="45" y="512" font-family="monospace" font-size="11" fill="#A0AEC0">{'}'}</text>

    {rounded_rect(380, 390, 320, 155, "#F0FFF4", "#38A169", 8)}
    {text_left(395, 413, "✅ 正確的程式碼", "good")}
    <rect x="395" y="420" width="290" height="110" rx="4" fill="#2D3748"/>
    <text x="405" y="440" font-family="monospace" font-size="11" fill="#68D391">let eventType = null; // 迴圈外宣告!</text>
    <text x="405" y="458" font-family="monospace" font-size="11" fill="#A0AEC0">while (true) {'{'}</text>
    <text x="405" y="476" font-family="monospace" font-size="11" fill="#A0AEC0">  const {'{'}done, value{'}'} = await reader.read();</text>
    <text x="405" y="494" font-family="monospace" font-size="11" fill="#A0AEC0">  // ... 處理 lines ...</text>
    <text x="405" y="512" font-family="monospace" font-size="11" fill="#A0AEC0">{'}'}</text>
    '''
    return svg_wrap(content, w, h)


def make_svg_snapshot_bug():
    """Diagram 7: Reference vs Snapshot - photo vs pointing"""
    w, h = 720, 480
    content = f'''
    {text_center(360, 30, "Bug #3：照片 vs 手指", "title")}
    {text_center(360, 52, "（Reference vs Snapshot 參考 vs 快照）", "label-sm")}

    <!-- Analogy -->
    {rounded_rect(20, 70, 680, 75, "#FAF5FF", "#805AD5", 10)}
    {text_left(40, 95, "🎯 想像你要把朋友的電話抄給別人：", "label")}
    {text_left(40, 118, "A. 你指著他的手機說「看那支」→ 但他換了號碼你就指到新號碼了（參考 Reference）", "label-sm")}
    {text_left(40, 136, "B. 你先拍張照片存起來 → 就算他換號碼，照片裡還是舊號碼（快照 Snapshot）", "label-sm")}

    <!-- Bad: Reference -->
    {rounded_rect(20, 160, 330, 35, "#FED7D7", "#E53E3E", 6)}
    {text_center(185, 183, "❌ 用「手指指」= 傳參考", "label-sm")}

    <!-- Timeline -->
    {rounded_rect(30, 208, 310, 55, "#FFF5F5", "#FEB2B2")}
    {text_left(40, 228, "步驟 1：工具查到結果", "label-xs")}
    <text x="40" y="248" font-family="monospace" font-size="10" fill="#2D3748">turnToolResults = [食物表格資料]</text>

    {rounded_rect(30, 275, 310, 55, "#FFF5F5", "#FEB2B2")}
    {text_left(40, 295, "步驟 2：存到訊息中（指著原陣列）", "label-xs")}
    <text x="40" y="315" font-family="monospace" font-size="10" fill="#E53E3E">tool_results: turnToolResults  // 指標!</text>

    {rounded_rect(30, 342, 310, 55, "#FFF5F5", "#FEB2B2")}
    {text_left(40, 362, "步驟 3：清空原陣列", "label-xs")}
    <text x="40" y="382" font-family="monospace" font-size="10" fill="#E53E3E">turnToolResults = []  // 指到的東西沒了!</text>
    {text_left(40, 393, "❌ 訊息裡的表格資料 → 空的！", "accent")}

    <!-- Good: Snapshot -->
    {rounded_rect(380, 160, 320, 35, "#C6F6D5", "#38A169", 6)}
    {text_center(540, 183, "✅ 用「拍照」= 做快照", "label-sm")}

    {rounded_rect(390, 208, 300, 55, "#F0FFF4", "#C6F6D5")}
    {text_left(400, 228, "步驟 1：工具查到結果", "label-xs")}
    <text x="400" y="248" font-family="monospace" font-size="10" fill="#2D3748">turnToolResults = [食物表格資料]</text>

    {rounded_rect(390, 275, 300, 55, "#F0FFF4", "#C6F6D5")}
    {text_left(400, 295, "步驟 2：先拍照，再存入訊息", "label-xs")}
    <text x="400" y="315" font-family="monospace" font-size="10" fill="#38A169">const snapshot = [...turnToolResults]</text>

    {rounded_rect(390, 342, 300, 55, "#F0FFF4", "#C6F6D5")}
    {text_left(400, 362, "步驟 3：清空原陣列 — 沒關係！", "label-xs")}
    <text x="400" y="382" font-family="monospace" font-size="10" fill="#38A169">turnToolResults = []  // 照片還在!</text>
    {text_left(400, 393, "✅ 訊息裡的資料完好無損！", "good")}

    <!-- Key insight -->
    {rounded_rect(20, 410, 680, 55, "#FFFFF0", "#D69E2E", 10)}
    {text_left(40, 433, "💡 在 JavaScript 中，陣列是「參考型別」(Reference Type)", "label")}
    {text_left(40, 453, "[...array] 叫做「展開運算子」，它會複製一份新的（就像拍照）", "label-sm")}
    '''
    return svg_wrap(content, w, h)


def make_svg_localhost_vs_external():
    """Diagram 8: Why localhost works but external doesn't"""
    w, h = 720, 420
    content = f'''
    {text_center(360, 30, "為什麼「在家」沒問題，「出門」就壞了？", "title")}

    <!-- localhost -->
    {rounded_rect(20, 60, 680, 155, "#F0FFF4", "#38A169", 10)}
    {text_left(40, 85, "🏠 在家測試 (localhost:3001)", "label")}

    {rounded_rect(40, 100, 120, 45, "#EBF8FF", "#3182CE")}
    {text_center(100, 127, "瀏覽器", "label-xs")}

    <line x1="160" y1="122" x2="198" y2="122" stroke="#38A169" stroke-width="2" marker-end="url(#ag)"/>

    {rounded_rect(200, 100, 120, 45, "#FEFCBF", "#D69E2E")}
    {text_center(260, 127, "前端Nginx", "label-xs")}

    <line x1="320" y1="122" x2="358" y2="122" stroke="#38A169" stroke-width="2" marker-end="url(#ag)"/>

    {rounded_rect(360, 100, 120, 45, "#F0FFF4", "#38A169")}
    {text_center(420, 127, "後端API", "label-xs")}

    {text_left(40, 170, "✅ 資料在同一台電腦內傳遞，速度極快，不會被切割", "label-sm")}
    {text_left(40, 192, "就像在同一間房子裡傳紙條 — 馬上就到，不會被風吹散", "label-xs")}
    {text_left(40, 207, "event: 和 data: 永遠在同一次讀取中到達", "label-xs")}

    <!-- External -->
    {rounded_rect(20, 230, 680, 175, "#FFF5F5", "#E53E3E", 10)}
    {text_left(40, 255, "🌐 出門測試 (aiia.csie.ncnu.edu.tw:9987)", "label")}

    {rounded_rect(30, 270, 90, 45, "#EBF8FF", "#3182CE")}
    {text_center(75, 297, "手機", "label-xs")}

    <line x1="120" y1="292" x2="143" y2="292" stroke="#E53E3E" stroke-width="1.5" marker-end="url(#ar)"/>

    {rounded_rect(145, 270, 100, 45, "#FED7D7", "#E53E3E")}
    {text_center(195, 292, "外部Nginx", "label-xs")}
    {text_center(195, 307, "(9987)", "label-xs")}

    <line x1="245" y1="292" x2="268" y2="292" stroke="#E53E3E" stroke-width="1.5" marker-end="url(#ar)"/>

    {rounded_rect(270, 270, 90, 45, "#FED7E2", "#D53F8C")}
    {text_center(315, 292, "網路", "label-xs")}
    {text_center(315, 307, "(TCP)", "label-xs")}

    <line x1="360" y1="292" x2="383" y2="292" stroke="#E53E3E" stroke-width="1.5" marker-end="url(#ar)"/>

    {rounded_rect(385, 270, 100, 45, "#FEFCBF", "#D69E2E")}
    {text_center(435, 292, "前端Nginx", "label-xs")}
    {text_center(435, 307, "(緩衝!)", "label-xs")}

    <line x1="485" y1="292" x2="508" y2="292" stroke="#E53E3E" stroke-width="1.5" marker-end="url(#ar)"/>

    {rounded_rect(510, 270, 90, 45, "#F0FFF4", "#38A169")}
    {text_center(555, 297, "後端API", "label-xs")}

    {text_left(40, 340, "❌ 資料穿越網路時可能被「切割」或「合併」", "accent")}
    {text_left(40, 360, "就像寄信 — 兩封信可能被郵差合成一包送，", "label-xs")}
    {text_left(40, 375, "或一封信被撕成兩半分批送到", "label-xs")}
    {text_left(40, 395, "event: 行可能在第一包，data: 行在第二包 → 程式讀第二包時已經忘記 event 了", "label-xs")}
    '''
    return svg_wrap(content, w, h)


def make_svg_tcp_chunking():
    """Diagram 9: TCP chunking visualization"""
    w, h = 720, 400
    content = f'''
    {text_center(360, 30, "TCP 封包切割：資料被「撕成兩半」", "title")}

    <!-- Original SSE events -->
    {text_left(30, 70, "後端發出的完整 SSE 資料：", "label")}

    <rect x="30" y="80" width="660" height="80" rx="6" fill="#2D3748"/>
    <text x="40" y="102" font-family="monospace" font-size="12" fill="#68D391">event: tool_result</text>
    <text x="40" y="120" font-family="monospace" font-size="12" fill="#63B3ED">data: {{"name":"search_food","data":{{"results":[{{"id":1,"name":"雞胸肉",</text>
    <text x="40" y="138" font-family="monospace" font-size="12" fill="#63B3ED">       "cal":121,"protein":23.3}},...]}}}}</text>
    <text x="40" y="155" font-family="monospace" font-size="12" fill="#718096">（空行）</text>

    <!-- Split point -->
    {text_center(360, 185, "經過網路傳輸後，可能被切成兩個封包：", "label")}

    <!-- Chunk 1 -->
    {rounded_rect(30, 200, 310, 90, "#FFF5F5", "#E53E3E", 6)}
    {text_left(45, 220, "📦 第 1 個封包", "label-sm")}
    <rect x="45" y="230" width="280" height="45" rx="4" fill="#2D3748"/>
    <text x="55" y="250" font-family="monospace" font-size="11" fill="#68D391">event: tool_result</text>
    <text x="55" y="268" font-family="monospace" font-size="11" fill="#63B3ED">data: {{"name":"sear</text>

    <!-- Chunk 2 -->
    {rounded_rect(380, 200, 310, 90, "#FFF5F5", "#E53E3E", 6)}
    {text_left(395, 220, "📦 第 2 個封包", "label-sm")}
    <rect x="395" y="230" width="280" height="45" rx="4" fill="#2D3748"/>
    <text x="405" y="250" font-family="monospace" font-size="11" fill="#63B3ED">ch_food","data":...}}}}</text>
    <text x="405" y="268" font-family="monospace" font-size="11" fill="#718096">（空行）</text>

    <!-- Arrow between -->
    <text x="345" y="255" font-family="{FONT_SVG}" font-size="24" fill="#E53E3E">✂️</text>

    <!-- Explanation -->
    {rounded_rect(30, 305, 660, 80, "#FFFFF0", "#D69E2E", 10)}
    {text_left(50, 330, "💡 如果 eventType 在迴圈內重設：", "label")}
    {text_left(50, 352, "讀完封包 1 → eventType = \"tool_result\" → 進入下一圈 → eventType 被重設為 null", "label-sm")}
    {text_left(50, 374, "讀封包 2 → 有 data 但 eventType 是 null → 不知道是什麼事件 → 丟棄 ❌", "label-sm")}
    '''
    return svg_wrap(content, w, h)


def make_svg_three_fixes():
    """Diagram 10: Summary of all three fixes"""
    w, h = 720, 500
    content = f'''
    {text_center(360, 30, "三個 Bug，三個修復，缺一不可", "title")}

    <!-- Fix 1 -->
    {rounded_rect(20, 60, 680, 120, "#FFF5F5", "#E53E3E", 10)}
    {rounded_rect(30, 65, 30, 30, "#E53E3E", "#E53E3E", 15)}
    <text x="45" y="86" text-anchor="middle" font-family="{FONT_SVG}" font-size="16" fill="white" font-weight="bold">1</text>

    {text_left(75, 82, "Nginx 代理緩衝（Proxy Buffering）", "label")}
    {text_left(75, 102, "問題：中間人把多則消息攢在一起才轉發", "label-sm")}
    {text_left(75, 120, "比喻：秘書把信攢成一疊才送來", "label-xs")}

    <rect x="75" y="130" width="610" height="35" rx="4" fill="#2D3748"/>
    <text x="85" y="152" font-family="monospace" font-size="12" fill="#68D391">修復：proxy_buffering off;  proxy_cache off;  chunked_transfer_encoding off;</text>

    <!-- Fix 2 -->
    {rounded_rect(20, 195, 680, 120, "#FFFFF0", "#D69E2E", 10)}
    {rounded_rect(30, 200, 30, 30, "#D69E2E", "#D69E2E", 15)}
    <text x="45" y="221" text-anchor="middle" font-family="{FONT_SVG}" font-size="16" fill="white" font-weight="bold">2</text>

    {text_left(75, 217, "變數作用域（Variable Scope）", "label")}
    {text_left(75, 237, "問題：紀錄事件類型的變數在迴圈內宣告，每圈被重設", "label-sm")}
    {text_left(75, 255, "比喻：白板每翻一頁就自動擦掉", "label-xs")}

    <rect x="75" y="265" width="610" height="35" rx="4" fill="#2D3748"/>
    <text x="85" y="287" font-family="monospace" font-size="12" fill="#68D391">修復：把 let eventType = null 從 while 迴圈內移到迴圈外</text>

    <!-- Fix 3 -->
    {rounded_rect(20, 330, 680, 120, "#EBF8FF", "#3182CE", 10)}
    {rounded_rect(30, 335, 30, 30, "#3182CE", "#3182CE", 15)}
    <text x="45" y="356" text-anchor="middle" font-family="{FONT_SVG}" font-size="16" fill="white" font-weight="bold">3</text>

    {text_left(75, 352, "陣列參考 vs 快照（Reference vs Snapshot）", "label")}
    {text_left(75, 372, "問題：存進訊息的是「指標」而非「副本」，清空後資料跟著消失", "label-sm")}
    {text_left(75, 390, "比喻：給別人看你手機而不是拍照傳給他", "label-xs")}

    <rect x="75" y="400" width="610" height="35" rx="4" fill="#2D3748"/>
    <text x="85" y="422" font-family="monospace" font-size="12" fill="#68D391">修復：const snapshot = [...turnToolResults] 先複製再存入</text>

    <!-- Bottom note -->
    {rounded_rect(20, 460, 680, 30, "#F0FFF4", "#38A169", 6)}
    {text_center(360, 480, "這三個問題同時存在，只修其中一兩個仍會有機率出錯 — 必須三個全修！", "good")}
    '''
    return svg_wrap(content, w, h)


def make_svg_industry_solutions():
    """Diagram 11: How experts solve SSE problems"""
    w, h = 720, 430
    content = f'''
    {text_center(360, 30, "業界常見的 SSE 問題解決方案", "title")}

    <!-- Solution 1 -->
    {rounded_rect(20, 60, 330, 100, "#EBF8FF", "#3182CE", 8)}
    {text_left(35, 85, "🔧 方案一：EventSource API", "label")}
    {text_left(35, 108, "瀏覽器內建的 SSE 接收器", "label-sm")}
    {text_left(35, 128, "自動處理 event/data 配對", "label-sm")}
    {text_left(35, 148, "⚠️ 不支援 POST 和自訂 Header", "label-xs")}

    <!-- Solution 2 -->
    {rounded_rect(370, 60, 330, 100, "#F0FFF4", "#38A169", 8)}
    {text_left(385, 85, "🔧 方案二：第三方 SSE 函式庫", "label")}
    {text_left(385, 108, "如 eventsource-parser", "label-sm")}
    {text_left(385, 128, "專門解析 SSE 格式，穩定可靠", "label-sm")}
    {text_left(385, 148, "需要額外安裝套件", "label-xs")}

    <!-- Solution 3 -->
    {rounded_rect(20, 175, 330, 100, "#FEFCBF", "#D69E2E", 8)}
    {text_left(35, 200, "🔧 方案三：WebSocket", "label")}
    {text_left(35, 223, "完全雙向的即時通信", "label-sm")}
    {text_left(35, 243, "比 SSE 更強大但更複雜", "label-sm")}
    {text_left(35, 263, "⚠️ 需要重寫整個通信架構", "label-xs")}

    <!-- Our solution -->
    {rounded_rect(370, 175, 330, 100, "#FAF5FF", "#805AD5", 8)}
    {text_left(385, 200, "⭐ 我們的方案：手動 fetch + 修 Bug", "label")}
    {text_left(385, 223, "保留 POST + JWT 認證支援", "label-sm")}
    {text_left(385, 243, "修正 3 個根本原因即可", "label-sm")}
    {text_left(385, 263, "不需額外套件，最小改動", "label-xs")}

    <!-- Comparison -->
    {rounded_rect(20, 290, 680, 125, "#F7FAFC", "#A0AEC0", 10)}
    {text_left(40, 315, "📊 方案比較", "label")}
    {text_left(40, 340, "EventSource：簡單但功能少（不能 POST）→ 不適合我們的 AI 顧問", "label-sm")}
    {text_left(40, 360, "第三方套件：穩定但增加依賴 → 適合大型專案", "label-sm")}
    {text_left(40, 380, "WebSocket：強大但改動大 → 過度設計（殺雞用牛刀）", "label-sm")}
    {text_left(40, 400, "修 Bug：最小改動、不增依賴、針對根因 → 我們選這個 ✅", "good")}
    '''
    return svg_wrap(content, w, h)


# ─────────────────────── Word Document Builder ───────────────────────

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    shading.append(shd)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Noto Sans CJK TC"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "Noto Sans CJK TC")
    return h

def add_para(doc, text, bold=False, italic=False, size=11, color=None, space_after=6, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Noto Sans CJK TC"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Noto Sans CJK TC")
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    return p

def add_image(doc, png_stream, width=Inches(6.2)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(png_stream, width=width)
    return p

def add_tip_box(doc, title, content):
    """Simulate a tip box with a table"""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, "F0F4FF")
    p = cell.paragraphs[0]
    run = p.add_run(f"💡 {title}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Noto Sans CJK TC"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Noto Sans CJK TC")
    p2 = cell.add_paragraph()
    run2 = p2.add_run(content)
    run2.font.size = Pt(10)
    run2.font.name = "Noto Sans CJK TC"
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), "Noto Sans CJK TC")
    doc.add_paragraph()  # spacer

def add_warn_box(doc, content):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, "FFF5F5")
    p = cell.paragraphs[0]
    run = p.add_run(f"⚠️ {content}")
    run.font.size = Pt(10)
    run.font.name = "Noto Sans CJK TC"
    run.font.color.rgb = RGBColor(0xC5, 0x30, 0x30)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Noto Sans CJK TC")
    doc.add_paragraph()


def build_document():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # ====== TITLE PAGE ======
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("當網頁「看不到」伺服器推來的資料")
    run.font.size = Pt(28)
    run.font.name = "Noto Sans CJK TC"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Noto Sans CJK TC")
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("一次 SSE 串流 Bug 的完整解析與修復")
    run.font.size = Pt(16)
    run.font.name = "Noto Sans CJK TC"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Noto Sans CJK TC")
    run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    for _ in range(2):
        doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("深入淺出 × 圖文並茂 × 從零開始理解")
    run.font.size = Pt(12)
    run.font.name = "Noto Sans CJK TC"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Noto Sans CJK TC")
    run.italic = True
    run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("2026 年 5 月")
    run.font.size = Pt(11)
    run.font.name = "Noto Sans CJK TC"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Noto Sans CJK TC")
    run.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)

    doc.add_page_break()

    # ====== TABLE OF CONTENTS ======
    add_heading(doc, "目錄", 1)
    toc_items = [
        "第 1 章　故事的開始：同一個網站，兩種結果",
        "第 2 章　網頁怎麼跟伺服器對話？",
        "第 3 章　SSE：讓伺服器「主動」說話",
        "第 4 章　中間人的角色：反向代理",
        "第 5 章　Bug #1 — 中間人把消息攢在手裡",
        "第 6 章　Bug #2 — 被自動擦掉的白板",
        "第 7 章　Bug #3 — 照片 vs 手指",
        "第 8 章　為什麼「在家」沒問題，「出門」就壞了？",
        "第 9 章　業界怎麼解決？我們怎麼解決？",
        "第 10 章　三個 Bug 的完整修復",
        "第 11 章　總結：一個 Bug 教會我們的事",
        "附錄 A　專有名詞對照表",
    ]
    for item in toc_items:
        add_para(doc, item, size=11, space_after=4)
    doc.add_page_break()

    # ====== CHAPTER 1 ======
    add_heading(doc, "第 1 章　故事的開始：同一個網站，兩種結果", 1)

    add_para(doc, "我們開發了一個「AI 營養顧問」網站。你可以問它：「雞胸肉有什麼營養？」，它會幫你查資料，然後把結果用漂亮的表格顯示出來。")
    add_para(doc, "在開發的過程中，我們遇到了一個詭異的現象：")

    add_para(doc, "✅ 在自己電腦上測試（輸入 localhost:3001）→ 表格正常顯示", bold=True, color=(0x38, 0xA1, 0x69))
    add_para(doc, "❌ 用外部網址測試（從手機或其他電腦連進來）→ 只有文字，表格消失了！", bold=True, color=(0xE5, 0x3E, 0x3E))

    add_para(doc, "同樣的程式碼、同樣的問題、同樣的伺服器，為什麼結果不一樣？")
    add_para(doc, "這個問題背後，其實牽涉到三個完全不同的技術層面。要理解它們，我們需要先從最基本的概念開始。")

    add_tip_box(doc, "閱讀指引",
        "本文會從日常生活的比喻開始，逐步引入技術細節。遇到專有名詞時，我們會先用白話解釋，"
        "再標上英文術語。你可以先跳過英文，等整體概念理解後再回來看。")

    doc.add_page_break()

    # ====== CHAPTER 2 ======
    add_heading(doc, "第 2 章　網頁怎麼跟伺服器對話？", 1)

    add_para(doc, "在講 Bug 之前，我們得先了解：當你打開一個網頁，背後發生了什麼事？")

    add_heading(doc, "2.1 一問一答的模式", 2)
    add_para(doc, "想像你走進一家餐廳。你的互動方式是：")
    add_para(doc, "1. 你看菜單，決定要點什麼（你 = 瀏覽器）")
    add_para(doc, "2. 你跟服務生說：「我要一份牛排」（發出請求 = Request）")
    add_para(doc, "3. 廚房做好後，服務生端過來（收到回應 = Response）")
    add_para(doc, "這就是網頁通信最基本的模式，叫做 HTTP（HyperText Transfer Protocol，超文本傳輸協議）。")

    add_tip_box(doc, "HTTP 是什麼？",
        "HTTP 就像是「網路上的共同語言」。你的瀏覽器（Chrome、Safari⋯）和伺服器"
        "都懂這個語言，所以能互相溝通。每次你打開網頁、點連結、送出表單，背後都是 HTTP 在運作。")

    # Insert diagram
    print("  Generating diagram: HTTP Basic...")
    svg = make_svg_http_basic()
    add_image(doc, svg_to_png(svg))

    add_heading(doc, "2.2 這個模式的限制", 2)
    add_para(doc, "傳統 HTTP 的問題是：伺服器不能「主動」跟你說話。")
    add_para(doc, "回到餐廳的比喻：如果你想知道「我的牛排做到哪了？」，你必須自己再舉手問一次。廚房不會主動跑來告訴你「正在煎第二面了喔」。")
    add_para(doc, "但我們的 AI 顧問需要的是「即時更新」——AI 在思考的時候，你要看到「思考中⋯」；AI 在查資料的時候，你要看到「正在搜尋食物⋯」；查完了，表格要立刻出現。")
    add_para(doc, "這就需要一種新的通信方式⋯⋯")

    doc.add_page_break()

    # ====== CHAPTER 3 ======
    add_heading(doc, "第 3 章　SSE：讓伺服器「主動」說話", 1)

    add_heading(doc, "3.1 從「一直問」到「主動通知」", 2)
    add_para(doc, "面對「即時更新」的需求，有兩種做法：")

    add_para(doc, "做法 A — 一直問（Polling）", bold=True)
    add_para(doc, "就像你每 5 秒鐘打一次電話問外送：「到了嗎？」「到了嗎？」「到了嗎？」")
    add_para(doc, "很煩，也很浪費電話費（網路流量）。")

    add_para(doc, "做法 B — 讓對方主動通知你（SSE）", bold=True)
    add_para(doc, "就像你開著外送 App，它會自動推播通知給你：「已接單」→「外送員出發」→「快到了」→「已送達」。")
    add_para(doc, "這就是 SSE 的概念。")

    print("  Generating diagram: SSE vs Polling...")
    svg = make_svg_sse_vs_polling()
    add_image(doc, svg_to_png(svg))

    add_heading(doc, "3.2 SSE 事件的格式", 2)
    add_para(doc, "SSE 的全名是 Server-Sent Events（伺服器傳送事件）。它的格式非常簡單，就像一張明信片，有兩個部分：")
    add_para(doc, "第一行：event: ○○○  →  這是什麼類型的消息（例如「查到食物資料了」）")
    add_para(doc, "第二行：data: { ... }   →  具體內容（JSON 格式的資料）")
    add_para(doc, "最後一個空行  →  代表「這則消息結束」")

    add_para(doc, "我們的 AI 顧問系統會發送 9 種不同類型的事件。其中最關鍵的是 tool_result，它帶著查詢結果的表格資料。如果這個事件漏掉了，表格就顯示不出來。", bold=True)

    print("  Generating diagram: SSE Format...")
    svg = make_svg_sse_format()
    add_image(doc, svg_to_png(svg))

    add_heading(doc, "3.3 瀏覽器怎麼接收 SSE？", 2)
    add_para(doc, "瀏覽器收到 SSE 資料後，需要用程式碼來解析。簡化的邏輯是：")
    add_para(doc, "1. 從伺服器讀一段資料")
    add_para(doc, "2. 把資料按行拆開")
    add_para(doc, "3. 如果看到 event: 開頭 → 記住事件類型")
    add_para(doc, "4. 如果看到 data: 開頭 → 根據事件類型來處理資料")
    add_para(doc, "5. 回到第 1 步，繼續讀下一段")
    add_para(doc, "看起來很簡單對吧？但問題就藏在「讀一段資料」這個步驟裡⋯⋯")

    doc.add_page_break()

    # ====== CHAPTER 4 ======
    add_heading(doc, "第 4 章　中間人的角色：反向代理", 1)

    add_heading(doc, "4.1 為什麼需要中間人？", 2)
    add_para(doc, "在現實生活中，大公司的 CEO 不會直接接你的電話。你得先打給前台，前台確認你的身份後，再把電話轉給秘書，秘書再轉給 CEO。")
    add_para(doc, "網站也一樣。出於安全和管理的考量，使用者的請求通常不會直接到達後端程式，而是先經過一個或多個「中間人」——技術上叫做「反向代理」（Reverse Proxy）。")

    add_tip_box(doc, "反向代理（Reverse Proxy）是什麼？",
        "「代理」就是「代替你做事的人」。「反向」代理是「代替伺服器接客」的角色。\n"
        "它站在伺服器前面，幫忙：1) 分配流量 2) 提供安全防護 3) 處理加密連線。\n"
        "最常用的反向代理軟體叫做 Nginx（唸作 engine-x）。")

    add_heading(doc, "4.2 我們系統的傳遞鏈", 2)
    add_para(doc, "我們的系統有兩種存取路徑，這就是問題的關鍵：")

    print("  Generating diagram: Proxy Chain...")
    svg = make_svg_proxy_chain()
    add_image(doc, svg_to_png(svg))

    add_para(doc, "在家測試時，資料只經過 1 個中間人（前端 Nginx），路徑短、速度快。")
    add_para(doc, "從外面連進來時，資料要經過 3 個中間人（外部 Nginx → 網路 → 前端 Nginx），每一層都可能對資料動手腳。")
    add_para(doc, "這就像傳話遊戲：傳越多人，訊息越容易走樣或漏掉。", bold=True)

    doc.add_page_break()

    # ====== CHAPTER 5 ======
    add_heading(doc, "第 5 章　Bug #1 — 中間人把消息攢在手裡", 1)
    add_para(doc, "現在我們來看第一個 Bug。", size=12)

    add_heading(doc, "5.1 問題描述", 2)
    add_para(doc, "Nginx 有一個預設行為，叫做「代理緩衝」（Proxy Buffering）。")
    add_para(doc, "什麼意思呢？想像一個秘書，她收到老闆寄來的信後，不會每封都立刻轉交給你。她會先攢在手裡，等收到好幾封後，一次全部交給你。")
    add_para(doc, "對普通的網頁請求來說，這樣做是好的——因為一次性傳送比較有效率。")
    add_para(doc, "但對 SSE 來說，這是災難性的！", bold=True, color=(0xE5, 0x3E, 0x3E))
    add_para(doc, "因為 SSE 的價值就在於「即時」。如果消息被攢在一起，「思考中」和「查到結果了」可能同時到達，瀏覽器就沒辦法依照正確的順序處理它們。")

    print("  Generating diagram: Buffering Problem...")
    svg = make_svg_buffering_problem()
    add_image(doc, svg_to_png(svg))

    add_heading(doc, "5.2 為什麼在本機沒問題？", 2)
    add_para(doc, "在本機測試時，資料從後端到前端 Nginx 的距離極短（都在同一台電腦裡），速度快到 Nginx 來不及攢——每收到一則消息就立刻轉發了。")
    add_para(doc, "但從外部連線時，多了一層外部 Nginx，這層 Nginx 的緩衝行為就可能把多則 SSE 事件打包在一起送出。")

    add_heading(doc, "5.3 修復方法", 2)
    add_para(doc, "在 Nginx 的設定檔中，對 SSE 使用的路徑加上三行指令：")
    add_para(doc, "proxy_buffering off;    → 關閉代理緩衝", size=10, color=(0x38, 0xA1, 0x69))
    add_para(doc, "proxy_cache off;        → 關閉代理快取", size=10, color=(0x38, 0xA1, 0x69))
    add_para(doc, "chunked_transfer_encoding off;  → 關閉分塊傳輸編碼", size=10, color=(0x38, 0xA1, 0x69))
    add_para(doc, "這三行告訴 Nginx：「收到什麼就立刻轉發，不要攢、不要存、不要重新包裝。」")

    add_tip_box(doc, "延伸知識：Chunked Transfer Encoding",
        "這是 HTTP 的一種傳輸方式，允許伺服器把一大段資料切成小塊（chunk）分批送出。\n"
        "Nginx 預設會把 SSE 事件「重新打包」成 chunk，可能改變原本的事件邊界。\n"
        "關閉它，就能確保 SSE 事件保持原樣通過。")

    doc.add_page_break()

    # ====== CHAPTER 6 ======
    add_heading(doc, "第 6 章　Bug #2 — 被自動擦掉的白板", 1)
    add_para(doc, "第二個 Bug 藏在 JavaScript 程式碼裡。", size=12)

    add_heading(doc, "6.1 什麼是「變數作用域」？", 2)
    add_para(doc, "在程式設計中，「變數」就像一個有名字的盒子，你可以把東西放進去。")
    add_para(doc, "「作用域」（Scope）決定了這個盒子的「有效範圍」——在什麼地方可以使用它，什麼時候會自動消失。")

    add_para(doc, "舉個生活例子：", bold=True)
    add_para(doc, "你在上課時用白板寫了一個公式。如果白板在「教室裡」，那整堂課你都能看到它。但如果白板在「你的筆記本某一頁」上，翻到下一頁就看不到了。")
    add_para(doc, "在 JavaScript（一種程式語言）中，用 let 宣告的變數如果寫在迴圈「裡面」，每次迴圈重新開始時，這個變數就會被重設為初始值——就像白板被自動擦掉。")

    add_heading(doc, "6.2 我們的程式碼出了什麼問題？", 2)
    add_para(doc, "我們的 SSE 解析器有一個 while 迴圈（不斷重複的結構），用來持續讀取伺服器送來的資料。程式碼裡有一個變數叫 eventType，負責記住「目前收到的是什麼類型的事件」。")

    add_para(doc, "錯誤的寫法：", bold=True, color=(0xE5, 0x3E, 0x3E))
    add_para(doc, "let eventType = null 寫在 while 迴圈「裡面」。結果每次迴圈重新開始，eventType 就被重設為 null（空白）。")

    add_para(doc, "正確的寫法：", bold=True, color=(0x38, 0xA1, 0x69))
    add_para(doc, "let eventType = null 寫在 while 迴圈「外面」。這樣不管迴圈跑幾圈，eventType 的值都會被保留。")

    print("  Generating diagram: Scope Bug...")
    svg = make_svg_scope_bug()
    add_image(doc, svg_to_png(svg))

    add_heading(doc, "6.3 這跟 TCP 切割有什麼關係？", 2)
    add_para(doc, "還記得第 4 章提到的嗎？資料經過網路傳輸時，可能被切成好幾個封包。")
    add_para(doc, "想像伺服器發出了這樣的 SSE 事件：")
    add_para(doc, "event: tool_result", size=10)
    add_para(doc, "data: {\"name\":\"search_food\", ...}", size=10)
    add_para(doc, "如果網路剛好在 event: 和 data: 之間把封包切開：")
    add_para(doc, "• 第一個封包帶著 event: tool_result")
    add_para(doc, "• 第二個封包帶著 data: {...}")
    add_para(doc, "在本機，因為傳輸極快，這兩行幾乎一定在同一個封包裡。但經過外部網路和多層代理後，它們可能被分到不同的封包中。")
    add_para(doc, "如果 eventType 在迴圈內重設，讀完第一個封包時 eventType 被設為 \"tool_result\"，但進入下一圈讀第二個封包時，eventType 已經被重設為 null 了！程式就不知道 data 屬於哪個事件，直接丟棄。", bold=True)

    print("  Generating diagram: TCP Chunking...")
    svg = make_svg_tcp_chunking()
    add_image(doc, svg_to_png(svg))

    add_warn_box(doc, "這就是為什麼本機完全正常、外部才出問題——不是程式碼「有時候」會錯，而是只有在資料被切割的情況下才會出錯。本機的快速傳輸恰好避開了這個問題。")

    doc.add_page_break()

    # ====== CHAPTER 7 ======
    add_heading(doc, "第 7 章　Bug #3 — 照片 vs 手指", 1)
    add_para(doc, "第三個 Bug 與 React（我們使用的前端框架）的狀態管理有關。", size=12)

    add_heading(doc, "7.1 「參考」和「快照」的差別", 2)
    add_para(doc, "想像這個情境：", bold=True)
    add_para(doc, "你的朋友手機上顯示著一個電話號碼。你要把這個號碼傳給另一個朋友。你有兩種做法：")
    add_para(doc, "A. 你說：「去看他手機」（參考 Reference）—— 如果他換了號碼，別人看到的就是新號碼")
    add_para(doc, "B. 你先拍張照片傳過去（快照 Snapshot）—— 就算他換號碼，照片裡還是原來那個")
    add_para(doc, "在程式設計中，JavaScript 的「陣列」（Array，一種存放多筆資料的容器）預設是用「參考」的方式傳遞的。也就是說，當你把一個陣列放進另一個地方，實際上放的是一根「手指」指向那個陣列，而不是把陣列複製一份。")

    add_heading(doc, "7.2 我們的程式碼出了什麼問題？", 2)
    add_para(doc, "我們的程式在 AI 查完資料後，需要：")
    add_para(doc, "步驟 1：把查到的表格資料存到一個叫 turnToolResults 的陣列中")
    add_para(doc, "步驟 2：把這個陣列放進「訊息」裡，準備顯示在畫面上")
    add_para(doc, "步驟 3：清空 turnToolResults，準備接收下一次查詢的結果")
    add_para(doc, "問題來了：步驟 2 存進去的是「手指」（參考），不是「照片」（副本）。所以當步驟 3 清空陣列時，訊息裡存的「手指」指到的地方也變空了——表格資料消失！")

    print("  Generating diagram: Snapshot Bug...")
    svg = make_svg_snapshot_bug()
    add_image(doc, svg_to_png(svg))

    add_heading(doc, "7.3 修復方法", 2)
    add_para(doc, "在步驟 2 之前，先「拍照」——用 JavaScript 的展開運算子 [...] 複製一份：")
    add_para(doc, "const snapshotResults = [...turnToolResults];  // 拍照！", size=10, color=(0x38, 0xA1, 0x69))
    add_para(doc, "然後把 snapshotResults（照片）存進訊息。這樣就算原本的陣列被清空，照片裡的資料還在。")

    add_tip_box(doc, "展開運算子 [...] 是什麼？",
        "三個點 ... 叫做「展開運算子」（Spread Operator）。它的功能是「把陣列裡的東西一個一個拿出來」。\n"
        "[...原陣列] 就是「建立一個新陣列，把原陣列的東西一個一個放進去」= 複製。\n"
        "就像你拍照：新照片跟本人長一樣，但是獨立的副本。")

    doc.add_page_break()

    # ====== CHAPTER 8 ======
    add_heading(doc, "第 8 章　為什麼「在家」沒問題，「出門」就壞了？", 1)

    add_para(doc, "到這裡，你可能會想：為什麼這三個 Bug 在本機測試時完全沒有問題？")
    add_para(doc, "答案是：它們都跟「網路傳輸路徑」有關。", bold=True)

    print("  Generating diagram: Localhost vs External...")
    svg = make_svg_localhost_vs_external()
    add_image(doc, svg_to_png(svg))

    add_heading(doc, "8.1 逐個分析", 2)

    add_para(doc, "Bug #1（Nginx 緩衝）：", bold=True)
    add_para(doc, "在本機，資料傳得太快，Nginx 來不及緩衝就已經轉發了。外部連線多了一層外部 Nginx + 網路延遲，緩衝行為才有時間「生效」。")

    add_para(doc, "Bug #2（變數作用域）：", bold=True)
    add_para(doc, "在本機，event: 和 data: 幾乎 100% 在同一個 TCP 封包裡，所以在同一次迴圈中就能配對成功。外部連線經過多層代理和網路傳輸，封包切割的機率大幅增加。")

    add_para(doc, "Bug #3（參考 vs 快照）：", bold=True)
    add_para(doc, "在本機，因為 Bug #1 和 #2 沒有觸發，工具結果事件都能正確收到並處理，所以就算用的是「參考」，時序上也沒問題。外部連線如果因為前兩個 Bug 導致事件遺失或錯亂，就可能觸發 React 狀態更新的競爭條件（Race Condition）。")

    add_tip_box(doc, "競爭條件（Race Condition）是什麼？",
        "想像兩個人同時要修改同一份文件：A 在第 3 行加了一句話，B 在第 3 行刪了一句話。\n"
        "最後文件變成什麼樣，取決於誰的動作「先到」。這種「結果取決於時機」的問題，就叫競爭條件。\n"
        "在我們的案例中，「存入訊息」和「清空陣列」如果時機不對，表格資料就會消失。")

    add_heading(doc, "8.2 重要觀念", 2)
    add_para(doc, "「在本機能跑不代表沒有 Bug」——這是軟體開發中非常重要的教訓。", bold=True)
    add_para(doc, "本機測試時，很多問題被「理想的環境」給掩蓋了。只有在接近真實的使用環境（不同網路、不同設備、不同頻寬）下測試，才能把隱藏的問題逼出來。")

    doc.add_page_break()

    # ====== CHAPTER 9 ======
    add_heading(doc, "第 9 章　業界怎麼解決？我們怎麼解決？", 1)

    add_para(doc, "SSE 解析問題並不罕見。業界有幾種常見的解決方案：")

    print("  Generating diagram: Industry Solutions...")
    svg = make_svg_industry_solutions()
    add_image(doc, svg_to_png(svg))

    add_heading(doc, "9.1 方案一：使用 EventSource API", 2)
    add_para(doc, "瀏覽器內建了一個叫 EventSource 的工具，專門用來接收 SSE。它會自動幫你處理事件配對的問題（event: 和 data: 的配對）。")
    add_para(doc, "但 EventSource 有一個致命缺點：它只支援 GET 請求。我們的 AI 顧問需要用 POST 請求（因為要傳送對話歷史等大量資料），而且需要附帶 JWT 認證 Header（身份驗證的令牌）。所以 EventSource 不適用。")

    add_heading(doc, "9.2 方案二：第三方 SSE 解析函式庫", 2)
    add_para(doc, "社群開發了一些專門解析 SSE 格式的工具包，例如 eventsource-parser。它們處理過各種邊界情況（edge case），穩定可靠。")
    add_para(doc, "缺點是增加了一個外部依賴（dependency）。對大型專案來說不是問題，但對我們的小專案來說，為了一個功能引入一整個套件有點殺雞用牛刀。")

    add_heading(doc, "9.3 方案三：改用 WebSocket", 2)
    add_para(doc, "WebSocket 是一種比 SSE 更強大的即時通信技術，支援雙向通信。但它需要完全不同的伺服器架構和前端程式碼。")
    add_para(doc, "對我們來說，SSE 已經滿足需求（伺服器→瀏覽器的單向推送），換成 WebSocket 是過度設計。")

    add_heading(doc, "9.4 我們的選擇：修正根本原因", 2)
    add_para(doc, "我們選擇保留現有的手動 fetch + SSE 解析架構，針對三個根本原因進行修復。理由：")
    add_para(doc, "1. 最小改動原則：只改出問題的地方，不動其他正常運作的部分")
    add_para(doc, "2. 不增加外部依賴：避免引入新的維護負擔")
    add_para(doc, "3. 針對根因：理解為什麼出錯比換一套方案更有價值")

    doc.add_page_break()

    # ====== CHAPTER 10 ======
    add_heading(doc, "第 10 章　三個 Bug 的完整修復", 1)

    print("  Generating diagram: Three Fixes...")
    svg = make_svg_three_fixes()
    add_image(doc, svg_to_png(svg))

    add_heading(doc, "10.1 修復 #1：關閉 Nginx 代理緩衝", 2)
    add_para(doc, "修改檔案：前端容器的 Dockerfile（Nginx 設定）")
    add_para(doc, "在 /api/ 路徑的 location 區塊加入：")
    add_para(doc, "proxy_buffering off;", size=10, bold=True, color=(0x38, 0xA1, 0x69))
    add_para(doc, "proxy_cache off;", size=10, bold=True, color=(0x38, 0xA1, 0x69))
    add_para(doc, "proxy_set_header Connection \"\";", size=10, bold=True, color=(0x38, 0xA1, 0x69))
    add_para(doc, "chunked_transfer_encoding off;", size=10, bold=True, color=(0x38, 0xA1, 0x69))
    add_para(doc, "效果：確保 SSE 事件收到後立即轉發，不緩衝、不快取、不重新打包。")

    add_heading(doc, "10.2 修復 #2：移動變數宣告位置", 2)
    add_para(doc, "修改檔案：advisorApi.js（SSE 解析器）")
    add_para(doc, "把 let eventType = null; 從 while 迴圈裡面移到外面：")
    add_para(doc, "修改前（錯誤）：變數在迴圈第 45 行，每次迴圈重設為 null", size=10, color=(0xE5, 0x3E, 0x3E))
    add_para(doc, "修改後（正確）：變數在迴圈前第 36 行，整個 SSE 接收過程中保持有效", size=10, color=(0x38, 0xA1, 0x69))
    add_para(doc, "效果：即使 event: 和 data: 跨越不同的 TCP 封包，eventType 的值也不會遺失。")

    add_heading(doc, "10.3 修復 #3：用快照取代參考", 2)
    add_para(doc, "修改檔案：AdvisorPage.jsx（前端頁面元件）")
    add_para(doc, "在 onReply 回呼函式中，存入訊息前先複製一份：")
    add_para(doc, "const snapshotResults = [...turnToolResults];", size=10, bold=True, color=(0x38, 0xA1, 0x69))
    add_para(doc, "const snapshotTools = [...turnTools];", size=10, bold=True, color=(0x38, 0xA1, 0x69))
    add_para(doc, "效果：訊息中存的是獨立的副本，不受之後清空陣列的影響。")

    add_warn_box(doc, "重要：這三個修復缺一不可。即使只修了其中兩個，在特定的網路條件下仍然可能出問題。必須三個全部修復才能完全解決。")

    doc.add_page_break()

    # ====== CHAPTER 11 ======
    add_heading(doc, "第 11 章　總結：一個 Bug 教會我們的事", 1)

    add_heading(doc, "11.1 技術面的啟示", 2)
    add_para(doc, "1. 分層架構的複雜性", bold=True)
    add_para(doc, "現代網頁應用通常有多層架構（瀏覽器 → 代理 → 後端）。每一層都有自己的預設行為和特性，必須確保每一層的設定都適合你的通信模式。")
    add_para(doc, "2. 環境差異", bold=True)
    add_para(doc, "本機開發環境和正式上線環境的差異可能導致同一段程式碼表現不同。這也是為什麼需要在盡量接近正式環境的條件下測試。")
    add_para(doc, "3. 多個小 Bug 的複合效應", bold=True)
    add_para(doc, "單獨來看，三個 Bug 中的每一個都可能不會造成問題。但它們同時存在時，就會在特定條件下引發顯著的故障。這種「複合 Bug」在實際開發中非常常見，也特別難以排查。")

    add_heading(doc, "11.2 方法論的啟示", 2)
    add_para(doc, "1. 追根溯源", bold=True)
    add_para(doc, "遇到 Bug 時，不要急著換方案或加 workaround。先理解「為什麼」會出問題，找到根本原因後再修復。")
    add_para(doc, "2. 最小改動原則", bold=True)
    add_para(doc, "修復 Bug 時，盡量只改出問題的地方，不要趁機大改架構。一次改太多東西，很難確定哪個改動真正修好了問題。")
    add_para(doc, "3. 從使用者的角度測試", bold=True)
    add_para(doc, "開發者習慣在本機測試，但使用者是從外部連進來的。記得站在使用者的角度，用他們會用的方式來測試。")

    add_heading(doc, "11.3 結語", 2)
    add_para(doc, "軟體開發中，看似簡單的問題背後往往隱藏著複雜的原因。「表格不見了」這個現象，牽涉到網路代理、程式語言特性、前端框架的狀態管理三個完全不同的領域。")
    add_para(doc, "能把這三個層面串在一起思考，找出它們的交互作用，就是從「寫程式的人」成長為「解決問題的工程師」的關鍵一步。", bold=True)

    doc.add_page_break()

    # ====== APPENDIX ======
    add_heading(doc, "附錄 A　專有名詞對照表", 1)

    terms = [
        ("HTTP", "HyperText Transfer Protocol", "超文本傳輸協議，瀏覽器和伺服器之間的通信規則"),
        ("SSE", "Server-Sent Events", "伺服器傳送事件，允許伺服器主動推送資料給瀏覽器"),
        ("Polling", "輪詢", "瀏覽器定期向伺服器查詢是否有新資料"),
        ("Reverse Proxy", "反向代理", "站在伺服器前面，代替伺服器接收請求的中間人"),
        ("Nginx", "Engine-X", "最常用的反向代理 / 網頁伺服器軟體"),
        ("Proxy Buffering", "代理緩衝", "代理伺服器先收集完整回應再一次轉發的預設行為"),
        ("TCP", "Transmission Control Protocol", "傳輸控制協議，確保網路資料可靠傳輸的底層規則"),
        ("Packet", "封包", "資料在網路上傳輸時的最小單位"),
        ("Variable Scope", "變數作用域", "變數在程式中可被存取的有效範圍"),
        ("Reference", "參考 / 參照", "指向原始資料的指標，不是獨立副本"),
        ("Snapshot", "快照", "在某個時間點對資料做的獨立副本"),
        ("Spread Operator", "展開運算子（...）", "JavaScript 語法，用於複製陣列或物件"),
        ("Race Condition", "競爭條件", "多個操作同時進行，結果取決於執行順序的問題"),
        ("JSON", "JavaScript Object Notation", "一種輕量的資料交換格式"),
        ("JWT", "JSON Web Token", "用於使用者身份驗證的加密令牌"),
        ("Docker", "容器化平台", "把應用程式和環境打包成獨立「容器」的工具"),
        ("React", "前端框架", "由 Meta 開發的 JavaScript UI 函式庫"),
        ("WebSocket", "雙向即時通信協議", "比 SSE 更強大，支援瀏覽器和伺服器雙向通信"),
    ]

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, txt in enumerate(["英文術語", "中文名稱", "簡單說明"]):
        hdr[i].text = txt
        for p in hdr[i].paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(10)
            p.runs[0].font.name = "Noto Sans CJK TC"
            p.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), "Noto Sans CJK TC")
        set_cell_shading(hdr[i], "EBF8FF")

    for eng, zh, desc in terms:
        row = table.add_row().cells
        for i, txt in enumerate([eng, zh, desc]):
            row[i].text = txt
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.name = "Noto Sans CJK TC"
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), "Noto Sans CJK TC")

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(4)
        row.cells[2].width = Cm(9.5)

    doc.add_paragraph()
    add_para(doc, "— 全文完 —", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, color=(0xA0, 0xAE, 0xC0))

    return doc


# ─────────────────────── Main ───────────────────────

if __name__ == "__main__":
    print("Starting SSE Tutorial Document Generation...")
    print("=" * 50)

    # Test SVG text rendering first
    print("Testing SVG text rendering...")
    test_svg = svg_wrap(f'{text_center(100, 25, "測試中文 Test English", "label")}', 200, 50)
    test_png = cairosvg.svg2png(bytestring=test_svg.encode("utf-8"), scale=2)
    assert len(test_png) > 1000, f"SVG text rendering failed! PNG too small: {len(test_png)} bytes"
    print(f"  SVG text rendering OK ({len(test_png)} bytes)")

    print("\nBuilding document...")
    doc = build_document()

    output_path = os.path.join(OUT_DIR, "SSE_Bug_教學文件_深入淺出.docx")
    doc.save(output_path)
    print(f"\n{'=' * 50}")
    print(f"Document saved to: {output_path}")
    print(f"File size: {os.path.getsize(output_path) / 1024:.1f} KB")
