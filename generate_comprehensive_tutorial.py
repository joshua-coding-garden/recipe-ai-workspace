#!/usr/bin/env python3
"""
生成 Recipe-AI 系統完整技術教學 Word 文件
16 章 + 4 附錄 + 28 張 SVG 圖表
"""
import io, os, re, textwrap, json
import cairosvg
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

BASE = "/home/aiiauser/JM/0325new"
OUT = os.path.join(BASE, "Recipe-AI_完整技術教學.docx")
FONT_ZH = "Noto Sans CJK TC"
FONT_EN = "Arial, Helvetica, sans-serif"
FONT_SVG = f"{FONT_ZH}, {FONT_EN}"

# ═══════════════════════════════════════════════════════════
# Section A: SVG Helpers
# ═══════════════════════════════════════════════════════════

COMMON_DEFS = '''
<marker id="arr-blue" markerWidth="10" markerHeight="7" refX="9" refY="3.5" orient="auto">
  <polygon points="0 0, 10 3.5, 0 7" fill="#3182CE"/></marker>
<marker id="arr-green" markerWidth="10" markerHeight="7" refX="9" refY="3.5" orient="auto">
  <polygon points="0 0, 10 3.5, 0 7" fill="#38A169"/></marker>
<marker id="arr-red" markerWidth="10" markerHeight="7" refX="9" refY="3.5" orient="auto">
  <polygon points="0 0, 10 3.5, 0 7" fill="#E53E3E"/></marker>
<marker id="arr-orange" markerWidth="10" markerHeight="7" refX="9" refY="3.5" orient="auto">
  <polygon points="0 0, 10 3.5, 0 7" fill="#DD6B20"/></marker>
<marker id="arr-purple" markerWidth="10" markerHeight="7" refX="9" refY="3.5" orient="auto">
  <polygon points="0 0, 10 3.5, 0 7" fill="#805AD5"/></marker>
<marker id="arr-teal" markerWidth="10" markerHeight="7" refX="9" refY="3.5" orient="auto">
  <polygon points="0 0, 10 3.5, 0 7" fill="#319795"/></marker>
<marker id="arr-gray" markerWidth="10" markerHeight="7" refX="9" refY="3.5" orient="auto">
  <polygon points="0 0, 10 3.5, 0 7" fill="#718096"/></marker>
<marker id="arrd-blue" markerWidth="7" markerHeight="10" refX="3.5" refY="9" orient="auto">
  <polygon points="0 0, 7 0, 3.5 10" fill="#3182CE"/></marker>
<marker id="arrd-green" markerWidth="7" markerHeight="10" refX="3.5" refY="9" orient="auto">
  <polygon points="0 0, 7 0, 3.5 10" fill="#38A169"/></marker>
<marker id="arrd-red" markerWidth="7" markerHeight="10" refX="3.5" refY="9" orient="auto">
  <polygon points="0 0, 7 0, 3.5 10" fill="#E53E3E"/></marker>
<marker id="arrd-orange" markerWidth="7" markerHeight="10" refX="3.5" refY="9" orient="auto">
  <polygon points="0 0, 7 0, 3.5 10" fill="#DD6B20"/></marker>
<marker id="arrd-purple" markerWidth="7" markerHeight="10" refX="3.5" refY="9" orient="auto">
  <polygon points="0 0, 7 0, 3.5 10" fill="#805AD5"/></marker>
<marker id="arrd-teal" markerWidth="7" markerHeight="10" refX="3.5" refY="9" orient="auto">
  <polygon points="0 0, 7 0, 3.5 10" fill="#319795"/></marker>
<marker id="arrd-gray" markerWidth="7" markerHeight="10" refX="3.5" refY="9" orient="auto">
  <polygon points="0 0, 7 0, 3.5 10" fill="#718096"/></marker>
'''

def svg_wrap(content, w, h):
    return f'''<svg xmlns="http://www.w3.org/2000/svg" width="{w}" height="{h}" viewBox="0 0 {w} {h}">
<defs>{COMMON_DEFS}
<style>
text {{ font-family: {FONT_SVG}; }}
.title {{ font-size: 18px; font-weight: bold; fill: #1a202c; }}
.subtitle {{ font-size: 15px; font-weight: bold; fill: #2d3748; }}
.label {{ font-size: 13px; fill: #2d3748; }}
.label-bold {{ font-size: 13px; fill: #2d3748; font-weight: bold; }}
.label-sm {{ font-size: 11px; fill: #4a5568; }}
.label-xs {{ font-size: 10px; fill: #718096; }}
.code {{ font-family: monospace; font-size: 11px; fill: #2d3748; }}
.white {{ font-size: 13px; fill: #ffffff; font-weight: bold; }}
.white-sm {{ font-size: 11px; fill: #ffffff; }}
.accent {{ font-size: 13px; fill: #e53e3e; font-weight: bold; }}
.good {{ font-size: 13px; fill: #38a169; font-weight: bold; }}
.num {{ font-size: 22px; fill: #ffffff; font-weight: bold; }}
</style>
</defs>
{content}
</svg>'''

def rr(x, y, w, h, fill="#EBF8FF", stroke="#3182CE", rx=8):
    return f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="{rx}" fill="{fill}" stroke="{stroke}" stroke-width="1.5"/>'

def rr_bold(x, y, w, h, fill="#EBF8FF", stroke="#3182CE", rx=8):
    return f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="{rx}" fill="{fill}" stroke="{stroke}" stroke-width="2.5"/>'

def circle(cx, cy, r, fill="#3182CE"):
    return f'<circle cx="{cx}" cy="{cy}" r="{r}" fill="{fill}"/>'

def tc(x, y, txt, cls="label"):
    return f'<text x="{x}" y="{y}" text-anchor="middle" class="{cls}">{txt}</text>'

def tl(x, y, txt, cls="label"):
    return f'<text x="{x}" y="{y}" class="{cls}">{txt}</text>'

def tr(x, y, txt, cls="label"):
    return f'<text x="{x}" y="{y}" text-anchor="end" class="{cls}">{txt}</text>'

def _color_hex(color):
    return {"blue":"#3182CE","green":"#38A169","red":"#E53E3E","orange":"#DD6B20",
            "purple":"#805AD5","teal":"#319795","gray":"#718096"}.get(color,"#3182CE")

def arr_r(x1, y, x2, color="blue", label="", ly=-12):
    parts = [f'<line x1="{x1}" y1="{y}" x2="{x2}" y2="{y}" stroke="{_color_hex(color)}" stroke-width="2" marker-end="url(#arr-{color})"/>']
    if label:
        parts.append(tc((x1+x2)/2, y+ly, label, "label-sm"))
    return "\n".join(parts)

def arr_d(x, y1, y2, color="blue", label="", lx=10):
    parts = [f'<line x1="{x}" y1="{y1}" x2="{x}" y2="{y2}" stroke="{_color_hex(color)}" stroke-width="2" marker-end="url(#arrd-{color})"/>']
    if label:
        parts.append(tl(x+lx, (y1+y2)/2+4, label, "label-sm"))
    return "\n".join(parts)

def arr_l(x1, y, x2, color="green", label="", ly=15):
    parts = [f'<line x1="{x1}" y1="{y}" x2="{x2}" y2="{y}" stroke="{_color_hex(color)}" stroke-width="2" marker-end="url(#arr-{color})"/>']
    if label:
        parts.append(tc((x1+x2)/2, y+ly, label, "label-sm"))
    return "\n".join(parts)

def dash_d(x, y1, y2, color="#718096"):
    return f'<line x1="{x}" y1="{y1}" x2="{x}" y2="{y2}" stroke="{color}" stroke-width="1.5" stroke-dasharray="6,4"/>'

def dash_r(x1, y, x2, color="#718096"):
    return f'<line x1="{x1}" y1="{y}" x2="{x2}" y2="{y}" stroke="{color}" stroke-width="1.5" stroke-dasharray="6,4"/>'

def arr_u(x, y1, y2, color="blue", label="", lx=10):
    c = _color_hex(color)
    parts = [f'<line x1="{x}" y1="{y1}" x2="{x}" y2="{y2}" stroke="{c}" stroke-width="2" marker-end="url(#arrd-{color})"/>']
    if label:
        parts.append(tl(x+lx, (y1+y2)/2+4, label, "label-sm"))
    return "\n".join(parts)

def svg_to_png(svg_str, scale=2):
    return io.BytesIO(cairosvg.svg2png(bytestring=svg_str.encode("utf-8"), scale=scale))


# ═══════════════════════════════════════════════════════════
# Section B: Word Document Helpers
# ═══════════════════════════════════════════════════════════

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)

style_normal = doc.styles['Normal']
style_normal.font.name = FONT_ZH
style_normal.font.size = Pt(11)
style_normal.paragraph_format.space_after = Pt(6)
style_normal.paragraph_format.line_spacing = 1.4
style_normal.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_ZH)

for lvl, sz, clr in [('Heading 1', 22, '1A365D'), ('Heading 2', 16, '2B6CB0'), ('Heading 3', 13, '2C7A7B')]:
    s = doc.styles[lvl]
    s.font.size = Pt(sz)
    s.font.color.rgb = RGBColor.from_string(clr)
    s.font.bold = True
    s.font.name = FONT_ZH
    s.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_ZH)
    s.paragraph_format.space_before = Pt(20 if sz > 15 else 12)
    s.paragraph_format.space_after = Pt(8)

code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = 'Consolas'
code_style.font.size = Pt(9)
code_style.paragraph_format.space_before = Pt(2)
code_style.paragraph_format.space_after = Pt(2)
code_style.paragraph_format.left_indent = Cm(0.5)
code_style.paragraph_format.element.append(
    code_style.paragraph_format.element.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'F0F0F0'
    })
)

def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
    shading.append(shd)

def h1(text):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.name = FONT_ZH
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_ZH)

def h2(text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.name = FONT_ZH
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_ZH)

def h3(text):
    h = doc.add_heading(text, level=3)
    for run in h.runs:
        run.font.name = FONT_ZH
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_ZH)

def p(text, bold=False, italic=False, size=11, color=None, align=None, space_after=6):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.name = FONT_ZH
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_ZH)
    if bold: run.bold = True
    if italic: run.italic = True
    if color: run.font.color.rgb = RGBColor(*color)
    para.paragraph_format.space_after = Pt(space_after)
    if align: para.alignment = align
    return para

def pb(text): return p(text, bold=True)
def pi(text): return p(text, italic=True, size=10, color=(0x71, 0x80, 0x96))

def code(text):
    for line in text.strip().split('\n'):
        para = doc.add_paragraph(style='CodeBlock')
        run = para.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

def bullet(text):
    para = doc.add_paragraph(style='List Bullet')
    para.text = text
    for run in para.runs:
        run.font.name = FONT_ZH
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_ZH)
    return para

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = hd
        for par in cell.paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in par.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = FONT_ZH
                run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_ZH)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            cell.text = str(val)
            for par in cell.paragraphs:
                for run in par.runs:
                    run.font.size = Pt(10)
                    run.font.name = FONT_ZH
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_ZH)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

def add_image(svg_str, width=Inches(6.2)):
    png = svg_to_png(svg_str)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(png, width=width)
    doc.add_paragraph()

def tip_box(title, content):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, "EBF8FF")
    pp = cell.paragraphs[0]
    run = pp.add_run(f"💡 {title}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = FONT_ZH
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_ZH)
    p2 = cell.add_paragraph()
    run2 = p2.add_run(content)
    run2.font.size = Pt(10)
    run2.font.name = FONT_ZH
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_ZH)
    doc.add_paragraph()

def warn_box(content):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, "FFF5F5")
    pp = cell.paragraphs[0]
    run = pp.add_run(f"⚠️ {content}")
    run.font.size = Pt(10)
    run.font.name = FONT_ZH
    run.font.color.rgb = RGBColor(0xC5, 0x30, 0x30)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_ZH)
    doc.add_paragraph()

def analogy_box(title, content):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, "F0FFF4")
    pp = cell.paragraphs[0]
    run = pp.add_run(f"🍳 {title}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = FONT_ZH
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_ZH)
    p2 = cell.add_paragraph()
    run2 = p2.add_run(content)
    run2.font.size = Pt(10)
    run2.font.name = FONT_ZH
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_ZH)
    doc.add_paragraph()

def prompt_box(title, content):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, "FAF5FF")
    pp = cell.paragraphs[0]
    run = pp.add_run(f"📋 Prompt: {title}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = FONT_ZH
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_ZH)
    for line in content.strip().split('\n'):
        cp = cell.add_paragraph()
        cr = cp.add_run(line)
        cr.font.size = Pt(9)
        cr.font.name = 'Consolas'
    doc.add_paragraph()

def page_break():
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════
# Section C: 28 SVG Diagrams
# ═══════════════════════════════════════════════════════════

def make_svg_01_architecture():
    """圖 1: 七服務架構圖"""
    w, h = 920, 650
    content = f'''
    {tc(460, 30, "Recipe-AI 系統七服務架構圖", "title")}

    <!-- Frontend -->
    {rr_bold(30, 60, 200, 90, "#EBF8FF", "#3182CE")}
    {tc(130, 90, "Frontend", "label-bold")}
    {tc(130, 110, "React + Vite + nginx", "label-sm")}
    {tc(130, 130, "Port 3001", "label-xs")}

    <!-- Control -->
    {rr_bold(310, 60, 200, 90, "#F0FFF4", "#38A169")}
    {tc(410, 90, "Control (核心)", "label-bold")}
    {tc(410, 110, "FastAPI + Advisor", "label-sm")}
    {tc(410, 130, "Port 8000", "label-xs")}

    <!-- Backend -->
    {rr_bold(590, 60, 200, 90, "#FFFAF0", "#DD6B20")}
    {tc(690, 90, "Backend", "label-bold")}
    {tc(690, 110, "FastAPI + asyncpg", "label-sm")}
    {tc(690, 130, "Port 8001", "label-xs")}

    <!-- Arrows top row -->
    {arr_r(230, 100, 308, "blue", "API + SSE")}
    {arr_r(510, 100, 588, "green", "httpx proxy")}

    <!-- PostgreSQL -->
    {rr_bold(590, 210, 200, 80, "#FFF5F5", "#E53E3E")}
    {tc(690, 245, "PostgreSQL", "label-bold")}
    {tc(690, 265, "Port 5432", "label-xs")}

    <!-- Gemma LLM -->
    {rr_bold(310, 210, 200, 80, "#FAF5FF", "#805AD5")}
    {tc(410, 240, "Gemma LLM", "label-bold")}
    {tc(410, 265, "Port 8080 (llama-server)", "label-xs")}

    <!-- GraphRAG -->
    {rr_bold(30, 210, 200, 80, "#E6FFFA", "#319795")}
    {tc(130, 240, "GraphRAG API", "label-bold")}
    {tc(130, 265, "Port 8002", "label-xs")}

    <!-- Health Vector -->
    {rr_bold(30, 340, 200, 80, "#FEFCBF", "#D69E2E")}
    {tc(130, 370, "Health Vector", "label-bold")}
    {tc(130, 395, "Port 8003", "label-xs")}

    <!-- Arrows -->
    {arr_d(690, 150, 208, "orange", "asyncpg")}
    {arr_d(410, 150, 208, "purple", "LLM 呼叫")}
    {arr_r(230, 250, 308, "teal")}
    {arr_d(130, 290, 338, "gray", "向量搜尋")}

    <!-- Docker vs Manual -->
    {rr(30, 460, 860, 50, "#F7FAFC", "#CBD5E0")}
    {tc(460, 490, "Docker 容器：Frontend / Control / Backend / PostgreSQL　　手動進程：Gemma LLM / GraphRAG / Health Vector", "label-sm")}

    <!-- Restaurant analogy -->
    {rr(30, 540, 860, 80, "#F0FFF4", "#38A169")}
    {tl(50, 565, "餐廳比喻：Frontend = 前台（接待客人）　Control = 領班（調度一切）　Backend = 廚房（實際料理）", "label-sm")}
    {tl(50, 590, "PostgreSQL = 食材庫（儲存所有資料）　Gemma = 大廚（思考判斷）　GraphRAG = 圖書館　Health Vector = 營養師資料庫", "label-sm")}
    '''
    return svg_wrap(content, w, h)


def make_svg_02_pipeline():
    """圖 2: 六步驟 Pipeline 全景圖"""
    w, h = 920, 380
    steps = [
        ("Step 1", "Echo", "即時確認", "temp=0.5", "#3182CE"),
        ("Step 2", "Phase1", "意圖判定", "temp=0.1", "#38A169"),
        ("Step 3", "Phase2", "工具調用", "temp=0.1", "#DD6B20"),
        ("Step 4", "Execute", "執行工具", "—", "#E53E3E"),
        ("Step 5", "Analysis", "分析回覆", "temp=0.3", "#805AD5"),
        ("Step 6", "DualMem", "雙記憶", "temp=0.2", "#319795"),
    ]
    parts = [tc(460, 30, "六步驟 LLM Pipeline 全景圖", "title")]
    for i, (step, name, desc, temp, clr) in enumerate(steps):
        x = 30 + i * 148
        parts.append(rr_bold(x, 60, 130, 110, "#FFFFFF", clr))
        parts.append(f'{circle(x+65, 85, 18, clr)}')
        parts.append(tc(x+65, 92, str(i+1), "num"))
        parts.append(tc(x+65, 125, name, "label-bold"))
        parts.append(tc(x+65, 145, desc, "label-sm"))
        parts.append(tc(x+65, 162, temp, "label-xs"))
        if i < 5:
            parts.append(arr_r(x+130, 115, x+148, "gray"))
    parts.append(rr(30, 210, 860, 50, "#EBF8FF", "#3182CE"))
    parts.append(tc(460, 240, "使用者輸入一句話 → 經過 6 個步驟 → 得到最終回覆 + 記憶更新", "label"))
    parts.append(rr(30, 280, 860, 70, "#F0FFF4", "#38A169"))
    parts.append(tl(50, 305, "比喻：你走進餐廳說「我想吃健康的」→ 服務生確認（Echo）→ 領班判斷你要什麼（Phase1）", "label-sm"))
    parts.append(tl(50, 325, "→ 廚師決定用什麼工具和食材（Phase2）→ 開始料理（Execute）→ 上菜時解說（Analysis）→ 記住你的喜好（DualMem）", "label-sm"))
    return svg_wrap("\n".join(parts), w, h)


def make_svg_03_phase1_intent():
    """圖 3: Phase1 意圖決策樹（15 種 intent）"""
    w, h = 900, 720
    intents = [
        ("search_food", "搜尋食物營養"),
        ("browse_food_database", "瀏覽食品資料庫"),
        ("analyze_recipe", "分析食譜營養"),
        ("search_symptom", "症狀搜尋"),
        ("search_graphrag", "學術論文問答"),
        ("nutrient_ranking", "營養素排名"),
        ("get_recipes", "查看已存食譜"),
        ("delete_recipe", "刪除食譜"),
        ("get_calendar", "查看行事曆"),
        ("add_to_calendar", "加入行事曆"),
        ("delete_calendar", "刪除行事曆"),
        ("get_profile", "查看個人資料"),
        ("update_profile", "更新個人資料"),
        ("browse_literature", "瀏覽學術論文"),
        ("synthesize_advice", "綜合營養建議"),
    ]
    parts = [tc(450, 30, "Phase1 意圖決策樹 — 15 種意圖", "title")]
    parts.append(rr_bold(330, 50, 240, 50, "#EBF8FF", "#3182CE"))
    parts.append(tc(450, 82, "使用者訊息", "label-bold"))
    parts.append(arr_d(450, 100, 130, "blue"))
    parts.append(rr_bold(330, 130, 240, 50, "#FAF5FF", "#805AD5"))
    parts.append(tc(450, 162, "Phase1 LLM 判定", "label-bold"))
    colors = ["#3182CE", "#38A169", "#DD6B20", "#E53E3E", "#805AD5",
              "#319795", "#D69E2E", "#C53030", "#2B6CB0", "#276749",
              "#9B2C2C", "#2C7A7B", "#6B46C1", "#B7791F", "#4A5568"]
    col1_x, col2_x, col3_x = 30, 330, 630
    for i, (intent, desc) in enumerate(intents):
        if i < 5:
            x, row = col1_x, i
        elif i < 10:
            x, row = col2_x, i - 5
        else:
            x, row = col3_x, i - 10
        y = 220 + row * 60
        parts.append(rr(x, y, 240, 45, "#FFFFFF", colors[i % len(colors)]))
        parts.append(tl(x + 10, y + 20, intent, "code"))
        parts.append(tl(x + 10, y + 36, desc, "label-xs"))
    parts.append(rr(30, 640, 840, 50, "#F0FFF4", "#38A169"))
    parts.append(tc(450, 670, "LLM 根據訊息內容、對話歷史、系統說明書，判斷應該呼叫哪個（或哪些）工具", "label-sm"))
    return svg_wrap("\n".join(parts), w, h)


def make_svg_04_phase2_waterfall():
    """圖 4: Phase2 系統優先 → LLM 備援瀑布圖"""
    w, h = 850, 420
    parts = [tc(425, 30, "Phase2: 系統優先 → LLM 備援", "title")]
    parts.append(rr_bold(30, 60, 240, 70, "#F0FFF4", "#38A169"))
    parts.append(tc(150, 90, "系統生成 tool_calls", "label-bold"))
    parts.append(tc(150, 110, "_system_generate_toolcalls()", "code"))
    parts.append(arr_d(150, 130, 160, "green"))
    parts.append(rr(70, 165, 160, 40, "#EBF8FF", "#3182CE"))
    parts.append(tc(150, 190, "成功？", "label-bold"))
    parts.append(arr_r(230, 185, 370, "green", "成功 → 直接用"))
    parts.append(rr_bold(370, 160, 200, 50, "#F0FFF4", "#38A169"))
    parts.append(tc(470, 190, "執行工具", "label-bold"))
    parts.append(arr_d(150, 205, 250, "red", "失敗"))
    parts.append(rr_bold(30, 250, 240, 70, "#FAF5FF", "#805AD5"))
    parts.append(tc(150, 280, "LLM 備援生成", "label-bold"))
    parts.append(tc(150, 300, "PROMPT_PHASE2", "code"))
    parts.append(arr_r(270, 285, 370, "purple", "LLM 生成"))
    parts.append(rr_bold(370, 260, 200, 50, "#FAF5FF", "#805AD5"))
    parts.append(tc(470, 290, "執行工具", "label-bold"))
    parts.append(rr(30, 350, 790, 50, "#FFFAF0", "#DD6B20"))
    parts.append(tc(425, 380, "為什麼先用系統？系統生成更快、更準確，不需等 LLM 回應。只有系統無法處理時才呼叫 LLM。", "label-sm"))
    return svg_wrap("\n".join(parts), w, h)


def make_svg_05_tool_categories():
    """圖 5: 15 個工具分類總覽"""
    w, h = 920, 620
    parts = [tc(460, 30, "15 個工具分類總覽", "title")]
    cats = [
        ("搜尋類", "#3182CE", "#EBF8FF", [
            "search_food — 食物營養搜尋",
            "search_symptom — 症狀搜尋",
            "search_graphrag — 學術問答",
            "get_nutrient_ranking — 營養素排名",
            "browse_taiwan_food — 瀏覽食品庫",
            "get_literature_papers — 瀏覽論文",
        ]),
        ("分析類", "#38A169", "#F0FFF4", [
            "analyze_recipe — 食譜分析",
            "synthesize_advice — 綜合建議",
        ]),
        ("CRUD 類", "#DD6B20", "#FFFAF0", [
            "get_saved_recipes — 查食譜",
            "delete_saved_recipe — 刪食譜",
            "get_calendar_entries — 查行事曆",
            "add_to_calendar — 加行事曆",
            "delete_calendar_entry — 刪行事曆",
        ]),
        ("個人資料類", "#805AD5", "#FAF5FF", [
            "get_user_profile — 查個人資料",
            "update_user_profile — 更新資料",
        ]),
    ]
    for ci, (cat_name, stroke, fill, tools) in enumerate(cats):
        x = 30 + ci * 220
        bh = 40 + len(tools) * 26
        parts.append(rr_bold(x, 60, 200, bh, fill, stroke))
        parts.append(tc(x + 100, 85, cat_name, "label-bold"))
        for ti, tool in enumerate(tools):
            parts.append(tl(x + 12, 110 + ti * 26, tool, "label-sm"))
    parts.append(rr(30, 520, 860, 70, "#F0FFF4", "#38A169"))
    parts.append(tl(50, 545, "比喻：搜尋類 = 查詢菜單和食材　分析類 = 營養師計算　CRUD 類 = 記帳本操作　個人資料 = 會員卡", "label-sm"))
    parts.append(tl(50, 565, "所有工具都由 Control 層的 _TOOL_DISPATCH 字典統一調度，前端不直接呼叫 Backend", "label-sm"))
    return svg_wrap("\n".join(parts), w, h)


def make_svg_06_search_food_chain():
    """圖 6: search_food 通信鏈"""
    w, h = 880, 360
    parts = [tc(440, 25, "search_food 完整通信鏈", "title")]
    layers = [
        (30, "Frontend", "使用者輸入食物名", "#3182CE"),
        (240, "Control", "advisor → _execute_tool", "#38A169"),
        (480, "Backend", "GET /food/search?q=雞蛋", "#DD6B20"),
        (690, "PostgreSQL", "SELECT FROM taiwan_foods", "#E53E3E"),
    ]
    for x, name, desc, clr in layers:
        parts.append(rr_bold(x, 55, 170, 70, "#FFFFFF", clr))
        parts.append(tc(x+85, 82, name, "label-bold"))
        parts.append(tc(x+85, 102, desc, "label-xs"))
    parts.append(arr_r(200, 85, 238, "blue", "SSE"))
    parts.append(arr_r(410, 85, 478, "green", "httpx"))
    parts.append(arr_r(650, 85, 688, "orange", "SQL"))
    parts.append(arr_l(688, 155, 650, "orange", "結果 rows"))
    parts.append(arr_l(478, 155, 410, "green", "JSON"))
    parts.append(arr_l(238, 155, 200, "blue", "tool_result SSE"))
    parts.append(rr(30, 190, 820, 70, "#EBF8FF", "#3182CE"))
    parts.append(tl(50, 215, "回傳資料：食物名、每百克營養素（蛋白質、鈣、鐵、維生素等 110+ 欄位）、來源（taiwan/foodb）", "label-sm"))
    parts.append(tl(50, 240, "搜尋策略：先精確匹配 → 再模糊搜尋 → 雙資料庫（taiwan_foods 2213 筆 + FooDB 992 筆）", "label-sm"))
    parts.append(rr(30, 280, 820, 50, "#F0FFF4", "#38A169"))
    parts.append(tl(50, 310, "比喻：你問服務生「有沒有雞蛋？」→ 服務生問廚房 → 廚房查食材庫 → 回報所有雞蛋相關食材及營養", "label-sm"))
    return svg_wrap("\n".join(parts), w, h)


def make_svg_07_symptom_pipeline():
    """圖 7: 症狀搜尋五階段管道"""
    w, h = 920, 480
    stages = [
        ("1. LLM 解析", "自由輸入→關鍵字", "#805AD5"),
        ("2. 雙路搜尋", "對照表+向量語義", "#3182CE"),
        ("3. LLM 相關性", "篩選 top 10", "#38A169"),
        ("4. 化合物→食物", "batch 查詢", "#DD6B20"),
        ("5. 去重彙總", "合併結果", "#E53E3E"),
    ]
    parts = [tc(460, 25, "症狀搜尋五階段管道", "title")]
    for i, (name, desc, clr) in enumerate(stages):
        x = 20 + i * 178
        parts.append(rr_bold(x, 55, 160, 80, "#FFFFFF", clr))
        parts.append(tc(x+80, 85, name, "label-bold"))
        parts.append(tc(x+80, 110, desc, "label-sm"))
        if i < 4:
            parts.append(arr_r(x+160, 95, x+178, "gray"))
    parts.append(rr(20, 160, 880, 90, "#EBF8FF", "#3182CE"))
    parts.append(tl(40, 185, "階段 1：使用者說「最近很累又頭痛」→ Gemma LLM 拆解為 [\"疲勞\", \"頭痛\"]", "label-sm"))
    parts.append(tl(40, 205, "階段 2：每個關鍵字同時查 symptom_mapping.json（87筆靜態）+ Health Vector 向量搜尋（1435筆語義）", "label-sm"))
    parts.append(tl(40, 225, "階段 3：候選合併後，LLM 篩選最相關的 10 個 health effects（若 LLM 掛，用 score 排序 fallback）", "label-sm"))
    parts.append(tl(40, 245, "階段 4-5：查 compound→food 關聯，去重彙總，回傳推薦食物", "label-sm"))
    parts.append(rr(20, 270, 880, 80, "#F0FFF4", "#38A169"))
    parts.append(tl(40, 300, "比喻：你跟藥師說「我好累又頭痛」→ 藥師先聽懂（LLM 解析）→ 翻查兩本參考書（雙路搜尋）", "label-sm"))
    parts.append(tl(40, 325, "→ 藥師判斷哪些最相關（LLM 篩選）→ 查哪些食物含這些營養素（compound→food）→ 整理成建議", "label-sm"))
    parts.append(rr(20, 370, 880, 80, "#FFFAF0", "#DD6B20"))
    parts.append(tl(40, 395, "重要：整個管道有 graceful degradation 機制——LLM 掛了也能運作", "label-bold"))
    parts.append(tl(40, 420, "階段 1 若 LLM 失敗 → 直接用原文當關鍵字；階段 3 若 LLM 失敗 → 用 score 排序", "label-sm"))
    return svg_wrap("\n".join(parts), w, h)


def make_svg_08_recipe_analysis():
    """圖 8: 食譜分析二階段流程"""
    w, h = 920, 500
    parts = [tc(460, 25, "食譜分析完整流程", "title")]
    stage1 = [("食譜文字", 40), ("提取食材", 200), ("雙 DB 匹配", 380), ("確認食材", 560)]
    stage2 = [("計算營養", 40), ("DRI 比對", 240), ("缺口分析", 440), ("推薦食物", 640)]
    parts.append(tl(40, 60, "第一階段：提取與確認", "subtitle"))
    for name, x in stage1:
        parts.append(rr(x, 75, 140, 50, "#EBF8FF", "#3182CE"))
        parts.append(tc(x+70, 105, name, "label-bold"))
    for i in range(3):
        parts.append(arr_r(stage1[i][1]+140, 100, stage1[i+1][1], "blue"))
    parts.append(tl(40, 155, "第二階段：分析與建議", "subtitle"))
    for name, x in stage2:
        parts.append(rr(x, 170, 160, 50, "#F0FFF4", "#38A169"))
        parts.append(tc(x+80, 200, name, "label-bold"))
    for i in range(3):
        parts.append(arr_r(stage2[i][1]+160, 195, stage2[i+1][1], "green"))
    parts.append(rr(30, 250, 860, 100, "#EBF8FF", "#3182CE"))
    parts.append(tl(50, 275, "通信鏈：Control 呼叫 Backend POST /recipe/calculate-nutrition", "label-bold"))
    parts.append(tl(50, 300, "Backend 做的事：提取食材名 → 雙 DB（taiwan_foods + FooDB）匹配 → 按用量換算營養素 → 比對 HPA DRI v8", "label-sm"))
    parts.append(tl(50, 320, "→ 找出缺口（33 項營養素，其中 27 項可計算）→ 推薦補足食物", "label-sm"))
    parts.append(tl(50, 340, "回傳：每個食材的營養、總營養、DRI 缺口百分比、推薦食物", "label-sm"))
    parts.append(rr(30, 370, 860, 100, "#F0FFF4", "#38A169"))
    parts.append(tl(50, 395, "比喻：你把食譜交給營養師 →", "label-sm"))
    parts.append(tl(50, 415, "營養師先認出每個食材（提取）→ 查每個食材的營養成分表（雙 DB）→ 按你用的份量算總營養（計算）", "label-sm"))
    parts.append(tl(50, 435, "→ 對照國家建議攝取量（DRI）→ 告訴你哪些不夠（缺口）→ 建議吃什麼補（推薦）", "label-sm"))
    return svg_wrap("\n".join(parts), w, h)


def make_svg_09_graphrag_chain():
    """圖 9: GraphRAG 通信鏈"""
    w, h = 880, 350
    parts = [tc(440, 25, "GraphRAG 學術問答通信鏈", "title")]
    layers = [
        (30, "Control", "search_graphrag", "#38A169"),
        (250, "GraphRAG API", "POST /retrieve-only", "#319795"),
        (500, "ChromaDB", "向量搜尋", "#805AD5"),
        (700, "NetworkX", "圖遍歷", "#DD6B20"),
    ]
    for x, name, desc, clr in layers:
        parts.append(rr_bold(x, 55, 190, 70, "#FFFFFF", clr))
        parts.append(tc(x+95, 82, name, "label-bold"))
        parts.append(tc(x+95, 104, desc, "label-sm"))
    parts.append(arr_r(220, 90, 248, "green", "httpx 60s"))
    parts.append(arr_r(440, 90, 498, "teal", "查詢"))
    parts.append(arr_r(690, 90, 698, "purple"))
    parts.append(rr(30, 155, 820, 80, "#EBF8FF", "#3182CE"))
    parts.append(tl(50, 180, "四種搜尋方法：A=純向量 B=Local(entity→graph traversal) C=Global(community reports) D=Drift(sub-queries)", "label-sm"))
    parts.append(tl(50, 205, "所有方法最終都回到 entity_chunk_map 找原始文本 chunk → 確保有根據（grounded）", "label-sm"))
    parts.append(tl(50, 225, "回傳：每個方法的 answer + evidence + 參考文獻（IF/SJR quartile 色標、DOI 連結）", "label-sm"))
    parts.append(rr(30, 255, 820, 60, "#F0FFF4", "#38A169"))
    parts.append(tl(50, 285, "比喻：你問圖書館員學術問題 → 他用四種方法找資料 → 每種方法都附上原文出處，確保不是瞎說的", "label-sm"))
    return svg_wrap("\n".join(parts), w, h)


def make_svg_10_delete_recipe():
    """圖 10: 刪除食譜確認流程"""
    w, h = 850, 350
    parts = [tc(425, 25, "刪除食譜確認流程（兩步驟安全機制）", "title")]
    parts.append(rr_bold(30, 60, 180, 60, "#EBF8FF", "#3182CE"))
    parts.append(tc(120, 95, "使用者：刪食譜", "label-bold"))
    parts.append(arr_r(210, 90, 270, "blue"))
    parts.append(rr_bold(270, 60, 220, 60, "#FFFAF0", "#DD6B20"))
    parts.append(tc(380, 85, "第一次呼叫", "label-bold"))
    parts.append(tc(380, 105, "confirmed=false", "code"))
    parts.append(arr_r(490, 90, 550, "orange"))
    parts.append(rr_bold(550, 60, 260, 60, "#FFF5F5", "#E53E3E"))
    parts.append(tc(680, 85, "回傳：預覽匹配食譜", "label-bold"))
    parts.append(tc(680, 105, "pending_ids=[23,24,25]", "code"))
    parts.append(arr_d(680, 120, 165, "red"))
    parts.append(rr_bold(550, 165, 260, 60, "#EBF8FF", "#3182CE"))
    parts.append(tc(680, 190, "使用者確認", "label-bold"))
    parts.append(tc(680, 210, "「全部刪除」/「刪第 1 個」", "label-xs"))
    parts.append(arr_d(680, 225, 265, "blue"))
    parts.append(rr_bold(550, 265, 260, 60, "#F0FFF4", "#38A169"))
    parts.append(tc(680, 290, "第二次：confirmed=true", "label-bold"))
    parts.append(tc(680, 308, "recipe_ids=[23,24,25]", "code"))
    return svg_wrap("\n".join(parts), w, h)


def make_svg_11_prompt_lifecycle():
    """圖 11: 9 個 Prompt 在 Pipeline 上的位置"""
    w, h = 920, 400
    parts = [tc(460, 25, "9 個 Prompt 生命週期圖", "title")]
    steps = ["Echo", "Phase1", "Phase2", "Execute", "Analysis", "DualMem"]
    for i, s in enumerate(steps):
        x = 30 + i * 148
        parts.append(rr(x, 50, 130, 40, "#F7FAFC", "#CBD5E0"))
        parts.append(tc(x+65, 76, s, "label-bold"))
    prompts = [
        (0, "PROMPT_ECHO", "L51"),
        (1, "PROMPT_PHASE1", "L92"),
        (2, "PROMPT_PHASE2", "L196"),
        (4, "PROMPT_ANALYSIS", "L262"),
        (5, "PROMPT_NARRATIVE", "L317"),
        (5, "PROMPT_STRUCTURED", "L341"),
    ]
    for pi, (step_idx, name, line) in enumerate(prompts):
        x = 30 + step_idx * 148
        y = 110 + (pi % 3) * 50
        if step_idx == 5 and pi > 4:
            y = 160
        parts.append(rr(x-10, y, 150, 35, "#FAF5FF", "#805AD5"))
        parts.append(tc(x+65, y+15, name, "code"))
        parts.append(tc(x+65, y+30, line, "label-xs"))
    extra_prompts = [
        ("PROMPT_CLASSIFY_REPLY", "L3318", "中斷偵測"),
        ("PROMPT_CLASSIFY_RESUME", "L3328", "恢復偵測"),
        ("PROMPT_EXTRACT_PARAMS", "L3638", "參數提取"),
    ]
    for i, (name, line, desc) in enumerate(extra_prompts):
        x = 30 + i * 300
        y = 290
        parts.append(rr(x, y, 280, 35, "#FFFAF0", "#DD6B20"))
        parts.append(tl(x+10, y+15, f"{name} ({line})", "code"))
        parts.append(tr(x+270, y+15, desc, "label-sm"))
    parts.append(rr(30, 340, 860, 40, "#EBF8FF", "#3182CE"))
    parts.append(tc(460, 365, "上方 6 個是 Pipeline 主線 Prompt，下方 3 個是工作流狀態機的輔助 Prompt", "label-sm"))
    return svg_wrap("\n".join(parts), w, h)


def make_svg_12_sse_timeline():
    """圖 12: SSE 事件時間軸"""
    w, h = 920, 620
    parts = [tc(460, 25, "SSE 事件時間軸 — 一次查詢的完整事件流", "title")]
    events = [
        ("echo", "即時確認回應", "#3182CE", 60),
        ("thinking_start", "開始思考", "#805AD5", 90),
        ("thinking_delta", "思考進度", "#805AD5", 115),
        ("thinking_end", "思考完成", "#805AD5", 140),
        ("phase_status", "Pipeline 階段", "#319795", 170),
        ("workflow_state", "工作流狀態", "#319795", 195),
        ("collecting_state", "收集參數狀態", "#319795", 220),
        ("tool_start", "開始執行工具", "#DD6B20", 250),
        ("tool_progress", "工具進度", "#DD6B20", 275),
        ("tool_result", "工具原始結果", "#DD6B20", 300),
        ("tool_done", "工具完成", "#DD6B20", 325),
        ("recovery_attempt", "自癒嘗試", "#E53E3E", 355),
        ("recovery_result", "自癒結果", "#E53E3E", 380),
        ("reply", "最終分析回覆", "#38A169", 410),
        ("memory", "記憶壓縮", "#38A169", 435),
        ("narrative_memory", "敘事記憶", "#38A169", 460),
        ("structured_store", "結構化記憶", "#38A169", 485),
        ("session_id", "Session ID", "#718096", 515),
        ("checkpoint_saved", "檢查點存檔", "#718096", 540),
        ("error", "錯誤", "#E53E3E", 565),
    ]
    parts.append(f'<line x1="200" y1="55" x2="200" y2="580" stroke="#CBD5E0" stroke-width="2"/>')
    for name, desc, clr, y in events:
        parts.append(f'<circle cx="200" cy="{y}" r="5" fill="{clr}"/>')
        parts.append(tl(220, y+4, name, "code"))
        parts.append(tl(450, y+4, desc, "label-sm"))
    parts.append(rr(30, 590, 860, 20, "#F7FAFC", "#CBD5E0"))
    parts.append(tc(460, 604, "時間由上到下，每個事件都是 SSE 格式：event: xxx\\ndata: {...}\\n\\n", "label-xs"))
    return svg_wrap("\n".join(parts), w, h)


def make_svg_13_six_schools():
    """圖 13: 中斷處理六派總覽"""
    w, h = 920, 530
    parts = [tc(460, 25, "中斷處理六個派別", "title")]
    schools = [
        ("智慧派", "_detect_interruption_smart", "LLM 三態判斷", "#805AD5", "LLM + Regex 交叉驗證"),
        ("規則派", "_detect_interruption_regex", "正則比對", "#3182CE", "三把武器：問句/回歸/新意圖"),
        ("態度派", "_classify_collecting_response", "回覆分類", "#38A169", "answer/soft_rejection/hard_abandon"),
        ("回歸派", "_detect_resume_smart", "恢復偵測", "#DD6B20", "隱式（參數提取）+ 顯式（LLM）"),
        ("快捷派", "[RECIPE_CONFIRM]", "食譜確認", "#E53E3E", "add/remove/confirm/abort"),
        ("前線派", "AbortController", "前端取消", "#319795", "使用者按取消 → abort SSE"),
    ]
    for i, (name, func, weapon, clr, desc) in enumerate(schools):
        row, col = i // 3, i % 3
        x = 30 + col * 300
        y = 60 + row * 210
        parts.append(rr_bold(x, y, 280, 180, "#FFFFFF", clr))
        parts.append(f'<rect x="{x}" y="{y}" width="280" height="40" rx="8" fill="{clr}"/>')
        parts.append(tc(x+140, y+26, name, "white"))
        parts.append(tl(x+15, y+65, f"函數：{func}", "code"))
        parts.append(tl(x+15, y+90, f"武器：{weapon}", "label-sm"))
        parts.append(tl(x+15, y+115, desc, "label-sm"))
        parts.append(tl(x+15, y+145, f"位置：advisor.py", "label-xs"))
    parts.append(rr(30, 490, 860, 30, "#F0FFF4", "#38A169"))
    parts.append(tc(460, 510, "六派協同運作：使用者說的每句話都經過多層判斷，確保不會誤中斷也不會漏中斷", "label-sm"))
    return svg_wrap("\n".join(parts), w, h)


def make_svg_14_tristate():
    """圖 14: 智慧派三態決策樹"""
    w, h = 850, 420
    parts = [tc(425, 25, "智慧派：LLM × Regex 三態決策", "title")]
    parts.append(rr_bold(310, 55, 230, 50, "#FAF5FF", "#805AD5"))
    parts.append(tc(425, 85, "使用者回覆", "label-bold"))
    parts.append(arr_d(350, 105, 145, "purple", "LLM 判定"))
    parts.append(arr_d(500, 105, 145, "blue", "Regex 判定"))
    parts.append(rr(260, 150, 150, 40, "#FAF5FF", "#805AD5"))
    parts.append(tc(335, 175, "LLM: A 或 B", "label"))
    parts.append(rr(440, 150, 150, 40, "#EBF8FF", "#3182CE"))
    parts.append(tc(515, 175, "Regex: T/F", "label"))
    parts.append(arr_d(425, 195, 240, "gray"))
    parts.append(rr(320, 245, 210, 40, "#F7FAFC", "#CBD5E0"))
    parts.append(tc(425, 270, "兩者一致？", "label-bold"))
    parts.append(arr_r(530, 265, 600, "green", "一致"))
    parts.append(arr_d(425, 285, 340, "red", "不一致"))
    parts.append(rr_bold(600, 240, 180, 50, "#F0FFF4", "#38A169"))
    parts.append(tc(690, 270, "answer 或 interrupt", "label-bold"))
    parts.append(rr_bold(320, 345, 210, 50, "#FFFAF0", "#DD6B20"))
    parts.append(tc(425, 375, "ambiguous（模糊）", "label-bold"))
    parts.append(rr(30, 245, 220, 130, "#EBF8FF", "#3182CE"))
    parts.append(tl(45, 270, "LLM A = 在回答問題", "label-sm"))
    parts.append(tl(45, 290, "LLM B = 在說別的事", "label-sm"))
    parts.append(tl(45, 315, "Regex T = 偵測到中斷", "label-sm"))
    parts.append(tl(45, 335, "Regex F = 未偵測到中斷", "label-sm"))
    parts.append(tl(45, 360, "兩者不一致 → ambiguous", "accent"))
    return svg_wrap("\n".join(parts), w, h)


def make_svg_15_collecting_state():
    """圖 15: 收集回覆分類流轉"""
    w, h = 850, 320
    parts = [tc(425, 25, "態度派：收集回覆三種分類", "title")]
    states = [
        (50, "answer", "提供了資訊", "#38A169"),
        (320, "soft_rejection", "猶豫/不確定", "#DD6B20"),
        (600, "hard_abandon", "明確放棄", "#E53E3E"),
    ]
    for x, name, desc, clr in states:
        parts.append(rr_bold(x, 60, 200, 80, "#FFFFFF", clr))
        parts.append(tc(x+100, 90, name, "label-bold"))
        parts.append(tc(x+100, 115, desc, "label-sm"))
    parts.append(rr(50, 170, 200, 50, "#F0FFF4", "#38A169"))
    parts.append(tc(150, 200, "繼續收集 → ready", "label-sm"))
    parts.append(rr(320, 170, 200, 50, "#FFFAF0", "#DD6B20"))
    parts.append(tc(420, 195, "寬容：改說法再問一次", "label-sm"))
    parts.append(tc(420, 210, "(最多寬容 1 次)", "label-xs"))
    parts.append(rr(600, 170, 200, 50, "#FFF5F5", "#E53E3E"))
    parts.append(tc(700, 200, "直接重置為 idle", "label-sm"))
    parts.append(rr(30, 250, 790, 50, "#EBF8FF", "#3182CE"))
    parts.append(tl(50, 275, "判斷依據：_SOFT_REJECTION_RE（猶豫模式）+ _SOFT_REJECTION_EXACT（精確匹配：「不知道」「隨便」「都行」等）", "label-sm"))
    parts.append(tl(50, 295, "hard_abandon 關鍵字：「算了」「不要了」「取消」「不用了」「不查了」", "label-sm"))
    return svg_wrap("\n".join(parts), w, h)


def make_svg_16_resume_flow():
    """圖 16: 回歸派恢復流程"""
    w, h = 850, 420
    parts = [tc(425, 25, "回歸派：中斷恢復偵測流程", "title")]
    parts.append(rr_bold(310, 55, 230, 50, "#EBF8FF", "#3182CE"))
    parts.append(tc(425, 85, "有中斷紀錄？", "label-bold"))
    parts.append(arr_d(425, 105, 145, "blue", "有"))
    parts.append(rr_bold(180, 150, 230, 50, "#F0FFF4", "#38A169"))
    parts.append(tc(295, 180, "隱式探測", "label-bold"))
    parts.append(tl(190, 195, "_extract_params_deterministic", "label-xs"))
    parts.append(rr_bold(440, 150, 230, 50, "#FAF5FF", "#805AD5"))
    parts.append(tc(555, 180, "顯式 LLM", "label-bold"))
    parts.append(tl(450, 195, "PROMPT_CLASSIFY_RESUME", "label-xs"))
    parts.append(arr_r(295, 200, 370, "green", "找到參數 → resume"))
    parts.append(arr_d(555, 200, 250, "purple", "LLM 判斷"))
    parts.append(rr(440, 255, 230, 40, "#FFFFFF", "#805AD5"))
    parts.append(tc(555, 280, "A = 回到舊任務", "label"))
    parts.append(rr(440, 305, 230, 40, "#FFFFFF", "#718096"))
    parts.append(tc(555, 330, "B = 新的事情", "label"))
    parts.append(rr(30, 360, 790, 40, "#FFFAF0", "#DD6B20"))
    parts.append(tc(425, 385, "隱式優先：使用者直接給出之前缺的資訊 → 不用問就恢復；否則才問 LLM「你是想繼續嗎？」", "label-sm"))
    return svg_wrap("\n".join(parts), w, h)


def make_svg_17_recipe_confirm():
    """圖 17: 食譜確認協議"""
    w, h = 850, 380
    parts = [tc(425, 25, "快捷派：[RECIPE_CONFIRM] 食譜確認協議", "title")]
    actions = [
        (50, "add", "新增食材", "#38A169"),
        (260, "remove", "移除食材", "#E53E3E"),
        (470, "confirm", "確認清單", "#3182CE"),
        (680, "abort", "放棄分析", "#718096"),
    ]
    for x, act, desc, clr in actions:
        parts.append(rr_bold(x, 60, 170, 70, "#FFFFFF", clr))
        parts.append(tc(x+85, 90, act, "label-bold"))
        parts.append(tc(x+85, 115, desc, "label-sm"))
    parts.append(rr(50, 160, 750, 90, "#EBF8FF", "#3182CE"))
    parts.append(tl(70, 185, "工具回傳格式：[RECIPE_CONFIRM] {\"action\":\"confirm\",\"ingredients\":[...]}", "code"))
    parts.append(tl(70, 210, "前端偵測到 [RECIPE_CONFIRM] → 顯示食材清單 UI → 使用者可以 add/remove/confirm/abort", "label-sm"))
    parts.append(tl(70, 235, "使用者操作後，前端組裝新訊息送回 Control → 根據 action 決定下一步", "label-sm"))
    parts.append(rr(50, 270, 750, 80, "#F0FFF4", "#38A169"))
    parts.append(tl(70, 295, "比喻：廚師列出要用的食材清單，你可以說：", "label-sm"))
    parts.append(tl(70, 315, "「加番茄」(add)、「不要洋蔥」(remove)、「就這些」(confirm)、「算了不做了」(abort)", "label-sm"))
    parts.append(tl(70, 335, "這是唯一一個不經過 LLM 判斷的快速通道，直接靠前端 UI 和後端正則處理", "label-sm"))
    return svg_wrap("\n".join(parts), w, h)


def make_svg_18_compound_dag():
    """圖 18: 複合任務 DAG 圖"""
    w, h = 880, 420
    parts = [tc(440, 25, "複合任務 DAG 依賴圖", "title")]
    parts.append(rr_bold(340, 55, 200, 50, "#EBF8FF", "#3182CE"))
    parts.append(tc(440, 85, "多意圖偵測", "label-bold"))
    parts.append(arr_d(380, 105, 145, "blue"))
    parts.append(arr_d(500, 105, 145, "blue"))
    parts.append(rr_bold(250, 150, 180, 50, "#F0FFF4", "#38A169"))
    parts.append(tc(340, 180, "analyze_recipe", "label-bold"))
    parts.append(rr_bold(470, 150, 180, 50, "#FFFAF0", "#DD6B20"))
    parts.append(tc(560, 180, "synthesize_advice", "label-bold"))
    parts.append(arr_d(340, 200, 250, "green", "依賴：先分析"))
    parts.append(rr_bold(250, 255, 180, 50, "#38A169", "#276749"))
    parts.append(tc(340, 285, "add_to_calendar", "white"))
    parts.append(dash_r(430, 275, 470, "#718096"))
    parts.append(tl(475, 280, "可平行", "label-xs"))
    parts.append(rr(50, 340, 780, 50, "#EBF8FF", "#3182CE"))
    parts.append(tc(440, 370, "DAG 規則：analyze_recipe 必須在 add_to_calendar 之前（需要食譜資料）；其餘可平行執行", "label-sm"))
    return svg_wrap("\n".join(parts), w, h)


def make_svg_19_recovery_engine():
    """圖 19: 自癒引擎流程"""
    w, h = 850, 430
    parts = [tc(425, 25, "自癒引擎 (ReIn) 三步驟流程", "title")]
    steps = [
        (60, "diagnose()", "診斷錯誤類型", "#E53E3E"),
        (320, "plan_recovery()", "制定恢復策略", "#DD6B20"),
        (580, "execute_recovery()", "執行恢復", "#38A169"),
    ]
    for x, func, desc, clr in steps:
        parts.append(rr_bold(x, 55, 210, 70, "#FFFFFF", clr))
        parts.append(tc(x+105, 85, func, "label-bold"))
        parts.append(tc(x+105, 110, desc, "label-sm"))
    parts.append(arr_r(270, 90, 318, "red"))
    parts.append(arr_r(530, 90, 578, "orange"))
    errors = ["empty_result（查無結果）", "tool_timeout（工具超時）",
              "param_mismatch（參數不匹配）", "backend_error（後端錯誤）"]
    strategies = ["synonym_expand（同義詞擴展）", "retry_timeout（重試）",
                  "field_remap（欄位重映射）", "skip（跳過）"]
    parts.append(rr(60, 150, 210, 130, "#FFF5F5", "#E53E3E"))
    parts.append(tc(165, 175, "四種錯誤", "label-bold"))
    for i, e in enumerate(errors):
        parts.append(tl(75, 200 + i*25, e, "label-sm"))
    parts.append(rr(320, 150, 210, 130, "#FFFAF0", "#DD6B20"))
    parts.append(tc(425, 175, "四種策略", "label-bold"))
    for i, s in enumerate(strategies):
        parts.append(tl(335, 200 + i*25, s, "label-sm"))
    parts.append(rr(580, 150, 210, 130, "#F0FFF4", "#38A169"))
    parts.append(tc(685, 175, "執行方式", "label-bold"))
    parts.append(tl(595, 200, "synonym: 嘗試替代詞", "label-sm"))
    parts.append(tl(595, 225, "retry: 直接重試一次", "label-sm"))
    parts.append(tl(595, 250, "remap: 換欄位名重試", "label-sm"))
    parts.append(tl(595, 275, "skip: 跳過並告知使用者", "label-sm"))
    parts.append(rr(60, 310, 730, 90, "#F0FFF4", "#38A169"))
    parts.append(tl(80, 340, "比喻：你問餐廳「有起司嗎？」→ 廚房查不到（empty_result）→ 自動幫你換個說法「芝士」再查（synonym_expand）", "label-sm"))
    parts.append(tl(80, 365, "→ 還是沒有就再試「乾酪」「cheese」（alternatives 最多試 5 個替代詞）", "label-sm"))
    parts.append(tl(80, 390, "同義詞庫：60+ 組常見台灣食物同義詞（芝士/起司、番茄/西紅柿、鳳梨/菠蘿...）", "label-sm"))
    return svg_wrap("\n".join(parts), w, h)


def make_svg_20_state_machine():
    """圖 20: 工作流狀態機"""
    w, h = 850, 420
    parts = [tc(425, 25, "工作流狀態機轉換圖", "title")]
    states = [
        (50, 100, "idle", "等待輸入", "#718096"),
        (250, 100, "collecting", "收集參數", "#DD6B20"),
        (480, 100, "ready", "參數齊全", "#3182CE"),
        (680, 100, "executing", "執行中", "#805AD5"),
        (480, 230, "done", "完成", "#38A169"),
    ]
    for x, y, name, desc, clr in states:
        parts.append(rr_bold(x, y, 150, 60, "#FFFFFF", clr))
        parts.append(tc(x+75, y+28, name, "label-bold"))
        parts.append(tc(x+75, y+48, desc, "label-sm"))
    parts.append(arr_r(200, 130, 248, "orange", "參數不足"))
    parts.append(arr_r(400, 130, 478, "blue", "參數齊全"))
    parts.append(arr_r(630, 130, 678, "purple", "開始執行"))
    parts.append(arr_d(555, 160, 228, "green", "完成"))
    parts.append(f'<path d="M250,160 Q150,200 150,160" stroke="#E53E3E" stroke-width="1.5" fill="none" stroke-dasharray="4,3"/>')
    parts.append(tl(130, 200, "中斷 → idle", "accent"))
    parts.append(rr(50, 310, 750, 80, "#EBF8FF", "#3182CE"))
    parts.append(tl(70, 335, "collecting 階段：最多追問 3 輪（_MAX_ROUNDS=3）；遇到 soft_rejection 會寬容一次改說法再問", "label-sm"))
    parts.append(tl(70, 360, "中斷發生時：當前狀態存入 interrupted_state，切回 idle 處理新意圖", "label-sm"))
    parts.append(tl(70, 380, "恢復時：從 interrupted_state 讀回來繼續", "label-sm"))
    return svg_wrap("\n".join(parts), w, h)


def make_svg_21_validation_pipeline():
    """圖 21: 六層驗證管道"""
    w, h = 880, 380
    layers = [
        ("1. 意圖別名修正", "_INTENT_ALIASES", "#3182CE"),
        ("2. 意圖合法性", "_VALID_INTENTS", "#38A169"),
        ("3. 工具參數驗證", "_validate_toolcalls", "#DD6B20"),
        ("4. 營養素映射", "_NUTRIENT_ZH_TO_FIELD", "#805AD5"),
        ("5. 品質過濾", "結果檢查", "#E53E3E"),
        ("6. JSON 解析", "_parse_llm_json", "#319795"),
    ]
    parts = [tc(440, 25, "六層驗證與自修復管道", "title")]
    for i, (name, func, clr) in enumerate(layers):
        x = 20 + i * 143
        parts.append(rr_bold(x, 55, 128, 80, "#FFFFFF", clr))
        parts.append(tc(x+64, 82, name, "label-bold"))
        parts.append(tc(x+64, 105, func, "code"))
        if i < 5:
            parts.append(arr_r(x+128, 95, x+143, "gray"))
    parts.append(rr(20, 160, 840, 100, "#EBF8FF", "#3182CE"))
    parts.append(tl(40, 185, "層 1：修正 LLM 常見錯誤，如 get_user_profile → get_profile、get_nutrient_ranking → nutrient_ranking", "label-sm"))
    parts.append(tl(40, 205, "層 2：過濾不在白名單的意圖名，防止 LLM 亂發明工具", "label-sm"))
    parts.append(tl(40, 225, "層 3：檢查必要參數是否存在、型別是否正確", "label-sm"))
    parts.append(tl(40, 245, "層 4：中文營養素名 → 資料庫欄位名（蛋白質→protein_per_100g、鈣→calcium_per_100g 等）", "label-sm"))
    parts.append(rr(20, 280, 840, 70, "#F0FFF4", "#38A169"))
    parts.append(tl(40, 305, "層 5：檢查工具回傳的結果是否合理（空結果觸發自癒引擎）", "label-sm"))
    parts.append(tl(40, 330, "層 6：三級 JSON fallback — 直接 parse → regex 提取 → 去掉 markdown 再試", "label-sm"))
    return svg_wrap("\n".join(parts), w, h)


def make_svg_22_memory_tree():
    """圖 22: 記憶樹架構圖"""
    w, h = 880, 480
    parts = [tc(440, 25, "雙記憶系統 + 記憶樹架構", "title")]
    parts.append(rr_bold(30, 55, 380, 80, "#FAF5FF", "#805AD5"))
    parts.append(tc(220, 80, "敘事記憶 (Narrative Memory)", "label-bold"))
    parts.append(tc(220, 105, "對話脈絡摘要，記對話流程不記具體數據", "label-sm"))
    parts.append(rr_bold(460, 55, 380, 80, "#EBF8FF", "#3182CE"))
    parts.append(tc(650, 80, "結構化記憶 (Structured Store)", "label-bold"))
    parts.append(tc(650, 105, "JSON 結構：user_facts / food_context / ...", "label-sm"))
    parts.append(rr(30, 160, 810, 100, "#FFFAF0", "#DD6B20"))
    parts.append(tc(435, 185, "結構化記憶 11 種欄位", "label-bold"))
    fields = "user_facts / food_context / recipe_entry / symptom_entry / ranking_entry / calendar_update / academic_entry / papers_referenced / dri_awareness / action_entry / saved_recipes_summary"
    parts.append(tc(435, 215, fields, "label-xs"))
    parts.append(tc(435, 240, "每種都有固定 schema、timestamp、保留上限（20~30 筆）", "label-xs"))
    parts.append(rr_bold(250, 290, 380, 100, "#F0FFF4", "#38A169"))
    parts.append(tc(440, 315, "記憶樹 (Memory Tree)", "label-bold"))
    parts.append(tc(440, 340, "Leaf → Session → User", "label"))
    parts.append(tc(440, 360, "三層架構 + embedding 向量相似度", "label-sm"))
    parts.append(rr(30, 410, 810, 50, "#EBF8FF", "#3182CE"))
    parts.append(tc(435, 440, "Pipeline Step 6：兩個 Prompt 並行執行 — PROMPT_NARRATIVE_SUMMARY + PROMPT_STRUCTURED_EXTRACT", "label-sm"))
    return svg_wrap("\n".join(parts), w, h)


def make_svg_23_advisor_page():
    """圖 23: AdvisorPage 資料流圖"""
    w, h = 920, 520
    parts = [tc(460, 25, "AdvisorPage 元件資料流", "title")]
    components = [
        (50, 60, "使用者輸入", "textarea + send", "#3182CE"),
        (300, 60, "advisorApi.js", "SSE 串流連線", "#38A169"),
        (560, 60, "SSE Callbacks", "18 種事件處理", "#805AD5"),
        (50, 180, "ChatMessages", "對話氣泡渲染", "#DD6B20"),
        (300, 180, "ToolResultCard", "15 種工具卡片", "#E53E3E"),
        (560, 180, "ThinkingPanel", "思考進度面板", "#319795"),
        (50, 300, "IndexedDB", "本地持久化", "#718096"),
        (300, 300, "localStorage", "最近 30 則", "#718096"),
        (560, 300, "State 管理", "history/memory/ws", "#3182CE"),
    ]
    for x, y, name, desc, clr in components:
        parts.append(rr(x, y, 200, 70, "#FFFFFF", clr))
        parts.append(tc(x+100, y+30, name, "label-bold"))
        parts.append(tc(x+100, y+55, desc, "label-sm"))
    parts.append(arr_r(250, 95, 298, "blue"))
    parts.append(arr_r(500, 95, 558, "green"))
    parts.append(arr_d(150, 130, 178, "orange"))
    parts.append(arr_d(400, 130, 178, "red"))
    parts.append(arr_d(660, 130, 178, "teal"))
    parts.append(arr_d(150, 250, 298, "gray"))
    parts.append(arr_d(400, 250, 298, "gray"))
    parts.append(rr(30, 400, 860, 90, "#F0FFF4", "#38A169"))
    parts.append(tl(50, 425, "資料流：使用者輸入 → advisorApi.js 建立 SSE 連線 → 18 種 callback 分發事件", "label-sm"))
    parts.append(tl(50, 450, "→ 更新各元件狀態（對話、工具卡片、思考面板）→ 持久化到 IndexedDB/localStorage", "label-sm"))
    parts.append(tl(50, 475, "ToolResultCard 有 15 種渲染器，依工具類型顯示不同 UI（表格、卡片、清單、圖表等）", "label-sm"))
    return svg_wrap("\n".join(parts), w, h)


def make_svg_24_unified_analysis():
    """圖 24: UnifiedAnalysisPage 七步驟"""
    w, h = 920, 420
    parts = [tc(460, 25, "整合分析頁七步驟流程", "title")]
    steps = [
        ("1.貼食譜", "#3182CE"), ("2.提取", "#38A169"), ("3.確認", "#DD6B20"),
        ("4.查詢", "#805AD5"), ("5.計算", "#E53E3E"), ("6.缺口", "#319795"),
        ("7.推薦", "#38A169"),
    ]
    for i, (name, clr) in enumerate(steps):
        x = 20 + i * 128
        parts.append(rr_bold(x, 55, 112, 60, "#FFFFFF", clr))
        parts.append(tc(x+56, 90, name, "label-bold"))
        if i < 6:
            parts.append(arr_r(x+112, 85, x+128, "gray"))
    parts.append(rr(20, 140, 880, 120, "#EBF8FF", "#3182CE"))
    parts.append(tl(40, 165, "步驟 1：使用者貼上食譜文字（如「雞胸肉 200g、花椰菜 100g、白飯一碗」）", "label-sm"))
    parts.append(tl(40, 185, "步驟 2：Backend 用 ai_service 正則提取食材名和份量", "label-sm"))
    parts.append(tl(40, 205, "步驟 3：使用者確認食材清單（可拖曳到停用詞區移除、也可拖回來）", "label-sm"))
    parts.append(tl(40, 225, "步驟 4-5：雙 DB 匹配食材 → 按份量計算營養素", "label-sm"))
    parts.append(tl(40, 245, "步驟 6-7：對照 HPA DRI v8（33 項）找出缺口 → 推薦補足食物", "label-sm"))
    parts.append(rr(20, 280, 880, 110, "#F0FFF4", "#38A169"))
    parts.append(tl(40, 305, "重要 UI 互動：", "label-bold"))
    parts.append(tl(40, 325, "• 食材 tag 可拖曳到停用詞區（停用詞也可拖回來）— 雙向 DnD", "label-sm"))
    parts.append(tl(40, 345, "• Section 2「各食材營養明細」食材名可點擊 → 開啟 TaiwanFoodModal 顯示完整營養", "label-sm"))
    parts.append(tl(40, 365, "• Section 3 DRI 缺口表的缺乏列有「推薦食物」toggle → 展開 5 欄表顯示排名", "label-sm"))
    parts.append(tl(40, 385, "• 非 advisor 路由，不走 LLM Pipeline；直接呼叫 Backend API", "label-sm"))
    return svg_wrap("\n".join(parts), w, h)


def make_svg_25_recipe_e2e():
    """圖 25: 食譜分析端到端通信鏈"""
    w, h = 920, 620
    parts = [tc(460, 25, "食譜分析端到端通信鏈", "title")]
    layers = [
        ("使用者", 35), ("Frontend", 160), ("nginx", 280),
        ("Control", 400), ("Backend", 560), ("PostgreSQL", 720),
    ]
    for name, x in layers:
        parts.append(rr(x, 50, 110, 35, "#EBF8FF", "#3182CE"))
        parts.append(tc(x+55, 73, name, "label-bold"))
    y_start = 100
    msgs = [
        (35, 160, "輸入食譜文字", "blue", y_start),
        (160, 280, "/api/advisor/chat", "blue", y_start+30),
        (280, 400, "POST /advisor/chat", "green", y_start+60),
        (400, 400, "Step1 Echo (LLM)", "purple", y_start+90),
        (400, 400, "Step2 Phase1 (LLM)", "purple", y_start+110),
        (400, 400, "Step3 Phase2", "purple", y_start+130),
        (400, 560, "POST /recipe/calculate", "orange", y_start+160),
        (560, 720, "SELECT foods", "red", y_start+190),
        (720, 560, "rows", "red", y_start+220),
        (560, 400, "JSON nutrition", "orange", y_start+250),
        (400, 400, "Step5 Analysis (LLM)", "purple", y_start+280),
        (400, 400, "Step6 DualMemory", "teal", y_start+300),
        (400, 160, "SSE events", "green", y_start+330),
        (160, 35, "渲染結果卡片", "blue", y_start+360),
    ]
    for x1, x2, label, clr, y in msgs:
        if x1 == x2:
            parts.append(tl(x1+5, y, label, "label-xs"))
        elif x1 < x2:
            parts.append(arr_r(x1+55, y, x2+55, clr, label))
        else:
            parts.append(arr_l(x1+55, y, x2+55, clr, label))
    return svg_wrap("\n".join(parts), w, h)


def make_svg_26_symptom_e2e():
    """圖 26: 症狀搜尋端到端通信鏈"""
    w, h = 920, 520
    parts = [tc(460, 25, "症狀搜尋端到端通信鏈", "title")]
    layers = [
        ("Control", 50), ("Gemma LLM", 230), ("Health Vector", 430), ("Backend", 640),
    ]
    for name, x in layers:
        parts.append(rr(x, 50, 160, 35, "#EBF8FF", "#3182CE"))
        parts.append(tc(x+80, 73, name, "label-bold"))
    steps = [
        (50, 230, "1. 症狀解析", "purple", 110),
        (230, 50, "拆解結果", "purple", 135),
        (50, 430, "2. 向量搜尋", "teal", 170),
        (430, 50, "語義匹配", "teal", 195),
        (50, 230, "3. 相關性判斷", "purple", 230),
        (230, 50, "top 10", "purple", 255),
        (50, 640, "4. 查食物", "orange", 290),
        (640, 50, "推薦食物", "orange", 315),
    ]
    for x1, x2, label, clr, y in steps:
        if x1 < x2:
            parts.append(arr_r(x1+80, y, x2+80, clr, label))
        else:
            parts.append(arr_l(x1+80, y, x2+80, clr, label))
    parts.append(rr(30, 360, 860, 130, "#F0FFF4", "#38A169"))
    parts.append(tl(50, 385, "完整流程：使用者說「手肘肌肉酸痛」", "label-bold"))
    parts.append(tl(50, 410, "→ Gemma 解析為 [\"肌肉酸痛\"] → 查 symptom_mapping + Health Vector 語義搜尋", "label-sm"))
    parts.append(tl(50, 435, "→ Gemma 從候選中篩出最相關的 → Backend 查 compound→food 關聯", "label-sm"))
    parts.append(tl(50, 460, "→ 回傳：化合物列表（如 Magnesium、Vitamin B6）+ 推薦食物（如 香蕉、菠菜）+ 相似度分數", "label-sm"))
    return svg_wrap("\n".join(parts), w, h)


def make_svg_27_calendar_crud():
    """圖 27: 行事曆 CRUD 鏈"""
    w, h = 880, 420
    parts = [tc(440, 25, "行事曆 CRUD 通信鏈", "title")]
    ops = [
        ("查詢", "get_calendar_entries", "GET /calendar/entries?start=...&amp;end=...", "#3182CE", 60),
        ("新增", "add_to_calendar", "POST /recipe/save + /calendar/entries", "#38A169", 160),
        ("刪除", "delete_calendar_entry", "GET entries + DELETE entries/{id}", "#E53E3E", 260),
    ]
    for name, tool, api, clr, y in ops:
        parts.append(rr_bold(30, y, 100, 50, "#FFFFFF", clr))
        parts.append(tc(80, y+30, name, "label-bold"))
        parts.append(rr(160, y, 200, 50, "#EBF8FF", "#3182CE"))
        parts.append(tc(260, y+25, tool, "code"))
        parts.append(tc(260, y+42, "Control 工具", "label-xs"))
        parts.append(arr_r(360, y+25, 410, "blue"))
        parts.append(rr(410, y, 440, 50, "#FFFAF0", "#DD6B20"))
        parts.append(tl(420, y+30, api, "label-sm"))
    parts.append(rr(30, 340, 820, 60, "#F0FFF4", "#38A169"))
    parts.append(tl(50, 365, "新增特殊邏輯：先 POST /recipe/save 存食譜 → 再 POST /calendar/entries 加到行事曆", "label-sm"))
    parts.append(tl(50, 385, "刪除特殊邏輯：先 GET 查到 entry_id → 再 DELETE，可指定 recipe_name 篩選特定筆", "label-sm"))
    return svg_wrap("\n".join(parts), w, h)


def make_svg_28_food_search_e2e():
    """圖 28: 食物搜尋 E2E"""
    w, h = 880, 370
    parts = [tc(440, 25, "食物搜尋端到端通信鏈", "title")]
    layers = [
        (30, "Control", "#38A169"),
        (240, "Backend", "#DD6B20"),
        (460, "taiwan_foods", "#3182CE"),
        (680, "FooDB foods", "#805AD5"),
    ]
    for x, name, clr in layers:
        parts.append(rr_bold(x, 55, 170, 50, "#FFFFFF", clr))
        parts.append(tc(x+85, 85, name, "label-bold"))
    parts.append(arr_r(200, 80, 238, "green", "GET /food/search?q=雞蛋"))
    parts.append(arr_r(410, 70, 458, "blue", "ILIKE %雞蛋%"))
    parts.append(arr_r(410, 95, 678, "purple", "ILIKE %雞蛋%"))
    parts.append(rr(30, 130, 820, 100, "#EBF8FF", "#3182CE"))
    parts.append(tl(50, 155, "搜尋策略：同時查詢兩個資料庫", "label-bold"))
    parts.append(tl(50, 180, "taiwan_foods（2213 筆）：台灣食品營養成分資料庫，110+ 營養素欄位，中文食品名", "label-sm"))
    parts.append(tl(50, 200, "FooDB foods（992 筆）：國際食品資料庫，25 分類，英文食品名（有 name_zh 翻譯欄位）", "label-sm"))
    parts.append(tl(50, 220, "合併去重後回傳：每筆含 source 欄位（taiwan/foodb）標示來源", "label-sm"))
    parts.append(rr(30, 250, 820, 90, "#F0FFF4", "#38A169"))
    parts.append(tl(50, 275, "自癒機制：若「芝士」查不到 → recovery.py 自動嘗試「起司」「乾酪」「cheese」", "label-bold"))
    parts.append(tl(50, 300, "食物名簡化：「一碗新鮮的花椰菜」→ 去掉量詞和修飾詞 →「花椰菜」重新搜尋", "label-sm"))
    parts.append(tl(50, 325, "比喻：你問超市店員「有沒有芝士？」→ 店員查兩個倉庫 → 沒有就幫你換個說法再查", "label-sm"))
    return svg_wrap("\n".join(parts), w, h)


# ═══════════════════════════════════════════════════════════
# Section D: Chapter Content Functions
# ═══════════════════════════════════════════════════════════

def build_cover():
    p("", space_after=60)
    p("Recipe-AI 系統完整技術教學", bold=True, size=28, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    p("六步驟 LLM Pipeline · 15 個工具 · 9 個 Prompt", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    p("六派中斷處理 · 自癒引擎 · 雙記憶系統", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    p("SSE 串流通信 · 複合問題處理 · 完整通信鏈", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)
    p("適合對象：高中生以上，無需程式設計基礎", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    p("2026 年 5 月", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    page_break()


def build_ch00():
    h1("第 0 章：前言——這份文件在講什麼？")

    analogy_box("核心比喻：高級餐廳",
        "整個 Recipe-AI 系統就像一間高級餐廳：\n"
        "• Frontend（前台）= 餐廳入口，負責接待客人、展示菜單、傳遞點餐\n"
        "• Control（領班）= 餐廳的核心調度者，接到客人需求後決定要找誰\n"
        "• Backend（廚房）= 實際料理的地方，有食材庫（資料庫）\n"
        "• PostgreSQL（食材庫）= 儲存所有食物營養資料的大倉庫\n"
        "• Gemma LLM（大廚）= 會思考、判斷的人工智慧大腦\n"
        "• GraphRAG（圖書館）= 查學術論文的地方\n"
        "• Health Vector（營養師資料庫）= 存放健康效果和症狀資料的知識庫")

    p("這份文件會帶你從零開始，理解這個系統裡每一個零件怎麼運作。我們會先用生活中的比喻讓你「感覺到」這個概念，再慢慢帶入正式的技術名詞。")

    h2("閱讀指引")
    p("在這份文件中，你會看到以下幾種特殊格式的方框：")
    analogy_box("比喻框（綠色底）", "用生活中的例子來解釋技術概念，幫助你建立直覺。")
    tip_box("提示框（藍色底）", "補充說明、進階知識、或是需要特別注意的細節。")
    warn_box("警告框（紅色底）：標示容易犯錯或需要小心的地方。")
    prompt_box("Prompt 框（紫色底）", "顯示系統中實際使用的 Prompt（指令模板），這是給 AI 看的「劇本」。")

    p("此外，程式碼區塊會用灰色背景、等寬字體顯示：")
    code('result = await _call_gemma(messages, temperature=0.3)\n# 這行呼叫 Gemma LLM，temperature=0.3 表示回覆的隨機程度')

    page_break()


def build_ch01():
    h1("第 1 章：系統全貌——七個服務如何合作")

    p("在開始深入各種技術細節之前，讓我們先從最高的角度看這個系統長什麼樣子。")

    analogy_box("想像一間餐廳",
        "你走進一間餐廳，從點餐到上菜，中間其實經過了很多人的合作：\n"
        "接待員帶你入座（Frontend）→ 領班聽你的需求（Control）→ 廚師開始備料和烹飪（Backend + DB）\n"
        "→ 如果你問了學術問題，領班會跑去圖書館查資料（GraphRAG）\n"
        "→ 如果你說身體不舒服，領班會去找營養師（Health Vector）\n"
        "→ 所有思考判斷都由大廚的大腦完成（Gemma LLM）")

    add_image(make_svg_01_architecture())
    pi("圖 1：七個服務的架構圖。每個方塊是一個獨立運行的程式，透過網路互相溝通。")

    h2("七個服務分別做什麼？")
    add_table(
        ["服務", "Port", "角色", "餐廳比喻"],
        [
            ["Frontend", "3001", "使用者介面（網頁）", "前台：接待客人、展示菜單"],
            ["Control", "8000", "核心調度、LLM Pipeline", "領班：調度一切、做決策"],
            ["Backend", "8001", "資料庫操作、API", "廚房：實際處理食材資料"],
            ["PostgreSQL", "5432", "資料庫", "食材庫：儲存所有資料"],
            ["Gemma LLM", "8080", "語言模型（AI 大腦）", "大廚的大腦：思考判斷"],
            ["GraphRAG", "8002", "學術論文知識圖譜", "圖書館：查學術資料"],
            ["Health Vector", "8003", "健康效果向量搜尋", "營養師資料庫"],
        ]
    )

    tip_box("什麼是 Port（埠號）？",
        "Port 就像是同一棟大樓裡的門牌號碼。大樓的地址是 localhost（本機電腦），"
        "每個服務住在不同的門牌號碼。Frontend 在 3001 號門、Control 在 8000 號門，"
        "它們透過門牌號碼找到彼此。")

    h2("服務之間怎麼溝通？")
    p("這些服務之間用 HTTP 協定溝通，就像寄信一樣：")
    bullet("Frontend → Control：使用者的每個操作都透過 API 送到 Control（走 nginx 代理）")
    bullet("Control → Backend：Control 用 httpx（一個 Python 工具）把請求轉給 Backend")
    bullet("Control → Gemma：Control 呼叫 Gemma LLM 做思考判斷（temperature 參數控制隨機度）")
    bullet("Backend → PostgreSQL：Backend 用 asyncpg 查詢資料庫")

    tip_box("什麼是 API？",
        "API（Application Programming Interface）就像餐廳的點餐單。你不需要走進廚房告訴廚師要什麼，"
        "只要在點餐單上勾選，服務生就會幫你傳達。API 就是程式之間的「點餐單」——定義好格式，"
        "一方發送請求，另一方回傳結果。")

    page_break()


def build_ch02():
    h1("第 2 章：六步驟 LLM Pipeline——AI 顧問如何思考")

    p("當你在 AI 營養顧問頁面輸入一句話，系統不是直接回答你，而是經過六個精密的步驟來處理。"
      "這六個步驟就叫做「Pipeline」（管道），每一步都有特定的任務。")

    analogy_box("什麼是 Pipeline（管道）？",
        "想像一條工廠的生產線：原料（你的問題）從一端進去，經過六個加工站，"
        "最後出來成品（AI 的回覆）。每個加工站只負責自己的工作，"
        "做完就把半成品傳給下一站。這就是 Pipeline。")

    add_image(make_svg_02_pipeline())
    pi("圖 2：六步驟 Pipeline 全景圖。每一步都有對應的 temperature（隨機度）參數。")

    tip_box("什麼是 Temperature（溫度）？",
        "Temperature 是控制 AI 回答有多「有創意」的參數。\n"
        "• temp=0.0：完全固定，每次回答都一樣（像背書）\n"
        "• temp=0.1：幾乎固定，適合需要精確判斷的地方（像考試）\n"
        "• temp=0.5：有一點變化，適合自然對話（像聊天）\n"
        "• temp=1.0：非常隨機，很有創意但可能離題（像即興發揮）\n"
        "注意到 Pipeline 裡不同步驟用不同 temperature：判斷意圖要精確（0.1），"
        "回覆使用者要自然（0.3~0.5）。")

    h2("Step 1：Echo（即時確認）— temp=0.5")
    p("使用者說了一句話後，系統第一件事是立刻回應，讓使用者知道「我聽到了，正在處理」。"
      "這就像餐廳服務生聽到你的點餐後說「好的，我馬上幫您處理」。")

    prompt_box("PROMPT_ECHO（advisor.py L51-85）",
        '請你根據使用者的訊息和對話脈絡，用繁體中文回應。\n\n'
        '## 規則\n'
        '- 簡潔回應\n'
        '- 如果使用者的訊息需要查詢工具：概括需求，表達「已收到、正在處理」\n'
        '- 如果是一般對話（打招呼、感謝、閒聊）：直接給出友善回應\n'
        '- 不要回答營養問題本身，不要提供營養數據或建議\n'
        '- 語氣親切自然\n\n'
        '## 範例\n'
        '使用者：「營養夠嗎」→ 收到，我先查看你的個人資料...\n'
        '使用者：「最近常覺得疲勞」→ 了解，我來查詢與疲勞相關的營養素...\n'
        '使用者：「你好」→ 你好！我是你的 AI 營養顧問...')

    p("逐行解說：")
    bullet("「請你根據使用者的訊息和對話脈絡，用繁體中文回應」→ 告訴 AI 要用中文回覆")
    bullet("「不要回答營養問題本身」→ 這一步只是確認，真正的答案在後面的步驟才生成")
    bullet("「{memory_block}」→ 會被替換成記憶內容，讓 AI 知道之前聊過什麼")
    bullet("「{history}」→ 會被替換成最近的對話歷史")
    bullet("「{message}」→ 會被替換成使用者這次說的話")

    h2("Step 2：Phase1（意圖判定）— temp=0.1")
    p("第二步是最關鍵的——系統要搞清楚使用者到底想做什麼。這就像領班聽完你的需求後，"
      "要判斷你是「要點菜」「要退菜」「要結帳」還是「只是聊天」。")

    add_image(make_svg_03_phase1_intent())
    pi("圖 3：Phase1 可以判斷出 15 種不同的意圖。")

    p("系統會送一段很長的 Prompt 給 LLM，裡面包含：")
    bullet("所有 15 種意圖的說明表格（告訴 AI 有哪些選項）")
    bullet("營養素欄位名對照表（讓 AI 知道「鈣」對應 calcium_per_100g）")
    bullet("確認詞處理規則（使用者說「好」「可以」→ 確認上一輪的動作）")
    bullet("複合任務範例（如「分析食譜，加到行事曆」→ 同時兩個意圖）")
    bullet("對話歷史和記憶（讓 AI 有上下文）")

    tip_box("為什麼用 temp=0.1？",
        "意圖判定必須非常精確——如果 AI 誤判你要「刪除食譜」而你只是想「查詢食譜」，"
        "後果很嚴重。所以 temperature 設很低（0.1），讓 AI 盡量給出最確定的答案。")

    h2("Step 3：Phase2（工具調用生成）— temp=0.1")
    p("知道使用者想做什麼之後，接下來要決定具體怎麼做——呼叫哪個工具、傳什麼參數。")

    add_image(make_svg_04_phase2_waterfall())
    pi("圖 4：Phase2 先嘗試系統生成，失敗才用 LLM。")

    p("這一步有個聰明的設計：先讓系統自己嘗試生成工具呼叫（不用 LLM），"
      "只有系統搞不定時才呼叫 LLM。為什麼？因為系統生成更快、更穩定、不會出錯。")

    prompt_box("PROMPT_PHASE2（advisor.py L196-255）",
        '你是工具調用生成模組。根據意圖判定結果和已知的使用者資訊，生成具體的工具呼叫指令。\n\n'
        '### 意圖判定結論（來自上一步）\n{phase1_json}\n\n'
        '### 系統已儲存的使用者資訊（結構化資料）\n{structured_store_json}\n\n'
        '### 可用工具與參數\n'
        '| 工具名 | 必要參數 | 可選參數 | 說明 |\n'
        '| search_symptom | symptom(中文) | — | 症狀→化合物→推薦食物 |\n'
        '| search_food | food_name | — | 搜尋食物營養成分 |\n'
        '... (共 15 個工具)\n\n'
        '### 輸出格式\n'
        '{"tool_calls":[{"tool":"tool_name","args":{"param":"value"}}],\n'
        ' "still_missing":["param_name"],\n'
        ' "ask_user":"若有缺失參數，寫一句自然的追問語句"}')

    h2("Step 4：Execute（執行工具）")
    p("這一步是真正「做事」的步驟。系統根據 Phase2 的指令，呼叫對應的工具。")
    p("所有 15 個工具都在 _TOOL_DISPATCH 字典中註冊：")
    code('_TOOL_DISPATCH = {\n'
         '    "search_food": _tool_search_food,\n'
         '    "analyze_recipe": _tool_analyze_recipe,\n'
         '    "search_symptom": _tool_search_symptom,\n'
         '    ... # 共 15 個工具\n'
         '}')
    p("每個工具函數負責：(1) 組裝 HTTP 請求 (2) 呼叫 Backend API (3) 解析回傳資料 (4) 格式化結果")

    h2("Step 5：Analysis（分析回覆）— temp=0.3")
    p("工具執行完畢後，AI 要根據查到的資料寫一份分析報告給使用者。")

    prompt_box("PROMPT_ANALYSIS（advisor.py L262-310）",
        '請你根據工具查詢結果，用繁體中文提供專業營養分析回覆。\n\n'
        '## 核心規則\n'
        '- 所有數據必須來自下方的「工具執行結果」，禁止從訓練資料引用\n'
        '- 禁止使用「根據一般營養學原理」等包裝語句\n'
        '- 前端已自動渲染工具結果的表格和卡片，你不需要列出原始數據清單\n'
        '- 你的角色是「分析師」：解讀數據的意義，不是複述數據')

    h2("Step 6：Dual Memory（雙記憶）— temp=0.2")
    p("最後一步是記憶更新。系統會同時執行兩個記憶 Prompt：")
    bullet("敘事記憶（Narrative Memory）：記錄對話流程，不記具體數據")
    bullet("結構化記憶（Structured Store）：提取具體的資料欄位（年齡、查過的食物等）")

    p("這兩種記憶就像人的大腦：敘事記憶是「這個客人先問了疲勞、又問了鐵質、對補鐵很有興趣」（劇情），"
      "結構化記憶是「25歲、女性、體重55公斤、查過鈣和鐵」（事實）。兩者互補，讓下次對話時 AI 更了解你。")

    page_break()


def build_ch03():
    h1("第 3 章：十五個工具全解析")

    p("前一章提到，Pipeline 的 Step 4 會「執行工具」。那這些工具到底有哪些？每個工具做什麼？"
      "如果是人在操作，會看到什麼？讓我們一個一個來看。")

    add_image(make_svg_05_tool_categories())
    pi("圖 5：15 個工具分成四大類——搜尋類、分析類、CRUD 類、個人資料類。")

    tip_box("什麼是 CRUD？",
        "CRUD 是四個英文字的縮寫：Create（新增）、Read（讀取）、Update（更新）、Delete（刪除）。"
        "這四個動作就是資料庫操作的全部了——就像你管理一本通訊錄，只有加人、查人、改資料、刪人四種操作。")

    h2("3.1 search_food — 搜尋食物營養成分")
    add_image(make_svg_06_search_food_chain())
    pi("圖 6：search_food 的完整通信鏈。")

    add_table(
        ["項目", "如果是人操作", "系統自動化"],
        [
            ["輸入", "在搜尋框輸入「雞蛋」", "message 中提到雞蛋 → Phase1 判定 intent=search_food"],
            ["過程", "按搜尋按鈕，等結果", "Control 呼叫 Backend GET /food/search?q=雞蛋"],
            ["回傳", "看到一張營養成分表", "JSON 包含：食物名、110+ 營養素欄位、來源"],
            ["Recovery", "換個說法再查", "自癒引擎自動嘗試同義詞"],
        ]
    )

    h2("3.2 search_symptom — 症狀搜尋")
    add_image(make_svg_07_symptom_pipeline())
    pi("圖 7：症狀搜尋的五階段管道。")

    add_table(
        ["項目", "如果是人操作", "系統自動化"],
        [
            ["輸入", "輸入「最近很累又頭痛」", "LLM 自動拆解為多個症狀關鍵字"],
            ["過程", "查症狀對照表 + 語義搜尋", "5 階段管道：解析→雙路→篩選→查食物→彙總"],
            ["回傳", "相關化合物和推薦食物清單", "化合物、食物、相似度分數、來源徽章"],
        ]
    )

    h2("3.3 analyze_recipe — 食譜分析")
    add_image(make_svg_08_recipe_analysis())
    pi("圖 8：食譜分析的二階段流程。")

    add_table(
        ["項目", "如果是人操作", "系統自動化"],
        [
            ["輸入", "貼上食譜文字", "Phase1 判定 intent=analyze_recipe"],
            ["過程", "營養師手動查每個食材", "Backend 自動提取→匹配→計算→DRI 比對"],
            ["回傳", "營養分析報告", "每個食材營養、總營養、DRI 缺口、推薦食物"],
        ]
    )

    h2("3.4 search_graphrag — 學術論文問答")
    add_image(make_svg_09_graphrag_chain())
    pi("圖 9：GraphRAG 學術問答通信鏈。")

    h2("3.5 delete_saved_recipe — 刪除食譜")
    add_image(make_svg_10_delete_recipe())
    pi("圖 10：刪除食譜的兩步驟安全機制。")

    p("刪除食譜有特殊的安全設計：第一次呼叫只會「預覽」匹配的食譜（confirmed=false），"
      "使用者確認後第二次呼叫才會真正刪除（confirmed=true）。這是為了防止誤刪。")

    h2("3.6-3.15 其他工具概覽")
    add_table(
        ["工具", "功能", "Backend API", "特殊邏輯"],
        [
            ["get_nutrient_ranking", "營養素含量排名", "GET /health/nutrients/{field}/top-foods", "支援合計欄位"],
            ["browse_taiwan_food", "瀏覽食品資料庫", "GET /food/taiwan/browse", "可選搜尋/分類"],
            ["get_literature_papers", "瀏覽學術論文", "GET /food/literature/papers", "—"],
            ["synthesize_advice", "綜合營養建議", "—", "從已有食譜資料生成建議"],
            ["get_saved_recipes", "查已存食譜", "GET /recipe/saved", "—"],
            ["get_calendar_entries", "查行事曆", "GET /calendar/entries", "日期範圍查詢"],
            ["add_to_calendar", "加行事曆", "POST /recipe/save + /calendar", "先存食譜再加行事曆"],
            ["delete_calendar_entry", "刪行事曆", "GET + DELETE /calendar", "先查再刪"],
            ["get_user_profile", "查個人資料", "GET /auth/me", "—"],
            ["update_user_profile", "更新資料", "PUT /profile", "gender 需轉換"],
        ]
    )

    page_break()


def build_ch04():
    h1("第 4 章：九個 Prompt 範本——AI 的劇本")

    p("Prompt 是什麼？簡單來說，Prompt 就是我們寫給 AI 的「指令」或「劇本」。"
      "AI 本身不會自動知道該怎麼做，它需要我們告訴它：你的角色是什麼、規則是什麼、該怎麼回答。")

    analogy_box("Prompt 就像演員的劇本",
        "AI 就像一個萬能的演員——它能演任何角色，但前提是你要給它劇本。\n"
        "Prompt 就是這份劇本：告訴 AI「你現在是營養顧問」「你要用中文回答」「你不能編造數據」等等。\n"
        "劇本寫得越清楚，演員的表現就越好。")

    add_image(make_svg_11_prompt_lifecycle())
    pi("圖 11：9 個 Prompt 在 Pipeline 中的位置。")

    h2("4.1 PROMPT_ECHO（即時確認 — L51）")
    p("用途：收到使用者訊息後，立刻回一句確認。溫度 0.5，讓回覆自然但不偏離。")
    p("變數：{memory_block}（記憶）、{history}（對話歷史）、{message}（使用者訊息）")

    h2("4.2 PROMPT_PHASE1（意圖判定 — L92）")
    p("這是最長的一個 Prompt（約 100 行），因為它要告訴 AI 所有可能的意圖、"
      "所有的規則、所有的參數格式。溫度 0.1，要求精確判斷。")
    p("關鍵設計：")
    bullet("15 種意圖表格 + 每個意圖的必要參數")
    bullet("營養素欄位名對照（中文→英文欄位名）")
    bullet("確認詞處理（「好」「可以」→ 上一輪動作）")
    bullet("複合任務範例（多意圖同時觸發）")
    bullet("輸出格式嚴格要求 JSON")

    h2("4.3 PROMPT_PHASE2（工具調用 — L196）")
    p("把 Phase1 的意圖轉成具體的工具呼叫指令。會參考 structured_store（記憶中的已知資訊），"
      "盡量減少追問使用者。")

    h2("4.4 PROMPT_ANALYSIS（分析回覆 — L262）")
    p("根據工具結果寫分析。核心規則是「只引用工具結果中的數據，禁止從訓練資料編造」。")

    h2("4.5 PROMPT_NARRATIVE_SUMMARY（敘事記憶 — L317）")
    p("壓縮對話歷史為摘要。規則是只記「流程和態度」，不記「具體數字」。")

    h2("4.6 PROMPT_STRUCTURED_EXTRACT（結構化提取 — L341）")
    p("從對話中提取 11 種固定欄位的值。輸出嚴格 JSON。")

    h2("4.7 PROMPT_CLASSIFY_REPLY（中斷偵測 — L3318）")
    prompt_box("PROMPT_CLASSIFY_REPLY",
        'AI 正在執行「{intent_zh}」，等使用者提供「{missing_desc}」。\n\n'
        '使用者說：「{message}」\n\n'
        '這句話是在回答上述問題嗎？\n'
        'A = 是（提供了所需的資訊）\n'
        'B = 否（在說別的事、問新問題、放棄、或換話題）\n\n'
        '只輸出 A 或 B：')
    p("只需要 LLM 回一個字母（A 或 B），非常快。用於判斷使用者是否中斷了當前任務。")

    h2("4.8 PROMPT_CLASSIFY_RESUME（恢復偵測 — L3328）")
    prompt_box("PROMPT_CLASSIFY_RESUME",
        '使用者之前有未完成的任務：「{intent_zh}」（還需要：{missing_desc}）。\n\n'
        '使用者現在說：「{message}」\n\n'
        '這句話是想回到那個未完成的任務嗎？\nA = 是\nB = 否\n\n只輸出 A 或 B：')

    h2("4.9 PROMPT_EXTRACT_PARAMS（參數提取 — L3638）")
    prompt_box("PROMPT_EXTRACT_PARAMS",
        '從使用者訊息中提取以下欄位的值。\n'
        '只輸出嚴格 JSON，不要加其他文字。找不到的欄位填 null。\n\n'
        '需要提取的欄位：\n{fields_desc}\n\n'
        '使用者訊息：{message}\n\nJSON：')
    p("這是 regex 提取失敗時的 LLM fallback。先嘗試用正則表達式提取（更快更準），失敗才用 LLM。")

    page_break()


def build_ch05():
    h1("第 5 章：SSE 通信協議——即時推送的秘密")

    p("當你在 AI 顧問頁面輸入問題後，你會看到回覆一點一點地「跑」出來——"
      "這不是因為系統很慢，而是因為它在用一種叫做 SSE 的技術，把結果即時推送給你。")

    analogy_box("SSE 就像廣播電台",
        "傳統的網路請求像打電話：你問一個問題，對方想好了才一次回答完。\n"
        "SSE（Server-Sent Events）像聽廣播：你接上頻道後，電台持續播送內容。\n"
        "這樣你就不用乾等——AI 想到什麼就立刻推過來，你邊看邊等。")

    add_image(make_svg_12_sse_timeline())
    pi("圖 12：一次查詢中 SSE 事件的完整時間軸。")

    h2("SSE 格式長什麼樣？")
    p("每個事件由兩行組成：")
    code('event: reply\ndata: {"text": "根據查詢結果..."}\n')
    p("第一行 event: 告訴前端「這是什麼類型的事件」，第二行 data: 是具體的資料（JSON 格式）。"
      "兩個事件之間用一個空行隔開。")

    h2("18 種 SSE 事件詳解")
    add_table(
        ["事件", "用途", "資料內容"],
        [
            ["echo", "即時確認回應", '{"text": "收到，正在處理..."}'],
            ["thinking_start", "開始思考", '{"step": "phase1"}'],
            ["thinking_delta", "思考進度", '{"text": "分析中..."}'],
            ["thinking_end", "思考完成", '{"step": "phase1"}'],
            ["phase_status", "Pipeline 階段", '{"phase": "intent", "status": "done"}'],
            ["tool_start", "開始執行工具", '{"tool": "search_food", "args": {...}}'],
            ["tool_progress", "工具進度", '{"tool": "search_food", "progress": 50}'],
            ["tool_result", "工具原始結果", '{"tool": "search_food", "result": {...}}'],
            ["tool_done", "工具完成", '{"tool": "search_food"}'],
            ["reply", "最終回覆", '{"text": "分析結果是..."}'],
            ["workflow_state", "工作流狀態", '{"phase": "collecting", "missing": [...]}'],
            ["collecting_state", "收集參數狀態", '{"intent": "...", "missing": [...]}'],
            ["memory", "記憶更新", '{"text": "壓縮後的記憶"}'],
            ["narrative_memory", "敘事記憶", '{"text": "對話摘要"}'],
            ["structured_store", "結構化記憶", '{"user_facts": {...}}'],
            ["session_id", "Session 識別", '{"session_id": "abc123"}'],
            ["checkpoint_saved", "檢查點存檔", '{"checkpoint_id": "..."}'],
            ["error", "錯誤", '{"message": "發生錯誤"}'],
        ]
    )

    h2("前端如何接收？advisorApi.js 程式碼走讀")
    p("前端用原生的 fetch API 建立 SSE 連線，然後用 ReadableStream 逐行讀取：")
    code('const response = await fetch("/api/advisor/chat", {\n'
         '  method: "POST",\n'
         '  headers: { "Content-Type": "application/json",\n'
         '             Authorization: `Bearer ${token}` },\n'
         '  body: JSON.stringify({ message, history, ... }),\n'
         '  signal: controller.signal,  // AbortController\n'
         '});\n'
         '\n'
         'const reader = response.body.getReader();\n'
         'const decoder = new TextDecoder();\n'
         'let buffer = "";\n'
         '\n'
         'while (true) {\n'
         '  const { done, value } = await reader.read();\n'
         '  if (done) break;\n'
         '  buffer += decoder.decode(value, { stream: true });\n'
         '  // 解析 buffer 中的 event: / data: 行...\n'
         '}')

    p("逐行解說：")
    bullet("fetch 建立 POST 請求，帶上 JWT token 做身份驗證")
    bullet("signal: controller.signal → 讓使用者可以按取消按鈕中斷連線")
    bullet("response.body.getReader() → 取得串流讀取器，不等全部完成就開始讀")
    bullet("decoder.decode(value, { stream: true }) → 把位元組轉成文字，stream:true 表示還有更多")
    bullet("解析邏輯：逐行掃描，遇到 event: 記住事件類型，遇到 data: 呼叫對應的 callback")

    page_break()


def build_ch06():
    h1("第 6 章：中斷處理六個派別——AI 如何應對使用者突然換話題")

    p("想像你正在跟朋友說一件事，說到一半朋友突然問了另一個問題——你會怎麼辦？"
      "是繼續說原來的？還是先回答新問題？這就是「中斷處理」要解決的問題。")
    p("在 AI 顧問系統中，使用者隨時可能換話題、放棄、或回到之前的問題。"
      "系統用了六種不同的「派別」來處理這些情況：")

    add_image(make_svg_13_six_schools())
    pi("圖 13：六個派別的總覽。每派有自己的武器和戰場。")

    h2("6.1 智慧派：_detect_interruption_smart — LLM 三態判斷")

    p("智慧派是最高層的判斷者。它同時使用 LLM 和正則表達式，比較兩者的結果：")

    add_image(make_svg_14_tristate())
    pi("圖 14：LLM 和 Regex 交叉驗證的三態決策。")

    p("為什麼要交叉驗證？")
    bullet("LLM 理解語意但偶爾判斷錯誤")
    bullet("Regex 速度快但只看表面模式")
    bullet("兩者一致 → 高信心度，直接採用")
    bullet("兩者不一致 → ambiguous（模糊），需要更謹慎的處理")

    code('async def _detect_interruption_smart(message, ws):\n'
         '    # 確認詞直接視為回答\n'
         '    if text in ("好", "好啊", "可以", ...):\n'
         '        return "answer"\n'
         '    \n'
         '    # LLM 判定\n'
         '    result = await _call_gemma([prompt], temperature=0.0)\n'
         '    llm_interrupt = (letter == "B")\n'
         '    \n'
         '    # Regex 判定\n'
         '    regex_interrupt = _detect_interruption_regex(message, ws)\n'
         '    \n'
         '    # 交叉比對\n'
         '    if llm_interrupt != regex_interrupt:\n'
         '        return "ambiguous"  # 不一致 → 模糊\n'
         '    return "interrupt" if llm_interrupt else "answer"')

    h2("6.2 規則派：_detect_interruption_regex — 正則比對")
    p("規則派用三把「武器」（正則表達式）來偵測中斷：")
    bullet("_QUESTION_RE：偵測問句（「？」「多少」「是什麼」「怎麼」等）")
    bullet("_RETURN_RE：偵測想回去的意圖（「回到剛才」「繼續之前」「那個食譜」）")
    bullet("_NEW_INTENT_RE：偵測新指令（「幫我查」「幫我記錄」「吃了」「加到行事曆」）")

    p("規則派還有一個精巧的設計——它會根據當前等待的參數類型來判斷。例如：")
    bullet("等 age（年齡）：使用者輸入「25 歲」→ 不是中斷")
    bullet("等 food_name（食物名）：使用者輸入「雞蛋」→ 不是中斷")
    bullet("等 recipe_text（食譜）：使用者輸入帶份量的文字 → 不是中斷")

    h2("6.3 態度派：_classify_collecting_response — 回覆分類")
    add_image(make_svg_15_collecting_state())
    pi("圖 15：三種回覆分類。")

    p("當系統在收集參數時，使用者的回覆可能不是在回答問題：")
    bullet("answer：正常回答（「鈣」「25歲」「今天」）→ 繼續收集")
    bullet("soft_rejection：猶豫（「不知道」「隨便」「都行」）→ 寬容一次，改說法再問")
    bullet("hard_abandon：明確放棄（「算了」「不要了」「取消」）→ 直接結束")

    h2("6.4 回歸派：_detect_resume_smart — 中斷恢復偵測")
    add_image(make_svg_16_resume_flow())
    pi("圖 16：恢復偵測流程——隱式優先於顯式。")

    p("如果使用者之前中斷了一個任務，回歸派負責偵測使用者是否想回來。")
    p("兩種恢復方式：")
    bullet("隱式恢復：使用者直接給出之前缺的參數（如之前問你年齡，你現在說「25」→ 自動恢復）")
    bullet("顯式恢復：使用者明確說「回到剛才」「繼續之前的」→ LLM 判斷是否是恢復")

    h2("6.5 快捷派：[RECIPE_CONFIRM] 食譜確認協議")
    add_image(make_svg_17_recipe_confirm())
    pi("圖 17：四種操作路徑。")

    p("食譜確認是唯一不經過 LLM 的快速通道。工具回傳的訊息中包含 [RECIPE_CONFIRM] 標記，"
      "前端偵測到後直接顯示食材確認 UI。")

    h2("6.6 前線派：Frontend AbortController")
    p("最後一道防線在前端。當使用者按下取消按鈕時：")
    code('const controller = new AbortController();\n'
         '// ... 建立連線時傳入 signal ...\n'
         'controller.abort(); // 使用者按取消 → 中斷 SSE 連線')
    p("前端呼叫 abort() → 瀏覽器中斷 HTTP 連線 → 後端收到斷線信號 → 停止處理。"
      "這是使用者能主動控制的唯一機制。")

    page_break()


def build_ch07():
    h1("第 7 章：複合問題處理——同時處理多個需求")

    p("使用者經常一句話裡包含多個需求：「幫我分析這個食譜，加到明天的行事曆，"
      "然後告訴我還缺什麼營養素」。這一句話裡有三個意圖！")

    add_image(make_svg_18_compound_dag())
    pi("圖 18：複合任務的 DAG 依賴圖。")

    analogy_box("什麼是 DAG（有向無環圖）？",
        "DAG 就像一份「做事順序表」。有些事情必須先做（先分析食譜才能加到行事曆），"
        "有些事情可以同時做（加行事曆和生成建議可以平行）。"
        "DAG 確保不會搞錯順序，也不會有「A 等 B、B 又等 A」的死結。")

    h2("多意圖偵測")
    p("Phase1 的輸出可以包含多個意圖：")
    code('{\n'
         '  "intents": ["analyze_recipe", "add_to_calendar", "synthesize_advice"],\n'
         '  "entities": {\n'
         '    "recipe_text": "雞胸肉 200g、花椰菜 100g",\n'
         '    "entry_date": "2026-05-21"\n'
         '  }\n'
         '}')

    h2("DAG 任務規劃")
    p("系統根據依賴規則建立 DAG：")
    bullet("analyze_recipe → add_to_calendar（行事曆需要食譜資料）")
    bullet("analyze_recipe → synthesize_advice（建議需要營養分析結果）")
    bullet("add_to_calendar 和 synthesize_advice 可以平行")

    h2("DAG 執行")
    p("按依賴層級執行：先執行沒有依賴的任務，完成後再執行依賴它的任務。同一層的任務可以平行。")

    h2("副作用傳播")
    p("前一個工具的結果可能影響後面的工具。例如 analyze_recipe 的結果中包含 recipe_id，"
      "add_to_calendar 需要用到這個 recipe_id。系統透過 _apply_side_effects() 自動傳播。")

    page_break()


def build_ch08():
    h1("第 8 章：自癒引擎——工具失敗了怎麼辦？")

    p("工具執行不一定每次都成功。最常見的問題是：使用者說「芝士」但資料庫裡叫「起司」，"
      "或者網路暫時斷線，或者參數格式不對。自癒引擎就是處理這些問題的「急救醫生」。")

    add_image(make_svg_19_recovery_engine())
    pi("圖 19：自癒引擎的三步驟流程。")

    analogy_box("自癒引擎像急救醫生",
        "工具出問題就像病人送急診：先診斷（什麼病？）→ 制定治療方案 → 執行治療。\n"
        "不同的病用不同的藥：查不到用同義詞、超時就重試、參數錯就修正。")

    h2("四種錯誤類型")
    add_table(
        ["錯誤類型", "情境", "信心度"],
        [
            ["empty_result", "搜尋食物查無結果", "0.9"],
            ["tool_timeout", "工具執行超時", "0.9"],
            ["param_mismatch", "營養素欄位名不匹配", "0.7"],
            ["backend_error", "後端服務 500 錯誤", "0.6-0.8"],
        ]
    )

    h2("四種恢復策略")
    add_table(
        ["策略", "做法", "適用情境"],
        [
            ["synonym_expand", "嘗試同義詞（最多 5 個）", "查食物查不到"],
            ["retry_timeout", "直接重試一次", "暫時性超時"],
            ["field_remap", "換營養素欄位名再試", "欄位名不匹配"],
            ["skip", "跳過此工具", "後端持續錯誤"],
        ]
    )

    h2("同義詞擴展機制")
    p("recovery.py 內建了 60+ 組常見台灣食物同義詞：")
    code('_COMMON_FOOD_SYNONYMS = {\n'
         '    "芝士": ["起司", "乾酪", "cheese"],\n'
         '    "番茄": ["西紅柿", "蕃茄"],\n'
         '    "馬鈴薯": ["洋芋", "土豆", "potato"],\n'
         '    "鳳梨": ["菠蘿"],\n'
         '    "高麗菜": ["甘藍菜", "包心菜", "捲心菜"],\n'
         '    "鮭魚": ["三文魚", "salmon"],\n'
         '    ... # 共 60+ 組\n'
         '}')

    p("還有食物名簡化功能——去掉量詞和修飾詞：")
    code('# 「一碗新鮮的花椰菜」\n'
         '# → 去掉量詞「一碗」\n'
         '# → 去掉修飾詞「新鮮的」\n'
         '# → 得到「花椰菜」')

    h2("SSE 通知")
    p("自癒過程中，前端會收到兩個 SSE 事件：")
    bullet("recovery_attempt：「正在嘗試恢復：芝士 → 起司」")
    bullet("recovery_result：「恢復成功」或「恢復失敗」")
    p("讓使用者知道系統在努力處理，不是卡住了。")

    page_break()


def build_ch09():
    h1("第 9 章：工作流狀態機——參數收集與追問機制")

    p("有些工具需要特定的參數才能執行。例如，搜尋食物需要知道「食物名稱」，"
      "查行事曆需要知道「日期範圍」。如果使用者沒提供這些資訊，系統就要追問。")

    add_image(make_svg_20_state_machine())
    pi("圖 20：工作流狀態機的轉換圖。")

    analogy_box("狀態機像點餐流程",
        "idle（等待）= 你坐在位子上還沒點餐\n"
        "collecting（收集）= 服務生在問你「要幾分熟？要什麼醬？」\n"
        "ready（準備好）= 所有選項都確認了\n"
        "executing（執行中）= 廚師正在做\n"
        "done（完成）= 菜上桌了")

    h2("WorkflowState 物件")
    code('class WorkflowState(BaseModel):\n'
         '    phase: str = "idle"          # 當前狀態\n'
         '    active_intent: str = ""      # 正在處理的意圖\n'
         '    collected: dict = {}         # 已收集的參數\n'
         '    missing: list[str] = []      # 還缺的參數\n'
         '    turn_count: int = 0          # 追問輪數\n'
         '    interrupted_from: str = ""   # 被中斷的來源\n'
         '    interrupted_state: dict = {} # 中斷時的狀態快照')

    h2("參數收集機制")
    p("每個意圖都有對應的 ParamSpec 定義，列出需要哪些參數：")
    code('WORKFLOW_DEFS = {\n'
         '    "search_symptom": [\n'
         '        ParamSpec("symptom", str, required=True,\n'
         '                  ask_zh="請問你有什麼症狀？"),\n'
         '    ],\n'
         '    "get_calendar": [\n'
         '        ParamSpec("start_date", str, required=True,\n'
         '                  ask_zh="請問你要查看哪段時間的行事曆？"),\n'
         '        ParamSpec("end_date", str, required=True),\n'
         '    ],\n'
         '    ... # 每個意圖都有定義\n'
         '}')

    h2("智慧填充")
    p("系統不會笨笨地追問——它會先嘗試自動填充：")
    bullet("_prefill_from_entities()：從 Phase1 提取的 entities 中取值")
    bullet("_fill_from_structured_store()：從記憶中的已知資訊取值（例如之前記住了年齡）")
    bullet("_prefill_from_profile()：從使用者個人資料取值")
    p("只有真的找不到時，才會追問使用者。")

    h2("追問限制")
    bullet("最多追問 3 輪（_MAX_ROUNDS=3），避免無限追問")
    bullet("遇到 soft_rejection（猶豫）→ 寬容一次，改說法再問")
    bullet("遇到 hard_abandon（放棄）→ 直接結束，回到 idle")
    bullet("遇到中斷（新問題）→ 存檔當前狀態到 interrupted_state，切去處理新問題")

    page_break()


def build_ch10():
    h1("第 10 章：驗證與自修復管道")

    p("AI（LLM）的回覆不一定是正確格式的。它可能把工具名打錯、用了不存在的參數、"
      "或者回傳的 JSON 格式有問題。驗證管道就是一連串的「品質檢查站」，確保 LLM 的輸出可用。")

    add_image(make_svg_21_validation_pipeline())
    pi("圖 21：六層驗證管道。")

    h2("層 1：意圖別名修正")
    p("LLM 經常用錯名字。例如它可能輸出 get_user_profile，但正確名字是 get_profile。")
    code('_INTENT_ALIASES = {\n'
         '    "get_user_profile": "get_profile",\n'
         '    "get_nutrient_ranking": "nutrient_ranking",\n'
         '    "browse_taiwan_food": "browse_food_database",\n'
         '    "get_saved_recipes": "get_recipes",\n'
         '    ... # 常見錯誤映射\n'
         '}')

    h2("層 2：意圖合法性過濾")
    p("過濾掉不在白名單中的意圖名，防止 LLM 發明不存在的工具。")

    h2("層 3：工具參數驗證")
    p("檢查必要參數是否存在、型別是否正確。缺少必要參數 → 進入 collecting 狀態追問。")

    h2("層 4：營養素欄位映射")
    p("使用者可能說「鈣」，但資料庫欄位名是 calcium_per_100g。系統有一張中英對照表自動轉換。")

    h2("層 5：品質過濾")
    p("檢查工具回傳結果是否為空 → 空結果觸發自癒引擎（Chapter 8）。")

    h2("層 6：JSON 三級 fallback")
    p("LLM 回傳的 JSON 可能格式不標準。系統有三級嘗試：")
    bullet("直接 json.loads() 解析")
    bullet("正則表達式提取 JSON 區塊")
    bullet("去掉 markdown 標記（```json ...```）再試")

    page_break()


def build_ch11():
    h1("第 11 章：雙記憶系統深入")

    p("人有短期記憶和長期記憶。AI 顧問也有兩種記憶，讓它能記住你是誰、你喜歡什麼、之前聊過什麼。")

    add_image(make_svg_22_memory_tree())
    pi("圖 22：雙記憶系統 + 記憶樹架構。")

    h2("敘事記憶（Narrative Memory）")
    p("記錄對話的「劇情」：")
    code('PROMPT_NARRATIVE_SUMMARY:\n'
         '"只記對話流程（誰問了什麼、AI 怎麼回的、話題轉折）"\n'
         '"不記具體數據（年齡、營養素數值等已由系統另存）"\n'
         '"記錄使用者的關注點"\n'
         '"記錄對話的邏輯鏈"')
    p("例如：「使用者先問了疲勞相關營養→AI建議補鐵→使用者又問哪些食物鐵最多→"
      "AI列出前10名→使用者對菠菜特別有興趣」")

    h2("結構化記憶（Structured Store）")
    p("記錄具體的「事實」，有 11 種固定欄位：")
    add_table(
        ["欄位", "記什麼", "保留上限"],
        [
            ["user_facts", "年齡、性別、體重、活動量", "覆蓋更新"],
            ["food_context", "搜尋過的食物、提到的食物", "追加"],
            ["recipe_entry", "分析過的食譜和營養", "—"],
            ["symptom_entry", "搜尋過的症狀", "最近 20 筆"],
            ["ranking_entry", "查過的營養素排名", "最近 20 筆"],
            ["calendar_update", "查過的行事曆日期", "最近 10 筆"],
            ["academic_entry", "查過的學術問題", "最近 20 筆"],
            ["papers_referenced", "引用過的論文", "去重追加"],
            ["dri_awareness", "提到的營養缺乏/過量", "覆蓋更新"],
            ["action_entry", "所有操作紀錄", "最近 30 筆"],
            ["saved_recipes_summary", "已存食譜摘要", "覆蓋更新"],
        ]
    )

    h2("記憶樹（Memory Tree）")
    p("三層架構：")
    bullet("Leaf（葉子）：單次對話的記憶片段")
    bullet("Session（會話）：一次對話的完整記憶")
    bullet("User（使用者）：跨越多次對話的長期記憶")
    p("使用 embedding 向量相似度來找到最相關的記憶，確保 AI 顧問記得你，即使隔了好幾天。")

    page_break()


def build_ch12():
    h1("第 12 章：前端頁面與元件")

    add_image(make_svg_23_advisor_page())
    pi("圖 23：AdvisorPage 的完整資料流。")

    h2("AdvisorPage — AI 營養顧問主頁")
    p("這是系統的預設首頁（/advisor），也是最複雜的前端元件。")
    bullet("SSE 連線：使用 advisorApi.js 建立，18 種事件各有 callback")
    bullet("IndexedDB：大量資料（完整對話歷史）存在瀏覽器的 IndexedDB")
    bullet("localStorage：最近 30 則對話存在 localStorage")
    bullet("ToolResultCard：根據工具類型渲染不同的 UI（表格、卡片、清單等）")
    bullet("ThinkingPanel：顯示 AI 的思考過程（step1~step6 進度）")

    add_image(make_svg_24_unified_analysis())
    pi("圖 24：整合分析頁的七步驟流程。")

    h2("UnifiedAnalysisPage — 整合分析頁")
    p("獨立的分析頁面（/analysis），不走 LLM Pipeline：")
    bullet("七步驟流程：貼食譜→提取→確認→查詢→計算→缺口→推薦")
    bullet("雙向拖曳：食材可拖到停用詞區，停用詞也可拖回來")
    bullet("食材點擊：點食材名開啟 TaiwanFoodModal 顯示完整營養")
    bullet("DRI 缺口表：缺乏列有「推薦食物」toggle")

    h2("其他頁面簡介")
    add_table(
        ["頁面", "路由", "功能"],
        [
            ["SymptomSearchPage", "/symptom-search", "症狀搜尋（5 階段管道）"],
            ["NutrientRankingPage", "/nutrient-ranking", "營養素排名（110+ 欄分 13 組）"],
            ["GraphRAGPage", "/graphrag", "學術論文 GraphRAG 問答"],
            ["DriGapPage", "/dri-gap", "DRI 缺口分析"],
            ["CalendarPage", "/calendar", "每日分析日曆"],
            ["SavedRecipesPage", "/recipes", "已儲存食譜"],
            ["FoodSearchPage", "/food-search", "食物搜尋"],
        ]
    )

    page_break()


def build_ch13():
    h1("第 13 章：完整通信鏈追蹤")

    p("前面的章節分別介紹了各個元件。現在讓我們追蹤一次完整的操作，"
      "從使用者點擊到資料庫查詢到畫面渲染，一步都不跳。")

    add_image(make_svg_25_recipe_e2e())
    pi("圖 25：食譜分析端到端完整通信鏈。")

    h2("範例：使用者輸入「幫我分析雞胸肉 200g 花椰菜 100g」")
    p("完整步驟：")
    bullet("1. 使用者在 AdvisorPage 輸入文字，按送出")
    bullet("2. advisorApi.js 建立 SSE 連線 → POST /api/advisor/chat")
    bullet("3. nginx 代理到 Control:8000 → POST /advisor/chat")
    bullet("4. Step1 Echo：LLM 回「收到，正在分析食譜」→ SSE echo 事件")
    bullet("5. Step2 Phase1：LLM 判定 intent=analyze_recipe → SSE phase_status")
    bullet("6. Step3 Phase2：系統生成 tool_calls=[{tool:analyze_recipe, args:{recipe_text:...}}]")
    bullet("7. Step4 Execute：Control 呼叫 Backend POST /recipe/calculate-nutrition")
    bullet("8. Backend 提取食材 → 查 taiwan_foods + FooDB → 按份量計算 → 比對 DRI")
    bullet("9. Backend 回傳 JSON（食材營養、總營養、DRI 缺口、推薦食物）")
    bullet("10. SSE tool_result 事件 → 前端渲染 ToolResultCard")
    bullet("11. Step5 Analysis：LLM 根據結果寫分析報告 → SSE reply 事件")
    bullet("12. Step6 DualMemory：並行更新敘事記憶 + 結構化記憶")
    bullet("13. 前端顯示完整結果：分析文字 + 營養表格 + DRI 缺口圖")

    add_image(make_svg_26_symptom_e2e())
    pi("圖 26：症狀搜尋端到端通信鏈。")

    add_image(make_svg_27_calendar_crud())
    pi("圖 27：行事曆 CRUD 通信鏈。")

    add_image(make_svg_28_food_search_e2e())
    pi("圖 28：食物搜尋端到端通信鏈。")

    page_break()


def build_ch14():
    h1("第 14 章：食材同義詞擴展 Pipeline")

    p("台灣用「番茄」，大陸用「西紅柿」，美國叫「tomato」。"
      "為了讓搜尋不受用詞差異影響，系統建立了一套自動同義詞擴展管道。")

    h2("中英雙向擴展流程")
    p("中文路徑：")
    bullet("1. MoeDict 中文字典查同義詞")
    bullet("2. Argos 翻譯成英文")
    bullet("3. WordNet 查英文同義詞和下位詞")
    bullet("4. 翻譯回中文")
    bullet("5. 去重 + 食物相關性驗證")

    p("英文路徑：")
    bullet("1. WordNet 查同義詞")
    bullet("2. 翻譯成中文")
    bullet("3. MoeDict 查中文同義詞")
    bullet("4. 翻譯回英文")
    bullet("5. 去重")

    h2("三路融合")
    p("MoeDict（中文字典）+ WordNet（英文字典）+ Argos（翻譯橋樑），三者融合產出同義詞。")
    p("目前成果：3,172 個食物 canonical，平均每個有 2.73 個同義詞，共 8,646 個同義詞。")

    warn_box("已知問題：Argos 翻譯對單字/短詞容易出錯（「小米」→「Mi」→ WordNet 吐 myocardial infarction），"
             "有 _validate_food_relevance 擋一部分但不能全擋。")

    page_break()


def build_ch15():
    h1("第 15 章：資料安全與部署")

    h2("JWT 身份驗證")
    p("JWT（JSON Web Token）就像一張「通行證」。使用者登入後拿到一張 token，"
      "之後每次操作都帶著這張 token 證明身份。")
    bullet("Token 有效期：7 天")
    bullet("密碼加密：bcrypt（不可逆加密）")
    bullet("所有 API 都需要帶 token（除了登入和註冊）")

    h2("審計日誌")
    p("所有涉及資料修改的操作都會記錄審計日誌：")
    bullet("誰做了什麼（JWT 中的使用者 ID）")
    bullet("修改前後的值")
    bullet("時間戳")
    bullet("日誌保留 90 天")

    h2("備份機制")
    bullet("stopword_store 每次寫入前自動快照（保留最近 10 份）")
    bullet("啟動時自動備份 food_synonym_dict.json")
    bullet("JSON 損毀時自動 rename .bak + 建空檔")

    h2("Docker 容器部署")
    add_table(
        ["服務", "容器化", "啟動方式"],
        [
            ["Frontend", "Docker", "docker run --network host recipe-frontend"],
            ["Control", "Docker", "docker run --network host + bind-mount"],
            ["Backend", "Docker", "docker run --network host recipe-backend"],
            ["PostgreSQL", "Docker", "docker run --network host recipe_ai_db"],
            ["Gemma LLM", "手動進程", "llama-server -m model.gguf --jinja"],
            ["GraphRAG", "手動進程", "python src/api_server.py"],
            ["Health Vector", "手動進程", "python server.py"],
        ]
    )

    page_break()


def build_appendix_a():
    h1("附錄 A：專有名詞中英對照表")
    add_table(
        ["英文", "中文", "說明"],
        [
            ["API", "應用程式介面", "程式之間的溝通規格"],
            ["SSE", "伺服器推送事件", "伺服器即時推送資料的技術"],
            ["LLM", "大型語言模型", "AI 的大腦，能理解和生成文字"],
            ["Pipeline", "管道/流水線", "一連串依序執行的步驟"],
            ["Temperature", "溫度", "控制 AI 回覆隨機度的參數"],
            ["Token", "令牌", "身份驗證的通行證（JWT）"],
            ["Prompt", "提示/指令", "給 AI 的劇本"],
            ["Intent", "意圖", "使用者想做什麼"],
            ["DAG", "有向無環圖", "表示依賴關係的圖"],
            ["CRUD", "增刪改查", "Create/Read/Update/Delete"],
            ["Webhook", "網路鉤子", "事件觸發的回呼 URL"],
            ["JSON", "JavaScript 物件標記", "結構化資料格式"],
            ["HTTP", "超文本傳輸協定", "網路上傳輸資料的標準方式"],
            ["REST", "表現層狀態轉換", "API 的設計風格"],
            ["SQL", "結構化查詢語言", "查詢資料庫用的語言"],
            ["DRI", "膳食營養素參考攝取量", "國家建議的每日營養攝取標準"],
            ["HPA", "國民健康署", "台灣衛生福利部的機構"],
            ["GraphRAG", "圖增強檢索生成", "結合知識圖譜的 AI 問答"],
            ["Embedding", "嵌入向量", "把文字轉成數字向量的技術"],
            ["ChromaDB", "向量資料庫", "存放和搜尋向量的資料庫"],
            ["Regex", "正則表達式", "文字模式匹配的工具"],
            ["Fallback", "備援/後備", "主要方式失敗時的替代方案"],
            ["Checkpoint", "檢查點", "存檔用的快照"],
            ["AbortController", "中止控制器", "取消網路請求的瀏覽器 API"],
            ["Graceful Degradation", "優雅降級", "部分功能壞了但整體還能用"],
        ]
    )


def build_appendix_b():
    h1("附錄 B：所有 18 種 SSE 事件完整表格")
    add_table(
        ["事件名", "觸發時機", "資料欄位", "前端處理"],
        [
            ["echo", "Step1 Echo 完成", "text", "顯示即時確認"],
            ["thinking_start", "開始 LLM 呼叫", "step", "顯示思考面板"],
            ["thinking_delta", "LLM 生成中", "text", "更新思考文字"],
            ["thinking_end", "LLM 呼叫完成", "step", "關閉思考面板"],
            ["phase_status", "Pipeline 階段切換", "phase, status", "更新進度條"],
            ["workflow_state", "狀態機變化", "phase, missing, ...", "更新 UI 狀態"],
            ["collecting_state", "收集參數中", "intent, missing", "顯示追問 UI"],
            ["tool_start", "開始執行工具", "tool, args", "顯示載入動畫"],
            ["tool_progress", "工具執行進度", "tool, progress", "更新進度"],
            ["tool_result", "工具結果", "tool, result", "渲染 ToolResultCard"],
            ["tool_done", "工具完成", "tool", "結束載入動畫"],
            ["recovery_attempt", "開始自癒", "tool, strategy", "顯示恢復提示"],
            ["recovery_result", "自癒結果", "success, explanation", "更新提示"],
            ["reply", "最終回覆", "text", "顯示分析文字"],
            ["memory", "記憶更新", "text", "更新 localStorage"],
            ["narrative_memory", "敘事記憶", "text", "更新 state"],
            ["structured_store", "結構化記憶", "{...}", "更新 state"],
            ["session_id", "Session ID", "session_id", "存 state"],
            ["checkpoint_saved", "檢查點", "checkpoint_id", "存 state"],
            ["error", "錯誤", "message", "顯示錯誤訊息"],
        ]
    )


def build_appendix_c():
    h1("附錄 C：所有 API 端點一覽")
    add_table(
        ["方法", "路徑", "功能", "需要 JWT"],
        [
            ["POST", "/auth/register", "註冊", "否"],
            ["POST", "/auth/login", "登入", "否"],
            ["GET", "/auth/me", "取得個人資料", "是"],
            ["POST", "/advisor/chat", "AI 顧問 SSE", "是"],
            ["POST", "/recipe/extract", "提取食材", "是"],
            ["POST", "/recipe/extract-with-amounts", "提取食材+份量", "是"],
            ["POST", "/recipe/reverse-search", "jieba 斷詞搜尋", "是"],
            ["POST", "/recipe/lookup", "食材匹配", "是"],
            ["POST", "/recipe/calculate-nutrition", "計算營養", "是"],
            ["POST", "/recipe/save", "儲存食譜", "是"],
            ["POST", "/recipe/analyze", "完整分析", "是"],
            ["GET", "/recipe/saved", "已存食譜", "是"],
            ["GET", "/food/search", "搜尋食物", "是"],
            ["GET", "/food/taiwan/browse", "瀏覽食品庫", "是"],
            ["GET", "/food/literature/papers", "學術論文", "是"],
            ["GET", "/food/symptom/{symptom}", "症狀搜尋", "是"],
            ["GET", "/health/nutrients/{field}/top-foods", "營養素排名", "是"],
            ["POST", "/health/dri-gap", "DRI 缺口", "是"],
            ["POST", "/health/effects-to-foods", "效果→食物", "是"],
            ["GET", "/dri/fields", "DRI 欄位", "是"],
            ["GET", "/calendar/entries", "行事曆查詢", "是"],
            ["POST", "/calendar/entries", "新增行事曆", "是"],
            ["DELETE", "/calendar/entries/{id}", "刪除行事曆", "是"],
            ["GET", "/stopwords", "停用詞列表", "是"],
            ["POST", "/stopwords", "新增停用詞", "是"],
            ["DELETE", "/stopwords/{word}", "刪除停用詞", "是"],
            ["PUT", "/profile", "更新個人資料", "是"],
        ]
    )


def build_appendix_d():
    h1("附錄 D：HPA 第八版 DRI 33 項營養素")
    add_table(
        ["營養素", "DB 欄位名", "可計算缺口", "說明"],
        [
            ["熱量", "cal_per_100g", "是", ""],
            ["蛋白質", "protein_per_100g", "是", ""],
            ["碳水化合物", "carbon_per_100g", "是", ""],
            ["脂肪", "fats_per_100g", "是", ""],
            ["膳食纖維", "dietary_fiber_per_100g", "是", ""],
            ["鈣", "calcium_per_100g", "是", ""],
            ["鐵", "iron_per_100g", "是", ""],
            ["鋅", "zinc_per_100g", "是", ""],
            ["鎂", "magnesium_per_100g", "是", ""],
            ["鉀", "potassium_per_100g", "是", ""],
            ["鈉", "sodium_per_100g", "是", ""],
            ["磷", "phosphorus_per_100g", "是", ""],
            ["錳", "manganese_per_100g", "是", ""],
            ["維生素 A", "retinol_ug", "是", ""],
            ["維生素 D", "vitamin_d_ug", "是", ""],
            ["維生素 E", "vitamin_e_mg", "是", ""],
            ["維生素 K", "vitamin_k_ug", "是", "合計 K1+K2MK4+K2MK7"],
            ["維生素 C", "vitamin_c_mg", "是", ""],
            ["維生素 B1", "thiamine_mg", "是", ""],
            ["維生素 B2", "riboflavin_mg", "是", ""],
            ["維生素 B6", "vitamin_b6_mg", "是", ""],
            ["維生素 B12", "vitamin_b12_ug", "是", ""],
            ["菸鹼素", "niacin_mg", "是", ""],
            ["葉酸", "folate_ug", "是", ""],
            ["n-3 多元不飽和", "n3_pufa_total_mg", "是", "合計 ALA+EPA+DHA"],
            ["n-6 多元不飽和", "n6_pufa_total_mg", "是", ""],
            ["反式脂肪", "trans_fat_per_100g", "是", ""],
            ["膽素", "—", "否", "DB 無欄位"],
            ["生物素", "—", "否", "DB 無欄位"],
            ["泛酸", "—", "否", "DB 無欄位"],
            ["碘", "—", "否", "DB 無欄位"],
            ["硒", "—", "否", "DB 無欄位"],
            ["氟", "—", "否", "DB 無欄位"],
        ]
    )


# ═══════════════════════════════════════════════════════════
# Section E: Build & Main
# ═══════════════════════════════════════════════════════════

def build():
    print("Building cover...")
    build_cover()

    chapters = [
        ("ch00", build_ch00),
        ("ch01", build_ch01),
        ("ch02", build_ch02),
        ("ch03", build_ch03),
        ("ch04", build_ch04),
        ("ch05", build_ch05),
        ("ch06", build_ch06),
        ("ch07", build_ch07),
        ("ch08", build_ch08),
        ("ch09", build_ch09),
        ("ch10", build_ch10),
        ("ch11", build_ch11),
        ("ch12", build_ch12),
        ("ch13", build_ch13),
        ("ch14", build_ch14),
        ("ch15", build_ch15),
    ]
    for name, fn in chapters:
        print(f"Building {name}...")
        fn()

    appendices = [
        ("appendix_a", build_appendix_a),
        ("appendix_b", build_appendix_b),
        ("appendix_c", build_appendix_c),
        ("appendix_d", build_appendix_d),
    ]
    for name, fn in appendices:
        print(f"Building {name}...")
        fn()

    print(f"Saving to {OUT}...")
    doc.save(OUT)
    size_mb = os.path.getsize(OUT) / 1024 / 1024
    print(f"Done! File size: {size_mb:.1f} MB")


if __name__ == "__main__":
    build()
