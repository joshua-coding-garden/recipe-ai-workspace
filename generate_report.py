#!/usr/bin/env python3
"""生成 AI 營養顧問 Tool Calling 技術報告 (.docx)"""

import os, glob
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import cairosvg, io

BASE = Path("/home/aiiauser/JM/0325new")
SVG_DIR = BASE / "報告圖表"
OUT = BASE / "AI營養顧問_Tool_Calling_技術報告.docx"

doc = Document()

# ─── Styles ───
style = doc.styles['Normal']
style.font.name = 'Microsoft JhengHei'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.35
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')

for lvl, sz, clr in [('Heading 1', 22, '1A5276'), ('Heading 2', 16, '2E86C1'), ('Heading 3', 13, '2874A6')]:
    s = doc.styles[lvl]
    s.font.size = Pt(sz)
    s.font.color.rgb = RGBColor.from_string(clr)
    s.font.bold = True
    s.font.name = 'Microsoft JhengHei'
    s.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
    s.paragraph_format.space_before = Pt(18 if sz > 15 else 12)
    s.paragraph_format.space_after = Pt(8)

code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = 'Consolas'
code_style.font.size = Pt(9)
code_style.paragraph_format.space_before = Pt(4)
code_style.paragraph_format.space_after = Pt(4)
code_style.paragraph_format.left_indent = Cm(0.5)
fmt = code_style.element
shd = fmt.makeelement(qn('w:pPr'), {})
code_style.paragraph_format.element.append(
    code_style.paragraph_format.element.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'F5F5F5'
    })
)

# ─── Helpers ───
def h1(text): doc.add_heading(text, level=1)
def h2(text): doc.add_heading(text, level=2)
def h3(text): doc.add_heading(text, level=3)
def p(text, bold=False, italic=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    return para
def pb(text): return p(text, bold=True)
def code(text):
    for line in text.strip().split('\n'):
        para = doc.add_paragraph(style='CodeBlock')
        run = para.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
def bullet(text):
    para = doc.add_paragraph(style='List Bullet')
    para.text = text
    return para
def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for par in cell.paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in par.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            cell.text = str(val)
            for par in cell.paragraphs:
                for run in par.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

def add_svg(filename, width_inches=6.0):
    svg_path = SVG_DIR / filename
    if not svg_path.exists():
        p(f"[圖表未找到: {filename}]", italic=True)
        return
    try:
        png_data = cairosvg.svg2png(url=str(svg_path), output_width=int(width_inches*150))
        doc.add_picture(io.BytesIO(png_data), width=Inches(width_inches))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        p(f"[圖表載入失敗: {filename} — {e}]", italic=True)

def page_break():
    doc.add_page_break()


# ══════════════════════════════════════════════
# 封面
# ══════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('AI 營養顧問\nTool Calling 機制\n完整技術報告')
run.font.size = Pt(32)
run.bold = True
run.font.color.rgb = RGBColor.from_string('1A5276')

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Recipe AI 系統 — advisor.py 完整解析\n涵蓋 13 個工具 × 6 個 Prompt × 4 步 LLM Pipeline\n所有通信協議、程式碼對應、逐行解釋')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor.from_string('5D6D7E')

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = date_p.add_run('2026 年 5 月 11 日')
r.font.size = Pt(12)

page_break()

# ══════════════════════════════════════════════
# 第 1 章：導論
# ══════════════════════════════════════════════
h1('第一章：導論')
p('這份報告的主角是一個「AI 營養顧問聊天機器人」。你可以把它想成：你在手機上跟一個很懂營養的朋友聊天，你說「今天午餐吃了雞腿便當」，它就幫你記錄下來，還告訴你吃了多少熱量、蛋白質夠不夠。')
p('這個系統不只是一個聊天視窗。它的背後連接了 5 個不同的服務程式，就像一家餐廳裡有前台接待、廚房、倉庫、圖書館和翻譯員，各司其職。當你說一句話，系統會：')
bullet('第一步：先回一句「收到，正在幫你處理」（讓你知道它在忙）')
bullet('第二步：分析你到底想做什麼（記錄飲食？查營養？問症狀？）')
bullet('第三步：決定要呼叫哪些「工具」去查資料')
bullet('第四步：把查到的結果整理成你看得懂的分析報告')

p('這份報告會完整拆解這個系統的每一個零件：它有哪些工具、每個工具怎麼運作、程式碼每一行在做什麼、如果是真人手動操作會怎麼做。所有的專有名詞，我們都會先用白話解釋，再進入技術細節。')

h2('1.1 報告涵蓋範圍')
bullet('13 個工具（Tool）的完整調用機制')
bullet('6 個 Prompt（給 AI 的指令模板）')
bullet('4 步 LLM Pipeline（AI 處理流程）')
bullet('完整通信協議（從你按 Enter 到畫面顯示結果的每一步）')
bullet('每次 AI 被呼叫時的參數限制（能處理多少文字、多「隨機」）')
bullet('工作流程狀態機（系統怎麼追蹤對話進度）')

h2('1.2 核心程式檔案')
add_table(
    ['檔案', '行數', '角色'],
    [
        ['advisor.py', '1950', '核心：6 個 Prompt + 13 個工具 + 4 步 Pipeline + 狀態機'],
        ['advisorApi.js', '99', '前端 SSE 連線客戶端'],
        ['AdvisorPage.jsx', '350+', '前端頁面：聊天介面 + 狀態管理'],
        ['ToolResultCard.jsx', '695', '13 種工具結果的專屬渲染元件'],
        ['config.py', '41', 'Gemma LLM 連線設定'],
        ['dri_gap_service.py', '200+', 'DRI 缺口分析服務'],
    ],
    col_widths=[2.5, 0.6, 3.5]
)

page_break()

# ══════════════════════════════════════════════
# 第 2 章：專有名詞解釋
# ══════════════════════════════════════════════
h1('第二章：專有名詞解釋')
p('在進入技術內容之前，我們先把所有會用到的「術語」用最白話的方式解釋一遍。每個名詞都會附上生活中的比喻，讓你一看就懂。')

terms = [
    ('LLM（大型語言模型）', 'Large Language Model',
     '一種經過大量文字訓練的 AI 程式，能理解人話並生成回應。',
     '想像一個讀過全世界所有書的超級秘書，你跟它說話，它能理解意思並回答。本系統用的是 Google 的 Gemma 模型。'),
    ('SSE（伺服器推送事件）', 'Server-Sent Events',
     '一種讓伺服器「主動」把資料一段一段送到瀏覽器的技術。',
     '像廣播電台：你調好頻道（建立連線），電台就持續播送消息給你，不需要你每次都打電話去問。普通的網頁請求像寄信（一來一往），SSE 像聽廣播（單向持續串流）。'),
    ('JWT（JSON Web Token）', 'JSON Web Token',
     '一串加密過的文字，用來證明「你是誰」。',
     '像電影院的票根：你買票（登入）後拿到一張票（JWT），之後進場、買爆米花都給人看這張票就好，不用每次重新買。這張票有效期 7 天。'),
    ('REST API', 'Representational State Transfer API',
     '程式之間溝通的標準方式，用 HTTP 動詞（GET/POST/PUT/DELETE）操作資源。',
     '像餐廳點餐：GET = 看菜單、POST = 點新菜、PUT = 改訂單、DELETE = 取消訂單。每次都附上桌號（URL）和你的需求（參數）。'),
    ('JSON', 'JavaScript Object Notation',
     '一種把資料組織成「名稱:值」格式的文字標準。',
     '像填寫表格：{"姓名": "小明", "年齡": 25}。機器和人都看得懂，是程式之間交換資料最常用的格式。'),
    ('Port（連接埠）', 'Port',
     '同一台電腦上，不同服務程式各自佔用的「門牌號碼」。',
     '一棟大樓有很多房間：8000 號是控制中心、8001 號是資料庫、8080 號是 AI 大腦。訪客（請求）要說清楚去幾號房。'),
    ('Token（語言模型用語）', 'Token',
     'AI 處理文字的最小單位。一個中文字大約等於 1-2 個 token。',
     '像拼圖的碎片：AI 不是一個字一個字讀，而是把文字切成小碎片（token）來處理。「雞腿便當」可能被切成「雞」「腿」「便」「當」四個 token。max_tokens=128 表示 AI 最多只能回答約 64 個中文字。'),
    ('Temperature（溫度參數）', 'Temperature',
     '控制 AI 回答「隨機程度」的數值，範圍 0~1。',
     '0 = 每次都給一樣的答案（像考試答案），1 = 天馬行空（像腦力激盪）。本系統根據不同步驟用不同溫度：分析意圖用 0.1（要精確），寫回覆用 0.3（稍有變化）。'),
    ('Prompt（提示詞）', 'Prompt',
     '給 AI 的「工作指令」，告訴它該怎麼回答。',
     '像給新員工的工作手冊：「你的角色是營養顧問，使用者問你○○就查△△，回答要用繁體中文，不超過 80 字。」本系統有 6 份不同的 Prompt，用在不同步驟。'),
    ('State Machine（狀態機）', 'State Machine',
     '一個追蹤「現在在什麼階段」的機制，根據事件在不同狀態之間切換。',
     '像自動販賣機：待機 → 投幣中 → 選擇商品 → 出貨 → 回到待機。每個狀態決定了接下來能做什麼。本系統有 5 個狀態：idle → collecting → ready → executing → done。'),
    ('GraphRAG', 'Graph Retrieval-Augmented Generation',
     '一種結合「知識圖譜」和「文獻檢索」的技術，讓 AI 能引用學術論文來回答問題。',
     '像一個聰明的圖書館員：不只幫你找書，還知道哪些書之間有關聯，能從 65 篇營養學論文中找到最相關的段落。'),
    ('DRI（每日建議攝取量）', 'Dietary Reference Intakes',
     '衛福部（HPA）公布的「每天應該吃多少各種營養素」的標準。',
     '像身體的「營養預算」：30 歲男性每天需要鈣 1000mg、鐵 10mg 等。系統用的是第八版（2022 年），包含 33 項營養素。'),
    ('Context Window（上下文視窗）', 'Context Window',
     'AI 一次能「看到」的文字量上限。超過就會「忘記」前面的內容。',
     '像一張桌子：桌子只有這麼大，放太多紙就會掉下去。Gemma 的桌子大約能放 4096-5500 個 token（約 2000-3000 個中文字）。'),
    ('Streaming（串流）', 'Streaming',
     '資料像水流一樣持續送出，而不是等全部準備好才一次送。',
     '像看 YouTube 影片：不是等整部下載完才播，而是邊下載邊播。本系統的回覆也是這樣，先送「收到」，再送工具結果，最後送分析。'),
    ('FastAPI', 'FastAPI',
     '一個用 Python 寫的網頁伺服器框架，負責接收和處理請求。',
     '像餐廳的點餐系統：客人（前端）下單，FastAPI 接單後交給廚房（處理邏輯）去做。本系統有 3 個 FastAPI 服務。'),
    ('React', 'React',
     '一個用 JavaScript 寫的前端框架，負責畫出你在瀏覽器看到的介面。',
     '像建築的室內裝潢：決定按鈕放哪裡、文字怎麼排列、聊天氣泡長什麼樣。'),
]

for zh, en, definition, analogy in terms:
    h3(f'{zh}')
    p(f'英文：{en}', italic=True)
    p(f'定義：{definition}')
    p(f'白話比喻：{analogy}')

page_break()

# ══════════════════════════════════════════════
# 第 3 章：系統架構總覽
# ══════════════════════════════════════════════
h1('第三章：系統架構總覽')
p('整個系統由 5 個獨立的服務程式組成，它們各自負責不同的工作，但共同運作。你可以把它們想成一家公司的 5 個部門：')

add_svg('01_系統架構圖.svg', 6.5)

h2('3.1 五大服務')
add_table(
    ['服務', '技術', 'Port', '角色', '比喻'],
    [
        ['Frontend\n前端', 'React + Vite\n+ nginx', '3001', '使用者看到的網頁介面', '櫃台接待'],
        ['Control\n控制中心', 'FastAPI\n(Python)', '8000', '業務邏輯、AI 顧問、\n認證、工具調度', '經理/調度員'],
        ['Backend\n後端資料', 'FastAPI\n+ asyncpg', '8001', '資料庫查詢\n(台灣食品/FooDB/\n飲食記錄/食譜)', '倉庫管理員'],
        ['GraphRAG\nAPI', 'FastAPI\n+ chromadb', '8002', '學術文獻檢索\n(65 篇營養學論文)', '圖書館員'],
        ['Gemma LLM', 'llama-server\n(llama.cpp)', '8080', 'AI 語言理解\n(Gemma-4 模型)', 'AI 大腦'],
    ],
    col_widths=[1.0, 1.2, 0.6, 1.8, 1.0]
)

h2('3.2 服務之間怎麼溝通')
p('所有服務都跑在同一台電腦上（localhost），透過不同的 Port 區分彼此。通訊路徑如下：')
bullet('使用者的瀏覽器 → nginx 把 /api/* 開頭的請求轉給 Control（Port 8000）')
bullet('Control 是「調度中心」，它會根據需要去呼叫 Backend（Port 8001）、GraphRAG（Port 8002）、Gemma（Port 8080）')
bullet('使用者永遠不會直接接觸到 Backend/GraphRAG/Gemma，一切都透過 Control 中轉')

h2('3.3 認證流程（JWT）')
p('使用者登入時，系統會發給他一張「通行證」（JWT Token），有效期 7 天。之後每次發送訊息，前端都會附上這張通行證，Control 會驗證它是否有效。')
code('''// 前端發送請求時附上 JWT
fetch("/api/advisor/chat", {
  method: "POST",
  headers: {
    "Content-Type": "application/json",
    Authorization: `Bearer ${token}`,     // ← JWT 通行證
  },
  body: JSON.stringify({ message, history, memory, workflow_state }),
});''')

page_break()

# ══════════════════════════════════════════════
# 第 4 章：完整通訊鏈
# ══════════════════════════════════════════════
h1('第四章：完整通訊鏈')
p('這一章我們要追蹤：從你在聊天框打字、按下 Enter，到畫面上出現 AI 回覆——中間到底發生了什麼事。')

add_svg('02_完整通訊鏈.svg', 6.5)

h2('4.1 一句話的旅程')
p('當你輸入「今天午餐吃了雞腿便當」，以下是這句話經歷的完整旅程：')
steps = [
    ('1. 使用者按 Enter', '前端 AdvisorPage.jsx 的 handleSend() 函式被觸發'),
    ('2. 前端發送請求', 'advisorApi.js 用 fetch() 發送 POST /api/advisor/chat，開始 SSE 串流'),
    ('3. nginx 轉發', '前端的 /api/* 被 nginx 轉發到 Control 服務（Port 8000）'),
    ('4. Control 驗證身分', 'require_auth 中介軟體驗證 JWT Token'),
    ('5. Step 1 — Echo', 'Control 呼叫 Gemma LLM，生成「收到」確認 → 推送 SSE echo 事件'),
    ('6. Step 2 — Intent', 'Control 呼叫 Gemma LLM，分析使用者意圖 → JSON 結構'),
    ('7. Step 3 — ToolCall', 'Control 用系統規則（或 LLM）產生工具呼叫指令'),
    ('8. 工具執行', 'Control 呼叫對應的後端 API → 推送 SSE tool_start/result/done 事件'),
    ('9. Step 4 — Analysis', 'Control 呼叫 Gemma LLM，把結果整理成分析 → 推送 SSE reply 事件'),
    ('10. 記憶壓縮', 'Control 呼叫 Gemma LLM，壓縮對話記憶 → 推送 SSE memory 事件'),
    ('11. 狀態更新', 'Control 推送 SSE workflow_state 事件'),
    ('12. 前端渲染', '前端收到所有事件，顯示工具卡片和分析回覆'),
]
for title, desc in steps:
    pb(title)
    p(f'　　{desc}')

h2('4.2 SSE 事件協議')
p('在這趟旅程中，Control 透過 SSE 串流推送了 9 種不同類型的事件。每個事件都有固定的格式：')
code('''event: 事件類型
data: {"JSON資料"}

// 例如：
event: echo
data: {"content": "好的，我來幫你記錄午餐..."}''')

add_svg('04_SSE事件流.svg', 6.5)

add_table(
    ['事件類型', '方向', '資料格式', '用途', '前端處理'],
    [
        ['thinking', '伺服器→瀏覽器', '{"status": "understanding"}', '顯示「思考中」動畫', 'setIsThinking(true)'],
        ['echo', '伺服器→瀏覽器', '{"content": "收到..."}', '立即顯示確認氣泡', '加入暫時訊息 (_isEcho=true)'],
        ['tool_start', '伺服器→瀏覽器', '{"name": "search_food",\n"args": {...}}', '顯示「正在搜尋食物」', 'ThinkingPanel 加入步驟'],
        ['tool_result', '伺服器→瀏覽器', '{"name": "search_food",\n"data": {...}}', '完整工具結果\n（用於渲染卡片）', '累積到 turnToolResults'],
        ['tool_done', '伺服器→瀏覽器', '{"name": "search_food",\n"summary": "找到8筆"}', '標記工具完成', 'ThinkingPanel 標記 done'],
        ['reply', '伺服器→瀏覽器', '{"content": "分析..."}', '最終分析回覆', '移除 echo，顯示最終訊息'],
        ['memory', '伺服器→瀏覽器', '{"summary": "25歲女性..."}', '壓縮後的對話記憶', '儲存到 React state'],
        ['workflow_state', '伺服器→瀏覽器', '{"phase": "done", ...}', '狀態機快照', '更新 workflowState'],
        ['error', '伺服器→瀏覽器', '{"message": "錯誤訊息"}', '錯誤通知', '顯示錯誤氣泡'],
    ],
    col_widths=[1.0, 0.9, 1.5, 1.3, 1.5]
)

h2('4.3 前端如何處理 SSE')
p('前端的 advisorApi.js（共 99 行）是 SSE 串流的「收音機」。它用瀏覽器原生的 fetch API 建立連線，然後逐行讀取伺服器推送的資料：')
code('''// advisorApi.js — 核心處理邏輯（第 33-86 行）
const reader = response.body.getReader();    // 取得串流讀取器
const decoder = new TextDecoder();           // 文字解碼器
let buffer = "";                             // 暫存區

while (true) {
  const { done, value } = await reader.read();  // 讀取下一段資料
  if (done) break;                               // 串流結束

  buffer += decoder.decode(value, { stream: true });  // 解碼
  const lines = buffer.split("\\n");                   // 按行分割
  buffer = lines.pop() || "";                          // 最後不完整的行放回暫存

  let eventType = null;
  for (const line of lines) {
    if (line.startsWith("event: ")) {
      eventType = line.slice(7).trim();      // 提取事件類型
    } else if (line.startsWith("data: ") && eventType) {
      const data = JSON.parse(line.slice(6)); // 解析 JSON 資料
      switch (eventType) {                    // 根據類型分派
        case "echo":    callbacks.onEcho?.(data);    break;
        case "reply":   callbacks.onReply?.(data);   break;
        case "memory":  callbacks.onMemory?.(data);  break;
        // ... 其他 6 種事件
      }
    }
  }
}''')

h2('4.4 HTTP 請求格式')
p('前端每次發送訊息時，會把以下資料打包成 JSON 送給 Control：')
code('''// 前端送出的 HTTP 請求 body
{
  "message": "今天午餐吃了雞腿便當",       // 使用者這次說的話
  "history": [                              // 之前的對話紀錄（最多 30 則）
    {"role": "user",      "content": "..."},
    {"role": "assistant", "content": "...", "tools_used": [...]}
  ],
  "user_profile": {},                       // 使用者個人資料（目前未用）
  "memory": "25歲女性;今日午餐已記錄...",   // 上一輪的壓縮記憶
  "workflow_state": {                       // 狀態機的當前狀態
    "phase": "idle",
    "active_intent": "",
    "collected": {},
    "missing": []
  }
}''')
p('程式碼位置：AdvisorRequest 類別定義在 advisor.py 第 355-360 行。')

page_break()

# ══════════════════════════════════════════════
# 第 5 章：四步驟 LLM Pipeline
# ══════════════════════════════════════════════
h1('第五章：四步驟 LLM Pipeline 詳解')
p('「Pipeline」的意思是「流水線」——就像工廠的生產線，每一站做不同的事。AI 營養顧問有 4 個站（Step），加上記憶壓縮和參數提取共 6 次 LLM 呼叫。')

add_svg('03_四步驟Pipeline.svg', 6.5)

# --- Step 1 ---
h2('5.1 Step 1: Echo — 即時確認回應')
p('目的：讓使用者知道「我收到你的訊息了，正在處理」。就像客服人員會先說「好的，請稍等」。')

h3('LLM 呼叫參數')
add_table(
    ['參數', '值', '說明'],
    [
        ['Temperature', '0.5', '中等隨機度，讓語氣自然不死板'],
        ['Max Tokens', '128', '最多約 64 個中文字（夠了，只需 1-2 句）'],
        ['嘗試次數', '1 次', '失敗就用預設回覆'],
        ['預設回覆', '"收到，正在處理你的問題..."', 'LLM 掛掉時的備用方案'],
    ],
    col_widths=[1.2, 1.5, 3.5]
)

h3('Prompt 全文（PROMPT_ECHO）')
p('這是系統給 Gemma LLM 的指令，告訴它「收到使用者訊息後，要怎麼回應」：', italic=True)
p('程式碼位置：advisor.py 第 38-78 行')
code('''請你根據使用者的訊息和對話脈絡，用繁體中文回應。

## 規則
- 一到兩句話，30-80 字
- 如果使用者的訊息需要查詢工具（記錄飲食、查營養、查症狀等）：
  概括需求，表達「已收到、正在處理」
- 如果是一般對話（打招呼、感謝、閒聊）：直接給出友善回應
- 不要回答營養問題本身，不要提供營養數據或建議
- 語氣親切自然

## 範例
使用者：「今天午餐吃了雞腿便當和一杯豆漿」
→ 好的，我來幫你記錄午餐的雞腿便當和豆漿，並計算營養成分。

使用者：「營養夠嗎」
→ 收到，我先查看你的個人資料，再幫你分析今天的營養攝取是否達標。

使用者：「最近常覺得疲勞」
→ 了解，我來查詢與疲勞相關的營養素和推薦食物。

使用者：「你好」
→ 你好！我是你的 AI 營養顧問，有什麼飲食或營養問題可以幫你的嗎？

---
{memory_block}        ← 上一輪的對話記憶
對話歷史：{history}    ← 之前的對話紀錄（最多 8 則，每則 ≤400 字）
使用者訊息：{message}  ← 這次使用者說的話

請直接輸出回應（不要加引號或其他格式）：''')

h3('SSE 輸出與前端渲染')
p('Echo 生成後，Control 會推送兩個 SSE 事件：')
code('''event: thinking
data: {"status": "understanding"}    ← 前端顯示思考動畫

event: echo
data: {"content": "好的，我來幫你記錄午餐..."}  ← 前端顯示暫時氣泡''')
p('前端收到 echo 後，會先顯示一個帶有 _isEcho: true 標記的暫時氣泡。等到 Step 4 的 reply 事件來了，這個暫時氣泡就會被替換成最終的分析回覆。')

h3('程式碼解讀')
code('''# advisor.py 第 790-805 行
async def _step_echo(message, history, memory):
    # 組合記憶區塊
    memory_block = f"使用者記憶：{memory}" if memory else "使用者記憶：（新對話）"

    # 套入 Prompt 模板
    prompt = PROMPT_ECHO.format(
        message=message,                        # 使用者這次說的話
        history=_format_history(history),        # 最近 8 則對話
        memory_block=memory_block,               # 壓縮記憶
    )

    # 呼叫 Gemma LLM
    result = await _call_gemma(
        [{"role": "user", "content": prompt}],
        temperature=0.5,    # 中等隨機
        max_tokens=128,     # 最多 128 token（約 64 字）
    )
    return result or "收到，正在處理你的問題..."  # 失敗備案''')

h3('如果是人手動操作')
p('想像你是客服人員：客戶打電話來說「我要查我的營養夠不夠」，你會先回答「好的，我幫您查一下，請稍等」。這就是 Echo 在做的事——先安撫使用者，讓他知道系統沒有當掉。')

page_break()

# --- Step 2 ---
h2('5.2 Step 2: Intent — 意圖偵測')
p('目的：搞清楚使用者到底想做什麼。這一步的結果不會顯示在畫面上（後台分析），但它決定了接下來要用哪些工具。')

h3('LLM 呼叫參數')
add_table(
    ['參數', '值', '說明'],
    [
        ['Temperature', '0.1', '極低隨機度（要精確分析，不要亂猜）'],
        ['Max Tokens', '512', '最多約 256 字（JSON 結構不長）'],
        ['嘗試次數', '2 次', '第一次 JSON 解析失敗會用簡化 prompt 再試'],
        ['歷史截斷', '最近 8 則 × 400 字', '避免超出 Context Window'],
        ['前輪結果截斷', '5000 字', '供多輪對話參考'],
    ],
    col_widths=[1.2, 1.8, 3.2]
)

h3('Prompt 全文（PROMPT_INTENT）')
p('這是系統最複雜的 Prompt，它要教 LLM 辨認 14 種不同的意圖：', italic=True)
p('程式碼位置：advisor.py 第 85-204 行')
code('''請你根據對話內容，分析使用者的意圖，輸出嚴格 JSON。

## 可辨識的意圖類型

| intent          | 觸發條件                         | 需要的 entities          |
|-----------------|----------------------------------|--------------------------|
| record_diet     | 提到吃了什麼、要記錄飲食         | food_items: [{name, amount_g, meal_type}] |
| query_dri       | 問營養是否充足、DRI 缺口         | age, gender, activity_level |
| search_food     | 問特定食物的營養                 | food_name                |
| search_symptom  | 提到健康症狀                     | symptom_en（英文）       |
| search_graphrag | 問「為什麼」、要學術根據         | query_en（英文學術查詢） |
| nutrient_ranking| 問某營養素哪些食物最多           | nutrient_field           |
| get_diet_logs   | 問過去吃了什麼                   | —                        |
| update_diet_log | 要修改飲食記錄                   | log_id                   |
| delete_diet_log | 要刪除飲食記錄                   | log_id                   |
| get_recipes     | 查看儲存的食譜                   | —                        |
| get_calendar    | 查看行事曆                       | start_date, end_date     |
| get_profile     | 查看個人資料                     | —                        |
| update_profile  | 更新個人資料                     | profile_fields           |
| general_chat    | 不需要工具的一般對話             | —                        |

## 份量估算規則
便當≈500g、雞排≈200g、煎餃一份≈180g、一碗飯≈200g、一杯豆漿≈300ml、
一顆蛋≈60g、一片吐司≈30g、一碗麵≈300g、排骨≈150g、魚排≈120g、
雞胸肉≈200g、牛排≈250g、沙拉≈150g、水果一份≈150g、一杯牛奶≈240ml

## 症狀翻譯對照
疲勞=fatigue、失眠=insomnia、頭痛=headache、掉髮=hair loss、
便秘=constipation、抽筋=muscle cramp、皮膚差=poor skin health、...

## 營養素欄位名對照
蛋白質=protein_per_100g、鈣=calcium_per_100g、鐵=iron_per_100g、...

## 確認詞處理
使用者說「好」「可以」「對」= 確認上一輪 AI 提議的動作。
請從對話歷史中找出上一輪 AI 建議了什麼，將其轉為對應的 intent。

## 輸出格式
嚴格 JSON（不要用 ``` 包裹）：
{
  "intents": ["record_diet"],
  "entities": {
    "food_items": [{"name": "雞腿便當", "amount_g": 500, "meal_type": "lunch"}],
    "food_name": "", "symptom_en": "", "query_en": "",
    "nutrient_field": "", "profile_fields": {},
    "log_id": null, "start_date": "", "end_date": "",
    "age": null, "gender": "", "activity_level": ""
  },
  "needs_profile": false,     ← true 表示需要先查個人資料
  "needs_diet_logs": false    ← true 表示需要先查飲食記錄
}''')

h3('輸出範例')
code('''使用者：「今天午餐吃了雞腿便當、一杯豆漿和一顆蛋」
→ {"intents":["record_diet"],
   "entities":{"food_items":[
     {"name":"雞腿便當","amount_g":500,"meal_type":"lunch"},
     {"name":"豆漿","amount_g":300,"meal_type":"lunch"},
     {"name":"蛋","amount_g":60,"meal_type":"lunch"}
   ]}}

使用者：「最近常覺得疲勞」
→ {"intents":["search_symptom"],
   "entities":{"symptom_en":"fatigue"}}

使用者：「我25歲女性，幫我看看營養是否充足」
→ {"intents":["update_profile","query_dri"],
   "entities":{"profile_fields":{"age":25,"gender":"female"},"age":25,"gender":"female"}}''')

h3('如果是人手動操作')
p('想像你是一個表單填寫員：客戶說了一段話，你要聽完後在表格上勾選「這是要記錄飲食」，然後填上「食物：雞腿便當、份量：500g、餐別：午餐」。Intent Detection 就是自動化的表單填寫。')

page_break()

# --- Step 3 ---
h2('5.3 Step 3: ToolCall — 工具指令生成')
p('目的：根據意圖分析的結果，決定要呼叫哪些工具，以及給這些工具什麼參數。')
p('這一步有一個重要的設計：系統優先，LLM 備援。', bold=True)
p('什麼意思呢？系統會先嘗試用純 Python 程式碼來生成工具指令（快速、可靠），只有當程式碼處理不了的複雜情況，才會去問 LLM。')

h3('路徑 A：系統生成（優先）')
p('函式：_system_generate_toolcalls()（第 714-783 行）')
p('純 Python 邏輯，不需要 LLM，速度快且穩定。能處理大部分常見情況：')
code('''# 系統生成的邏輯（簡化版）
def _system_generate_toolcalls(intent, profile):
    intents = intent.get("intents", [])
    entities = intent.get("entities", {})

    # 記錄飲食：為每項食物各生成一個 add_diet_log
    if "record_diet" in intents:
        return [{"tool": "add_diet_log", "args": {
            "food_name": item["name"],
            "amount_g": item.get("amount_g", 200),
            "meal_type": item.get("meal_type", "lunch"),
        }} for item in entities.get("food_items", [])]

    # 查 DRI：需要 age + gender
    if "query_dri" in intents:
        if intent.get("needs_profile"):
            return [{"tool": "get_user_profile", "args": {}}]
        return [{"tool": "query_dri", "args": {
            "age": entities.get("age", 30),
            "gender": entities.get("gender", "male"),
        }}]

    # ... 其他意圖的規則''')

h3('路徑 B：LLM 生成（備援）')
p('只有當路徑 A 回傳 None（處理不了）時才會啟用。')
add_table(
    ['參數', '值', '說明'],
    [
        ['Temperature', '0.1', '要精確，不能亂生成工具呼叫'],
        ['Max Tokens', '512', '工具指令 JSON 不長'],
        ['嘗試次數', '2 次', '第一次 JSON 解析失敗會重試'],
    ],
    col_widths=[1.2, 1.5, 3.5]
)

h3('Prompt 全文（PROMPT_TOOLCALL）')
p('程式碼位置：advisor.py 第 211-278 行')
code('''請你根據需求分析結果，生成工具呼叫指令。

## 可用工具
| 工具名              | 參數                           | 說明           |
|---------------------|--------------------------------|----------------|
| add_diet_log        | food_name, amount_g, meal_type | 新增飲食記錄   |
| query_dri           | age, gender, activity_level    | DRI 缺口分析   |
| search_food         | food_name                      | 搜尋食物營養   |
| search_symptom      | symptom（英文）                | 症狀相關營養素 |
| search_graphrag     | query（英文學術詞）            | 學術文獻搜尋   |
| get_nutrient_ranking| nutrient_field, top_n          | 營養素排名     |
| get_diet_logs       | （無參數）                     | 查詢飲食記錄   |
| update_diet_log     | log_id, food_name, amount_g    | 修改記錄       |
| delete_diet_log     | log_id                         | 刪除記錄       |
| get_saved_recipes   | （無參數）                     | 查詢食譜       |
| get_calendar_entries| start_date, end_date           | 行事曆查詢     |
| get_user_profile    | （無參數）                     | 讀取個人資料   |
| update_user_profile | age, gender, height_cm, ...    | 更新個人資料   |

## 規則
- needs_profile 為 true → 輸出 [{"tool":"get_user_profile","args":{}}]
- record_diet → 為每項食物各生成一個 add_diet_log
- 不需要工具 → 輸出空陣列 []

## 輸出格式
[{"tool":"工具名","args":{"參數":"值"}}]''')

h3('驗證層')
p('無論是系統生成還是 LLM 生成的工具呼叫，都會經過 _validate_toolcalls() 驗證（第 633-695 行）：')
bullet('檢查工具名是否在 _TOOL_DISPATCH 註冊表中（13 個合法工具名）')
bullet('檢查必填參數是否齊全（例如 query_dri 必須有 age 和 gender）')
bullet('自動填入缺少的參數（從 entities 或使用者 profile）')
bullet('強制型別轉換（例如 age 的字串 "25" 轉成數字 25）')
bullet('去除重複的工具呼叫')

page_break()

# --- Step 4 ---
h2('5.4 Step 4: Analysis — 分析回覆')
p('目的：把工具查到的結果，整理成使用者看得懂的分析報告。這是使用者最終看到的回覆文字。')

h3('LLM 呼叫參數')
add_table(
    ['參數', '值', '說明'],
    [
        ['Temperature', '0.3', '平衡準確與自然語氣'],
        ['Max Tokens', '2048', '最多約 1024 字（最長的回覆）'],
        ['嘗試次數', '2 次', '第一次太短（< 50 字）會要求更詳細'],
        ['工具結果截斷', '12000 字', '避免超出 Context Window'],
    ],
    col_widths=[1.2, 1.5, 3.5]
)

h3('Prompt 全文（PROMPT_ANALYSIS）')
p('程式碼位置：advisor.py 第 285-331 行')
code('''請你根據工具查詢結果，用繁體中文提供專業營養分析回覆。

## 核心規則
- 所有數據必須來自下方的「工具執行結果」，禁止從訓練資料引用任何營養數字
- 沒有查到相關數據就明確說「目前沒有查到相關數據」
- 禁止使用「根據一般營養學原理」「通常建議」等包裝語句
- 前端已自動渲染工具結果的表格和卡片，你不需要列出原始數據清單
- 你的角色是「分析師」：解讀數據的意義，不是複述數據

## 回覆結構（每次必須包含三個部分）

### ① 數據摘要
簡述關鍵數字。例：
- 記錄飲食：「本次記錄了 3 項食物，估算總熱量約 850 kcal」
- DRI 分析：「共 X 項達標、Y 項不足、Z 項過量」

### ② 分析觀點
指出值得注意的發現。例：
- 「今天攝取偏向高碳水低蛋白，蛋白質僅達建議量的 40%」
- 「鐵和鈣是最嚴重的缺口，分別只達建議量的 30% 和 45%」

### ③ 下一步建議
具體可執行的行動。例：
- 「建議晚餐加一份深綠蔬菜（如菠菜）來補充鐵質和葉酸」
- 「要不要我查詢今天的 DRI 達成率？」

---
使用者原始問題：{message}
{memory_block}
工具執行結果：{tool_results}    ← 截斷至 12000 字

請用繁體中文輸出分析回覆：''')

h3('如果是人手動操作')
p('想像你是一位營養師：助理幫你查好了客戶的飲食記錄和 DRI 標準，放在桌上。你看完這些數據後，寫一份簡報給客戶：「你今天蛋白質不夠，建議晚餐加個蛋」。Analysis 就是自動化的營養師簡報。')

page_break()

# ══════════════════════════════════════════════
# 第 6 章：記憶壓縮
# ══════════════════════════════════════════════
h1('第六章：記憶壓縮系統')
p('AI 的「記憶力」是有限的（Context Window）。如果把所有對話歷史都塞進去，很快就會滿了。所以系統用一個「壓縮器」，把每一輪的對話精華壓縮成一小段記憶，傳給下一輪。')

add_svg('14_記憶壓縮流程.svg', 5.5)

h2('6.1 Prompt 全文（MEMORY_PROMPT）')
p('程式碼位置：advisor.py 第 338-348 行')
code('''你是記憶壓縮器。根據舊記憶和這一輪對話，輸出更新後的使用者記憶。

規則：
- 只列關於使用者的事實（年齡、性別、活動量、目標、飲食內容、健康狀況等）
- 用頓號分隔，一行，不換行
- 新資訊覆蓋舊資訊
- 工具查到的關鍵數據也記（如「DRI 熱量目標 2150kcal」）
- 不記 AI 的建議，只記使用者的事實和查詢結果
- 最多 200 字''')

h2('6.2 LLM 呼叫參數')
add_table(
    ['參數', '值', '說明'],
    [
        ['Temperature', '0.2', '低隨機度（事實壓縮要準確）'],
        ['Max Tokens', '1024', '壓縮結果不需要太長'],
        ['輸出截斷', '1500 字', '程式碼會截斷超過 1500 字的輸出'],
        ['使用者訊息截斷', '500 字', '只取使用者這輪說的前 500 字'],
        ['AI 回覆截斷', '800 字', '只取分析回覆的前 800 字'],
    ],
    col_widths=[1.2, 1.5, 3.5]
)

h2('6.3 何時觸發、何時跳過')
bullet('觸發：工具執行完畢後、或狀態機回覆後（且不在 collecting 階段）')
bullet('跳過：collecting 階段（正在追問參數時），節省 LLM 呼叫')
p('程式碼位置：advisor.py 第 1424-1445 行、第 1904-1912 行')

page_break()

# ══════════════════════════════════════════════
# 第 7 章：第六個 Prompt
# ══════════════════════════════════════════════
h1('第七章：第六個 Prompt — 參數提取')
p('當系統在「collecting」階段（追問使用者缺少的參數），使用者回答了一個值（例如「25歲」），系統需要從這句話中提取出年齡。')
p('提取分兩步：先用程式規則（regex），不行的話才請 LLM 幫忙。')

h2('7.1 第一步：程式規則提取')
p('函式：_extract_params_deterministic()（第 1479-1550 行）')
p('用正規表達式（regex）從文字中找出特定模式：')
code('''# 例如提取年齡
if "age" in missing:
    m = re.search(r"(\\d{1,3})\\s*歲", text)   # 找「25歲」這種格式
    if m:
        result["age"] = int(m.group(1))     # 提取出 25

# 提取性別
if "gender" in missing:
    if re.search(r"女性|女生|女", text):
        result["gender"] = "female"
    elif re.search(r"男性|男生|男", text):
        result["gender"] = "male"

# 提取症狀（中文→英文查表）
if "symptom" in missing:
    for zh, en in _SYMPTOM_ZH_TO_EN.items():  # {"疲勞": "fatigue", ...}
        if zh in text:
            result["symptom"] = en
            break''')

h2('7.2 第二步：LLM 提取（備援）')
p('當 regex 提取不到（使用者的表述太自由），才用 LLM：')

h3('Prompt 全文（PROMPT_EXTRACT_PARAMS）')
p('程式碼位置：advisor.py 第 1553-1561 行')
code('''從使用者訊息中提取以下欄位的值。
只輸出嚴格 JSON，不要加其他文字。找不到的欄位填 null。

需要提取的欄位：
- age: 年齡（整數）
- gender: 性別（輸出 male 或 female）
- activity_level: 活動量（high / moderate / low）
- food_name: 食物名稱（中文）
- symptom: 症狀（英文，如 fatigue, insomnia, headache）
- query: 學術查詢（英文）
- nutrient_field: 營養素欄位名（如 protein_per_100g）
- log_id: 飲食記錄編號（整數）
- start_date: 開始日期（YYYY-MM-DD 格式）
- end_date: 結束日期（YYYY-MM-DD 格式）

使用者訊息：{message}

JSON：''')

add_table(
    ['參數', '值', '說明'],
    [
        ['Temperature', '0.1', '要精確提取，不要猜'],
        ['Max Tokens', '128', '輸出就是一小段 JSON'],
    ],
    col_widths=[1.2, 1.5, 3.5]
)

page_break()

# ══════════════════════════════════════════════
# 第 8 章：13 個工具
# ══════════════════════════════════════════════
h1('第八章：13 個工具完整解析')
p('這是本報告最核心的一章。我們會逐一拆解每個工具：它做什麼、人手動操作長什麼樣、需要什麼參數、呼叫哪個後端 API、回傳什麼資料、前端怎麼顯示。')

add_svg('05_工具分發機制.svg', 6.0)

h2('8.0 工具分發機制')
p('所有工具都透過一個「分發字典」來找到對應的處理函式：')
code('''# advisor.py 第 1403-1417 行
_TOOL_DISPATCH = {
    "query_dri":           _tool_query_dri,          # → DRI 缺口
    "get_nutrient_ranking": _tool_get_nutrient_ranking, # → 營養素排名
    "search_food":         _tool_search_food,         # → 食物搜尋
    "search_symptom":      _tool_search_symptom,      # → 症狀搜尋
    "search_graphrag":     _tool_search_graphrag,     # → 學術文獻
    "get_diet_logs":       _tool_get_diet_logs,       # → 飲食記錄
    "add_diet_log":        _tool_add_diet_log,        # → 新增飲食
    "update_diet_log":     _tool_update_diet_log,     # → 修改飲食
    "delete_diet_log":     _tool_delete_diet_log,     # → 刪除飲食
    "get_saved_recipes":   _tool_get_saved_recipes,   # → 查食譜
    "get_calendar_entries": _tool_get_calendar_entries, # → 行事曆
    "get_user_profile":    _tool_get_user_profile,    # → 讀個人資料
    "update_user_profile": _tool_update_user_profile, # → 改個人資料
}

# 第 945-954 行 — 分發執行
async def _execute_tool(name, arguments, auth_token, user_message):
    fn = _TOOL_DISPATCH.get(name)        # 用工具名查字典
    if not fn:
        return {"error": f"Unknown tool: {name}"}
    return await fn(arguments, auth_token, user_message)  # 呼叫函式''')
p('每次執行工具時，Control 會發出 3 個 SSE 事件：tool_start → tool_result → tool_done。')

# Tool 1: query_dri
tools_data = [
    {
        'num': '1', 'name': 'query_dri', 'zh': 'DRI 缺口分析',
        'label': '查詢 DRI 建議攝取量',
        'human': '打開 DRI 缺口分析頁面 → 輸入年齡和性別 → 點擊「分析」→ 看到每項營養素的達成率和缺口',
        'params': [('age', '整數', '必填', '年齡'), ('gender', '"male"/"female"', '必填', '性別'), ('activity_level', '字串', '選填（預設 "moderate"）', '活動量')],
        'calls': '1. 自動取今天的飲食記錄：GET http://localhost:8000/diet/logs\n2. 呼叫 dri_gap_service.analyze_dri_gap()\n3. 載入 HPA v8 JSON → 取得目標值 → 比較實際 vs 目標\n4. 對每個缺口：GET http://localhost:8001/health/nutrients/{field}/top-foods',
        'response': '{"profile": "Female, 19-30 years",\n "gaps": [{\n   "nutrient_zh": "鐵", "unit": "mg",\n   "target": 15, "actual": 5.2, "ratio": 0.35,\n   "deficit": 9.8, "status": "deficient",\n   "recommended_foods": [{"name": "牡蠣", "content_per_100g": 28.5}]\n }]}',
        'renderer': 'DriGapTable — 紅/橙/綠色標的達成率表格，可展開看推薦食物',
        'page': 'DriGapPage（/dri-gap）',
        'lines': '957-1003',
        'svg': '08_DRI缺口分析流程.svg',
    },
    {
        'num': '2', 'name': 'get_nutrient_ranking', 'zh': '營養素排名',
        'label': '查詢營養素排名',
        'human': '打開營養素排名頁面 → 從下拉選單選「鎂」→ 看到鎂含量最高的前 10 名食物',
        'params': [('nutrient_field', '字串 (如 magnesium_per_100g)', '必填', '營養素欄位名'), ('top_n', '整數', '選填（預設 10）', '要幾名')],
        'calls': 'GET http://localhost:8000/health/nutrients/{field}/top-foods?limit={top_n}',
        'response': '{"nutrient": "magnesium_per_100g",\n "top_foods": [\n   {"name": "南瓜子", "category": "堅果", "content_per_100g": 550, "source": "taiwan_foods"},\n   {"name": "Pumpkin seeds", "content_per_100g": 535, "source": "foodb"}\n ]}',
        'renderer': 'NutrientRankingTable — 排名表格（名次/食物/類別/含量/來源）',
        'page': 'NutrientRankingPage（/nutrient-ranking）',
        'lines': '1118-1126',
        'svg': None,
    },
    {
        'num': '3', 'name': 'search_food', 'zh': '食物搜尋',
        'label': '搜尋食物資訊',
        'human': '打開食物搜尋頁面 → 輸入「雞胸肉」→ 看到台灣食品資料庫和 FooDB 的搜尋結果',
        'params': [('food_name', '字串', '必填', '食物名稱')],
        'calls': 'GET http://localhost:8000/food/search?q={food_name}&limit=8',
        'response': '{"food_name": "雞胸肉",\n "results": [{\n   "id": 123, "name": "雞胸肉-去皮",\n   "category": "禽肉類", "source": "taiwan_foods",\n   "nutrition": {"calories": 117, "protein": 24.2, "carbs": 0, "fat": 1.9, "fiber": 0}\n }]}',
        'renderer': 'SearchFoodCard — 食物卡片（名稱/類別/來源標籤/營養網格）',
        'page': 'FoodSearchPage（/food-search）',
        'lines': '1129-1156',
        'svg': None,
    },
    {
        'num': '4', 'name': 'search_symptom', 'zh': '症狀搜尋',
        'label': '搜尋症狀相關營養素',
        'human': '打開症狀搜尋頁面 → 輸入「疲勞」→ 看到相關化合物和推薦食物',
        'params': [('symptom', '英文字串 (如 fatigue)', '必填', '症狀英文名')],
        'calls': '對使用者訊息中的中文詞 + 英文症狀名，逐一嘗試：\nPOST http://localhost:8000/health/symptom-search {"keyword": "..."}\n取得分最高的結果（effects 數 + recommendations 數）',
        'response': '{"symptom": "fatigue",\n "effects": [{"name": "Reduced Fatigue"}],\n "recommendations": [{\n   "compound": {"name": "Iron", "description": "..."},\n   "health_effect": "Reduced Fatigue",\n   "top_foods": [{"name": "牡蠣", "food_group": "Seafood", "content_per_100g": 28.5}]\n }]}',
        'renderer': 'SymptomResultCard — 效果標籤 + 推薦卡片（化合物/食物表格）',
        'page': 'SymptomSearchPage（/symptom-search）',
        'lines': '1159-1189',
        'svg': '10_症狀搜尋流程.svg',
    },
    {
        'num': '5', 'name': 'search_graphrag', 'zh': '學術文獻搜尋',
        'label': '查詢學術文獻',
        'human': '打開 GraphRAG 頁面 → 輸入「蛋白質對肌肉修復的重要性」→ 看到相關論文和原文片段',
        'params': [('query', '英文字串', '必填', '學術查詢詞')],
        'calls': 'POST http://localhost:8002/retrieve-only\n{"query": "protein muscle repair", "methods": ["B", "C"]}\n方法 B = 局部搜尋（entity → graph traversal → chunks）\n方法 C = 全域搜尋（community reports → key entities → chunks）\nTimeout: 60 秒（比其他工具的 30 秒長一倍）',
        'response': '{"query": "protein muscle repair",\n "sources": [{\n   "type": "paper", "doc_id": "doc_001",\n   "paper_meta": {\n     "title": "Dietary Protein and Muscle...",\n     "authors": ["Smith J."], "journal": "Nutrients",\n     "year": 2023, "doi": "10.3390/...",\n     "impact_factor": 5.9, "sjr_quartile": "Q1"\n   },\n   "text_preview": "Protein intake of 1.6-2.2g/kg/day..."\n }]}',
        'renderer': 'GraphRagResultCard → 使用 ReferenceList 元件\n（IF/SJR 色標、DOI 連結、study type 徽章、PDF 檢視按鈕）',
        'page': 'GraphRAGPage（/graphrag）',
        'lines': '1192-1237',
        'svg': '11_學術文獻流程.svg',
    },
    {
        'num': '6', 'name': 'get_diet_logs', 'zh': '查詢飲食記錄',
        'label': '查詢飲食日誌',
        'human': '打開飲食日誌頁面 → 看到所有記錄（食物名、克數、餐別、時間、營養素）',
        'params': [],
        'calls': 'GET http://localhost:8000/diet/logs（附 JWT Token）',
        'response': '{"logs": [{\n   "id": 42, "food_name": "雞腿便當",\n   "amount_g": 500, "meal_type": "lunch",\n   "logged_at": "2026-05-11T12:30:00",\n   "calories": 750, "protein": 35, "carbs": 80, "fat": 25\n }], "total": 15}',
        'renderer': 'DietLogsTable — 完整表格（ID/食物/克數/餐別/時間/營養素）',
        'page': 'DietLogPage（/diet）',
        'lines': '1242-1263',
        'svg': None,
    },
    {
        'num': '7', 'name': 'add_diet_log', 'zh': '新增飲食記錄',
        'label': '新增飲食記錄',
        'human': '打開飲食日誌頁面 → 點「新增」→ 輸入食物名稱和克數 → 儲存',
        'params': [('food_name', '字串', '必填', '食物名稱'), ('amount_g', '數字', '選填（預設 200）', '克數'), ('meal_type', '字串', '選填（預設 "lunch"）', '餐別')],
        'calls': '1. 自動搜尋食物：GET /food/search?q={food_name}&limit=1\n   → 取得 food_id 和 source（用於連結營養資料）\n2. 新增記錄：POST /diet/logs\n   payload: {food_name, amount_g, meal_type, food_id, food_source, logged_at}',
        'response': '{"success": true, "id": 43,\n "food_name": "雞腿便當", "amount_g": 500,\n "nutrition": {"calories": 750, "protein": 35, "carbs": 80, "fat": 25}}',
        'renderer': 'DietLogActionCard (action="add") — 綠色確認卡片，顯示自動計算的營養',
        'page': 'DietLogPage（/diet）',
        'lines': '1266-1321',
        'svg': '09_飲食記錄流程.svg',
    },
    {
        'num': '8', 'name': 'update_diet_log', 'zh': '修改飲食記錄',
        'label': '修改飲食記錄',
        'human': '打開飲食日誌 → 點某筆記錄的「編輯」→ 改食物名或克數 → 儲存',
        'params': [('log_id', '整數', '必填', '記錄 ID'), ('food_name', '字串', '選填', '新食物名'), ('amount_g', '數字', '選填', '新克數'), ('meal_type', '字串', '選填', '新餐別')],
        'calls': 'PUT http://localhost:8000/diet/logs/{log_id}（附 JWT）\npayload: 只包含要修改的欄位',
        'response': '{"success": true, "id": 42, "food_name": "雞胸肉便當"}',
        'renderer': 'DietLogActionCard (action="update")',
        'page': 'DietLogPage（/diet）',
        'lines': '1324-1333',
        'svg': None,
    },
    {
        'num': '9', 'name': 'delete_diet_log', 'zh': '刪除飲食記錄',
        'label': '刪除飲食記錄',
        'human': '打開飲食日誌 → 點某筆記錄的「刪除」→ 確認刪除',
        'params': [('log_id', '整數', '必填', '記錄 ID')],
        'calls': 'DELETE http://localhost:8000/diet/logs/{log_id}（附 JWT）',
        'response': '{"success": true, "message": "已刪除記錄 #42"}',
        'renderer': 'DietLogActionCard (action="delete")',
        'page': 'DietLogPage（/diet）',
        'lines': '1336-1342',
        'svg': None,
    },
    {
        'num': '10', 'name': 'get_saved_recipes', 'zh': '查詢已儲存食譜',
        'label': '查詢已儲存食譜',
        'human': '打開食譜頁面 → 看到所有儲存的食譜清單',
        'params': [],
        'calls': 'GET http://localhost:8000/recipe/list（附 JWT）',
        'response': '{"recipes": [{\n   "id": 5, "name": "番茄炒蛋",\n   "servings": 2, "created_at": "2026-05-10",\n   "ingredients": ["番茄", "雞蛋", "蔥"],\n   "nutrition": {"calories": 180, "protein": 12}\n }], "total": 8}',
        'renderer': 'SavedRecipesCard — 食譜卡片（食材標籤 + 營養網格）',
        'page': 'SavedRecipesPage（/recipes）',
        'lines': '1347-1372',
        'svg': None,
    },
    {
        'num': '11', 'name': 'get_calendar_entries', 'zh': '行事曆查詢',
        'label': '查詢行事曆',
        'human': '打開行事曆頁面 → 選擇日期範圍 → 看到該時段的記錄',
        'params': [('start_date', 'YYYY-MM-DD', '必填', '開始日期'), ('end_date', 'YYYY-MM-DD', '必填', '結束日期')],
        'calls': 'GET http://localhost:8000/calendar/entries?start_date={}&end_date={}（附 JWT）',
        'response': '{"entries": [{"id": 1, "entry_date": "2026-05-11", "recipe_id": 5}]}',
        'renderer': 'CalendarEntriesCard — 日期列表',
        'page': 'DailyAnalysisCalendarPage（/calendar）',
        'lines': '1375-1383',
        'svg': None,
    },
    {
        'num': '12', 'name': 'get_user_profile', 'zh': '讀取個人資料',
        'label': '讀取個人資料',
        'human': '打開個人資料頁面 → 看到自己的年齡、性別、身高、體重等',
        'params': [],
        'calls': 'GET http://localhost:8000/profile（附 JWT）',
        'response': '{"username": "小明", "email": "user@example.com",\n "gender": "female", "age": 25,\n "height_cm": 165, "weight_kg": 55,\n "activity_level": "moderate", "goal": "maintain"}',
        'renderer': 'UserProfileCard — 兩欄顯示個人欄位',
        'page': 'ProfilePage（/profile）',
        'lines': '1386-1391',
        'svg': None,
    },
    {
        'num': '13', 'name': 'update_user_profile', 'zh': '更新個人資料',
        'label': '更新個人資料',
        'human': '打開個人資料頁面 → 修改欄位 → 點儲存',
        'params': [('age', '整數', '選填', '年齡'), ('gender', '字串', '選填', '性別'), ('height_cm', '數字', '選填', '身高'), ('weight_kg', '數字', '選填', '體重'), ('activity_level', '字串', '選填', '活動量'), ('goal', '字串', '選填', '目標')],
        'calls': 'PUT http://localhost:8000/profile（附 JWT）\npayload: 只包含要修改的欄位',
        'response': '{"success": true, "profile": {完整 profile 物件}}',
        'renderer': 'ProfileUpdatedCard — 顯示更新後的資料',
        'page': 'ProfilePage（/profile）',
        'lines': '1394-1400',
        'svg': None,
    },
]

for tool in tools_data:
    h2(f"8.{tool['num']} {tool['name']} — {tool['zh']}")
    p(f"前端標籤：{tool['label']}", italic=True)

    h3('如果是人手動操作')
    p(tool['human'])

    if tool['params']:
        h3('參數')
        add_table(
            ['參數名', '型別', '必填/選填', '說明'],
            [[p[0], p[1], p[2], p[3]] for p in tool['params']],
            col_widths=[1.5, 2.0, 1.5, 1.5]
        )
    else:
        h3('參數')
        p('無參數（只需要 JWT Token 認證）')

    h3('後端 HTTP 呼叫')
    code(tool['calls'])

    h3('回傳資料格式')
    code(tool['response'])

    h3('前端渲染')
    p(tool['renderer'])
    p(f"對應頁面：{tool['page']}")
    p(f"程式碼位置：advisor.py 第 {tool['lines']} 行")

    if tool.get('svg'):
        add_svg(tool['svg'], 5.5)

page_break()

# ══════════════════════════════════════════════
# 第 9 章：工具↔功能對照表
# ══════════════════════════════════════════════
h1('第九章：工具與功能對照表')
p('這張表讓你一目了然：使用者在前端頁面做的每個操作，對應到 AI 顧問的哪個工具，以及那個工具去呼叫了哪個後端 API。')

add_svg('07_意圖到工具映射.svg', 6.5)
add_svg('17_工具功能對照表.svg', 6.5)

add_table(
    ['使用者功能', '前端頁面', 'AI 工具', '後端端點'],
    [
        ['整合分析\n（DRI 缺口）', 'DriGapPage\n/dri-gap', 'query_dri\n+ get_diet_logs（自動）', 'dri_gap_service\n+ /diet/logs'],
        ['營養素排名', 'NutrientRankingPage\n/nutrient-ranking', 'get_nutrient_ranking', '/health/nutrients/\n{field}/top-foods'],
        ['食物搜尋', 'FoodSearchPage\n/food-search', 'search_food', '/food/search'],
        ['症狀搜尋', 'SymptomSearchPage\n/symptom-search', 'search_symptom', '/health/\nsymptom-search'],
        ['學術文獻', 'GraphRAGPage\n/graphrag', 'search_graphrag', 'GraphRAG:8002\n/retrieve-only'],
        ['飲食記錄', 'DietLogPage\n/diet', 'add/get/update/\ndelete_diet_log', '/diet/logs\n(CRUD)'],
        ['食譜管理', 'SavedRecipesPage\n/recipes', 'get_saved_recipes', '/recipe/list'],
        ['行事曆', 'CalendarPage\n/calendar', 'get_calendar_entries', '/calendar/entries'],
        ['個人資料', 'ProfilePage\n/profile', 'get/update_\nuser_profile', '/profile'],
    ],
    col_widths=[1.3, 1.5, 1.5, 1.5]
)

page_break()

# ══════════════════════════════════════════════
# 第 10 章：狀態機
# ══════════════════════════════════════════════
h1('第十章：工作流程狀態機')
p('前面我們說過「State Machine（狀態機）」就像自動販賣機，有不同的階段。AI 營養顧問的狀態機負責處理一個重要問題：「如果使用者沒有一次給齊所有資訊怎麼辦？」')
p('例如使用者說「營養夠嗎」，系統需要知道年齡和性別才能查 DRI，但使用者沒說。這時候狀態機就會進入「收集資訊」階段，一步一步追問。')

add_svg('06_狀態機流程.svg', 6.0)

h2('10.1 五個狀態')
add_table(
    ['狀態', '英文', '說明', '比喻'],
    [
        ['閒置', 'idle', '等待使用者新訊息', '櫃台空閒，等客人來'],
        ['收集中', 'collecting', '缺少必要參數，追問使用者', '問客人：「請問幾位？」'],
        ['就緒', 'ready', '所有參數齊全，準備執行', '客人都說清楚了，準備上菜'],
        ['執行中', 'executing', '工具正在呼叫中', '廚房正在做菜'],
        ['完成', 'done', '工具執行完畢，生成分析', '菜做好了，端上桌'],
    ],
    col_widths=[0.7, 0.9, 2.0, 2.0]
)

h2('10.2 WorkflowState 資料模型')
p('程式碼位置：advisor.py 第 363-371 行')
code('''class WorkflowState(BaseModel):
    phase: str = "idle"               # 當前階段
    active_intent: str = ""           # 正在處理的意圖
    collected: dict = {}              # 已收集到的參數 {"age": 25, "gender": "female"}
    missing: list[str] = []           # 還缺少的參數 ["activity_level"]
    pending_intents: list[str] = []   # 等待處理的其他意圖（多意圖佇列）
    auto_fetch_done: list[str] = []   # 已自動取過的資料 ["profile"]
    turn_count: int = 0               # 目前在 collecting 第幾輪
    interrupted_from: str = ""        # 被中斷的上一個意圖''')

h2('10.3 參數收集流程')
add_svg('12_參數收集流程.svg', 6.0)

p('每個意圖需要哪些參數，定義在 WORKFLOW_DEFS（第 384-427 行）：')
add_table(
    ['意圖', '必填參數', '可從 Profile 自動填', '追問訊息'],
    [
        ['query_dri', 'age, gender', 'age, gender,\nactivity_level', '「請問你幾歲？」\n「你的性別是？」'],
        ['search_food', 'food_name', '—', '「想查哪種食物？」'],
        ['search_symptom', 'symptom', '—', '「你有什麼症狀？」'],
        ['search_graphrag', 'query', '—', '「想查什麼主題？」'],
        ['nutrient_ranking', 'nutrient_field', '—', '「想查哪種營養素？」'],
        ['update_diet_log', 'log_id', '—', '「要修改哪筆記錄？」'],
        ['delete_diet_log', 'log_id', '—', '「要刪除哪筆記錄？」'],
        ['get_calendar', 'start_date,\nend_date', '—', '「要看哪段時間？」'],
    ],
    col_widths=[1.3, 1.3, 1.3, 2.3]
)

h2('10.4 中斷偵測')
p('如果使用者在 collecting 階段突然轉換話題，系統會偵測到「中斷」：')
p('函式：_detect_interruption()（第 1452-1476 行）')
bullet('不算中斷的：「好」「25歲」「女性」「疲勞」（是在回答追問的問題）')
bullet('算中斷的：「算了」「取消」「不要了」（放棄語）')
bullet('算中斷的：「幫我記錄」「我吃了」（新的動作指令）')
p('偵測到中斷後，系統會把 interrupted_from 記在記憶中，然後重置狀態機到 idle，開始處理新的意圖。')

h2('10.5 多意圖佇列')
p('使用者可能一句話包含多個意圖，例如「我 25 歲女性，幫我看看營養」→ 同時包含 update_profile + query_dri。')
p('系統會把第一個意圖設為 active_intent，其餘放進 pending_intents 佇列。第一個處理完後，自動處理下一個（前提是參數齊全）。')

page_break()

# ══════════════════════════════════════════════
# 第 11 章：LLM 呼叫參數完整清單
# ══════════════════════════════════════════════
h1('第十一章：LLM 呼叫參數完整清單')
p('這章匯總所有 LLM 呼叫的參數，讓你一次看清整個系統在什麼時候呼叫 AI、給它多大的空間、設什麼限制。')

h2('11.1 六次 LLM 呼叫')
add_table(
    ['步驟', 'Prompt', 'Temperature', 'Max Tokens', '嘗試次數', '觸發條件'],
    [
        ['Step 1\nEcho', 'PROMPT_ECHO\n(第 38-78 行)', '0.5', '128\n(≈64 字)', '1', '每輪對話\n(collecting 跳過)'],
        ['Step 2\nIntent', 'PROMPT_INTENT\n(第 85-204 行)', '0.1', '512\n(≈256 字)', '2', '每輪對話\n(collecting 不走)'],
        ['Step 3\nToolCall', 'PROMPT_TOOLCALL\n(第 211-278 行)', '0.1', '512', '2', '僅系統生成\n失敗時啟用'],
        ['Step 4\nAnalysis', 'PROMPT_ANALYSIS\n(第 285-331 行)', '0.3', '2048\n(≈1024 字)', '2', '有工具結果時'],
        ['Memory\n壓縮', 'MEMORY_PROMPT\n(第 338-348 行)', '0.2', '1024', '1', '有結果且\n非 collecting'],
        ['Params\n提取', 'PROMPT_EXTRACT\n(第 1553-1561 行)', '0.1', '128', '1', 'regex 提取\n失敗時備援'],
    ],
    col_widths=[0.8, 1.3, 0.8, 0.8, 0.7, 1.0]
)

h2('11.2 輸入截斷限制')
add_table(
    ['項目', '截斷限制', '程式碼位置', '說明'],
    [
        ['對話歷史', '最近 8 則 × 每則 400 字', '第 495-501 行', '避免 Context Window 超載'],
        ['上輪工具結果', '5000 字', '第 825 行', '供 Intent Detection 參考'],
        ['本輪工具結果', '12000 字', '第 908 行', '供 Analysis 合成回覆'],
        ['記憶輸入 — 使用者訊息', '500 字', '第 1432 行', '壓縮用'],
        ['記憶輸入 — AI 回覆', '800 字', '第 1433 行', '壓縮用'],
        ['記憶輸出', '1500 字', '第 1442 行', '壓縮結果截斷'],
    ],
    col_widths=[1.5, 1.5, 1.0, 2.5]
)

h2('11.3 Timeout 設定')
add_table(
    ['連線對象', 'Timeout', '程式碼位置'],
    [
        ['Gemma LLM', '120 秒（connect 10 秒）', '第 27 行 _GEMMA_TIMEOUT'],
        ['Backend API', '30 秒', '第 28 行 _BACKEND_TIMEOUT'],
        ['GraphRAG API', '60 秒', '第 1194 行'],
    ],
    col_widths=[2.0, 2.5, 2.0]
)

h2('11.4 Gemma LLM 設定')
p('程式碼位置：config.py')
add_table(
    ['設定項', '值', '說明'],
    [
        ['gemma_url', 'http://localhost:8080/v1/chat/completions', 'OpenAI 相容 API 端點'],
        ['gemma_model_name', 'gemma-4', '模型名稱'],
        ['gemma_enabled', 'true', '可停用 LLM 功能'],
        ['Context Window', '4096-5500 tokens', '由 llama-server -c 參數決定'],
    ],
    col_widths=[1.5, 3.0, 2.0]
)

page_break()

# ══════════════════════════════════════════════
# 第 12 章：前端渲染
# ══════════════════════════════════════════════
h1('第十二章：前端渲染機制')
p('前端負責把 SSE 串流事件轉換成使用者看到的畫面。核心元件架構如下：')

add_svg('13_前端渲染架構.svg', 6.0)

h2('12.1 AdvisorPage.jsx — 主頁面')
p('這是 AI 營養顧問的主要頁面，管理所有的聊天狀態：')
add_table(
    ['React State', '型別', '用途'],
    [
        ['messages', 'array', '所有聊天訊息（使用者 + AI），存 localStorage（最多 30 則）'],
        ['input', 'string', '輸入框的文字'],
        ['isLoading', 'boolean', '是否正在等待回應'],
        ['steps', 'array', '工具執行進度（給 ThinkingPanel）'],
        ['isThinking', 'boolean', '是否顯示思考動畫'],
        ['memory', 'string', '壓縮記憶（傳給下一輪）'],
        ['workflowState', 'object', '狀態機快照（傳給下一輪）'],
    ],
    col_widths=[1.3, 0.8, 4.5]
)

h2('12.2 Echo → Reply 替換機制')
p('當 echo 事件到來時，前端會先顯示一個帶 _isEcho: true 標記的暫時氣泡。當 reply 事件到來時：')
bullet('移除所有 _isEcho 訊息')
bullet('加入最終的 AI 回覆訊息（包含 tools_used 和 tool_results）')
p('這樣使用者看到的效果是：先出現「收到...」→ 被替換成完整的分析回覆。')

h2('12.3 ToolResultCard — 13 種結果渲染器')
p('每種工具的結果都有專屬的渲染元件，確保資料以最適合的方式呈現：')
add_table(
    ['工具', '渲染元件', '呈現方式'],
    [
        ['query_dri', 'DriGapTable', '紅/橙/綠達成率表格 + 推薦食物展開列'],
        ['get_nutrient_ranking', 'NutrientRankingTable', '排名表格（名次/食物/含量）'],
        ['search_food', 'SearchFoodCard', '食物卡片（營養五邊形）'],
        ['search_symptom', 'SymptomResultCard', '效果標籤 + 化合物推薦卡片'],
        ['search_graphrag', 'GraphRagResultCard', '論文卡片（IF/SJR/DOI）'],
        ['get_diet_logs', 'DietLogsTable', '完整飲食記錄表格'],
        ['add_diet_log', 'DietLogActionCard', '綠色確認：已記錄 + 營養'],
        ['update_diet_log', 'DietLogActionCard', '藍色確認：已更新'],
        ['delete_diet_log', 'DietLogActionCard', '紅色確認：已刪除'],
        ['get_saved_recipes', 'SavedRecipesCard', '食譜卡片 + 食材標籤'],
        ['get_calendar_entries', 'CalendarEntriesCard', '日期列表'],
        ['get_user_profile', 'UserProfileCard', '個人資料欄位'],
        ['update_user_profile', 'ProfileUpdatedCard', '更新後的資料'],
    ],
    col_widths=[1.5, 1.8, 3.0]
)

page_break()

# ══════════════════════════════════════════════
# 第 13 章：情境序列圖
# ══════════════════════════════════════════════
h1('第十三章：常見情境完整序列')
p('這章用完整的序列圖展示幾個典型使用場景。每個場景都標出了所有的 SSE 事件、HTTP 呼叫和 LLM 呼叫。')

h2('13.1 情境一：記錄飲食')
p('使用者說：「今天午餐吃了雞腿便當和一杯豆漿」')
p('這是最簡單的情境——參數齊全，不需要追問，直接執行。')
add_svg('15_情境_記錄飲食.svg', 6.5)

p('通信步驟解析：')
bullet('1. 使用者送出訊息 → 前端 POST /advisor/chat')
bullet('2. Echo: Gemma 回「好的，我來幫你記錄午餐...」(temp=0.5, max=128)')
bullet('3. Intent: Gemma 分析出 record_diet + 2 項 food_items (temp=0.1, max=512)')
bullet('4. ToolCall: 系統直接生成 2 個 add_diet_log（不需 LLM）')
bullet('5. 執行第一個工具：先 GET /food/search?q=雞腿便當 → 再 POST /diet/logs')
bullet('6. 執行第二個工具：先 GET /food/search?q=豆漿 → 再 POST /diet/logs')
bullet('7. Analysis: Gemma 分析兩項食物的營養 (temp=0.3, max=2048)')
bullet('8. Memory: 壓縮「午餐已記錄雞腿便當500g+豆漿300g」(temp=0.2, max=1024)')
p('本次對話共呼叫 Gemma LLM 4 次。', bold=True)

h2('13.2 情境二：DRI 分析（需要追問）')
p('使用者說：「營養夠嗎」→ 系統需要年齡和性別才能查 DRI。')
add_svg('16_情境_DRI分析.svg', 6.5)

p('通信步驟解析：')
pb('第一輪（需要追問）：')
bullet('1. Echo: 「收到，我先查看你的個人資料...」')
bullet('2. Intent: 偵測到 query_dri，needs_profile=true')
bullet('3. 狀態機：自動呼叫 get_user_profile → 取得 profile')
bullet('4. profile 有 gender 但缺 age → 進入 collecting 階段')
bullet('5. SSE reply:「請問你幾歲呢？」+ workflow_state:{phase:"collecting"}')
bullet('本輪呼叫 Gemma 2 次（Echo + Intent）')
pb('第二輪（提供年齡）：')
bullet('1. 使用者回答「25歲」→ 前端帶 workflow_state:{phase:"collecting"}')
bullet('2. 跳過 Echo（collecting 階段不做 Echo）')
bullet('3. regex 提取：25 → age=25')
bullet('4. missing=[] → phase→ready → 執行 query_dri')
bullet('5. Analysis: 分析 DRI 缺口')
bullet('6. Memory: 壓縮「25歲女性;DRI分析...」')
bullet('本輪呼叫 Gemma 2 次（Analysis + Memory）')
p('兩輪合計呼叫 Gemma LLM 4 次。', bold=True)

page_break()

# ══════════════════════════════════════════════
# 第 14 章：驗證機制
# ══════════════════════════════════════════════
h1('第十四章：工具分發與驗證機制')
p('AI 不是 100% 可靠的——它可能會產生不存在的工具名、缺少必填參數、或者回答太短。系統用三層驗證來防範這些問題。')

h2('14.1 第一層：意圖驗證')
p('函式：_validate_intent_result()（第 560-630 行）')
bullet('白名單檢查：intents 中的每一項都必須在 _VALID_INTENTS 中（14 個合法意圖）')
bullet('衝突消除：如果同時有 general_chat 和其他意圖，移除 general_chat')
bullet('症狀翻譯：用 _SYMPTOM_ZH_TO_EN 查表，把中文症狀轉成英文')
bullet('營養素轉換：用 _NUTRIENT_ZH_TO_FIELD 查表，把「鈣」轉成 calcium_per_100g')
bullet('份量填充：用 _PORTION_DEFAULTS 表，給沒有份量的食物加預設值')

h2('14.2 第二層：工具呼叫驗證')
p('函式：_validate_toolcalls()（第 633-695 行）')
bullet('工具名檢查：必須在 _TOOL_DISPATCH 的 13 個合法工具中')
bullet('必填參數檢查：根據 _TOOL_REQUIRED_ARGS 確認必填欄位')
bullet('自動填入：缺少的參數嘗試從 entities 或 profile 補上')
bullet('型別轉換：_coerce_type() 把字串轉數字（"25" → 25）')
bullet('去重：移除完全重複的工具呼叫')

h2('14.3 第三層：分析驗證')
p('函式：_validate_analysis()（第 700-707 行）')
bullet('最短長度：回覆少於 50 字會觸發重試')
bullet('重試策略：第二次嘗試會在 Prompt 後加「請用更詳細的方式分析」')

page_break()

# ══════════════════════════════════════════════
# 第 15 章：總結
# ══════════════════════════════════════════════
h1('第十五章：總結')
p('這份報告完整拆解了 AI 營養顧問的 Tool Calling 機制。以下是關鍵數字的總結：')

add_table(
    ['項目', '數量/內容'],
    [
        ['工具數量', '13 個（涵蓋飲食、營養、症狀、文獻、個人資料）'],
        ['Prompt 數量', '6 個（Echo / Intent / ToolCall / Analysis / Memory / Extract）'],
        ['LLM Pipeline 步驟', '4 步（Echo → Intent → ToolCall → Analysis）'],
        ['SSE 事件類型', '9 種（thinking / echo / tool_start / result / done / reply / memory / workflow_state / error）'],
        ['狀態機狀態', '5 個（idle → collecting → ready → executing → done）'],
        ['驗證層', '3 層（Intent / ToolCall / Analysis）'],
        ['服務數量', '5 個（Frontend / Control / Backend / GraphRAG / Gemma）'],
        ['程式碼總行數', '1950 行（advisor.py）'],
        ['一次簡單對話的 LLM 呼叫', '4 次（Echo + Intent + Analysis + Memory）'],
        ['一次需追問的對話', '4-6 次（多一輪收集 + 可能的 LLM 參數提取）'],
    ],
    col_widths=[2.5, 4.0]
)

h2('15.1 設計哲學')
bullet('系統優先，LLM 備援：能用程式碼解決的就不問 AI（ToolCall 的系統生成、參數的 regex 提取）')
bullet('串流回饋：用 SSE 讓使用者即時看到進度，不用乾等')
bullet('多輪收集：不強迫使用者一次說完所有資訊')
bullet('三層驗證：不盲信 LLM 的輸出，每一步都做驗證和修正')
bullet('記憶壓縮：在有限的 Context Window 中保留最重要的對話脈絡')

h2('15.2 完整通信鏈總覽')
p('從使用者按 Enter 到看到回覆，資料經過的完整路徑：')
code('''使用者瀏覽器
  → fetch POST /api/advisor/chat
    → nginx proxy (port 3001 → 8000)
      → Control FastAPI advisor_chat()
        → Gemma LLM (port 8080) × 4-6 次
        → Backend API (port 8001) × N 次（視工具數量）
        → GraphRAG API (port 8002) × 0-1 次（僅學術查詢）
      ← SSE stream: thinking → echo → tool_start → tool_result
                     → tool_done → reply → memory → workflow_state
    ← 前端逐一處理 SSE 事件
  ← 畫面更新：工具卡片 + 分析文字 + 記憶儲存''')

# ─── 儲存 ───
doc.save(str(OUT))
print(f"✅ 報告已生成: {OUT}")
print(f"   檔案大小: {OUT.stat().st_size / 1024:.0f} KB")
